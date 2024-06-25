from collections import defaultdict
import textwrap
from flask import session,Flask, render_template, request, redirect, url_for, session, flash,jsonify, make_response,Response,send_from_directory
from flask_cors import CORS
import pandas as pd
from datetime import datetime, date, timedelta
import calendar
import zipfile
import io
import re
import json
import numpy as np
import bcrypt
import os
import secrets
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
# from bson import ObjectId
import uuid 
from datetime import datetime, timedelta
from collections.abc import Iterable
import random
import warnings
from tabulate import tabulate
from itertools import product
from flask import send_file
import tempfile
import csv    
from random import sample   
import time
import threading  
import shutil
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Table, Spacer, PageBreak,TableStyle
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph,Image
import urllib
import schedule
import base64
from PIL import Image
from flask_cors import CORS, cross_origin
from functools import wraps
from flask_simple_captcha import CAPTCHA
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Spacer, PageBreak, Table
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from io import BytesIO
import textwrap
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, PageBreak, Image, PageTemplate, Frame
from reportlab.platypus.doctemplate import PageTemplate, BaseDocTemplate, Frame
from ldap3 import Server, Connection, ALL,SUBTREE
import pyodbc
import traceback
from docx import Document
from io import StringIO
import paramiko
import chardet
from werkzeug.utils import secure_filename
from dateutil.parser import parse as parse_datetime, UnknownTimezoneWarning
from fuzzywuzzy import fuzz
import operator
from collections import defaultdict
import regex as rx
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import time
from datetime import datetime






# configure email sending
smtp_server = "smtp.gmail.com"
smtp_port = 587  # Typically 587 for TLS
sender_email = "sruthi.k@pinacalabs.com"
sender_password = "mjnpmuzudfpjuybc"


app = Flask(__name__)
cors = CORS(app)

# this is merged data
# serverIP = '10.40.16.189'
# userSQL = 'OGBBANK'
# pwdSQL = 'root12345'



#this is user and scenarios DB
serverIP2 = '10.40.16.189'
userSQL2 = 'OGBBANK'
pwdSQL2 = 'root12345'



# # Establish connection to the SQL Server database
# # this is merged data
# conn = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP};Database=ticketid;UID={userSQL};PWD={pwdSQL}")
# print("DB CONNECTED Merged" ,conn)


#this is user and scenarios DB
# conn2 = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
conn2 = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
print("DB CONNECTED Users and Scenarios" ,conn2)

mysql2 = conn2.cursor()  # this is merged data


# print("DB CONNECTED" ,conn2)


# mysql2 = conn2.cursor() #this is user and scenarios DB

# # =============== Kamal Data Connetion Starts ==============

# server = 'tcp:Charan\\MSSQLSERVER04,49172'
# database = 'Kamal_upload'
# username = 'triveni_amlcft'
# password = 'Triv@amlcft1234'
# conn_str = f'DRIVER={{SQL Server}};SERVER={server};DATABASE={database};UID={username};PWD={password}'
conn_str = 'Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=kamal_upload; MARS_Connection=YES'

# Function to establish a database connection   
def get_db_connection():
    sql_conn = None
    try:
        sql_conn = pyodbc.connect(conn_str)
        print("Connection to SQL Server established successfully.")
    except Exception as e:
        print("Error connecting to SQL Server:", e)
    
    return sql_conn

# SQL query to retrieve table names
conn_kamal = get_db_connection()

# Initialize tables variable to an empty list


# Check if the connection was successful before proceeding
if conn_kamal:
    kamal_tables = []
    try:
        tables_query = "SELECT TABLE_NAME FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_TYPE = 'BASE TABLE'"
        cursor = conn_kamal.cursor()
        kamal_tables = cursor.execute(tables_query).fetchall()
        print("Tables in database:")
        for table_info in kamal_tables:
            print(table_info.TABLE_NAME)
    except Exception as e:
        print("Error executing query:", e)

else:
    print("Failed to establish a database connection.")

# # =============== Kamal Data Connetion Ends ==============






# =================== Active Directotry ================================

AD_SERVER = 'ldap://172.32.15.9'

For_AD_USERNAME_update = 'pinaca.pavan@ogb.com'
For_AD_PASSWORD_update = '9701004594Pa@'

AD_BASE_DN = 'OU=AML,DC=ogb,DC=com'



def SSO(AD_USERNAME,AD_PASSWORD):
    try:
        server = Server(AD_SERVER)
        conn = Connection(server, user=AD_USERNAME, password=AD_PASSWORD,auto_bind=True) 
        print("conn status : ",conn)

        status = conn.result['result']
        if status == 0:


            conn.search(AD_BASE_DN,'(objectClass=*)',attributes=['member','CN'])

            users=[]


            for doc in conn.entries:
                role = str(doc['cn'])
                userDetailsAD = list(doc['member'])



                if userDetailsAD:
                    for Aduser in userDetailsAD:

                        conn.search(Aduser,'(objectClass=*)',attributes=['description','company','countryCode','displayName','distinguishedName','mail','mobile','physicalDeliveryOfficeName','primaryGroupID','sAMAccountName','sAMAccountType','telephoneNumber','title','userAccountControl','userPrincipalName'])
                        
                        
                        for user_entry in conn.entries:
                            user_data = {'role': role}
                            entry_attributes = user_entry.entry_attributes_as_dict
                            for attribute, values in entry_attributes.items():
                                if len(values) > 0:
                                    user_data[attribute] = values[0]
                                else:
                                    user_data[attribute] = None
                            
                            users.append(user_data)
            

            conn.unbind()

            
            userPrincipalName = []
            for doc in users:
                userPrincipalName.append(doc['userPrincipalName'])

            
            if any(AD_USERNAME.lower() == username.lower() for username in userPrincipalName):
                return AD_USERNAME
        else:
            return "locked"
    except:
        return "invaliedCred"
    




def detltaADupdate(USERNAME,PASSWORD):
    serverIPAD = '10.40.16.189'
    userSQLAD = 'OGBBANK'
    pwdSQLAD = 'root12345'

    #this is user and scenarios DB
    connAD = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIPAD};Database=ticketid;UID={userSQLAD};PWD={pwdSQLAD}")
    print("DB CONNECTED for AD User Update" ,connAD)

    # Create a cursor object
    mysqlAD = connAD.cursor()


    server = Server(AD_SERVER)
  
    conn = Connection(server, user=USERNAME, password=PASSWORD,auto_bind=True) 
    

    conn.search(AD_BASE_DN,'(objectClass=*)',attributes=['member','CN'])

    users=[]


    for doc in conn.entries:
        role = str(doc['cn'])
        userDetailsAD = list(doc['member'])


        if userDetailsAD:
            for Aduser in userDetailsAD:

                conn.search(Aduser,'(objectClass=*)',attributes=['description','company','countryCode','displayName','distinguishedName','mail','mobile','physicalDeliveryOfficeName','primaryGroupID','sAMAccountName','sAMAccountType','telephoneNumber','title','userAccountControl','userPrincipalName'])
                # conn.search(Aduser,'(whenChanged >= {})'.format(last_changed_formatted),attributes=['company','countryCode','displayName','distinguishedName','mail','mobile','physicalDeliveryOfficeName','primaryGroupID','sAMAccountName','sAMAccountType','telephoneNumber','title','userAccountControl','userPrincipalName'])
                
                for user_entry in conn.entries:
                    user_data = {'role': role}
                    entry_attributes = user_entry.entry_attributes_as_dict
                    for attribute, values in entry_attributes.items():
                        if len(values) > 0:
                                user_data[attribute] = values[0]
                        else:
                                user_data[attribute] = None
                    
                    users.append(user_data)

    conn.unbind()


    curAD = mysqlAD.connection.cursor()
        
    curAD.execute("SELECT 1 FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'user' ")
    table_exists = curAD.fetchone()

    if not table_exists:
        create_table_sql = """
        CREATE TABLE [user] (
            id INT IDENTITY(1,1) PRIMARY KEY,
            EmpId NVARCHAR(500),
            UserName NVARCHAR(500),
            EmailId NVARCHAR(500),
            MobileNo NVARCHAR(500),
            Address NVARCHAR(500),
            Role NVARCHAR(500),
            Image NVARCHAR(500),
            LeaveStatus NVARCHAR(500),
            Status NVARCHAR(500),
            [File] NVARCHAR(500),
            Alerts_generated NVARCHAR(500),
            Assigned_to NVARCHAR(500),
            telephoneNumber NVARCHAR(500),
            userPrincipalName NVARCHAR(500),
            distinguishedName NVARCHAR(500),
            sAMAccountType NVARCHAR(500),
            physicalDeliveryOfficeName NVARCHAR(500),
            title NVARCHAR(500),
            primaryGroupID NVARCHAR(500),
            userAccountControl NVARCHAR(500)
            )
        """
        curAD.execute(create_table_sql)
        mysqlAD.connection.commit()

    

    removedUsers = []
    for doc in users:
        curAD.execute("SELECT userPrincipalName FROM [user] WHERE userPrincipalName = ?", (doc['userPrincipalName'],))
        existing_email = curAD.fetchone()
        removedUsers.append(doc['userPrincipalName'])

        if existing_email:
            update_data_sql = f"""
            UPDATE [user]
            SET
            EmpId = ?,
            UserName = ?,
            MobileNo = ?,
            Address = ?,
            Role = ?,
            Image = ?,
            LeaveStatus = ?,
            Status = ?,
            [File] = ?,
            telephoneNumber = ?,
            userPrincipalName = ?,
            distinguishedName = ?,
            sAMAccountType = ?,
            physicalDeliveryOfficeName = ?,
            title = ?,
            primaryGroupID = ?,
            userAccountControl = ?
            WHERE userPrincipalName = ?
            """
            if doc['role'] == 'AML MLRO':
                definedRole = 'MLRO'
            elif doc['role'] == 'AML CM':
                definedRole = 'CM/SM'
            elif doc['role'] == 'AML GM':
                definedRole = 'DGM/PO'
            elif doc['role'] == 'AML Admin':
                definedRole = 'ADMIN'
            elif doc['role'] == 'AML HO Admin':
                definedRole = 'IT OFFICER'
            elif doc['role'] == 'AML Branch Maker':
                definedRole = 'BranchMakers'
            elif doc['role'] == 'AML Branch Checker':
                definedRole = 'ROS'
            elif doc['role'] == 'AML PINACA_ADMIN':
                definedRole = 'PINACA_ADMIN'
            elif doc['role'] == 'SDN/USER':
                definedRole = 'SDN/USER'
            values = (doc['description'],doc['displayName'],doc['mobile'],None,definedRole,None,'Working','Approved',None,doc['telephoneNumber'],doc['userPrincipalName'],doc['distinguishedName'],doc['sAMAccountType'],doc['physicalDeliveryOfficeName'],doc['title'],doc['primaryGroupID'],doc['userAccountControl'],doc['userPrincipalName'])
            curAD.execute(update_data_sql, values)
            mysqlAD.connection.commit()

        else:
            insert_data_sql = """
            INSERT INTO [user] (
            
                EmpId,
                UserName,
                EmailId,
                MobileNo,
                Address,
                Role,
                Image,
                LeaveStatus,
                Status,
                [File],
                Alerts_generated,
                Assigned_to,
                telephoneNumber,
                userPrincipalName,
                distinguishedName,
                sAMAccountType,
                physicalDeliveryOfficeName,
                title,
                primaryGroupID,
                userAccountControl
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """
            if doc['role'] == 'AML MLRO':
                definedRole = 'MLRO'
            elif doc['role'] == 'AML CM':
                definedRole = 'CM/SM'
            elif doc['role'] == 'AML GM':
                definedRole = 'DGM/PO'
            elif doc['role'] == 'AML Admin':
                definedRole = 'ADMIN'
            elif doc['role'] == 'AML HO Admin':
                definedRole = 'IT OFFICER'
            elif doc['role'] == 'AML Branch Maker':
                definedRole = 'BranchMakers'
            elif doc['role'] == 'AML Branch Checker':
                definedRole = 'ROS'
            elif doc['role'] == 'AML PINACA_ADMIN':
                definedRole = 'PINACA_ADMIN'
            elif doc['role'] == 'SDN/USER':
                definedRole = 'SDN/USER'
            values = [doc['description'],doc['displayName'],doc['mail'],doc['mobile'],None,definedRole,None,'Working','Approved',None,None,None,doc['telephoneNumber'],doc['userPrincipalName'],doc['distinguishedName'],doc['sAMAccountType'],doc['physicalDeliveryOfficeName'],doc['title'],doc['primaryGroupID'],doc['userAccountControl']]
            curAD.execute(insert_data_sql, values)
            mysqlAD.connection.commit()
    

    removed_users_str = ', '.join([f"'{user}'" for user in removedUsers])

    removeSQL = f"DELETE FROM [user] WHERE userPrincipalName NOT IN ({removed_users_str})"

    curAD.execute(removeSQL)
    mysqlAD.connection.commit()
    mysqlAD.close()

# ======================================================================



secret_key = secrets.token_hex(16)
app.secret_key = secret_key
YOUR_CONFIG = {
    'SECRET_CAPTCHA_KEY': secret_key,
    'CAPTCHA_LENGTH': 6,
    'CAPTCHA_DIGITS': False,
    'EXPIRE_SECONDS': 600,
}
SIMPLE_CAPTCHA = CAPTCHA(config=YOUR_CONFIG)
app = SIMPLE_CAPTCHA.init_app(app)
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_FILE_DIR'] = '/temp/flask_session'

app.permanent_session_lifetime = timedelta(minutes=30)
# session(app)


# @app.after_request
# def add_security_headers(response):
#     csp_directives = {
#         'default-src': "'self' data:",
#         'script-src': "'self' 'unsafe-inline' https://cdn.jsdelivr.net https://cdn.datatables.net https://ajax.googleapis.com https://cdnjs.cloudflare.com https://gyrocode.github.io https://maxcdn.bootstrapcdn.com https://unpkg.com https://d3js.org",
#         'style-src': "'self' 'unsafe-inline' https://cdn.jsdelivr.net https://fonts.googleapis.com https://maxcdn.bootstrapcdn.com https://cdn.datatables.net https://gyrocode.github.io https://stackpath.bootstrapcdn.com",
#         'font-src': "'self' https://fonts.gstatic.com https://stackpath.bootstrapcdn.com https://cdn.jsdelivr.net",
#         'img-src': "'self' data: https://cdn.datatables.net",
#         'frame-ancestors': "'none'"
#     }
#     csp_header = "; ".join([f"{directive} {csp_directives[directive]}" for directive in csp_directives])
#     response.headers['Content-Security-Policy'] = csp_header
#     return response




# ==================== Email Part ========================

def send_email(receiver_email, subject, body):
    # configure email sending
    smtp_server = "smtp.gmail.com"
    smtp_port = 587  # Typically 587 for TLS
    sender_email = "sruthi.k@pinacalabs.com"
    password = "mjnpmuzudfpjuybc"
    # Create a MIME multipart message
    message = MIMEMultipart()
    message['From'] = sender_email
    message['To'] = receiver_email
    message['Subject'] = subject
    # Attach the email body
    message.attach(MIMEText(body, 'plain'))
    # Create a SMTP session
    with smtplib.SMTP(smtp_server, smtp_port) as server:
        server.starttls()
        server.login(sender_email, password)
        text = message.as_string()
        server.sendmail(sender_email, receiver_email, text)

# ========================================================
# ================================================= Socket IO Usage =====================================================================

# def notification(mails):

   

#      userinfo = users_collection.find_one({'emailid':mails})

#     # =========== Taking the array's of the tickets which are storing the cases ========================== 
     
#      if "allocated_tickets" in userinfo:
#         pendingAlertsNotfy = len(userinfo['allocated_tickets'])
#      else:
#          pendingAlertsNotfy = 0
     
#      if "Sent_Back_Case_Alerts" in userinfo:
#         unsuffeccientAlerst = len(userinfo['Sent_Back_Case_Alerts'])
#      else:
#          unsuffeccientAlerst = 0
     
#      if "rised_closed_tickets" in userinfo:
#         rised_closed_tickets = len(userinfo['rised_closed_tickets'])
#      else:
#          rised_closed_tickets = 0
     
#      if "Sent_Back_Alerts" in userinfo:
#         Sent_Back_Alerts = len(userinfo['Sent_Back_Alerts'])
#      else:
#          Sent_Back_Alerts = 0
     
#      if "Offline_assigned_tickets" in userinfo:
#         offlineCases = len(userinfo['Offline_assigned_tickets'])
#      else:
#          offlineCases = 0

# # ============================
#      if userinfo['role'] == 'CM/SM':
#         users_collection.update_one({'emailid':mails},{"$set":{'pendingAlertsNotfy':pendingAlertsNotfy,'unsufecientAlertsNotfy':unsuffeccientAlerst,"risedClosedCount":rised_closed_tickets}})
#      elif userinfo['role'] == 'MLRO':
#         users_collection.update_one({'emailid':mails},{"$set":{'pendingAlertsNotfy':pendingAlertsNotfy,'unsufecientAlertsNotfy':unsuffeccientAlerst,"sentBackClosedCount":Sent_Back_Alerts}})  
#      else:
#         users_collection.update_one({'emailid':mails},{"$set":{'pendingAlertsNotfy':pendingAlertsNotfy,'unsufecientAlertsNotfy':unsuffeccientAlerst,"offlineCasesCount":offlineCases}})
    
#     # ============================
#      if "pendingAlertsNotfy" in userinfo:
#         prvPendingAlertsNotfy = userinfo['pendingAlertsNotfy']
#      else:
#         prvPendingAlertsNotfy = 0  

#      if "unsufecientAlertsNotfy" in userinfo:
#         prvunsufecientAlertsNotfy = userinfo['unsufecientAlertsNotfy']
#      else:
#          prvunsufecientAlertsNotfy = 0
     
#      if "risedClosedCount" in userinfo:
#         prvrisedClosedCount = userinfo['risedClosedCount']
#      else:
#          prvrisedClosedCount = 0
     
#      if "sentBackClosedCount" in userinfo:
#         prvsentClosedCount = userinfo['sentBackClosedCount']
#      else:
#          prvsentClosedCount = 0
     
#      if "offlineCasesCount" in userinfo:
#         prvofflineCount = userinfo['offlineCasesCount']
#      else:
#          prvofflineCount = 0

# # ===================================================
         
#      mainInfo = users_collection.find_one({'emailid':mails})

#      presentpen = mainInfo['pendingAlertsNotfy']
#      if "prvPendingAlertsNotfy" in mainInfo:
#         prevpen = mainInfo['prvPendingAlertsNotfy']
#      else:
#          prevpen = 0
#          users_collection.update_one({'emailid':mails},{"$set":{'prvPendingAlertsNotfy':prevpen}})

#      presentuns = mainInfo['unsufecientAlertsNotfy']
#      if "prvunsufecientAlertsNotfy" in mainInfo:
#         prvuns = mainInfo['prvunsufecientAlertsNotfy']
#      else:
#          prvuns = 0
#          users_collection.update_one({'emailid':mails},{"$set":{'prvunsufecientAlertsNotfy':prvuns}})
    
#      if "risedClosedCount" in mainInfo:
#         presentclosed = mainInfo['risedClosedCount']
#         if "prvrisedClosedCount" in mainInfo:
#             prvclosed = mainInfo['prvrisedClosedCount']
#         else:
#             prvclosed = 0
#             users_collection.update_one({'emailid':mails},{"$set":{'prvrisedClosedCount':prvclosed}})
     
#      if "sentBackClosedCount" in mainInfo:
#         presentsentclosed = mainInfo['sentBackClosedCount']
#         if "prvsentClosedCount" in mainInfo:
#             prvsentclosed = mainInfo['prvsentClosedCount']
#         else:
#             prvsentclosed = 0
#             users_collection.update_one({'emailid':mails},{"$set":{'prvsentClosedCount':prvsentclosed}})
     
#      if "offlineCasesCount" in mainInfo:
#         presentofflineclosed = mainInfo['offlineCasesCount']
#         if "prvofflineCount" in mainInfo:
#             prvofflineclosed = mainInfo['prvofflineCount']
#         else:
#             prvofflineclosed = 0
#             users_collection.update_one({'emailid':mails},{"$set":{'prvofflineCount':prvofflineclosed}})

#     #  ================================
#      print("prevpen : ",prevpen)
#      print("presentpen : ",presentpen)
     
#      if prevpen > presentpen:
#          users_collection.update_one({'emailid':mails},{"$set":{'prvPendingAlertsNotfy':prvPendingAlertsNotfy}})
#      if prvuns > presentuns:
#          users_collection.update_one({'emailid':mails},{"$set":{'prvunsufecientAlertsNotfy':prvunsufecientAlertsNotfy}})
     
#      if userinfo['role'] == 'AGM' or userinfo['role'] == 'DGM/PO' or userinfo['role'] == 'ROS':
#         if prvofflineclosed > presentofflineclosed:
#             users_collection.update_one({'emailid':mails},{"$set":{'prvofflineCount':prvofflineCount}})
#      if userinfo['role'] == 'CM/SM':
#         if prvclosed > presentclosed:
#             users_collection.update_one({'emailid':mails},{"$set":{'prvrisedClosedCount':prvrisedClosedCount}})
#      if userinfo['role'] == 'MLRO':
#         if prvsentclosed > presentsentclosed:
#             users_collection.update_one({'emailid':mails},{"$set":{'prvsentClosedCount':prvsentClosedCount}})
     
     
# # ===================================================


#      main = users_collection.find_one({'emailid':mails})
     
#      if "pendingAlertsNotfy" in main and "prvPendingAlertsNotfy" in main:
#         presentPendingAlertsNotfy  = main['pendingAlertsNotfy']
#         prvPendingAlertsNotfy  = main['prvPendingAlertsNotfy']
#      else:
#         presentPendingAlertsNotfy  = 0
#         prvPendingAlertsNotfy  = 0

#      print("presentPendingAlertsNotfy : ",presentPendingAlertsNotfy)
#      print("prvPendingAlertsNotfy : ",prvPendingAlertsNotfy)

#      if "unsufecientAlertsNotfy" in main and "prvunsufecientAlertsNotfy" in main:
#         presentunsufecientAlertsNotfy  = main['unsufecientAlertsNotfy']
#         prvunsufecientAlertsNotfy  = main['prvunsufecientAlertsNotfy']
#      else:
#         presentunsufecientAlertsNotfy  = 0
#         prvunsufecientAlertsNotfy  = 0
     
#      if "risedClosedCount" in main and "prvrisedClosedCount" in main:
#         presentclosedNotfy  = main['risedClosedCount']
#         prvclosedNotfy  = main['prvrisedClosedCount']
#      else:
#         presentclosedNotfy  = 0
#         prvclosedNotfy  = 0
     
#      if "sentBackClosedCount" in main and "prvsentClosedCount" in main:
#         presentsentclosedNotfy  = main['sentBackClosedCount']
#         prvsentclosedNotfy  = main['prvsentClosedCount']
#      else:
#         presentsentclosedNotfy  = 0
#         prvsentclosedNotfy  = 0
     
#      if "offlineCasesCount" in main and "prvofflineCount" in main:
#         presentofflineNotfy  = main['offlineCasesCount']
#         prvofflineNotfy  = main['prvofflineCount']
#      else:
#         presentofflineNotfy  = 0
#         prvofflineNotfy  = 0


# # ===================================================
        
#      if presentPendingAlertsNotfy and (presentPendingAlertsNotfy > prvPendingAlertsNotfy):
#          pendingcount = presentPendingAlertsNotfy - prvPendingAlertsNotfy
#      else:
#          pendingcount = 0
#      print("pendingcount : ",pendingcount)    
    
#      if presentunsufecientAlertsNotfy and (presentunsufecientAlertsNotfy > prvunsufecientAlertsNotfy):
#          unsuffeccientcount = presentunsufecientAlertsNotfy - prvunsufecientAlertsNotfy
#      else:
#          unsuffeccientcount = 0
     
#      if presentclosedNotfy and (presentclosedNotfy > prvclosedNotfy):
#          verifyClosedcount = presentclosedNotfy - prvclosedNotfy
#      else:
#          verifyClosedcount = 0
    
#      if presentsentclosedNotfy and (presentsentclosedNotfy > prvsentclosedNotfy):
#          sentClosedcount = presentsentclosedNotfy - prvsentclosedNotfy
#      else:
#          sentClosedcount = 0
     
#      if presentofflineNotfy and (presentofflineNotfy > prvofflineNotfy):
#          risedofflinecount = presentofflineNotfy - prvofflineNotfy
#      else:
#          risedofflinecount = 0

         
# ===============================================================================

# ============================ DashBoard Per Day Logic ===========================================

    #  current_datetime = datetime.now()
     # Extract only the date and set the time to midnight
    #  current_date = str(current_datetime.date())
    #  current_date = '2024-01-23'
    #  print('current_date : ',current_date)

    #  if pendingcount != 0:

    #                 perDay = users_collection.find_one({"emailid":mails,"pendingAlerts_perDay":{"$exists":True}})

    #                 if perDay:
    #                     print("holaa.................")
    #                     dateExists = users_collection.find_one({"emailid": mails, f'pendingAlerts_perDay.{current_date}': {"$exists": True}})
    #                     if dateExists:
    #                         users_collection.update_one(
    #                                 {"emailid": mails, f'pendingAlerts_perDay.{current_date}': {"$exists": True}},
    #                                 {"$inc": {f'pendingAlerts_perDay.$.{current_date}': pendingcount}}
    #                             )
    #                     else:
    #                         users_collection.update_one({"emailid":mails},{"$push":{"pendingAlerts_perDay":{current_date:pendingcount}}})

    #                 else:
    #                     users_collection.update_one({"emailid":mails},{"$set":{"pendingAlerts_perDay":[{current_date:pendingcount}]}})

   

# ================================================================================================

#      return {"pendingcount":pendingcount,"unsuffeccientcount":unsuffeccientcount,"verifyClosedcount":verifyClosedcount,'sentBackClosed':sentClosedcount,"offlineCases":risedofflinecount}

# # ===============Clearing the notification function ===========================

# def clearnotification(mails,endpoint):
#      userinfo = users_collection.find_one({'emailid':mails})

# # ====================================
  

     
#      if "allocated_tickets" in userinfo:
#         pendingAlertsNotfy = len(userinfo['allocated_tickets'])
#      else:
#          pendingAlertsNotfy = 0
     
#      if "Sent_Back_Case_Alerts" in userinfo:
#         unsuffeccientAlerst = len(userinfo['Sent_Back_Case_Alerts'])
#      else:
#          unsuffeccientAlerst = 0
     
#      if "rised_closed_tickets" in userinfo:
#         rised_closed_tickets = len(userinfo['rised_closed_tickets'])
#      else:
#          rised_closed_tickets = 0
     
#      if "Sent_Back_Alerts" in userinfo:
#         Sent_Back_Alerts = len(userinfo['Sent_Back_Alerts'])
#      else:
#          Sent_Back_Alerts = 0
     
#      if "Offline_assigned_tickets" in userinfo:
#         offlineCases = len(userinfo['Offline_assigned_tickets'])
#      else:
#          offlineCases = 0
    
# # =====================================================
         
#      if endpoint == 'nextLevel':
#         users_collection.update_one({'emailid':mails},{"$set":{'pendingAlertsNotfy':pendingAlertsNotfy}})
#      if endpoint == 'sentBack':
#         users_collection.update_one({'emailid':mails},{"$set":{'unsufecientAlertsNotfy':unsuffeccientAlerst}})
#      if endpoint == 'closedVerify':
#         users_collection.update_one({'emailid':mails},{"$set":{'risedClosedCount':rised_closed_tickets}})
#      if endpoint == 'sentBackClosed':
#         users_collection.update_one({'emailid':mails},{"$set":{'sentBackClosedCount':Sent_Back_Alerts}})
#      if endpoint == 'offline':
#         users_collection.update_one({'emailid':mails},{"$set":{'offlineCasesCount':offlineCases}})
     

#     # ==========================================================
        
    
#      if "pendingAlertsNotfy" in userinfo:
#         prvPendingAlertsNotfy = userinfo['pendingAlertsNotfy']
#      else:
#         prvPendingAlertsNotfy = 0  

#      if "unsufecientAlertsNotfy" in userinfo:
#         prvunsufecientAlertsNotfy = userinfo['unsufecientAlertsNotfy']
#      else:
#          prvunsufecientAlertsNotfy = 0
     
#      if "risedClosedCount" in userinfo:
#         prvclosedNotfy = userinfo['risedClosedCount']
#      else:
#          prvclosedNotfy = 0
     
#      if "sentBackClosedCount" in userinfo:
#         prvsentclosedNotfy = userinfo['sentBackClosedCount']
#      else:
#          prvsentclosedNotfy = 0
     
#      if "offlineCasesCount" in userinfo:
#         prvofflineNotfy = userinfo['offlineCasesCount']
#      else:
#          prvofflineNotfy = 0

# # ==============================================
         
#      if endpoint == 'nextLevel':
#         users_collection.update_one({'emailid':mails},{"$set":{'prvPendingAlertsNotfy':prvPendingAlertsNotfy}})
    
#      if endpoint == 'sentBack':
#         users_collection.update_one({'emailid':mails},{"$set":{'prvunsufecientAlertsNotfy':prvunsufecientAlertsNotfy}})
    
#      if endpoint == 'closedVerify':
#         users_collection.update_one({'emailid':mails},{"$set":{'prvrisedClosedCount':prvclosedNotfy}})
    
#      if endpoint == 'sentBackClosed':
#         users_collection.update_one({'emailid':mails},{"$set":{'prvsentClosedCount':prvsentclosedNotfy}})
    
#      if endpoint == 'offline':
#         users_collection.update_one({'emailid':mails},{"$set":{'prvofflineCount':prvofflineNotfy}})
    
     
         
# # ====================================================================



#      main = users_collection.find_one({'emailid':mails})
     
#      if "pendingAlertsNotfy" in main and "prvPendingAlertsNotfy" in main:
#         presentPendingAlertsNotfy  = main['pendingAlertsNotfy']
#         prvPendingAlertsNotfy  = main['prvPendingAlertsNotfy']
#      else:
#         presentPendingAlertsNotfy  = 0
#         prvPendingAlertsNotfy  = 0

#      if "unsufecientAlertsNotfy" in main and "prvunsufecientAlertsNotfy" in main:
#         presentunsufecientAlertsNotfy  = main['unsufecientAlertsNotfy']
#         prvunsufecientAlertsNotfy  = main['prvunsufecientAlertsNotfy']
#      else:
#         presentunsufecientAlertsNotfy  = 0
#         prvunsufecientAlertsNotfy  = 0

#      if "risedClosedCount" in main and "prvrisedClosedCount" in main:
#         presentclosedNotfy  = main['risedClosedCount']
#         prvclosedNotfy  = main['prvrisedClosedCount']
#      else:
#         presentclosedNotfy  = 0
#         prvclosedNotfy  = 0
     
#      if "sentBackClosedCount" in main and "prvsentClosedCount" in main:
#         presentsentclosedNotfy  = main['sentBackClosedCount']
#         prvsentclosedNotfy  = main['prvsentClosedCount']
#      else:
#         presentsentclosedNotfy  = 0
#         prvsentclosedNotfy  = 0
    
#      if "offlineCasesCount" in main and "prvofflineCount" in main:
#         presentofflineNotfy  = main['offlineCasesCount']
#         prvofflineNotfy  = main['prvofflineCount']
#      else:
#         presentofflineNotfy  = 0
#         prvofflineNotfy  = 0

# # =====================================================
     
#      if presentPendingAlertsNotfy and (presentPendingAlertsNotfy > prvPendingAlertsNotfy):
#          pendingcount = presentPendingAlertsNotfy - prvPendingAlertsNotfy
#      else:
#          pendingcount = 0
    
#      if presentunsufecientAlertsNotfy and (presentunsufecientAlertsNotfy > prvunsufecientAlertsNotfy):
#          unsuffeccientcount = presentunsufecientAlertsNotfy - prvunsufecientAlertsNotfy
#      else:
#          unsuffeccientcount = 0

#      if presentclosedNotfy and (presentclosedNotfy > prvclosedNotfy):
#          verifyClosedcount = presentclosedNotfy - prvclosedNotfy
#      else:
#          verifyClosedcount = 0
     
#      if presentsentclosedNotfy and (presentsentclosedNotfy > prvsentclosedNotfy):
#          sentClosedcount = presentsentclosedNotfy - prvsentclosedNotfy
#      else:
#          sentClosedcount = 0
     
#      if presentofflineNotfy and (presentofflineNotfy > prvofflineNotfy):
#          risedofflinecount = presentofflineNotfy - prvofflineNotfy
#      else:
#          risedofflinecount = 0


#     # ============================

#      return {"pendingcount":pendingcount,"unsuffeccientcount":unsuffeccientcount,"verifyClosedcount":verifyClosedcount,'sentBackClosed':sentClosedcount,"offlineCases":risedofflinecount}


# ============================ END OF THE NOTIFICATION FUNTION'S BOTH GETTING AND CLEARING===============================================

def fetch_current_values(query,code,Threshold_Type):
    with conn2.cursor() as cursor:
        cursor.execute(query,(code,Threshold_Type))
        data = cursor.fetchall()
        rip_thresholds_ind={}
        rip_thresholds_legal={}
        json_data = json.loads(data[0][0])
        for key,value in json_data.items():
            if 'individual' in key.lower():
                if 'low' in key.lower():
                    rip_thresholds_ind[1]=value
                if 'mid' in key.lower() or 'med' in key.lower():
                    rip_thresholds_ind[2]=value
                if 'high' in key.lower():
                    rip_thresholds_ind[3]=value
            if 'legal' in key.lower():
                if 'low' in key.lower():
                    rip_thresholds_legal[1]=value
                if 'mid' in key.lower() or 'med' in key.lower():
                    rip_thresholds_legal[2]=value
                if 'high' in key.lower():
                    rip_thresholds_legal[3]=value
    return rip_thresholds_ind,rip_thresholds_legal


def TM11():
        
    try:
        current_datetime = datetime.now()
        current_date = current_datetime.date()

        onDate1 = datetime.strptime('29-05-2024', '%d-%m-%Y')
        # ofDate = datetime.strptime('31-03-2024', '%d-%m-%Y')

        query = """select Current_values from Thresholds where code=? AND Threshold_Type=?;"""
        rip_thresholds_ind,rip_thresholds_legal = fetch_current_values(query,'TM_1_1','RFI')

        print(rip_thresholds_ind)
        print(rip_thresholds_legal)

       
        # while onDate1 <= ofDate:
        date = onDate1 - timedelta(days=1)
        date = date.strftime('%d-%m-%Y')
        
        # onDate1 += timedelta(days=1)
        date = str(date)
        print(date)

        current_date = datetime.now().date()
        raisedDate = datetime.combine(current_date, datetime.min.time())

        quireTM11 = f"""

                    WITH AggregatedTransactions AS (
                    SELECT
                        MAX('TM_1_1') AS scenario_code,
                        MAX('--') AS ticketid,
                        MAX('{raisedDate}') AS alert_created_on,
                        NULL AS alert_allocated_on,
                        mc.ACCTNO, 
                        mc.TXDATE,
                        SUM(CAST(mc.TXAMT AS FLOAT)) AS TXAMT,
                        MAX(mc.TXTYPE) AS TXTYPE,
                        MAX(mc.CASHFLOWTYPE) AS CASHFLOWTYPE,
                        MAX(mc.TRNFLOWTYPE) AS TRNFLOWTYPE,
                        MAX(mc.CURRCD) AS CURRCD,
                        MAX(mc.TXFRCURRCD) AS TXFRCURRCD,
                        MAX(mc.TXFRCOUNTRY) AS TXFRCOUNTRY
                    FROM                     
                        TRANSACTIONS mc
                    JOIN
                        CUSTOMERS cd ON mc.ACCTNO = cd.ACCTNO
                    WHERE 
                        mc.TXTYPE = 'C' 
                        AND mc.TXDATE = '{date}' 
                        AND mc.TRNFLOWTYPE = 'C' 
                        AND mc.TXFRCURRCD = 'INR'
                    GROUP BY 
                        mc.ACCTNO, mc.TXDATE
                    HAVING 
                        (((MAX(cd.RIP) = '1' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[1]}) OR
                            (MAX(cd.RIP) = '2' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[2]}) OR
                            (MAX(cd.RIP) = '3' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[2]}) OR
                            (MAX(cd.RIP) IS NULL AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[3]})
                        )  AND MAX(cd.CUST_TYPE_CODE) IN ('IND', 'INDIV'))
                        OR 
                        (((MAX(cd.RIP) = '1' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[1]}) OR
                            (MAX(cd.RIP) = '2' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[2]}) OR
                            (MAX(cd.RIP) = '3' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[2]}) OR
                            (MAX(cd.RIP) IS NULL AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[3]})
                        )  AND MAX(cd.CUST_TYPE_CODE) NOT IN ('IND', 'INDIV'))
                )
                SELECT
                    at.scenario_code,
                    at.ticketid,
                    at.alert_created_on,
                    at.alert_allocated_on,
                    at.TXDATE,
                    at.TXAMT,
                    at.TXTYPE,
                    at.CASHFLOWTYPE,
                    at.TRNFLOWTYPE,
                    at.CURRCD,
                    at.TXFRCURRCD,
                    at.TXFRCOUNTRY,
                    cd.*
                FROM 
                    AggregatedTransactions at
                JOIN 
                    CUSTOMERS cd ON at.ACCTNO = cd.ACCTNO

                    """
        

        res = mysql2.execute(quireTM11)




        finalRes = res.fetchall()
        column_names = [desc[0] for desc in res.description]
        Acccounts = []

        if finalRes:
            flat_data = []
            for row in finalRes:
                flat_row = []
                for index, sublist in enumerate(row):
                        if sublist is None:
                            flat_row.append(None)
                        else:
                            flat_row.append(str(sublist))

                flat_row[1] = f"ARM-VRV-{current_date}-{uuid.uuid4()}"
                Acccounts.append(flat_row[21])

                
                flat_data.append(tuple(flat_row))

            df = pd.DataFrame(flat_data)

            print(df)

            
            tableInfo = mysql2.execute("SELECT 1 FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'scenarios'")
            table_exists = tableInfo.fetchone()

            if not table_exists:
                column_definitions = ", ".join([f"{col} NVARCHAR(MAX)" for col in column_names])
                create_table_query = f"CREATE TABLE scenarios ({column_definitions})"
                mysql2.execute(create_table_query)
                mysql2.commit()



            placeholders = ", ".join(["?"] * len(column_names))
            insert_query = f"INSERT INTO scenarios ({', '.join(column_names)}) VALUES ({placeholders})"
        


            data = [tuple(row) for row in df.to_numpy()]



            batch_size = 1000

            for i in range(0, len(data), batch_size):
                batch = data[i:i+batch_size]
                mysql2.executemany(insert_query, batch)
                mysql2.commit()
            
            



    except Exception as e:
        mysql2.connection.rollback()
        print(f"Somthing went wrong {e}")

        


def TM12():


    try:
        current_datetime = datetime.now()
        current_date = current_datetime.date()

        onDate1 = datetime.strptime('29-05-2024', '%d-%m-%Y')
        # ofDate = datetime.strptime('31-03-2024', '%d-%m-%Y')

        query = """select Current_values from Thresholds where code=? AND Threshold_Type=?;"""
        rip_thresholds_ind,rip_thresholds_legal = fetch_current_values(query,'TM_1_2','RFI')

        # while onDate1 <= ofDate:
        date = onDate1 - timedelta(days=1)
        date = date.strftime('%d-%m-%Y')
        
        # onDate1 += timedelta(days=1)
        date = str(date)
        print(date)

        current_date = datetime.now().date()
        raisedDate = datetime.combine(current_date, datetime.min.time())

        quireTM11 = f"""

                    WITH AggregatedTransactions AS (
                    SELECT
                        MAX('TM_1_2') AS scenario_code,
                        MAX('--') AS ticketid,
                        MAX('{raisedDate}') AS alert_created_on,
                        NULL AS alert_allocated_on,
                        mc.ACCTNO, 
                        mc.TXDATE,
                        SUM(CAST(mc.TXAMT AS FLOAT)) AS TXAMT,
                        MAX(mc.TXTYPE) AS TXTYPE,
                        MAX(mc.CASHFLOWTYPE) AS CASHFLOWTYPE,
                        MAX(mc.TRNFLOWTYPE) AS TRNFLOWTYPE,
                        MAX(mc.CURRCD) AS CURRCD,
                        MAX(mc.TXFRCURRCD) AS TXFRCURRCD,
                        MAX(mc.TXFRCOUNTRY) AS TXFRCOUNTRY
                    FROM                     
                        TRANSACTIONS mc
                    JOIN
                        CUSTOMERS cd ON mc.ACCTNO = cd.ACCTNO
                    WHERE 
                        mc.TXTYPE = 'C' 
                        AND mc.TXDATE = '{date}' 
                        AND mc.TRNFLOWTYPE = 'D' 
                        AND mc.TXFRCURRCD = 'INR'
                    GROUP BY 
                        mc.ACCTNO, mc.TXDATE
                    HAVING 
                            (((MAX(cd.RIP) = '1' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[1]}) OR
                            (MAX(cd.RIP) = '2' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[2]}) OR
                            (MAX(cd.RIP) = '3' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[2]}) OR
                            (MAX(cd.RIP) IS NULL AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[3]})
                        )  AND MAX(cd.CUST_TYPE_CODE) IN ('IND', 'INDIV'))
                        OR 
                        (((MAX(cd.RIP) = '1' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[1]}) OR
                            (MAX(cd.RIP) = '2' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[2]}) OR
                            (MAX(cd.RIP) = '3' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[2]}) OR
                            (MAX(cd.RIP) IS NULL AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[3]})
                        )  AND MAX(cd.CUST_TYPE_CODE) NOT IN ('IND', 'INDIV'))
                )
                SELECT
                    at.scenario_code,
                    at.ticketid,
                    at.alert_created_on,
                    at.alert_allocated_on,
                    at.TXDATE,
                    at.TXAMT,
                    at.TXTYPE,
                    at.CASHFLOWTYPE,
                    at.TRNFLOWTYPE,
                    at.CURRCD,
                    at.TXFRCURRCD,
                    at.TXFRCOUNTRY,
                    cd.*
                FROM 
                    AggregatedTransactions at
                JOIN 
                    CUSTOMERS cd ON at.ACCTNO = cd.ACCTNO

                    """

        res = mysql2.execute(quireTM11)

        finalRes = res.fetchall()
        print(finalRes)

        column_names = [desc[0] for desc in res.description]

        if finalRes:
            flat_data = []
            for row in finalRes:
                flat_row = []
                for index, sublist in enumerate(row):
                        if sublist is None:
                            flat_row.append(None)  # Keep None as None
                        else:
                            flat_row.append(str(sublist))

                flat_row[1] = f"ARM-VRV-{current_date}-{uuid.uuid4()}"

                
                flat_data.append(tuple(flat_row))

            df = pd.DataFrame(flat_data)

            
            tableInfo = mysql2.execute("SELECT 1 FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'scenarios'")
            table_exists = tableInfo.fetchone()

            if not table_exists:
                column_definitions = ", ".join([f"{col} NVARCHAR(MAX)" for col in column_names])
                create_table_query = f"CREATE TABLE scenarios ({column_definitions})"
                mysql2.execute(create_table_query)
                mysql2.commit()



            placeholders = ", ".join(["?"] * len(column_names))
            insert_query = f"INSERT INTO scenarios ({', '.join(column_names)}) VALUES ({placeholders})"
        


            data = [tuple(row) for row in df.to_numpy()]



            batch_size = 1000;

            for i in range(0, len(data), batch_size):
                batch = data[i:i+batch_size]
                mysql2.executemany(insert_query, batch)
                mysql2.commit()
    

    except Exception as e:
        mysql2.connection.rollback()
        print(f"Somthing went wrong {e}")

   

def TM13():

    try:
        current_datetime = datetime.now()
        current_date = current_datetime.date()

        onDate1 = datetime.strptime('29-05-2024', '%d-%m-%Y')
        # ofDate = datetime.strptime('31-03-2024', '%d-%m-%Y')

        query = """select Current_values from Thresholds where code=? AND Threshold_Type=?;"""
        rip_thresholds_ind,rip_thresholds_legal = fetch_current_values(query,'TM_1_3','RFI')
        # while onDate1 <= ofDate:
        date = onDate1 - timedelta(days=1)
        date = date.strftime('%d-%m-%Y')
        
        # onDate1 += timedelta(days=1)
        date = str(date)
        print(date)

        current_date = datetime.now().date()
        raisedDate = datetime.combine(current_date, datetime.min.time())

        quireTM11 = f"""

                    WITH AggregatedTransactions AS (
                    SELECT
                        MAX('TM_1_3') AS scenario_code,
                        MAX('--') AS ticketid,
                        MAX('{raisedDate}') AS alert_created_on,
                        NULL AS alert_allocated_on,
                        mc.ACCTNO, 
                        mc.TXDATE,
                        SUM(CAST(mc.TXAMT AS FLOAT)) AS TXAMT,
                        MAX(mc.TXTYPE) AS TXTYPE,
                        MAX(mc.CASHFLOWTYPE) AS CASHFLOWTYPE,
                        MAX(mc.TRNFLOWTYPE) AS TRNFLOWTYPE,
                        MAX(mc.CURRCD) AS CURRCD,
                        MAX(mc.TXFRCURRCD) AS TXFRCURRCD,
                        MAX(mc.TXFRCOUNTRY) AS TXFRCOUNTRY
                    FROM                     
                        TRANSACTIONS mc
                    JOIN
                        CUSTOMERS cd ON mc.ACCTNO = cd.ACCTNO
                    WHERE 
                        mc.TXTYPE = 'T' 
                        AND mc.TXDATE = '{date}' 
                        AND mc.TRNFLOWTYPE = 'C' 
                        AND mc.TXFRCURRCD = 'INR'
                    GROUP BY 
                        mc.ACCTNO, mc.TXDATE
                    HAVING 
                            (((MAX(cd.RIP) = '1' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[1]}) OR
                            (MAX(cd.RIP) = '2' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[2]}) OR
                            (MAX(cd.RIP) = '3' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[2]}) OR
                            (MAX(cd.RIP) IS NULL AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[3]})
                        )  AND MAX(cd.CUST_TYPE_CODE) IN ('IND', 'INDIV'))
                        OR 
                        (((MAX(cd.RIP) = '1' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[1]}) OR
                            (MAX(cd.RIP) = '2' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[2]}) OR
                            (MAX(cd.RIP) = '3' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[2]}) OR
                            (MAX(cd.RIP) IS NULL AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[3]})
                        )  AND MAX(cd.CUST_TYPE_CODE) NOT IN ('IND', 'INDIV'))
                )
                SELECT
                    at.scenario_code,
                    at.ticketid,
                    at.alert_created_on,
                    at.alert_allocated_on,
                    at.TXDATE,
                    at.TXAMT,
                    at.TXTYPE,
                    at.CASHFLOWTYPE,
                    at.TRNFLOWTYPE,
                    at.CURRCD,
                    at.TXFRCURRCD,
                    at.TXFRCOUNTRY,
                    cd.*
                FROM 
                    AggregatedTransactions at
                JOIN 
                    CUSTOMERS cd ON at.ACCTNO = cd.ACCTNO

                    """

        res = mysql2.execute(quireTM11)

        finalRes = res.fetchall()
        print(finalRes)

        column_names = [desc[0] for desc in res.description]

        if finalRes:
            flat_data = []
            for row in finalRes:
                flat_row = []
                for index, sublist in enumerate(row):
                        if sublist is None:
                            flat_row.append(None)  # Keep None as None
                        else:
                            flat_row.append(str(sublist)) 

                flat_row[1] = f"ARM-VRV-{current_date}-{uuid.uuid4()}"

                
                flat_data.append(tuple(flat_row))

            df = pd.DataFrame(flat_data)

            
            tableInfo = mysql2.execute("SELECT 1 FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'scenarios'")
            table_exists = tableInfo.fetchone()

            if not table_exists:
                column_definitions = ", ".join([f"{col} NVARCHAR(MAX)" for col in column_names])
                create_table_query = f"CREATE TABLE scenarios ({column_definitions})"
                mysql2.execute(create_table_query)
                mysql2.commit()



            placeholders = ", ".join(["?"] * len(column_names))
            insert_query = f"INSERT INTO scenarios ({', '.join(column_names)}) VALUES ({placeholders})"
        


            data = [tuple(row) for row in df.to_numpy()]



            batch_size = 1000;

            for i in range(0, len(data), batch_size):
                batch = data[i:i+batch_size]
                mysql2.executemany(insert_query, batch)
                mysql2.commit()
    

    except Exception as e:
        mysql2.connection.rollback()
        print(f"Somthing went wrong {e}")


def TM14():
    try:
        current_datetime = datetime.now()
        current_date = current_datetime.date()

        onDate1 = datetime.strptime('29-05-2024', '%d-%m-%Y')
        # ofDate = datetime.strptime('31-03-2024', '%d-%m-%Y')

        query = """select Current_values from Thresholds where code=? AND Threshold_Type=?;"""
        rip_thresholds_ind,rip_thresholds_legal = fetch_current_values(query,'TM_1_4','RFI')
        # while onDate1 <= ofDate:
        date = onDate1 - timedelta(days=1)
        date = date.strftime('%d-%m-%Y')
        
        # onDate1 += timedelta(days=1)
        date = str(date)
        print(date)

        current_date = datetime.now().date()
        raisedDate = datetime.combine(current_date, datetime.min.time())

        quireTM11 = f"""

                    WITH AggregatedTransactions AS (
                    SELECT
                        MAX('TM_1_4') AS scenario_code,
                        MAX('--') AS ticketid,
                        MAX('{raisedDate}') AS alert_created_on,
                        NULL AS alert_allocated_on,
                        mc.ACCTNO, 
                        mc.TXDATE,
                        SUM(CAST(mc.TXAMT AS FLOAT)) AS TXAMT,
                        MAX(mc.TXTYPE) AS TXTYPE,
                        MAX(mc.CASHFLOWTYPE) AS CASHFLOWTYPE,
                        MAX(mc.TRNFLOWTYPE) AS TRNFLOWTYPE,
                        MAX(mc.CURRCD) AS CURRCD,
                        MAX(mc.TXFRCURRCD) AS TXFRCURRCD,
                        MAX(mc.TXFRCOUNTRY) AS TXFRCOUNTRY
                    FROM                     
                        TRANSACTIONS mc
                    JOIN
                        CUSTOMERS cd ON mc.ACCTNO = cd.ACCTNO
                    WHERE 
                        mc.TXTYPE = 'T' 
                        AND mc.TXDATE = '{date}' 
                        AND mc.TRNFLOWTYPE = 'D' 
                        AND mc.TXFRCURRCD = 'INR'
                    GROUP BY 
                        mc.ACCTNO, mc.TXDATE
                    HAVING 
                            (((MAX(cd.RIP) = '1' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[1]}) OR
                            (MAX(cd.RIP) = '2' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[2]}) OR
                            (MAX(cd.RIP) = '3' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[2]}) OR
                            (MAX(cd.RIP) IS NULL AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[3]})
                        )  AND MAX(cd.CUST_TYPE_CODE) IN ('IND', 'INDIV'))
                        OR 
                        (((MAX(cd.RIP) = '1' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[1]}) OR
                            (MAX(cd.RIP) = '2' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[2]}) OR
                            (MAX(cd.RIP) = '3' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[2]}) OR
                            (MAX(cd.RIP) IS NULL AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[3]})
                        )  AND MAX(cd.CUST_TYPE_CODE) NOT IN ('IND', 'INDIV'))
                )
                SELECT
                    at.scenario_code,
                    at.ticketid,
                    at.alert_created_on,
                    at.alert_allocated_on,
                    at.TXDATE,
                    at.TXAMT,
                    at.TXTYPE,
                    at.CASHFLOWTYPE,
                    at.TRNFLOWTYPE,
                    at.CURRCD,
                    at.TXFRCURRCD,
                    at.TXFRCOUNTRY,
                    cd.*
                FROM 
                    AggregatedTransactions at
                JOIN 
                    CUSTOMERS cd ON at.ACCTNO = cd.ACCTNO

                    """

        res = mysql2.execute(quireTM11)

        finalRes = res.fetchall()
        print(finalRes)

        column_names = [desc[0] for desc in res.description]

        if finalRes:
            flat_data = []
            for row in finalRes:
                flat_row = []
                for index, sublist in enumerate(row):
                        if sublist is None:
                            flat_row.append(None)  # Keep None as None
                        else:
                            flat_row.append(str(sublist)) 

                flat_row[1] = f"ARM-VRV-{current_date}-{uuid.uuid4()}"

                
                flat_data.append(tuple(flat_row))

            df = pd.DataFrame(flat_data)

            
            tableInfo = mysql2.execute("SELECT 1 FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'scenarios'")
            table_exists = tableInfo.fetchone()

            if not table_exists:
                column_definitions = ", ".join([f"{col} NVARCHAR(MAX)" for col in column_names])
                create_table_query = f"CREATE TABLE scenarios ({column_definitions})"
                mysql2.execute(create_table_query)
                mysql2.commit()



            placeholders = ", ".join(["?"] * len(column_names))
            insert_query = f"INSERT INTO scenarios ({', '.join(column_names)}) VALUES ({placeholders})"
        


            data = [tuple(row) for row in df.to_numpy()]



            batch_size = 1000;

            for i in range(0, len(data), batch_size):
                batch = data[i:i+batch_size]
                mysql2.executemany(insert_query, batch)
                mysql2.commit()
    

    except Exception as e:
        mysql2.connection.rollback()
        print(f"Somthing went wrong {e}")





def TM21():


    try:
        onDate1 = datetime.strptime('29-05-2024', '%d-%m-%Y')
        # end_date = datetime.strptime('01-04-2024', '%d-%m-%Y')

        query = """select Current_values from Thresholds where code=? AND Threshold_Type=?;"""
        rip_thresholds_ind,rip_thresholds_legal = fetch_current_values(query,'TM_2_1','RFI')

        # while onDate1 <= end_date:

        date = onDate1.strftime('%d-%m-%Y')
        # date = onDate1 - timedelta(days=1)
        # date = date.strftime('%d-%m-%Y')
        # print(date)
        

        presentDate = datetime.strptime(date, '%d-%m-%Y')
        pastOneDate = presentDate - timedelta(days=30)
        pastOneTimestamp = pastOneDate.strftime('%Y-%m-%d')


        presentDate = presentDate - timedelta(days=1)
        presentDate = presentDate.strftime('%Y-%m-%d') 

        current_date = datetime.now().date()
        raisedDate = datetime.combine(current_date, datetime.min.time())
        print(pastOneTimestamp)
        print(presentDate)




        
        quire = f"""


                WITH AggregatedTransactions AS (
                    SELECT
                        MAX('TM_2_1') AS scenario_code,
                        MAX('--') AS ticketid,
                        MAX('{raisedDate}') AS alert_created_on,
                        NULL AS alert_allocated_on,
                        mc.ACCTNO, 
                        MAX('{str(pastOneTimestamp)} to {str(presentDate)}') AS TXDATE,
                        SUM(CAST(mc.TXAMT AS FLOAT)) AS TXAMT,
                        MAX(mc.TXTYPE) AS TXTYPE,
                        MAX(mc.CASHFLOWTYPE) AS CASHFLOWTYPE,
                        MAX(mc.TRNFLOWTYPE) AS TRNFLOWTYPE,
                        MAX(mc.CURRCD) AS CURRCD,
                        MAX(mc.TXFRCURRCD) AS TXFRCURRCD,
                        MAX(mc.TXFRCOUNTRY) AS TXFRCOUNTRY
                    FROM                     
                        TRANSACTIONS mc
                    JOIN
                        CUSTOMERS cd ON mc.ACCTNO = cd.ACCTNO
                    WHERE 
                        mc.TXTYPE IN ('C') 
                        AND TRY_CONVERT(DATE, mc.TXDATE, 105) BETWEEN '{pastOneTimestamp}' AND '{presentDate}'
                        AND TRY_CONVERT(DATE, cd.ACCT_OPENDATE, 105) <= '{pastOneTimestamp}'  
                        AND mc.TRNFLOWTYPE = 'C' 
                        AND mc.TXFRCURRCD = 'INR'
                    GROUP BY 
                        mc.ACCTNO
                    HAVING 
                        (((MAX(cd.RIP) = '1' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[1]}) OR
                            (MAX(cd.RIP) = '2' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[2]}) OR
                            (MAX(cd.RIP) = '3' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[2]}) OR
                            (MAX(cd.RIP) IS NULL AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[3]})
                        )  AND MAX(cd.CUST_TYPE_CODE) IN ('IND', 'INDIV'))
                        OR 
                        (((MAX(cd.RIP) = '1' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[1]}) OR
                            (MAX(cd.RIP) = '2' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[2]}) OR
                            (MAX(cd.RIP) = '3' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[2]}) OR
                            (MAX(cd.RIP) IS NULL AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[3]})
                        )  AND MAX(cd.CUST_TYPE_CODE) NOT IN ('IND', 'INDIV'))
                )
                SELECT
                    at.scenario_code,
                    at.ticketid,
                    at.alert_created_on,
                    at.alert_allocated_on,
                    at.TXDATE,
                    at.TXAMT,
                    at.TXTYPE,
                    at.CASHFLOWTYPE,
                    at.TRNFLOWTYPE,
                    at.CURRCD,
                    at.TXFRCURRCD,
                    at.TXFRCOUNTRY,
                    cd.*
                FROM 
                    AggregatedTransactions at
                JOIN 
                    CUSTOMERS cd ON at.ACCTNO = cd.ACCTNO
                
                

    """


        res = mysql2.execute(quire)

        finalRes = res.fetchall()
        print(finalRes)

        column_names = [desc[0] for desc in res.description]
        
        if finalRes:
            flat_data = []
            for row in finalRes:
                flat_row = []
                for index, sublist in enumerate(row):
                        if sublist is None:
                            flat_row.append(None)  # Keep None as None
                        else:
                            flat_row.append(str(sublist)) 

                flat_row[1] = f"ARM-VRV-{current_date}-{uuid.uuid4()}"

                
                flat_data.append(tuple(flat_row))

            df = pd.DataFrame(flat_data)
            
            tableInfo = mysql2.execute("SELECT 1 FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'scenarios'")
            table_exists = tableInfo.fetchone()

            if not table_exists:
                column_definitions = ", ".join([f"{col} NVARCHAR(MAX)" for col in column_names])
                create_table_query = f"CREATE TABLE scenarios ({column_definitions})"
                mysql2.execute(create_table_query)
                mysql2.commit()



            placeholders = ", ".join(["?"] * len(column_names))
            insert_query = f"INSERT INTO scenarios ({', '.join(column_names)}) VALUES ({placeholders})"
        


            data = [tuple(row) for row in df.to_numpy()]



            batch_size = 1000;

            for i in range(0, len(data), batch_size):
                batch = data[i:i+batch_size]
                mysql2.executemany(insert_query, batch)
                mysql2.commit()

        # onDate1 += timedelta(days=1)

    except Exception as e:
        mysql2.connection.rollback()
        print(f"Somthing went wrong {e}")



    


def TM22():


    try:
        onDate1 = datetime.strptime('29-05-2024', '%d-%m-%Y')
        # end_date = datetime.strptime('01-04-2024', '%d-%m-%Y')

        query = """select Current_values from Thresholds where code=? AND Threshold_Type=?;"""
        rip_thresholds_ind,rip_thresholds_legal = fetch_current_values(query,'TM_2_2','RFI')

        # while onDate1 <= end_date:

        date = onDate1.strftime('%d-%m-%Y')
        # date = onDate1 - timedelta(days=1)
        # date = date.strftime('%d-%m-%Y')
        # print(date)
        

        presentDate = datetime.strptime(date, '%d-%m-%Y')
        pastOneDate = presentDate - timedelta(days=30)
        pastOneTimestamp = pastOneDate.strftime('%Y-%m-%d')


        presentDate = presentDate - timedelta(days=1)
        presentDate = presentDate.strftime('%Y-%m-%d') 

        current_date = datetime.now().date()
        raisedDate = datetime.combine(current_date, datetime.min.time())

        print(pastOneTimestamp)
        print(presentDate)


        
        quire = f"""


                WITH AggregatedTransactions AS (
                    SELECT
                        MAX('TM_2_2') AS scenario_code,
                        MAX('--') AS ticketid,
                        MAX('{raisedDate}') AS alert_created_on,
                        NULL AS alert_allocated_on,
                        mc.ACCTNO, 
                        MAX('{str(pastOneTimestamp)} to {str(presentDate)}') AS TXDATE,
                        SUM(CAST(mc.TXAMT AS FLOAT)) AS TXAMT,
                        MAX(mc.TXTYPE) AS TXTYPE,
                        MAX(mc.CASHFLOWTYPE) AS CASHFLOWTYPE,
                        MAX(mc.TRNFLOWTYPE) AS TRNFLOWTYPE,
                        MAX(mc.CURRCD) AS CURRCD,
                        MAX(mc.TXFRCURRCD) AS TXFRCURRCD,
                        MAX(mc.TXFRCOUNTRY) AS TXFRCOUNTRY
                    FROM                     
                        TRANSACTIONS mc
                    JOIN
                        CUSTOMERS cd ON mc.ACCTNO = cd.ACCTNO
                    WHERE 
                        mc.TXTYPE IN ('C') 
                        AND TRY_CONVERT(DATE, mc.TXDATE, 105) BETWEEN '{pastOneTimestamp}' AND '{presentDate}'
                        AND TRY_CONVERT(DATE, cd.ACCT_OPENDATE, 105) <= '{pastOneTimestamp}'  
                        AND mc.TRNFLOWTYPE = 'D' 
                        AND mc.TXFRCURRCD = 'INR'
                    GROUP BY 
                        mc.ACCTNO
                    HAVING 
                            (((MAX(cd.RIP) = '1' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[1]}) OR
                            (MAX(cd.RIP) = '2' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[2]}) OR
                            (MAX(cd.RIP) = '3' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[2]}) OR
                            (MAX(cd.RIP) IS NULL AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[3]})
                        )  AND MAX(cd.CUST_TYPE_CODE) IN ('IND', 'INDIV'))
                        OR 
                        (((MAX(cd.RIP) = '1' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[1]}) OR
                            (MAX(cd.RIP) = '2' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[2]}) OR
                            (MAX(cd.RIP) = '3' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[2]}) OR
                            (MAX(cd.RIP) IS NULL AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[3]})
                        )  AND MAX(cd.CUST_TYPE_CODE) NOT IN ('IND', 'INDIV'))
                )
                SELECT
                    at.scenario_code,
                    at.ticketid,
                    at.alert_created_on,
                    at.alert_allocated_on,
                    at.TXDATE,
                    at.TXAMT,
                    at.TXTYPE,
                    at.CASHFLOWTYPE,
                    at.TRNFLOWTYPE,
                    at.CURRCD,
                    at.TXFRCURRCD,
                    at.TXFRCOUNTRY,
                    cd.*
                FROM 
                    AggregatedTransactions at
                JOIN 
                    CUSTOMERS cd ON at.ACCTNO = cd.ACCTNO
                
                

    """


        res = mysql2.execute(quire)

        finalRes = res.fetchall()
        print(finalRes)

        column_names = [desc[0] for desc in res.description]
        
        if finalRes:
            flat_data = []
            for row in finalRes:
                flat_row = []
                for index, sublist in enumerate(row):
                        if sublist is None:
                            flat_row.append(None)  # Keep None as None
                        else:
                            flat_row.append(str(sublist)) 

                flat_row[1] = f"ARM-VRV-{current_date}-{uuid.uuid4()}"

                
                flat_data.append(tuple(flat_row))

            df = pd.DataFrame(flat_data)


            
            tableInfo = mysql2.execute("SELECT 1 FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'scenarios'")
            table_exists = tableInfo.fetchone()

            if not table_exists:
                column_definitions = ", ".join([f"{col} NVARCHAR(MAX)" for col in column_names])
                create_table_query = f"CREATE TABLE scenarios ({column_definitions})"
                mysql2.execute(create_table_query)
                mysql2.commit()



            placeholders = ", ".join(["?"] * len(column_names))
            insert_query = f"INSERT INTO scenarios ({', '.join(column_names)}) VALUES ({placeholders})"
        


            data = [tuple(row) for row in df.to_numpy()]



            batch_size = 1000;

            for i in range(0, len(data), batch_size):
                batch = data[i:i+batch_size]
                mysql2.executemany(insert_query, batch)
                mysql2.commit()

        # onDate1 += timedelta(days=1)

    except Exception as e:
        mysql2.connection.rollback()
        print(f"Somthing went wrong {e}")





def TM23():


    try:
        onDate1 = datetime.strptime('29-05-2024', '%d-%m-%Y')
        # end_date = datetime.strptime('01-04-2024', '%d-%m-%Y')

        query = """select Current_values from Thresholds where code=? AND Threshold_Type=?;"""
        rip_thresholds_ind,rip_thresholds_legal = fetch_current_values(query,'TM_2_3','RFI')

        # while onDate1 <= end_date:

        date = onDate1.strftime('%d-%m-%Y')
        # date = onDate1 - timedelta(days=1)
        # date = date.strftime('%d-%m-%Y')
        # print(date)
        

        presentDate = datetime.strptime(date, '%d-%m-%Y')
        pastOneDate = presentDate - timedelta(days=30)
        pastOneTimestamp = pastOneDate.strftime('%Y-%m-%d')


        presentDate = presentDate - timedelta(days=1)
        presentDate = presentDate.strftime('%Y-%m-%d') 

        current_date = datetime.now().date()
        raisedDate = datetime.combine(current_date, datetime.min.time())

        print(pastOneTimestamp)
        print(presentDate)



        
        quire = f"""


                WITH AggregatedTransactions AS (
                    SELECT
                        MAX('TM_2_3') AS scenario_code,
                        MAX('--') AS ticketid,
                        MAX('{raisedDate}') AS alert_created_on,
                        NULL AS alert_allocated_on,
                        mc.ACCTNO, 
                        MAX('{str(pastOneTimestamp)} to {str(presentDate)}') AS TXDATE,
                        SUM(CAST(mc.TXAMT AS FLOAT)) AS TXAMT,
                        MAX(mc.TXTYPE) AS TXTYPE,
                        MAX(mc.CASHFLOWTYPE) AS CASHFLOWTYPE,
                        MAX(mc.TRNFLOWTYPE) AS TRNFLOWTYPE,
                        MAX(mc.CURRCD) AS CURRCD,
                        MAX(mc.TXFRCURRCD) AS TXFRCURRCD,
                        MAX(mc.TXFRCOUNTRY) AS TXFRCOUNTRY
                    FROM                     
                        TRANSACTIONS mc
                    JOIN
                        CUSTOMERS cd ON mc.ACCTNO = cd.ACCTNO
                    WHERE 
                        mc.TXTYPE IN ('T','L') 
                        AND TRY_CONVERT(DATE, mc.TXDATE, 105) BETWEEN '{pastOneTimestamp}' AND '{presentDate}'
                        AND TRY_CONVERT(DATE, cd.ACCT_OPENDATE, 105) <= '{pastOneTimestamp}'  
                        AND mc.TRNFLOWTYPE = 'C' 
                        AND mc.TXFRCURRCD = 'INR'
                    GROUP BY 
                        mc.ACCTNO
                    HAVING 
                            (((MAX(cd.RIP) = '1' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[1]}) OR
                            (MAX(cd.RIP) = '2' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[2]}) OR
                            (MAX(cd.RIP) = '3' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[2]}) OR
                            (MAX(cd.RIP) IS NULL AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[3]})
                        )  AND MAX(cd.CUST_TYPE_CODE) IN ('IND', 'INDIV'))
                        OR 
                        (((MAX(cd.RIP) = '1' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[1]}) OR
                            (MAX(cd.RIP) = '2' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[2]}) OR
                            (MAX(cd.RIP) = '3' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[2]}) OR
                            (MAX(cd.RIP) IS NULL AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[3]})
                        )  AND MAX(cd.CUST_TYPE_CODE) NOT IN ('IND', 'INDIV'))
                )
                SELECT
                    at.scenario_code,
                    at.ticketid,
                    at.alert_created_on,
                    at.alert_allocated_on,
                    at.TXDATE,
                    at.TXAMT,
                    at.TXTYPE,
                    at.CASHFLOWTYPE,
                    at.TRNFLOWTYPE,
                    at.CURRCD,
                    at.TXFRCURRCD,
                    at.TXFRCOUNTRY,
                    cd.*
                FROM 
                    AggregatedTransactions at
                JOIN 
                    CUSTOMERS cd ON at.ACCTNO = cd.ACCTNO
                
                

    """


        res = mysql2.execute(quire)

        finalRes = res.fetchall()
        print(finalRes)

        column_names = [desc[0] for desc in res.description]
        
        if finalRes:
            flat_data = []
            for row in finalRes:
                flat_row = []
                for index, sublist in enumerate(row):
                        if sublist is None:
                            flat_row.append(None)  # Keep None as None
                        else:
                            flat_row.append(str(sublist)) 

                flat_row[1] = f"ARM-VRV-{current_date}-{uuid.uuid4()}"

                
                flat_data.append(tuple(flat_row))

            df = pd.DataFrame(flat_data)

            
            tableInfo = mysql2.execute("SELECT 1 FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'scenarios'")
            table_exists = tableInfo.fetchone()

            if not table_exists:
                column_definitions = ", ".join([f"{col} NVARCHAR(MAX)" for col in column_names])
                create_table_query = f"CREATE TABLE scenarios ({column_definitions})"
                mysql2.execute(create_table_query)
                mysql2.commit()



            placeholders = ", ".join(["?"] * len(column_names))
            insert_query = f"INSERT INTO scenarios ({', '.join(column_names)}) VALUES ({placeholders})"
        


            data = [tuple(row) for row in df.to_numpy()]



            batch_size = 1000;

            for i in range(0, len(data), batch_size):
                batch = data[i:i+batch_size]
                mysql2.executemany(insert_query, batch)
                mysql2.commit()

        # onDate1 += timedelta(days=1)

    except Exception as e:
        mysql2.connection.rollback()
        print(f"Somthing went wrong {e}")




def TM24():


    try:
        onDate1 = datetime.strptime('29-05-2024', '%d-%m-%Y')
        # end_date = datetime.strptime('01-04-2024', '%d-%m-%Y')

        query = """select Current_values from Thresholds where code=? AND Threshold_Type=?;"""
        rip_thresholds_ind,rip_thresholds_legal = fetch_current_values(query,'TM_2_4','RFI')

        # while onDate1 <= end_date:

        date = onDate1.strftime('%d-%m-%Y')
        # date = onDate1 - timedelta(days=1)
        # date = date.strftime('%d-%m-%Y')
        # print(date)
        

        presentDate = datetime.strptime(date, '%d-%m-%Y')
        pastOneDate = presentDate - timedelta(days=30)
        pastOneTimestamp = pastOneDate.strftime('%Y-%m-%d')


        presentDate = presentDate - timedelta(days=1)
        presentDate = presentDate.strftime('%Y-%m-%d') 

        current_date = datetime.now().date()
        raisedDate = datetime.combine(current_date, datetime.min.time())

        print(pastOneTimestamp)
        print(presentDate)


        
        quire = f"""


                WITH AggregatedTransactions AS (
                    SELECT
                        MAX('TM_2_4') AS scenario_code,
                        MAX('--') AS ticketid,
                        MAX('{raisedDate}') AS alert_created_on,
                        NULL AS alert_allocated_on,
                        mc.ACCTNO, 
                        MAX('{str(pastOneTimestamp)} to {str(presentDate)}') AS TXDATE,
                        SUM(CAST(mc.TXAMT AS FLOAT)) AS TXAMT,
                        MAX(mc.TXTYPE) AS TXTYPE,
                        MAX(mc.CASHFLOWTYPE) AS CASHFLOWTYPE,
                        MAX(mc.TRNFLOWTYPE) AS TRNFLOWTYPE,
                        MAX(mc.CURRCD) AS CURRCD,
                        MAX(mc.TXFRCURRCD) AS TXFRCURRCD,
                        MAX(mc.TXFRCOUNTRY) AS TXFRCOUNTRY
                    FROM                     
                        TRANSACTIONS mc
                    JOIN
                        CUSTOMERS cd ON mc.ACCTNO = cd.ACCTNO
                    WHERE 
                        mc.TXTYPE IN ('T','L') 
                        AND TRY_CONVERT(DATE, mc.TXDATE, 105) BETWEEN '{pastOneTimestamp}' AND '{presentDate}'
                        AND TRY_CONVERT(DATE, cd.ACCT_OPENDATE, 105) <= '{pastOneTimestamp}'  
                        AND mc.TRNFLOWTYPE = 'D' 
                        AND mc.TXFRCURRCD = 'INR'
                    GROUP BY 
                        mc.ACCTNO
                    HAVING 
                            (((MAX(cd.RIP) = '1' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[1]}) OR
                            (MAX(cd.RIP) = '2' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[2]}) OR
                            (MAX(cd.RIP) = '3' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[2]}) OR
                            (MAX(cd.RIP) IS NULL AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_ind[3]})
                        )  AND MAX(cd.CUST_TYPE_CODE) IN ('IND', 'INDIV'))
                        OR 
                        (((MAX(cd.RIP) = '1' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[1]}) OR
                            (MAX(cd.RIP) = '2' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[2]}) OR
                            (MAX(cd.RIP) = '3' AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[2]}) OR
                            (MAX(cd.RIP) IS NULL AND SUM(CAST(mc.TXAMT AS FLOAT)) >= {rip_thresholds_legal[3]})
                        )  AND MAX(cd.CUST_TYPE_CODE) NOT IN ('IND', 'INDIV'))
                )
                SELECT
                    at.scenario_code,
                    at.ticketid,
                    at.alert_created_on,
                    at.alert_allocated_on,
                    at.TXDATE,
                    at.TXAMT,
                    at.TXTYPE,
                    at.CASHFLOWTYPE,
                    at.TRNFLOWTYPE,
                    at.CURRCD,
                    at.TXFRCURRCD,
                    at.TXFRCOUNTRY,
                    cd.*
                FROM 
                    AggregatedTransactions at
                JOIN 
                    CUSTOMERS cd ON at.ACCTNO = cd.ACCTNO
                
                

    """


        res = mysql2.execute(quire)

        finalRes = res.fetchall()
        print(finalRes)

        column_names = [desc[0] for desc in res.description]
        
        if finalRes:
            flat_data = []
            for row in finalRes:
                flat_row = []
                for index, sublist in enumerate(row):
                        if sublist is None:
                            flat_row.append(None)  # Keep None as None
                        else:
                            flat_row.append(str(sublist)) 

                flat_row[1] = f"ARM-VRV-{current_date}-{uuid.uuid4()}"

                
                flat_data.append(tuple(flat_row))

            df = pd.DataFrame(flat_data)


            
            tableInfo = mysql2.execute("SELECT 1 FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'scenarios'")
            table_exists = tableInfo.fetchone()

            if not table_exists:
                column_definitions = ", ".join([f"{col} NVARCHAR(MAX)" for col in column_names])
                create_table_query = f"CREATE TABLE scenarios ({column_definitions})"
                mysql2.execute(create_table_query)
                mysql2.commit()



            placeholders = ", ".join(["?"] * len(column_names))
            insert_query = f"INSERT INTO scenarios ({', '.join(column_names)}) VALUES ({placeholders})"
        


            data = [tuple(row) for row in df.to_numpy()]



            batch_size = 1000;

            for i in range(0, len(data), batch_size):
                batch = data[i:i+batch_size]
                mysql2.executemany(insert_query, batch)
                mysql2.commit()

        # onDate1 += timedelta(days=1)

    except Exception as e:
        mysql2.connection.rollback()
        print(f"Somthing went wrong {e}")





def insert_into_tickets():

    try:

        cur = mysql2.connection.cursor()

        insert_data_sql = ["ALTER TABLE scenarios ADD allocatedTicket INTEGER NULL",
        "ALTER TABLE scenarios ADD mlroCasesTicket INTEGER NULL",
        "ALTER TABLE scenarios ADD cmSMCasesTicket INTEGER NULL",
        "ALTER TABLE scenarios ADD gmCasesTicket INTEGER NULL",
        "ALTER TABLE scenarios ADD mlroClosedTicket INTEGER NULL",
        "ALTER TABLE scenarios ADD ten_percent_ticket INTEGER NULL",
        "ALTER TABLE scenarios ADD cmSmClosedTicket INTEGER NULL",
        "ALTER TABLE scenarios ADD sentBackClosedTicket INTEGER NULL",
        "ALTER TABLE scenarios ADD unsatisfiedTicket INTEGER NULL",
        "ALTER TABLE scenarios ADD deletedTicket INTEGER NULL",
        "ALTER TABLE scenarios ADD approved INTEGER NULL",
        "ALTER TABLE scenarios ADD rejected INTEGER NULL",
        "ALTER TABLE scenarios ADD currentDate NVARCHAR(500) NULL"]
    
        for sql in insert_data_sql:
            cur.execute(sql)


        indexes_sql = """
                        CREATE INDEX idx_allocatedTicket ON scenarios (allocatedTicket);
                        CREATE INDEX idx_mlroCasesTicket ON scenarios (mlroCasesTicket);
                        CREATE INDEX idx_cmSMCasesTicket ON scenarios (cmSMCasesTicket);
                        CREATE INDEX idx_gmCasesTicket ON scenarios (gmCasesTicket);
                        CREATE INDEX idx_mlroClosedTicket ON scenarios (mlroClosedTicket);
                        CREATE INDEX idx_ten_percent_ticket ON scenarios (ten_percent_ticket);
                        CREATE INDEX idx_cmSmClosedTicket ON scenarios (cmSmClosedTicket);
                        CREATE INDEX idx_sentBackClosedTicket ON scenarios (sentBackClosedTicket);
                        CREATE INDEX idx_unsatisfiedTicket ON scenarios (unsatisfiedTicket);
                        CREATE INDEX idx_deletedTicket ON scenarios (deletedTicket);
                        CREATE INDEX idx_approved ON scenarios (approved);
                        CREATE INDEX idx_rejected ON scenarios (rejected);
                    """
        cur.execute(indexes_sql)
        cur.commit()

        mysql2.connection.commit()

    except Exception as e:
        mysql2.connection.rollback()
        print(e)



def secure_route(required_role=None):
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            print("Middleware executing...")
            print(session)
            if 'email_id' not in session:
                session['pls_login'] = 'Unauthorized access to this page.'
                flash('Access denied. Please sign in.', 'error')
                return redirect(url_for('sign_in'))  # Redirect to the same page
            if isinstance(required_role, list):
                if required_role and 'user_role' in session and session['user_role'] not in required_role:
                            session['pls_login'] = 'Unauthorized access to this page.'
                            flash('Access denied. You do not have the required role.', 'error')
                            return redirect(url_for('sign_in'))  # Redirect to the same page
            else:
                if required_role and 'user_role' in session and session['user_role'] != required_role:
                    session['pls_login'] = 'Unauthorized access to this page.'
                    flash('Access denied. You do not have the required role.', 'error')
                    return redirect(url_for('sign_in'))  # Redirect to the same page
            return func(*args, **kwargs)
        return wrapper
    return decorator
def apply_secure_route_to_all_routes():
    for rule in app.url_map.iter_rules():
        endpoint = rule.endpoint
        if endpoint != 'static':
            view_func = app.view_functions[endpoint]
            app.view_functions[endpoint] = secure_route()(view_func)

apply_secure_route_to_all_routes()






@app.route('/', methods=['GET', 'POST'])
def sign_in():
    msg = None
    if "success_reg_msg" in session:
        msg = session.pop("success_reg_msg",None)
    if "Invalid_password" in session:
        msg = session.pop("Invalid_password",None)
    if "no_user_found" in session:
        msg = session.pop("no_user_found",None)
    if "pls_login" in session:
        msg = session.pop("pls_login",None)
    if "login_required_success" in session:
        msg = session.pop("login_required_success",None)
    if "File_Not_Uploaded" in session:
        msg = session.pop("File_Not_Uploaded",None)
    if "incorrect_captcha" in session:
        msg = session.pop("incorrect_captcha",None)
    session.clear()
    new_captcha_dict = SIMPLE_CAPTCHA.create()
    return render_template('sign_in.html',msg=msg,captcha=new_captcha_dict)
    
    
@app.route('/refresh_captcha')
def refresh_captcha():
    new_captcha_dict = SIMPLE_CAPTCHA.create()
    return {"captcha_html":new_captcha_dict}        


# @app.route("/post_login", methods=["POST", "GET"])
# def post_login():
#     if request.method == "POST":
#         # connPostLogin = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
#         connPostLogin = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
#         mysqlPostLogin = connPostLogin.cursor()
#         email = request.form.get('email')
#         try: 
#             res = email.lower().index("@ogb.com")
#             email = email
#         except:
#             email = f"{email}@ogb.com"
#         password = request.form.get('password')
#         c_hash = request.form.get('captcha-hash')
#         c_text = request.form.get('captcha-text')
#         if SIMPLE_CAPTCHA.verify(c_text, c_hash):
#             userExist = SSO(email,password)
#             print("userExist : ",userExist)

#             if userExist != 'locked' or userExist != 'invaliedCred':
#                 sql = """
#                     SELECT *
#                     FROM [user]
#                     WHERE LOWER(userPrincipalName) = LOWER(?)
#                 """
#                 # sql = """
#                 #     SELECT *
#                 #     FROM [user]
#                 #     WHERE EmailId = ?
#                 # """
#                 # try:
#                 mysqlPostLogin.execute(sql, (userExist))
#                 # except:
#                 #     session["no_user_found"] = "Somthing went Wrong please Re-login Again"
#                 #     return redirect(url_for("sign_in"))
#                 # cursor.execute(sql, (email))
#                 user_tuple = mysqlPostLogin.fetchone()
#                 user = None
#                 if user_tuple:
#                     user = {
#                         "EmailId": user_tuple[3],
#                         # "Password": user_tuple[7],# u can remove this
#                         "Role": user_tuple[6],
#                         "Status": user_tuple[9],
#                     }
                
#                 current_datetime = datetime.now()
#                 current_date = current_datetime.date()
#                 if user:
#                     email_column = user["EmailId"]
#                     # hashed_password = user["Password"]
#                     user_role = user["Role"]
#                     status = user["Status"]
                
#                     # if password == hashed_password: 
                
                
                
#                     if status == "Approved":
#                         session["email_id"] = str(email_column)
#                         session["user_role"] = user_role
#                         session.permanent = True
#                         if user_role == "IT OFFICER":
#                                 try:
#                                     # run_tm_functions(user_role,userExist)
#                                     # run_tm_functions(user_role,email)
#                                     return redirect(url_for("ITdashboard"))
#                                 except Exception as e:
#                                     return f'Something Went Wrong: {e} Try to Re-Login Again', 500                               
                                
#                         elif user_role == "MLRO":
#                             try:
#                                 return redirect(url_for("MLROdashboard"))
#                             except Exception as e:
#                                 return f'Something Went Wrong: {e} Try to Re-Login Again', 500
#                         elif user_role == "CM/SM":
#                             try:
#                                 return redirect(url_for("CM_SM_dashboard"))
#                             except Exception as e:
#                                 return f'Something Went Wrong: {e} Try to Re-Login Again', 500
#                         elif user_role == "DGM/PO":
#                             try:
#                                 return redirect(url_for("DGMdashboard"))
#                             except Exception as e:
#                                 return f'Something Went Wrong: {e} Try to Re-Login Again', 500
#                         elif user_role == "ADMIN":
#                             try:
#                                 return redirect(url_for("AdminAllUsers"))
#                             except Exception as e:
#                                 return f'Something Went Wrong: {e} Try to Re-Login Again', 500
#                         elif user_role == "BranchMakers":
#                             try:
#                                 return redirect(url_for("branchmakers"))
#                             except Exception as e:
#                                 return f'Something Went Wrong: {e} Try to Re-Login Again', 500
#                         elif user_role == "ROS":
#                             try:
#                                 return redirect(url_for("ROSDashboard"))
#                             except Exception as e:
#                                 return f'Something Went Wrong: {e} Try to Re-Login Again', 500
#                         elif user_role == "SDN/USER":
#                             try:
#                                 return redirect(url_for("SDN_user"))
#                             except Exception as e:
#                                 return f'Something Went Wrong: {e} Try to Re-Login Again', 500
#                         # else:
#                         #     new_captcha_dict = SIMPLE_CAPTCHA.create()
#                         #     return render_template('sign_in.html',active=True,captcha=new_captcha_dict)
                            
#                     elif status == "Rejected":
#                         session["pls_login"] = "Your account has been rejected. Please contact Admin for support."
#                         return redirect(url_for("sign_in"))
#                     elif status == "Created":
#                         session["pls_login"] = "Your account is not yet approved. Please wait for approval."
#                         return redirect(url_for("sign_in"))
#                     elif status == "Deleted":
#                         session["pls_login"] = "Your account has been Deleted. Please contact Admin for support."
#                         return redirect(url_for("sign_in"))
#                     else:
#                         session["Invalid_password"] = "Invalid Credentials"
#                         return redirect(url_for("sign_in"))
#                 else:
#                     session["no_user_found"] = "No user found, Check your Credentials or Contact Adminitrator ."
#                     return redirect(url_for("sign_in"))
#             else:
#                 session["no_user_found"] = "No user found, Check your Credentials  or Contact Adminitrator."
#                 return redirect(url_for("sign_in"))
#         else:
#             session["incorrect_captcha"] = "Incorrect Captcha, please enter correct captcha"
#             return redirect(url_for("sign_in"))
#     return redirect(url_for("sign_in"))

@app.route("/post_login", methods=["POST", "GET"])
def post_login():
    if request.method == "POST":

        connPostLogin = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

        mysqlPostLogin = connPostLogin.cursor()
        
        email = request.form.get('email')
        print(email)

        password = request.form.get('password')
        c_hash = request.form.get('captcha-hash')
        c_text = request.form.get('captcha-text')
        if SIMPLE_CAPTCHA.verify(c_text, c_hash):

       
                sql = """
                    SELECT *
                    FROM [user]
                    WHERE EmailId = ?
                """

                mysqlPostLogin.execute(sql, (email))

                user_tuple = mysqlPostLogin.fetchone()

                user = None
                if user_tuple:
                    
                    user = {
                        "EmailId": user_tuple[3],
                        "Password": user_tuple[7],# u can remove this
                        "Role": user_tuple[6],
                        "Status": user_tuple[10],
                    }
                    print(user,"user")
                
                current_datetime = datetime.now()
                current_date = current_datetime.date()

                if user:
                    email_column = user["EmailId"]
                    hashed_password = user["Password"]
                    user_role = user["Role"]
                    print(user_role)
                    status = user["Status"]
                
                    if password == hashed_password: 
                
                
                
                        if status == "Approved":
                            session["email_id"] = str(email_column)
                            session["user_role"] = user_role
                            session.permanent = True
                            if user_role == "IT OFFICER":
                                    try:
                                        
                                        return redirect(url_for("ITdashboard"))
                                    except Exception as e:
                                        return f'Something Went Wrong: {e} Try to Re-Login Again', 500                               
                                    
                            elif user_role == "MLRO":
                                try:
                                    return redirect(url_for("MLROdashboard"))
                                except Exception as e:
                                    return f'Something Went Wrong: {e} Try to Re-Login Again', 500
                            elif user_role == "CM/SM":
                                try:
                                    return redirect(url_for("CM_SM_dashboard"))
                                except Exception as e:
                                    return f'Something Went Wrong: {e} Try to Re-Login Again', 500
                            elif user_role == "DGM/PO":
                                try:
                                    return redirect(url_for("DGMdashboard"))
                                except Exception as e:
                                    return f'Something Went Wrong: {e} Try to Re-Login Again', 500
                            elif user_role == "ADMIN":
                                try:
                                    return redirect(url_for("AdminAllUsers"))
                                except Exception as e:
                                    return f'Something Went Wrong: {e} Try to Re-Login Again', 500
                            elif user_role == "BranchMakers":
                                try:
                                    return redirect(url_for("branchmakers"))
                                except Exception as e:
                                    return f'Something Went Wrong: {e} Try to Re-Login Again', 500
                            elif user_role == "ROS":
                                try:
                                    return redirect(url_for("ROSDashboard"))
                                except Exception as e:
                                    return f'Something Went Wrong: {e} Try to Re-Login Again', 500
                            elif user_role == "PINACA_ADMIN":
                                try:
                                    return redirect(url_for("SDN_USER"))
                                except Exception as e:
                                    return f'Something Went Wrong: {e} Try to Re-Login Again', 500
                            elif user_role == "SDN/USER":
                                try:
                                    return redirect(url_for("SDN_USER"))
                                except Exception as e:
                                    return f'Something Went Wrong: {e} Try to Re-Login Again', 500
                            # else:
                            #     new_captcha_dict = SIMPLE_CAPTCHA.create()
                            #     return render_template('sign_in.html',active=True,captcha=new_captcha_dict)
                                
                        elif status == "Rejected":
                            session["pls_login"] = "Your account has been rejected. Please contact Admin for support."
                            return redirect(url_for("sign_in"))
                        elif status == "Created":
                            session["pls_login"] = "Your account is not yet approved. Please wait for approval."
                            return redirect(url_for("sign_in"))
                        elif status == "Deleted":
                            session["pls_login"] = "Your account has been Deleted. Please contact Admin for support."
                            return redirect(url_for("sign_in"))
                        else:
                            session["Invalid_password"] = "Invalid Credentials"
                            return redirect(url_for("sign_in"))
                else:
                    session["no_user_found"] = "No user found, Check your Credentials or Contact Adminitrator ."
                    return redirect(url_for("sign_in"))
            # else:
            #     session["no_user_found"] = "No user found, Check your Credentials."
            #     return redirect(url_for("sign_in"))
        else:
            session["incorrect_captcha"] = "Incorrect Captcha, please enter correct captcha"
            return redirect(url_for("sign_in"))
    return redirect(url_for("sign_in"))



@app.after_request
def add_header(response):
    response.cache_control.no_store = True
    return response
@app.route('/logout')
def logout():
    try:
        del session
    except:
        pass
    response = make_response(redirect(url_for("sign_in")))

    response.cache_control.no_store = True

    return response



# ====================== ADMIN CODE START'S HERE ====================================================


# ------------------ ADMIN MAINLANDING PAGE -----------------------

@app.route("/AdminAllUsers", methods=['POST', 'GET'])
@secure_route(required_role=['ADMIN'])
def AdminAllUsers():
    
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    
    email = session['email_id']
    
    # connAdminAllUsers = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connAdminAllUsers = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlAdminAllUsers = connAdminAllUsers.cursor()
    try:
        query = "SELECT * FROM [user] WHERE EmailId = ?" 

        mysqlAdminAllUsers.execute(query, (email,))
        
        rows = mysqlAdminAllUsers.fetchall()

        columns = [col[0] for col in mysqlAdminAllUsers.description]

        mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

        mlro = mlro[0]

        
        if mlro is None:
            return "User data not found. Please log in again."
        
        if 'image' in mlro:
                    # Encode the image data as a base64 string
                    mlro['image'] = base64.b64encode(mlro['image']).decode('utf-8')
        
        data=[]

        select_query = "SELECT * FROM [user]"
        
        mysqlAdminAllUsers.execute(select_query)
        columns = [desc[0] for desc in mysqlAdminAllUsers.description]
        users_data = mysqlAdminAllUsers.fetchall()

        if users_data is None:
            return "User data not found. Please log in again."

        for row in users_data:
            scenario_object = {}
            for i, value in enumerate(row):
                scenario_object[columns[i]] = value
            data.append(scenario_object)
        
        # Assuming you have a column named 'image' in your table
        for user in data:
            if 'image' in user:

                user['image'] = base64.b64encode(user['image']).decode('utf-8')


        return render_template("AllUsers.html",  users=users_data, data=data, role='ADMIN', type='AllUsers')

    except Exception as e:
        mysqlAdminAllUsers.rollback()
        return f'Something Went Wrong: {e} Try to Re-Login Again', 500


# ----------------- MAKEING USERS ON LEAVE BY ADMIN  END - POINT -----------------

@app.route('/userLeaveStatus', methods=['POST','GET'])
@secure_route(required_role=['ADMIN'])
def userLeaveStatus():
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    
    # connAdminAllUsersLeave = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connAdminAllUsersLeave = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlAdminAllUsersLeave = connAdminAllUsersLeave.cursor()
    emailid = request.form.get('emailid')
    
    try:
        updateQuery = "UPDATE [user] SET LeaveStatus = 'Leave' WHERE EmailId = ?"
        mysqlAdminAllUsersLeave.execute(updateQuery, (emailid,))
        mysqlAdminAllUsersLeave.commit()

        connAdminAllUsersLeave.close()    
        
        return redirect(url_for('AdminAllUsers'))
    except Exception as e:
        mysqlAdminAllUsersLeave.rollback()
        connAdminAllUsersLeave.close()    
        return f"Somthing went wrong {e} please Re-Login",500


# ---------------- MAKEING USERS WORKING FROM LEAVE BY ADMIN END - POINT --------------------


@app.route('/userWorkingStatus', methods=['POST','GET'])
@secure_route(required_role=['ADMIN'])
def userWorkingStatus():
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    
    # mysqlAdminAllUsersWorking = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    mysqlAdminAllUsersWorking = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlAdminAllUsersWorking = connAdminAllUsersWorking.cursor()
    
    emailid = request.form.get('emailid')

    try:

        updateQuery = "UPDATE [user] SET LeaveStatus = 'Working' WHERE EmailId = ?"
        mysqlAdminAllUsersWorking.execute(updateQuery, (emailid,))
        mysqlAdminAllUsersWorking.commit()
        
        connAdminAllUsersWorking.close()

        return redirect(url_for('AdminAllUsers'))
    
    except Exception as e:
        mysqlAdminAllUsersWorking.rollback()
        connAdminAllUsersWorking.close()

        return f"Something Went Wrong {e} , Please Re-Login Again",500


# ===================== ADMINE PART END'S HERE ==================================================================


# -------------------------------------------------------------------------------------------------------------------------------------------



# ========================= IT OFFICER / HO ADMIN CODE START'S HERE  ==========================================================


# ------------------- HO ADMIN LANDING PAGE OR DASHBOARD -------------------------------------------


@app.route('/ITdashboard', methods=['GET'])
@secure_route(required_role='IT OFFICER')
def ITdashboard():
        
        if 'email_id' not in session:
            return redirect(url_for('post_login'))

        email = session['email_id']


        # connItDash = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
        connItDash = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

        mysqlItDash = connItDash.cursor()

        try:

            query = "SELECT * FROM [user] WHERE EmailId = ?"

            mysqlItDash.execute(query, (email,))
                
            rows = mysqlItDash.fetchall()

            columns = [col[0] for col in mysqlItDash.description]

            info = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

            user = info[0]


            countSubmited = 0
            
            
            mysqlItDash.execute("SELECT COUNT(*) FROM scenarios WHERE alert_allocated_on IS NOT NULL")



            temp = mysqlItDash.fetchone()[0]
            if temp != 0:
                    countSubmited = temp

            print(countSubmited)

            connItDash.close()

            return render_template("IT OFFICER.html",ituser=user,countSubmited=countSubmited,numUsers=0,mlroDetails=[],cmUsers=[],cmDetails=[],perdayDatait=[],role='IT OFFICER',type='ITdashboard')
        
        except Exception as e:
            mysqlItDash.rollback()
            connItDash.close()
            return f'Something Went Wrong: {e} ,Please Re-Login Again', 500


# ----------------- HO ADMIN OR IT OFFICER FINET REPORTS PAGE END - POINT -----------------------------------


@app.route('/FINnetReports', methods=['GET'])
@secure_route(required_role='IT OFFICER')
def FINnetReports():
    # user = users_collection.find_one({'role': 'IT OFFICER'})

    # ituser = {'image': ""}
    # if user:
    #     ituser = users_collection.find_one({'emailid': user.get('emailid')})

    #     if ituser and 'image' in ituser:
    #         ituser['image'] = base64.b64encode(ituser['image']).decode('utf-8')

    return render_template('FINnet_report.html',type='FINnetReports',role='IT OFFICER')



# ------------------ ALLOCATION FUNCTION WHICH IS AUTOMATED --------------------------



# @app.route('/allocate', methods=['GET','POST'])
# @secure_route(required_role='IT OFFICER')
# def allocate():


#     connAllocate = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
#     # connAllocate = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

#     mysqlAllocate = connAllocate.cursor()
#     try:
#             current_datetime = datetime.now()
#             current_date = current_datetime.date()
#             midnight_datetime = datetime.combine(current_date, datetime.min.time())


#             mlrosIdList = []
#             ticketId = []


#             mysqlAllocate.execute("SELECT id FROM scenarios WHERE alert_allocated_on IS NULL")

#             ticket_ids = mysqlAllocate.fetchall()

#             for ticket_id in ticket_ids:
#                 ticketId.append(ticket_id[0])

#             mysqlAllocate.execute("SELECT id FROM [user] WHERE Role = 'MLRO' ")

#             mlroId = mysqlAllocate.fetchall()

#             for id in mlroId:
#                 mlrosIdList.append(id[0])


#             while ticketId:
#                         ticket = random.choice(ticketId)
                        
#                         mlro = random.choice(mlrosIdList)

                        
#                         mysqlAllocate.execute("UPDATE scenarios SET allocatedTicket = ? WHERE alert_allocated_on IS NULL", (mlro, ticket))
#                         mysqlAllocate.commit()
                        
#                         mysqlAllocate.execute("UPDATE scenarios SET alert_allocated_on = ? WHERE allocatedTicket IS NOT NULL AND  alert_allocated_on IS NULL ", (midnight_datetime, ticket))
#                         mysqlAllocate.commit()
                        
                        
#                         ticketId.remove(ticket)

            



#             connAllocate.close()
     
#             return redirect(url_for('ITdashboard'))
#     except Exception as e:

#         mysqlAllocate.rollback()
#         connAllocate.close()

#         return f'Something Went Wrong {e} , Please Re-Login Again ', 500
   
def allocate():


    # connAllocate = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connAllocate = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlAllocate = connAllocate.cursor()
    try:
            current_datetime = datetime.now()
            current_date = current_datetime.date()
            midnight_datetime = datetime.combine(current_date, datetime.min.time())


            tickets = []
            allMlros = []

            mysqlAllocate.execute("SELECT * FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = 'scenarios' AND COLUMN_NAME = 'alert_allocated_on'")
            column_exists = mysqlAllocate.fetchone()

            # If the column doesn't exist, add it
            if not column_exists:
                mysqlAllocate.execute("ALTER TABLE scenarios ADD alert_allocated_on DATETIME")
       
            mysqlAllocate.execute("SELECT id FROM [user] WHERE Role = 'MLRO'")

            mlroId = mysqlAllocate.fetchone()[0]

            mysqlAllocate.execute("UPDATE scenarios SET allocatedTicket = ? WHERE alert_allocated_on IS NULL",(mlroId,))
            mysqlAllocate.commit()

            mysqlAllocate.execute("UPDATE scenarios SET alert_allocated_on = ? WHERE allocatedTicket = ?", (midnight_datetime, mlroId))
            mysqlAllocate.commit()


            



            connAllocate.close()
     
            # return redirect(url_for('ITdashboard'))
    except Exception as e:

        mysqlAllocate.rollback()
        connAllocate.close()

        return f'Something Went Wrong {e} , Please Re-Login Again '
   


# ------------------ """ ONLINE STR DOWNLOAD PAGE """ IN FINET REPORT PAGE TAB  -------------------------------



@app.route('/online_STR_download_page',methods=['GET'])
@secure_route(required_role='IT OFFICER')
def online_STR_download_page():
    

        if 'email_id' not in session:
            return redirect(url_for('post_login'))

        email = session['email_id']


        # connAdminonlineSTRDownload = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
        connAdminonlineSTRDownload = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

        mysqlonlineSTRDownload = connAdminonlineSTRDownload.cursor()

        try:
            query = "SELECT * FROM [user] WHERE EmailId = ?"

            mysqlonlineSTRDownload.execute(query, (email,))
                
            rows = mysqlonlineSTRDownload.fetchall()

            columns = [col[0] for col in mysqlonlineSTRDownload.description]

            info = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

            user = info[0]




            mysqlonlineSTRDownload.execute("SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = 'scenarios'")
            columns = [column[0] for column in mysqlonlineSTRDownload.fetchall()]

            data = []

            if 'reportedDate' in columns:

                mysqlonlineSTRDownload.execute("SELECT ACCTNO, reportedDate, ticketid,TXDATE,TXTYPE,TRNFLOWTYPE,CUSTCD,scenario_code FROM scenarios WHERE approved IS NOT NULL")
            
                for doc in mysqlonlineSTRDownload.fetchall():
                    account_number = doc[0]
                    date = doc[1].date()
                    ticket_id = doc[2]
                    TXDATE = doc[3]
                    TXTYPE = doc[4]
                    TRNFLOWTYPE = doc[5]
                    CUSTCD = doc[6]
                    scenario_code = doc[7]
            
                    obj = {
                        'AccountNumber': account_number,
                        'Date': date,
                        'TicketId': ticket_id,
                        'TXDATE':TXDATE,
                        'TXTYPE':TXTYPE,
                        'TRNFLOWTYPE':TRNFLOWTYPE,
                        'CUSTCD':CUSTCD,
                        'scenario_code':scenario_code
                    }
                    data.append(obj)
            
            connAdminonlineSTRDownload.close()
                
            return render_template('online_STR_Downloads.html',data=data,type='FINnetReports',ituser=user,role='IT OFFICER')
        except Exception as e:
            mysqlonlineSTRDownload.rollback()
            connAdminonlineSTRDownload.close()

            return f"Somthing went wrong {e} , please Re-Login Again",500



# ------------------ """ OFFLINE STR DOWNLOAD PAGE """ IN FINET REPORT PAGE TAB  -------------------------------



# @app.route('/offline_STR_download_page',methods=['GET'])
# @secure_route(required_role='IT OFFICER')
# def offline_STR_download_page():


#     if 'email_id' not in session:
#             return redirect(url_for('post_login'))

#     email = session['email_id']


#     try:
#         # connAdminofflineSTRDownload = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
#         connAdminofflineSTRDownload = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

#         mysqlofflineSTRDownload = connAdminofflineSTRDownload.cursor()


#         query = "SELECT * FROM [user] WHERE EmailId = ?"

#         mysqlofflineSTRDownload.execute(query, (email,))
                
#         rows = mysqlofflineSTRDownload.fetchall()

#         columns = [col[0] for col in mysqlofflineSTRDownload.description]

#         info = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

#         user = info[0]


#         mysqlofflineSTRDownload.execute("SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = 'offline_collection'")
#         columns = [column[0] for column in mysqlofflineSTRDownload.fetchall()]

#         if 'Created_Date' in columns:
        
#             mysqlofflineSTRDownload.execute("SELECT o1.[AccountNumber], o1.[Created_Date], o1.[ticket_id] FROM [offline_collection] o1, [offline_tickets] o2 WHERE o1.[ticket_id] = o2.[ticketId] AND o2.[approved] = 1 AND [ROS_cmt] IS NOT NULL AND [DGM_cmt] IS NOT NULL")
            
#             data = []
            
#             for doc in mysqlofflineSTRDownload.fetchall():
#                 account_number = doc[0]
#                 date = doc[1]
#                 ticket_id = doc[2]
#                 date_str = doc[1]  # Access the first element of the list
#                 date_obj = datetime.strptime(date_str, "%Y-%m-%d").date()

#                 obj = {
#                     'AccountNumber': account_number,
#                     'Date': date_obj,
#                     'TicketId': ticket_id
#                 }
#                 data.append(obj)
                

        
#         connAdminofflineSTRDownload.close()
        
#         return render_template('offline_STR_Downloads.html',data=data,type='FINnetReports',ituser=user,role='IT OFFICER')

#     except Exception as e:
#         mysqlofflineSTRDownload.rollback()
#         connAdminofflineSTRDownload.close()
        
#         return f"Somthong went Wrong {e} , Please Re-Login Again",500


@app.route('/offline_STR_download_page',methods=['POST','GET'])
@secure_route(required_role='IT OFFICER')
def offline_STR_download_page():
    data = []

    cur = mysql2.connection.cursor()

    #cur.execute("SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = 'offline_collection'")
    cur.execute("SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = 'offline_scenarios'")
    
    columns = [column[0] for column in cur.fetchall()]

    # Check if 'Created_Date' is in the list of columns
    if 'Created_Date' in columns:
        print("start")
        # Execute the SELECT query
        # cur.execute("SELECT o1.[AccountNumber], o1.[Created_Date], o1.[ticket_id] FROM [dbo].[offline_collection] o1, [dbo].[offline_tickets] o2 WHERE o1.[ticket_id] = o2.[ticketId] AND o2.[approved] = 1 AND [ROS_cmt] IS NOT NULL AND [DGM_cmt] IS NOT NULL")
        #MY QUR
        cur.execute("SELECT [AccountNumber], [Created_Date], [ticket_id] FROM [dbo].[offline_scenarios] WHERE [approved] = 1 AND [ROS_cmt] IS NOT NULL AND [DGM_cmt] IS NOT NULL")
        
        for doc in cur.fetchall():
            account_number = doc[0]
            date = doc[1]
            ticket_id = doc[2]
            date_str = doc[1]  # Access the first element of the list
            date_obj = datetime.strptime(date_str, "%Y-%m-%d").date()

            obj = {
                'AccountNumber': account_number,
                'Date': date_obj,
                'TicketId': ticket_id
            }
            data.append(obj)
        print(data)
            
    query = "SELECT * FROM [dbo].[user] WHERE [Role] = ?"

    cur.execute(query, ('IT OFFICER',))
            
    rows = cur.fetchall()

    columns = [col[0] for col in cur.description]

    info = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

    user = info[0]
    
    return render_template('offline_STR_Downloads.html',data=data,type='FINnetReports',ituser=user,role='IT OFFICER')




# ------------------ """ ONLINE STR DOWNLOAD ZIP PROCESS """   -------------------------------


# @app.route('/download_pdf_strr', methods=['POST'])
# @secure_route(required_role=['IT OFFICER'])
# def download_pdf_strr():
#     accNo = request.form.get('accNumber')
#     dateSubmited = request.form.get('date')
#     TicketId = request.form.get('TicketId')
#     txdate = request.form.get('TXDATE')
#     txtype = request.form.get('TXTYPE')
#     trnsflowtype = request.form.get('TRNFLOWTYPE')
#     CUSTCD = request.form.get('CUSTCD')
#     scenario_code = request.form.get('scenario_code')



#     connAdminonlineSTRpdf = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
#     # connAdminonlineSTRpdf = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

#     mysqlonlineSTRpdf = connAdminonlineSTRpdf.cursor()

#     try:

#         matching_object = []

#         mysqlonlineSTRpdf.execute("SELECT * FROM scenarios WHERE approved IS NOT NULL AND ticketid = ?", ( TicketId,))

#         columns = [desc[0] for desc in mysqlonlineSTRpdf.description]
#         rows = mysqlonlineSTRpdf.fetchall()

#         for row in rows:
#             matching_object.append(dict(zip(columns, row)))

#         if matching_object:
#             temp_dir = tempfile.mkdtemp()
#             zip_filename = os.path.join(temp_dir, "str_data.zip")

#             with zipfile.ZipFile(zip_filename, 'w') as zipf:
#                 for obj in matching_object:
#                     if obj["ticketid"] == TicketId:

#                         keys_to_remove = ['ticketid', 'allocatedTicket', 'mlroCasesTicket', 'cmSMCasesTicket', 'gmCasesTicket', 'approved', 'rejected', 'id','currentDate']
#                         for key in keys_to_remove:
#                             obj.pop(key, None)

#                         temp_csv_filename = os.path.join(temp_dir, "str_data.csv")
#                         temp_pdf_filename = os.path.join(temp_dir, "str_data.pdf")
#                         temp_txt_filename = os.path.join(temp_dir, "str_data.txt")
#                         temp_docx_filename = os.path.join(temp_dir, "str_data.docx")
#                         merged_collection_data = mysqlonlineSTRpdf.fetchall()
#                         try:
#                             divDate = txdate.split(' to ')
#                             presentDate1 = datetime.strptime(divDate[1], '%Y-%m-%d').strftime('%Y-%m-%d')
#                             pastDate1 = datetime.strptime(divDate[0], '%Y-%m-%d').strftime('%Y-%m-%d')


#                             query1 = """
#                                 SELECT mc.TXDATE, t.TRNFLOWTYPE, mc.TXAMT, mc.TXFRCURRCD, mc.TXNNO, mc.ACCTNO AS TRNFACCTNO
#                                 FROM TRANSACTIONS t
#                                 JOIN TRANSACTIONS mc ON mc.TRNFACCTNO = t.ACCTNO AND mc.TXNNO = t.TXNNO AND mc.TXDATE = t.TXDATE
#                                 WHERE t.ACCTNO = ? AND t.TXTYPE = ? AND t.TRNFLOWTYPE = ? AND CONVERT(DATE, t.TXDATE, 105) BETWEEN ? AND ?
#                                 UNION
#                                 SELECT TXDATE, TRNFLOWTYPE, TXAMT, TXFRCURRCD, TXNNO, TRNFACCTNO
#                                 FROM TRANSACTIONS
#                                 WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TRNFACCTNO IS NOT NULL AND TRNFACCTNO != 'M' AND CONVERT(DATE, TXDATE, 105) BETWEEN ? AND ?
#                             """
#                             mysqlonlineSTRpdf.execute(query1, (accNo,txtype,trnsflowtype,pastDate1, presentDate1, accNo,txtype,trnsflowtype,pastDate1, presentDate1))
#                             scenarioTransactionData = mysqlonlineSTRpdf.fetchall()
#                             scenarioTransactionData = [(accNo,txtype) + tuple(record) for record in scenarioTransactionData]



#                             scenariosmultiDrmultiCRquery = """
#                                 WITH UniqueTransactions AS (
#                                     SELECT DISTINCT mc.TXDATE, mc.TXNNO
#                                     FROM TRANSACTIONS mc
#                                     WHERE mc.ACCTNO = ? AND mc.TXTYPE = ? AND mc.TRNFACCTNO = 'M' AND mc.TRNFLOWTYPE = ? AND CONVERT(DATE, mc.TXDATE, 105) BETWEEN ? AND ?
#                                 )
#                                 SELECT t.ACCTNO,t.TXTYPE, t.TXDATE, t.TRNFLOWTYPE, t.TXAMT, t.TXFRCURRCD, t.TXNNO,NULL
#                                 FROM TRANSACTIONS t
#                                 JOIN UniqueTransactions ut ON t.TXDATE = ut.TXDATE AND t.TXNNO = ut.TXNNO
#                                 WHERE t.TRNFACCTNO = 'M'
#                                 ORDER BY t.TXNNO ASC
#                             """
#                             mysqlonlineSTRpdf.execute(scenariosmultiDrmultiCRquery, (accNo,txtype,trnsflowtype, pastDate1, presentDate1))
#                             scenarioTransactionDataMultiDrMultiCr = mysqlonlineSTRpdf.fetchall()
#                         except:
#                             query1 = """
#                                 SELECT mc.TXDATE, t.TRNFLOWTYPE, mc.TXAMT, mc.TXFRCURRCD, mc.TXNNO, mc.ACCTNO AS TRNFACCTNO
#                                 FROM TRANSACTIONS t
#                                 JOIN TRANSACTIONS mc ON mc.TRNFACCTNO = t.ACCTNO AND mc.TXNNO = t.TXNNO AND mc.TXDATE = t.TXDATE
#                                 WHERE t.ACCTNO = ? AND t.TXTYPE = ? AND t.TRNFLOWTYPE = ? AND t.TXDATE = ?
#                                 UNION
#                                 SELECT TXDATE, TRNFLOWTYPE, TXAMT, TXFRCURRCD, TXNNO, TRNFACCTNO
#                                 FROM TRANSACTIONS
#                                 WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TXDATE = ? AND TRNFACCTNO IS NOT NULL AND TRNFACCTNO != 'M'
#                             """
#                             mysqlonlineSTRpdf.execute(query1, (accNo,txtype,trnsflowtype,txdate, accNo,txtype,trnsflowtype,txdate))
#                             scenarioTransactionData = mysqlonlineSTRpdf.fetchall()
#                             scenarioTransactionData = [(accNo,txtype) + tuple(record) for record in scenarioTransactionData]


#                             scenariosmultiDrmultiCRquery = """
#                                 WITH UniqueTransactions AS (
#                                     SELECT DISTINCT mc.TXDATE, mc.TXNNO
#                                     FROM TRANSACTIONS mc
#                                     WHERE mc.ACCTNO = ? AND mc.TXTYPE = ? AND mc.TRNFACCTNO = 'M' AND mc.TRNFLOWTYPE = ? AND mc.TXDATE = ? 
#                                 )
#                                 SELECT t.ACCTNO,t.TXTYPE,t.TXDATE, t.TRNFLOWTYPE, t.TXAMT, t.TXFRCURRCD, t.TXNNO,NULL
#                                 FROM TRANSACTIONS t
#                                 JOIN UniqueTransactions ut ON t.TXDATE = ut.TXDATE AND t.TXNNO = ut.TXNNO
#                                 WHERE t.TRNFACCTNO = 'M'
#                                 ORDER BY t.TXNNO ASC
#                             """
#                             mysqlonlineSTRpdf.execute(scenariosmultiDrmultiCRquery, (accNo,txtype,trnsflowtype, txdate))
#                             scenarioTransactionDataMultiDrMultiCr = mysqlonlineSTRpdf.fetchall()

#                         mysqlonlineSTRpdf.close()
#                         scenarios_extened = scenarioTransactionData + scenarioTransactionDataMultiDrMultiCr
                        

#                         merged_collection_data = [list(row) for row in scenarios_extened]
#                         merged_collection_columns = ['ACCTNO','TXTYPE','TXDATE', 'TRNFLOWTYPE', 'TXAMT','TXFRCURRCD','TXNNO','TRNFACCTNO']
#                         df_merged_collection = pd.DataFrame(merged_collection_data, columns=merged_collection_columns)
#                         df_merged_collection['ACCTNO'] = "' " + df_merged_collection['ACCTNO'].astype(str)
#                         df_merged_collection['TRNFACCTNO'] = "' " + df_merged_collection['TRNFACCTNO'].astype(str)
#                         merged_collection_csv_filename = os.path.join(temp_dir, "Scenario_data.csv")
                        
#                         df_merged_collection.to_csv(merged_collection_csv_filename, index=False,quoting=csv.QUOTE_NONNUMERIC,encoding='utf-8')
                        
#                         flattened_data_list = [obj]
                        
#                         df = pd.DataFrame(flattened_data_list)


#                         df.to_csv(temp_csv_filename, index=False)
#                         list_of_dicts = df.to_dict(orient='records')

#                         generate_pdf_from_csv(temp_csv_filename, temp_pdf_filename, list_of_dicts)
#                         generate_txt_from_dataframe(df, temp_txt_filename)
#                         generate_docx_from_dataframe(df, temp_docx_filename)
                        
#                         zipf.write(temp_csv_filename, 'str_data.csv')
#                         zipf.write(temp_pdf_filename, 'str_data.pdf')
#                         zipf.write(temp_txt_filename, 'str_data.txt')
#                         zipf.write(temp_docx_filename, 'str_data.docx')

#                         zipf.write(merged_collection_csv_filename, 'Scenario_data.csv')

#             return send_from_directory(temp_dir, "str_data.zip", as_attachment=True)

#         connAdminonlineSTRpdf.close()
#         return render_template('FINnet_report.html', message='No approved cases found', type='FINnetReports', role='IT OFFICER')
#     except Exception as e:
#         mysqlonlineSTRpdf.rollback()
#         connAdminonlineSTRpdf.close()
#         return f"Something went wrong {e} , please Re-Login Again",500






@app.route('/download_pdf_strr', methods=['POST'])
@secure_route(required_role=['IT OFFICER'])
def download_pdf_strr():
    accNo = request.form.get('accNumber')
    dateSubmited = request.form.get('date')
    TicketId = request.form.get('TicketId')
    txdate = request.form.get('TXDATE')
    txtype = request.form.get('TXTYPE')
    trnsflowtype = request.form.get('TRNFLOWTYPE')
    CUSTCD = request.form.get('CUSTCD')
    scenario_code = request.form.get('scenario_code')


    try:

       
        temp_dir = tempfile.mkdtemp()
        zip_filename = os.path.join(temp_dir, f"STR_{accNo}.zip")

        with zipfile.ZipFile(zip_filename, 'w') as zipf:
            
                    AccountpersonalDetailsCSV = os.path.join(temp_dir, "Account_Personal.csv")
                    AccountDetailsCSV = os.path.join(temp_dir, "Account_Details.csv")
                    kcfile1 = os.path.join(temp_dir, "KC1.csv")
                    kcfile2 = os.path.join(temp_dir, "KC2.csv")
                    Gcfile1 = os.path.join(temp_dir, "GS1.csv")
                    

            

                    debit_details,credit_details,cash_transaction_count,cash_deposits,cash_withdrawals = communFunction(accNo,dateSubmited,txdate)
                    print("gone in..........")
                    
                    AccountPersonalDetails(accNo,dateSubmited,TicketId,txtype,trnsflowtype,CUSTCD,scenario_code,AccountpersonalDetailsCSV)
                    AccountDetails(accNo,dateSubmited,TicketId,txtype,trnsflowtype,CUSTCD,scenario_code,AccountDetailsCSV,debit_details,credit_details,cash_transaction_count,cash_deposits,cash_withdrawals)
                    kc1(accNo,dateSubmited,TicketId,txtype,trnsflowtype,CUSTCD,scenario_code,kcfile1)
                    kc2(accNo,dateSubmited,TicketId,txtype,trnsflowtype,CUSTCD,scenario_code,kcfile2)
                    Gs2(accNo,dateSubmited,TicketId,txtype,trnsflowtype,CUSTCD,scenario_code,Gcfile1,txdate,debit_details,credit_details,cash_transaction_count,cash_deposits,cash_withdrawals)
                    print("got out..........")
                        
                        
                    zipf.write(AccountpersonalDetailsCSV, 'Account_Personal.csv')
                    zipf.write(AccountDetailsCSV, 'Account_Details.csv')
                    zipf.write(kcfile1, 'KC1.csv')
                    zipf.write(kcfile2, 'KC2.csv')
                    zipf.write(Gcfile1, 'GS1.csv')

                    if txtype == 'C':
                        TCfile1 = os.path.join(temp_dir, "TC1.csv")
                        TC1(accNo,dateSubmited,TicketId,txtype,trnsflowtype,CUSTCD,scenario_code,TCfile1,txdate)
                        zipf.write(TCfile1, 'TC1.csv')

                    else:
                        pass
                    

        return send_from_directory(temp_dir, f"STR_{accNo}.zip", as_attachment=True)
    
    except:
        redirect(url_for('online_STR_download_page'))

    

# --------------- STR ZIP CSV FILES FUNCTIONS WHICH WILL BE CALLED IN UPPER END POINT ----------------------------------


def AccountPersonalDetails(ACCNO,DATE_of_SUBMITION_GM,TICKETID,TXTYPE,TRANSFLOWTYPE,CUSTCD,SCENARIO_CODE,filePath):

    # connAdminonlineSTRpdf = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connAdminonlineSTRpdf = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlonlineSTRpdf = connAdminonlineSTRpdf.cursor()

    quiryForAcc = "SELECT ACCTNO,PRIMARY_SOL_ID,CUST_TYPE_CODE FROM CUSTOMERS WHERE ACCTNO = ? "

    personalDetails = mysqlonlineSTRpdf.execute(quiryForAcc,(ACCNO,)).fetchall()


    colum = ["Account Number","Relationship Type","Individual/Non-Individual","Unique Reference Number","Name of Non-Customer"]

    personalPandas = pd.DataFrame.from_records(personalDetails,columns=["Account Number","Relationship Type","Individual/Non-Individual"])
    

    personalPandas["Individual/Non-Individual"] = personalPandas["Individual/Non-Individual"].apply(
    lambda x: "Individual" if x in ["IND", "INDI"] else "Non-Individual"
    )

    personalPandas["Unique Reference Number"] = ""

    personalPandas["Name of Non-Customer"] = ""

    personalPandas = personalPandas[colum]


    personalPandas.to_csv(filePath,index=False)

    connAdminonlineSTRpdf.close()
    



# ////////////////////////////////// Preceding TXDATE of the first Report ////////////////////////////////

#     WITH FirstTransactionDate AS (
#     SELECT 
#         MIN(TRY_CONVERT(datetime, TXDATE, 105)) AS FirstDate
#     FROM 
#         Transactions
#     WHERE 
#         AcctNo = '400632003000012'
# )
# SELECT 
#     COUNT(*) AS CashTransactionCount
# FROM 
#     Transactions
# WHERE 
#     TRNFLOWTYPE = 'C'
#     AND AcctNo = '400632003000012'
#     AND TRY_CONVERT(datetime, TXDATE, 105) >= DATEADD(MONTH, -12, (SELECT FirstDate FROM FirstTransactionDate))
#     AND TRY_CONVERT(datetime, TXDATE, 105) < (SELECT FirstDate FROM FirstTransactionDate);

# ////////////////////////////////// Preceding TXDATE of the first Report ////////////////////////////////

def communFunction(ACCNO,DATE_of_SUBMITION_GM,txdate):
    quiryForAccTRD = """SELECT 
                    COUNT(*) AS DebitCount,
                    SUM(CONVERT(FLOAT, TXAMT)) AS TotalDebitAmount
                FROM 
                    TRANSACTIONS
                WHERE 
                    TRNFLOWTYPE = 'D' AND ACCTNO = ?
                    AND TRY_CONVERT(datetime, TXDATE, 105) >= DATEADD(MONTH, -12, ?)
                GROUP BY 
                    ACCTNO;
                """
    
    quiryForAccTRC = """SELECT 
                    COUNT(*) AS CrditCount,
                    SUM(CONVERT(FLOAT, TXAMT)) AS TotalCreditAmount
                FROM 
                    TRANSACTIONS
                WHERE 
                    TRNFLOWTYPE = 'C' AND ACCTNO = ?
                    AND TRY_CONVERT(datetime, TXDATE, 105) >= DATEADD(MONTH, -12, ?)
                GROUP BY 
                    ACCTNO;
                """
    
    quiryForAccTRCcount = """SELECT 
                    COUNT(*) AS CashTransactionCount
                FROM 
                    TRANSACTIONS
                WHERE 
                    TXTYPE = 'C' AND ACCTNO = ?
                    AND TRY_CONVERT(datetime, TXDATE, 105) >= DATEADD(MONTH, -12, ?)
                GROUP BY 
                    ACCTNO;
                """
    

    quiryForAccTRCdeposits = """SELECT 
                    SUM(CONVERT(FLOAT, TXAMT)) AS TotalCashDeposits
                FROM 
                    TRANSACTIONS
                WHERE 
                    TXTYPE = 'C' AND TRNFLOWTYPE = 'C' AND ACCTNO = ?
                    AND TRY_CONVERT(datetime, TXDATE, 105) >= DATEADD(MONTH, -12, ?)
                GROUP BY 
                    ACCTNO;
                """
    

    quiryForAccTRCWith = """SELECT 
                    SUM(CONVERT(FLOAT, TXAMT)) AS TotalCashWithdrawals
                FROM 
                    TRANSACTIONS
                WHERE 
                    TXTYPE = 'C' AND TRNFLOWTYPE = 'D' AND ACCTNO = ?
                    AND TRY_CONVERT(datetime, TXDATE, 105) >= DATEADD(MONTH, -12, ?)
                GROUP BY 
                    ACCTNO;
                """
    

    try:
        divDate = txdate.split(' to ')
        presentDate1 = datetime.strptime(divDate[1], '%Y-%m-%d').strftime('%Y-%m-%d')
        pastDate1 = datetime.strptime(divDate[0], '%Y-%m-%d').strftime('%Y-%m-%d')
    except:
        
        parsed_date = datetime.strptime(txdate, '%d-%m-%Y')
        presentDate1 = parsed_date.strftime('%Y-%m-%d')

    print(presentDate1)



    debit_details = fetch_single_result(quiryForAccTRD, (ACCNO,presentDate1 ), (0, 0.0))
    credit_details = fetch_single_result(quiryForAccTRC, (ACCNO,presentDate1 ), (0, 0.0))
    cash_transaction_count = fetch_single_result(quiryForAccTRCcount, (ACCNO,presentDate1 ), (0,))
    cash_deposits = fetch_single_result(quiryForAccTRCdeposits, (ACCNO,presentDate1 ), (0.0,))
    cash_withdrawals = fetch_single_result(quiryForAccTRCWith, (ACCNO, presentDate1), (0.0,))

    return debit_details,credit_details,cash_transaction_count,cash_deposits,cash_withdrawals



def fetch_single_result(query, params, default):
    # connAdminonlineSTRpdf = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connAdminonlineSTRpdf = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlonlineSTRpdf = connAdminonlineSTRpdf.cursor()

    result = mysqlonlineSTRpdf.execute(query, params).fetchone()
    return result if result else default


def AccountDetails(ACCNO,DATE_of_SUBMITION_GM,TICKETID,TXTYPE,TRANSFLOWTYPE,CUSTCD,SCENARIO_CODE,filePath,debit_details,credit_details,cash_transaction_count,cash_deposits,cash_withdrawals):


    # connAdminonlineSTRpdf = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    # connAdminonlineSTRpdf = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connAdminonlineSTRpdf = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")


    quiryForAcc = "SELECT CUST_CONSTITUTION,ACCTNO,ACCT_STATUS,NULL AS Reason_for_Account_Freeze,ACCT_OPENDATE,ACCT_CLOSEDATE,PRIMARY_SOL_ID FROM CUSTOMERS WHERE ACCTNO = ? "

   
   
    
    account_details = fetch_single_result(quiryForAcc, (ACCNO,), (None, None, None, None, None, None, None))
   

    # Combine all details into one list
    combined_details = [
        account_details[0],  # Account Type
        account_details[1],  # Account Number
        account_details[2],  # Account Status
        account_details[3],  # Reason for Account Freeze
        account_details[4],  # Date of Account Opening
        account_details[5],  # Date of Account Closing
        account_details[6],  # Branch Code of Account
        debit_details[0],    # No Of Debits (In last 12 months)
        debit_details[1],    # Total Debit Amount (In last 12 months) Amount
        credit_details[0],   # No Of Credits (In last 12 months)
        credit_details[1],   # Total Credit Amount (In last 12 months) Amount
        cash_transaction_count[0],  # No Of Cash Transaction (In last 12 months)
        cash_deposits[0],    # Total Cash Deposit (In last 12 months) Amount
        cash_withdrawals[0]  # Total Cash Withdrawal (In last 12 months) Amount
    ]


    print("personalDetails : ",combined_details)

    colum = ["Account Type","Account Number","Account Status","Reason for Account Freeze","Date of Account Opening","Date of Account Closing","Branch Code of Account","No Of Debits (In last 12 months)","Total Debit Amount (In last 12 months) Amount","No Of Credits (In last 12 months)","Total Credit Amount (In last 12 months) Amount","No Of Cash Transaction (In last 12 months)","Total Cash Deposit (In last 12 months) Amount","Total Cash Withdrawal (In last 12 months) Amount"]

    personalPandas = pd.DataFrame([combined_details],columns=colum)
    


    personalPandas.to_csv(filePath,index=False)







def kc1(ACCNO,DATE_of_SUBMITION_GM,TICKETID,TXTYPE,TRANSFLOWTYPE,CUSTCD,SCENARIO_CODE,filePath):

    # connAdminonlineSTRpdf = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connAdminonlineSTRpdf = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlonlineSTRpdf = connAdminonlineSTRpdf.cursor()
    print("in........")


    quiryForAcc = "SELECT CustomerName,NULL AS MIDDLEnAME,NULL AS LASTNAME,'Full Name Is Provided On First Name' as decleration,Gender,'INDIAN' AS nationality,NULL AS FATHER,NULL AS MOTHERNAME,NULL AS SPOUSE, NULL AS MOBILE,NULL AS ALTERNATIVEMOBILE,NULL AS TELE,NULL AS EMAILID,DOB,Perm_ADDRESSRECTYPE,Permanent_Address,COUNTRY,PINCODE,STATE,NULL AS DISTRICT,NULL AS CITY,Comm_ADDRESSRECTYPE,Communication_Address,COUNTRY,PINCODE,STATE,NULL AS DISTSEC, NULL AS CITYSEC,CUSTCD,PAN,NULL AS DECLAIREPAN,CKYC,NULL AS CKYCDECLAIR,Othr_doc,VoterID,AADHARNO,'INDIVIDUAL' AS CUSTTYPE,NULL AS OTHERCUSTTYPE,CUSTCD,NULL AS EMPNAME , NULL AS EMPADDRESS,NULL AS EMPLOC,NULL AS EMPCOUNTRY,NULL AS EMPPIN , NULL AS EMPSTATE,NULL AS EMPDIST,NULL AS EMPCITY,Occupation,NULL AS OTHEROccupation,ACCT_OPENDATE,RIP,KYC_UpdateDate FROM CUSTOMERS WHERE ACCTNO = ? AND CUST_TYPE_CODE IN ('IND','INDI')"

    personalDetails = mysqlonlineSTRpdf.execute(quiryForAcc,(ACCNO,)).fetchall()

    print("out........")

    colum = ["First name","Middle name","Last name","Declaration (If Last name is not available)","Gender","Nationality","Name of Father","Name of Mother","Spouse/Partner Name","Mobile Number","Alternate Mobile Number","Telephone Number","Email ID","Date of Birth","Primary Address 1","Primary Address Locality","Primary Address Country","Primary Address Pin Code","Primary Address State","Primary Address District","Primary Address City / Village / Town","Secondary Address 1","Secondary Address Locality","Secondary Address Country","econdary Address Pin Code","Secondary Address State","Secondary Address District","Secondary Address City / Village / Town","UCIC","PAN","Declaration (If PAN is not available)","CKYC number","Declaration (If CKYC is not available)","OTHER DOC","Voter ID","Identity verified using Aadhaar ID","Customer Type","Other Customer Type","Customer ID","Employer Name","Employer Address 1","Employer Address Locality","Employer Address Country","Employer Address Pin Code","Employer Address State","Employer Address District","Employer Address City / Village / Town","Occupation","Other Occupation","Date of Customer On-boarding","Customer Risk Level","Date of last KYC / re-KYC"]

    personalPandas = pd.DataFrame.from_records(personalDetails,columns=colum)



    personalPandas.to_csv(filePath,index=False)

    connAdminonlineSTRpdf.close()




def kc2(ACCNO,DATE_of_SUBMITION_GM,TICKETID,TXTYPE,TRANSFLOWTYPE,CUSTCD,SCENARIO_CODE,filePath):

    # connAdminonlineSTRpdf = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connAdminonlineSTRpdf = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlonlineSTRpdf = connAdminonlineSTRpdf.cursor()


    quiryForAcc = "SELECT CUSTCD,CustomerName, NULL AS MOBILE,NULL AS TELE,NULL AS EMAILID,NULL AS WEBSIT,Perm_ADDRESSRECTYPE,Permanent_Address,COUNTRY,PINCODE,STATE,NULL AS DISTRICT,NULL AS CITY,CUST_TYPE_CODE,NULL AS OTHERSCUSTTYPE,RIP,ACCT_OPENDATE,KYC_UpdateDate,DateofRegistration,NULL AS REGADRESS,NULL AS REGLOC,CountryofRegistration,PINCODE,STATE,NULL AS REGDIST,NULL AS REGCITY,NULL AS FCRASTATUS,'Not Available' AS FCRA,NULL AS REGFCRADATE,STATE,Segment,NULL AS UBO,NULL AS CUSTIDTYPE,NULL AS CUSTNUMBER,GISTIN_NO,PAN,NULL AS DECLAIREPAN,NULL AS PEKRN,NULL AS TAN, NULL AS IEC,NULL AS TANDECL FROM CUSTOMERS WHERE ACCTNO = ? AND CUST_TYPE_CODE NOT IN ('IND','INDI')"

    personalDetails = mysqlonlineSTRpdf.execute(quiryForAcc,(ACCNO,)).fetchall()


    colum = ["UCIC","Entity Name","Mobile Number","Telephone Number","Email ID","Company website","Address Linel","Locality","Country","Pin Code","state","District","City / Village / Town","Customer Type","Other Customer Type","Customer Risk Level","On boarding Date","Last KYC Date","Date of Incorporation","Registered Address Line1","Registered Locality","Registered Country","Registered Pin Code","Registered State","Registered District","Registered City /Village / Town","FCRA Status","FCRA Registration Number","FCRA Registration","State","Lines of Business","Declaration (If UBO is not available)","Company ID Type","Company ID Number","GSTIN","PAN","Declaration (If PAN is not available)","PEKRN","TAN","IEC","Declaration (If IEC is not available)"]

    personalPandas = pd.DataFrame.from_records(personalDetails,columns=colum)
    

    personalPandas.to_csv(filePath,index=False)

    connAdminonlineSTRpdf.close()






def Gs2(ACCNO,DATE_of_SUBMITION_GM,TICKETID,TXTYPE,TRANSFLOWTYPE,CUSTCD,SCENARIO_CODE,filePath,txdate,debit_details,credit_details,cash_transaction_count,cash_deposits,cash_withdrawals):

    # connAdminonlineSTRpdf = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connAdminonlineSTRpdf = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlonlineSTRpdf = connAdminonlineSTRpdf.cursor()



    accQuires = """SELECT 
                    SUM(CONVERT(FLOAT, TXAMT)) AS TotalDebitAmount
                FROM 
                    TRANSACTIONS
                WHERE 
                    TRNFLOWTYPE = 'D' AND ACCTNO = ?
                    AND TRY_CONVERT(datetime, TXDATE, 105) >= DATEADD(MONTH, -12, TRY_CONVERT(datetime, TXDATE, 105))
                GROUP BY 
                    ACCTNO;
                """

    allDebits = mysqlonlineSTRpdf.execute(accQuires,(ACCNO,)).fetchone()[0]

    accQuirescrid = """SELECT 
                    SUM(CONVERT(FLOAT, TXAMT)) AS TotalCriditstAmount
                FROM 
                    TRANSACTIONS
                WHERE 
                    TRNFLOWTYPE = 'C' AND ACCTNO = ?
                    AND TRY_CONVERT(datetime, TXDATE, 105) >= DATEADD(MONTH, -12, TRY_CONVERT(datetime, TXDATE, 105))
                GROUP BY 
                    ACCTNO;
                """

    allCredit = mysqlonlineSTRpdf.execute(accQuirescrid,(ACCNO,)).fetchone()[0]


    totalIncomePerYear = allCredit - allDebits



    codes = ['TM_1_1','TM_1_2','TM_1_3','TM_1_4','TM_2_1','TM_2_2','TM_2_3','TM_2_4']

    if SCENARIO_CODE in codes:
        mainormulti = 'Main';
    else:
        mainormulti = "associate"


    debit_detailsCount  = debit_details[0]  # No Of Debits (In last 12 months)
    debit_detailsAmount =    debit_details[1]    # Total Debit Amount (In last 12 months) Amount
    credit_detailscount =  credit_details[0]   # No Of Credits (In last 12 months)
    credit_detailsamount = credit_details[1]   # Total Credit Amount (In last 12 months) Amount
    cash_transaction_count1 = cash_transaction_count[0]  # No Of Cash Transaction (In last 12 months)
    cash_deposits1 =    cash_deposits[0]    # Total Cash Deposit (In last 12 months) Amount
    cash_withdrawals1 =    cash_withdrawals[0]

    if TXTYPE == 'C':
        try:
            divDate = txdate.split(' to ')
            presentDate1 = datetime.strptime(divDate[1], '%Y-%m-%d').strftime('%Y-%m-%d')
            pastDate1 = datetime.strptime(divDate[0], '%Y-%m-%d').strftime('%Y-%m-%d')

            query1 = """
                SELECT COUNT(*) AS COUNTC,SUM(CONVERT(FLOAT, TXAMT)) AS TotalCashAmount
                FROM TRANSACTIONS
                WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TRY_CONVERT(DATE, TXDATE, 105) BETWEEN ? AND ? 
            """
            mysqlonlineSTRpdf.execute(query1, (ACCNO,TXTYPE,TRANSFLOWTYPE,pastDate1, presentDate1))
            scenarioTransactionData = mysqlonlineSTRpdf.fetchone()
            cashTotalD = None
            cashTotalC = None
            if TRANSFLOWTYPE == 'D':
                cashTotalD = scenarioTransactionData[0]
            if TRANSFLOWTYPE == 'C':
                cashTotalC = scenarioTransactionData[0]
            
            cashTotalcount = scenarioTransactionData[0]
            cashTotal = scenarioTransactionData[1]

            

        except:
            query1 = """
                SELECT COUNT(*) AS COUNTC,SUM(CONVERT(FLOAT, TXAMT)) AS TotalCashAmount 
                FROM TRANSACTIONS
                WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TXDATE = ? 
            """
            mysqlonlineSTRpdf.execute(query1, (ACCNO,TXTYPE,TRANSFLOWTYPE,txdate))
            scenarioTransactionData = mysqlonlineSTRpdf.fetchone()
            cashTotalD = None
            cashTotalC = None
            if TRANSFLOWTYPE == 'D':
                cashTotalD = scenarioTransactionData[0]
            if TRANSFLOWTYPE == 'C':
                cashTotalC = scenarioTransactionData[0]
            
            cashTotalcount = scenarioTransactionData[0]
            cashTotal = scenarioTransactionData[1]


    else:
        cashTotalD = None
        cashTotalC = None
        cashTotalcount = None
        cashTotal = None


    try:
        divDate = txdate.split(' to ')
        presentDate1 = datetime.strptime(divDate[1], '%Y-%m-%d').strftime('%Y-%m-%d')
        pastDate1 = datetime.strptime(divDate[0], '%Y-%m-%d').strftime('%Y-%m-%d')

        query1 = """
            SELECT COUNT(*) AS COUNTC,SUM(CONVERT(FLOAT, TXAMT)) AS TotalCashAmount
            FROM TRANSACTIONS
            WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TRY_CONVERT(DATE, TXDATE, 105) BETWEEN ? AND ? AND TRNFACCTNO LIKE '%ATM%'
        """
        mysqlonlineSTRpdf.execute(query1, (ACCNO,TXTYPE,TRANSFLOWTYPE,pastDate1, presentDate1))
        scenarioTransactionData = mysqlonlineSTRpdf.fetchone()
        cardTotalD=None
        cardTotalC=None
        if TRANSFLOWTYPE == 'D':
            cardTotalD = scenarioTransactionData[0]
        if TRANSFLOWTYPE == 'C':
            cardTotalC = scenarioTransactionData[0]
        
        cardTotalcount = scenarioTransactionData[0]
        cardTotal = scenarioTransactionData[1]

        

    except:
        query1 = """
            SELECT COUNT(*) AS COUNTC,SUM(CONVERT(FLOAT, TXAMT)) AS TotalCashAmount 
            FROM TRANSACTIONS
            WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TXDATE = ? AND TRNFACCTNO LIKE '%ATM%'
        """
        mysqlonlineSTRpdf.execute(query1, (ACCNO,TXTYPE,TRANSFLOWTYPE,txdate))
        scenarioTransactionData = mysqlonlineSTRpdf.fetchone()

        cardTotalD=None
        cardTotalC=None
        if TRANSFLOWTYPE == 'D':
            cardTotalD = scenarioTransactionData[0]
        if TRANSFLOWTYPE == 'C':
            cardTotalC = scenarioTransactionData[0]
        
        cardTotalcount = scenarioTransactionData[0]
        cardTotal = scenarioTransactionData[1]
     


    try:
        divDate = txdate.split(' to ')
        presentDate1 = datetime.strptime(divDate[1], '%Y-%m-%d').strftime('%Y-%m-%d')
        pastDate1 = datetime.strptime(divDate[0], '%Y-%m-%d').strftime('%Y-%m-%d')

        query1 = """
            SELECT COUNT(*) AS COUNTC,SUM(CONVERT(FLOAT, TXAMT)) AS TotalCashAmount
            FROM TRANSACTIONS
            WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TRY_CONVERT(DATE, TXDATE, 105) BETWEEN ? AND ? AND REMARKS LIKE '%IMPS%'
        """
        mysqlonlineSTRpdf.execute(query1, (ACCNO,TXTYPE,TRANSFLOWTYPE,pastDate1, presentDate1))
        scenarioTransactionData = mysqlonlineSTRpdf.fetchone()


        impsTotalD = None
        impsTotalC = None
        if TRANSFLOWTYPE == 'D':
            impsTotalD = scenarioTransactionData[0]
        if TRANSFLOWTYPE == 'C':
            impsTotalC = scenarioTransactionData[0]
        
        impsTotalcount = scenarioTransactionData[0]
        impsTotal = scenarioTransactionData[1]

        

    except:
        query1 = """
            SELECT COUNT(*) AS COUNTC,SUM(CONVERT(FLOAT, TXAMT)) AS TotalCashAmount 
            FROM TRANSACTIONS
            WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TXDATE = ? AND REMARKS LIKE '%IMPS%'
        """
        mysqlonlineSTRpdf.execute(query1, (ACCNO,TXTYPE,TRANSFLOWTYPE,txdate))
        scenarioTransactionData = mysqlonlineSTRpdf.fetchone()

        impsTotalD = None
        impsTotalC = None
        if TRANSFLOWTYPE == 'D':
            impsTotalD = scenarioTransactionData[0]
        if TRANSFLOWTYPE == 'C':
            impsTotalC = scenarioTransactionData[0]
        
        impsTotalcount = scenarioTransactionData[0]
        impsTotal = scenarioTransactionData[1]

    


    try:
        divDate = txdate.split(' to ')
        presentDate1 = datetime.strptime(divDate[1], '%Y-%m-%d').strftime('%Y-%m-%d')
        pastDate1 = datetime.strptime(divDate[0], '%Y-%m-%d').strftime('%Y-%m-%d')

        query1 = """
            SELECT COUNT(*) AS COUNTC,SUM(CONVERT(FLOAT, TXAMT)) AS TotalCashAmount
            FROM TRANSACTIONS
            WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TRY_CONVERT(DATE, TXDATE, 105) BETWEEN ? AND ? AND (REMARKS LIKE '%NEFT%' OR REMARKS LIKE '%RTGS%')
        """
        mysqlonlineSTRpdf.execute(query1, (ACCNO,TXTYPE,TRANSFLOWTYPE,pastDate1, presentDate1))
        scenarioTransactionData = mysqlonlineSTRpdf.fetchone()


        nftsTotalD = None
        nftsTotalC = None
        if TRANSFLOWTYPE == 'D':
            nftsTotalD = scenarioTransactionData[0]
        if TRANSFLOWTYPE == 'C':
            nftsTotalC = scenarioTransactionData[0]
        
        nftsTotalcount = scenarioTransactionData[0]
        nftsTotal = scenarioTransactionData[1]

        

    except:
        query1 = """
            SELECT COUNT(*) AS COUNTC,SUM(CONVERT(FLOAT, TXAMT)) AS TotalCashAmount 
            FROM TRANSACTIONS
            WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TXDATE = ? AND (REMARKS LIKE '%NEFT%' OR REMARKS LIKE '%RTGS%')
        """
        mysqlonlineSTRpdf.execute(query1, (ACCNO,TXTYPE,TRANSFLOWTYPE,txdate))
        scenarioTransactionData = mysqlonlineSTRpdf.fetchone()

        nftsTotalD = None
        nftsTotalC = None
        if TRANSFLOWTYPE == 'D':
            nftsTotalD = scenarioTransactionData[0]
        if TRANSFLOWTYPE == 'C':
            nftsTotalC = scenarioTransactionData[0]
        
        nftsTotalcount = scenarioTransactionData[0]
        nftsTotal = scenarioTransactionData[1]


    
    CBWTTotalD = None
    CBWTTotalC = None
    CBWTTotalcount = None
    CBWTTotal = None




    try:
        divDate = txdate.split(' to ')
        presentDate1 = datetime.strptime(divDate[1], '%Y-%m-%d').strftime('%Y-%m-%d')
        pastDate1 = datetime.strptime(divDate[0], '%Y-%m-%d').strftime('%Y-%m-%d')

        query1 = """
            SELECT COUNT(*) AS COUNTC,SUM(CONVERT(FLOAT, TXAMT)) AS TotalCashAmount
            FROM TRANSACTIONS
            WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TRY_CONVERT(DATE, TXDATE, 105) BETWEEN ? AND ? 
        """
        mysqlonlineSTRpdf.execute(query1, (ACCNO,TXTYPE,TRANSFLOWTYPE,pastDate1, presentDate1))
        scenarioTransactionData = mysqlonlineSTRpdf.fetchone()


        gtTotalD = None
        gtTotalC = None
        if TRANSFLOWTYPE == 'D':
            gtTotalD = scenarioTransactionData[0]
        if TRANSFLOWTYPE == 'C':
            gtTotalC = scenarioTransactionData[0]
        
        gtTotalcount = scenarioTransactionData[0]
        gtTotal = scenarioTransactionData[1]

        

    except:
        query1 = """
            SELECT COUNT(*) AS COUNTC,SUM(CONVERT(FLOAT, TXAMT)) AS TotalCashAmount 
            FROM TRANSACTIONS
            WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TXDATE = ? 
        """
        mysqlonlineSTRpdf.execute(query1, (ACCNO,TXTYPE,TRANSFLOWTYPE,txdate))
        scenarioTransactionData = mysqlonlineSTRpdf.fetchone()

        gtTotalD = None
        gtTotalC = None
        if TRANSFLOWTYPE == 'D':
            gtTotalD = scenarioTransactionData[0]
        if TRANSFLOWTYPE == 'C':
            gtTotalC = scenarioTransactionData[0]
        
        gtTotalcount = scenarioTransactionData[0]
        gtTotal = scenarioTransactionData[1]



    

    quire = f"SELECT CustomerName,Permanent_Address,PAN,CKYC,RIP,CUST_CONSTITUTION,'{totalIncomePerYear}' AS incomePerYear,'{mainormulti}' AS ROLE,'{cashTotalD}' AS cashTotalDebits,'{cashTotalC}' AS cashTotalDebits,'{cashTotalcount}' AS cashTotalCount,'{cashTotal}' AS cashTotalAMT,'{cardTotalD}' AS cardTotalDebits,'{cardTotalC}' AS cardTotalDebits,'{cardTotalcount}' AS cardTotalCount,'{cardTotal}' AS cardTotalAMT,'{impsTotalD}' AS impsTotalDebits,'{impsTotalC}' AS impsTotalDebits,'{impsTotalcount}' AS impsTotalCount,'{impsTotal}' AS impsTotalAMT,'{nftsTotalD}' AS nftsTotalDebits,'{nftsTotalC}' AS nftsTotalDebits,'{nftsTotalcount}' AS nftsTotalCount,'{nftsTotal}' AS nftsTotalAMT,'{CBWTTotalD}' AS CBWTTotalDebits,'{CBWTTotalC}' AS impsTotalDebits,'{CBWTTotalcount}' AS impsTotalCount,'{CBWTTotal}' AS impsTotalAMT,'{gtTotalD}' AS gtTotalDebits,'{gtTotalC}' AS gtTotalDebits,'{gtTotalcount}' AS impsTotalCount,'{gtTotal}' AS impsTotalAMT,'{debit_detailsCount}' AS DEBITCOUNT,'{debit_detailsAmount}' AS DEBITAMOUNT,'{credit_detailscount}' AS CREDCOUNT,'{credit_detailsamount}' AS CREDAMOUNT,'{cash_transaction_count1}' AS COUNTCASH,'{cash_deposits1}' AS CASHDEPO,'{cash_withdrawals1}' AS WITHCASH  FROM  CUSTOMERS  WHERE  ACCTNO = ? "

    suspisusGS = mysqlonlineSTRpdf.execute(quire,(ACCNO,)).fetchall()



    colum = ["Name","Address","PAN","Unique Id (CKYC/GSTN/CIN)","Risk","Profession/ LOB","Income (per year)","Role (Main / associate)","Cash Total Debits","Cash Total Credits","Cash Total Count","Cash Total Amount","Card Total Debits","Card Total Credits","Card Total Count","Card Total Amount","IMPS Total Debits","IMPS Total Credits","IMPS Total Count","IMPS Total Amount","NEFt/RTG Total Debits","NEFT/RTG Total Credits","NEFT/RTG Total Count","NEFT/RTG Total Amount","CDET Total Debits","CDWT Total Credits","CDWT Total Count","CDWT Total Amount","General Transaction Total Debits","General Transaction Total Credits","General Transaction Total Count","General Transaction Total Amount","No Of Debits (In last 12 months)","Total Debit Amount (In last 12 months) Amount","No Of Credits (In last 12 months)","Total Credit Amount (In last 12 months) Amount","No Of Cash Transaction (In last 12 months)","Total Cash Deposit (In last 12 months) Amount","Total Cash Withdrawal (In last 12 months) Amount"]

    personalPandas = pd.DataFrame.from_records(suspisusGS,columns=colum)


    second_query1 = f"""SELECT SuspeciousDuoCrime,SuspeciousDuoComplexTr,SuspeciousDuoNoeco,terrorisumFunding FROM scenarios WHERE ACCTNO = ? AND scenario_code = ? AND TXDATE = ?  """  
    second_queryRes = mysqlonlineSTRpdf.execute(second_query1, (ACCNO,SCENARIO_CODE,txdate)).fetchone()

    print(second_queryRes)

    susDueTo = None
    if second_queryRes[0] == 'Yes':
        susDueTo = "Suspicion Due to Proceeds of crime"
    if second_queryRes[1] == 'Yes':
        susDueTo = "Suspicion Due to Unusual or complex transactions"
    if second_queryRes[2] == 'Yes':
        susDueTo = "Suspicion Due to No economic rationale or Bonafide purposes"
    if second_queryRes[3] == 'Yes':
        susDueTo = "Suspicion Due to (Suspicion of) financing of terrorism"
    



    second_scenarios = f"""SELECT Alert_title,Current_values FROM Thresholds WHERE code = ? """  
    second_scenarios = mysqlonlineSTRpdf.execute(second_scenarios, (SCENARIO_CODE,)).fetchone()

    susSummery = f"{second_scenarios[0]},with the threshould limit for ,{second_scenarios[1]}"



    second_query = f"""SELECT TXAMT,TXTYPE,TXDATE,TRNFLOWTYPE,ACCTNO,CustomerName,NULL AS DESTINATION,'{susDueTo}' AS SUSDUE,Investigation,dgmprefilled,'{susSummery}' AS susSum,dgmprefilled ,dgmcomment,fileName FROM scenarios WHERE ACCTNO = ? AND scenario_code = ? AND TXDATE = ?  """  
    second_query_results = mysqlonlineSTRpdf.execute(second_query, (ACCNO,SCENARIO_CODE,txdate)).fetchall()

    second_columns = ["Total Violated Amount","TRANSACTION TYPE","TXDATE","Transaction Flow Type/Catogary","Account Number","Source of Funds", "Destination of funds", 'Suspicion due to','Source of alert','Red Flag Indicator','Type of Suspicion',"Queries (One or more)","Narration","Attachments"]  
    second_df = pd.DataFrame.from_records(second_query_results, columns=second_columns)

    # Concatenate the new DataFrame to the existing DataFrame
    combined_df = pd.concat([personalPandas, second_df], ignore_index=True)

    # Save the combined DataFrame to the CSV file
    combined_df.to_csv(filePath, index=False)
    
    connAdminonlineSTRpdf.close()





def TC1(ACCNO,DATE_of_SUBMITION_GM,TICKETID,TXTYPE,TRANSFLOWTYPE,CUSTCD,SCENARIO_CODE,filePath,txdate):

    # connAdminonlineSTRpdf = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connAdminonlineSTRpdf = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlonlineSTRpdf = connAdminonlineSTRpdf.cursor()


    if TXTYPE == 'C':
        try:
            divDate = txdate.split(' to ')
            presentDate1 = datetime.strptime(divDate[1], '%Y-%m-%d').strftime('%Y-%m-%d')
            pastDate1 = datetime.strptime(divDate[0], '%Y-%m-%d').strftime('%Y-%m-%d')

            query1 = """
                SELECT 'CUSTOMER' AS RELATIONFLAG,TXNNO,TXDATE,NULL AS TXTIME,ACCTNO,TXAMT,TRNFACCTNO,'Not Available' AS NONCUSTREF,ISNULL(CHEQNO, 'Cash') AS CHEQNO,NULL AS INSTID,TRF_SOL_ID,'Not Available' AS THIREDpARTYAN,NULL AS DECLEARTAION,TRF_SOL_ID 
                FROM TRANSACTIONS
                WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TRY_CONVERT(DATE, TXDATE, 105) BETWEEN ? AND ? 
            """
            mysqlonlineSTRpdf.execute(query1, (ACCNO,TXTYPE,TRANSFLOWTYPE,pastDate1, presentDate1))
            scenarioTransactionData = mysqlonlineSTRpdf.fetchall()

            second_columns = ['Relationship Flag','Transaction ID','Transaction Date','Transaction Time','Account Number','Deposit/Withdrawal','Non-Customer Reference Number','Instrument Type','Instrument ID','Transaction Branch Code','Thired Party PAN (Non-Customer)','Declaration (If Third Party PAn is not Available)','Branch Code of Account']

            second_df = pd.DataFrame.from_records(scenarioTransactionData, columns=second_columns)


        except:
            query1 = """
                SELECT 'CUSTOMER' AS RELATIONFLAG,TXNNO,TXDATE,NULL AS TXTIME,ACCTNO,TXAMT,TRNFACCTNO,'Not Available' AS NONCUSTREF,ISNULL(CHEQNO, 'Cash') AS CHEQNO,NULL AS INSTID,TRF_SOL_ID,'Not Available' AS THIREDpARTYAN,NULL AS DECLEARTAION,TRF_SOL_ID 
                FROM TRANSACTIONS
                WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TXDATE = ? 
            """
            mysqlonlineSTRpdf.execute(query1, (ACCNO,TXTYPE,TRANSFLOWTYPE,txdate))
            scenarioTransactionData = mysqlonlineSTRpdf.fetchall()

            second_columns = ['Relationship Flag','Transaction ID','Transaction Date','Transaction Time','Account Number','Transaction Amount','Deposit/Withdrawal','Non-Customer Reference Number','Instrument Type','Instrument ID','Transaction Branch Code','Thired Party PAN (Non-Customer)','Declaration (If Third Party PAn is not Available)','Branch Code of Account']

            second_df = pd.DataFrame.from_records(scenarioTransactionData, columns=second_columns)


        second_df.to_csv(filePath, index=False)
    else:
        # No operation for other TXTYPE values
        pass



# def TS1(ACCNO,DATE_of_SUBMITION_GM,TICKETID,TXTYPE,TRANSFLOWTYPE,CUSTCD,SCENARIO_CODE,filePath,txdate):

#      # connAdminonlineSTRpdf = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
#     connAdminonlineSTRpdf = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

#     mysqlonlineSTRpdf = connAdminonlineSTRpdf.cursor()



#     try:
#         divDate = txdate.split(' to ')
#         presentDate1 = datetime.strptime(divDate[1], '%Y-%m-%d').strftime('%Y-%m-%d')
#         pastDate1 = datetime.strptime(divDate[0], '%Y-%m-%d').strftime('%Y-%m-%d')

#         query1 = """
#             SELECT 'CUSTOMER' AS RELATIONFLAG,TXDATE,NULL AS TXTIME,TXNNO,SENDERNAME,BENEFICIARYNAME,NULL AS SENDEIFSC,ACCTNO,NULL AS BENFIIFSC,TRNFACCTNO,TXTYPE,TXAMT,NULL AS NARRATION 
#             FROM TRANSACTIONS
#             WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TRY_CONVERT(DATE, TXDATE, 105) BETWEEN ? AND ? 
#         """
#         mysqlonlineSTRpdf.execute(query1, (ACCNO,TXTYPE,TRANSFLOWTYPE,pastDate1, presentDate1))
#         scenarioTransactionData = mysqlonlineSTRpdf.fetchall()

#         second_columns = ['Relationship Flag','Transaction Date','Transaction Time','Transaction ID','Sender Name','Beneficiary Name','Sender IFSC','Sender Account Number','Beneficiary IFSC','Beneficiary Account Number','Transaction Type','Transaction Amount','Narration']

#         second_df = pd.DataFrame.from_records(scenarioTransactionData, columns=second_columns)

#         res = second_df['Beneficiary Account Number']

#         print(res)

#         if res.isnull().any():  # Check if there are any null values in the 'Beneficiary Account Number' column
#             new_rows = []  # List to hold new rows to be added

#             for index, row in second_df.iterrows():
#                 if pd.isnull(row['Beneficiary Account Number']):  # Check if the 'Beneficiary Account Number' is null for this row
#                     txId = row['Transaction ID']

#                     queryTo = """
#                         SELECT ACCTNO, TXNNO, SENDERNAME,TXAMT
#                         FROM TRANSACTIONS 
#                         WHERE TXNNO = ? AND TXDATE = ? AND TRY_CONVERT(DATE, TXDATE, 105) BETWEEN ? AND ? AND TRNFACCTNO = ?
#                     """

#                     mysqlonlineSTRpdf.execute(queryTo, (txId, pastDate1, presentDate1, TXTYPE, ACCNO))
#                     TrnfAcctno = mysqlonlineSTRpdf.fetchall()

#                     for acctno, txnno, sendername,txamt in TrnfAcctno:
#                         # Create a new row
#                         new_row = row.copy()

#                         # Swap Sender Account Number and Beneficiary Account Number
#                         new_row['Sender Account Number'], new_row['Beneficiary Account Number'] = ACCNO, acctno
                        
#                         # Change Transaction Type
#                         if new_row['Transaction Type'] == 'C':
#                             new_row['Transaction Type'] = 'D'
#                         elif new_row['Transaction Type'] == 'D':
#                             new_row['Transaction Type'] = 'C'

#                         # Set Beneficiary Name as SENDERNAME
#                         new_row['Beneficiary Name'] = sendername
#                         new_row['Transaction Amount'] = txamt

#                         # Append the new row to the list
#                         new_rows.append(new_row)

#                     # Remove the original null row
#                     second_df.drop(index, inplace=True)

#             # Concatenate the new rows to the original DataFrame
#             if new_rows:
#                 second_df = pd.concat([second_df, pd.DataFrame(new_rows)], ignore_index=True)

#         # Reset the index after appending rows
#         second_df.reset_index(drop=True, inplace=True)

#     except:
#         query1 = """
#             SELECT 'CUSTOMER' AS RELATIONFLAG,TXDATE,NULL AS TXTIME,TXNNO,SENDERNAME,BENEFICIARYNAME,NULL AS SENDEIFSC,ACCTNO,NULL AS BENFIIFSC,TRNFACCTNO,TXTYPE,TXAMT,NULL AS NARRATION 
#             FROM TRANSACTIONS
#             WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TXDATE = ? 
#         """
#         mysqlonlineSTRpdf.execute(query1, (ACCNO,TXTYPE,TRANSFLOWTYPE,txdate))
#         scenarioTransactionData = mysqlonlineSTRpdf.fetchall()

#         second_columns = ['Relationship Flag','Transaction Date','Transaction Time','Transaction ID','Sender Name','Beneficiary Name','Sender IFSC','Sender Account Number','Beneficiary IFSC','Beneficiary Account Number','Transaction Type','Transaction Amount','Narration']

#         second_df = pd.DataFrame.from_records(scenarioTransactionData, columns=second_columns)

#         res = second_df['Beneficiary Account Number']

#         if res.isnull().any():  # Check if there are any null values in the 'Beneficiary Account Number' column
#             new_rows = []  # List to hold new rows to be added

#             for index, row in second_df.iterrows():
#                 if pd.isnull(row['Beneficiary Account Number']):  # Check if the 'Beneficiary Account Number' is null for this row
#                     txId = row['Transaction ID']

#                     queryTo = """
#                         SELECT ACCTNO, TXNNO, SENDERNAME,TXAMT
#                         FROM TRANSACTIONS 
#                         WHERE TXNNO = ? AND TXDATE = ? AND TXTYPE = ? AND TRNFACCTNO = ?
#                     """

#                     mysqlonlineSTRpdf.execute(queryTo, (txId, txdate, TXTYPE, ACCNO))
#                     TrnfAcctno = mysqlonlineSTRpdf.fetchall()

#                     for acctno, txnno, sendername,txamt in TrnfAcctno:
#                         # Create a new row
#                         new_row = row.copy()

#                         # Swap Sender Account Number and Beneficiary Account Number
#                         new_row['Sender Account Number'], new_row['Beneficiary Account Number'] = ACCNO, acctno
                        
#                         # Change Transaction Type
#                         if new_row['Transaction Type'] == 'C':
#                             new_row['Transaction Type'] = 'D'
#                         elif new_row['Transaction Type'] == 'D':
#                             new_row['Transaction Type'] = 'C'

#                         # Set Beneficiary Name as SENDERNAME
#                         new_row['Beneficiary Name'] = sendername
#                         new_row['Transaction Amount'] = txamt

#                         # Append the new row to the list
#                         new_rows.append(new_row)

#                     # Remove the original null row
#                     second_df.drop(index, inplace=True)

#             # Concatenate the new rows to the original DataFrame
#             if new_rows:
#                 second_df = pd.concat([second_df, pd.DataFrame(new_rows)], ignore_index=True)

#         # Reset the index after appending rows
#         second_df.reset_index(drop=True, inplace=True)



#     second_df.to_csv(filePath, index=False)
















# ------------------ """ OFFLINE STR DOWNLOAD ZIP PROCESS """   -------------------------------


# @app.route('/download_pdf_offline_str', methods=['POST'])
# @secure_route(required_role=['IT OFFICER'])
# def download_pdf_offline_str():
    
#     accNo = request.form.get('accNumber')
#     dateSubmited = request.form.get('date')
#     TicketId = request.form.get('TicketId')
#     format = request.form.get('format')

#     # connAdminofflineSTRpdf = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
#     connAdminofflineSTRpdf = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

#     mysqlofflineSTRpdf = connAdminofflineSTRpdf.cursor()

#     try:
#         mysqlofflineSTRpdf.execute("SELECT o1.[ticket_id],o1.[Created_Date],o1.[Customerno],o1.[casename],o1.[scenario],o1.[Guidance],o1.[RuleScenario],o1.[personname],o1.[SourceofAlert],o1.[alertindicator],o1.[SuspiciousDueToproceedofCrime],o1.[SuspiciousDueToComplexTranscaction],o1.[SuspiciousDueToNoecoRational],o1.[SuspiciousDueToFinancingTerrorism],o1.[AttemptedTranscaction],o1.[LEAInformed],o1.[PriorityRating],o1.[ReportCoverage],o1.[leadetails],o1.[AdditionalDocument],o1.[Aroundofsuspision],o1.[DetailsofInvestigation],o1.[AccountNumber],o1.[AccountType],o1.[holdername],o1.[AccountHolderType],o1.[AccountStatus],o1.[DateofOpening],o1.[RiskRating],o1.[CummulativeCerditTurnover],o1.[CummulativeDebitTurnover],o1.[CummulativeCashDepositTurnover],o1.[CummulativeCashWithdrawalTurnover],o1.[NoOfTransactionsToBeReported],o1.[TransactionDate],o1.[TransactionsID],o1.[TransactionMode],o1.[DebitCredit],o1.[amount],o1.[TransactionsCurrency],o1.[ProductType],o1.[ProductIdentifiers],o1.[TransactionType],o1.[unit],o1.[Date],o1.[DispositionOfFunds],o1.[RelatedAccountNumber],o1.[RelatedInstitutionName],o1.[Remark],o1.[ROS_cmt],o1.[Created_By],o1.[DGM_cmt] FROM [offline_collection] o1, [offline_tickets] o2 WHERE o1.[ticket_id] = o2.[ticketId] AND o2.[approved] = 1 AND [ROS_cmt] IS NOT NULL AND [DGM_cmt] IS NOT NULL AND o1.[ticket_id] = ?", (TicketId,))
        
#         rows = mysqlofflineSTRpdf.fetchall()

#         if rows:
#             temp_dir = tempfile.mkdtemp()
#             zip_file_path = os.path.join(temp_dir, "offline_data.zip")

#             with zipfile.ZipFile(zip_file_path, 'w') as zipf:
#                 for row in rows:
#                     flattened_data = flatten_data(row)
#                     temp_csv_filename = os.path.join(temp_dir, "str_offline_data.csv")
#                     temp_pdf_filename = os.path.join(temp_dir, "str_offline_data.pdf")
#                     temp_docx_filename = os.path.join(temp_dir, "str_offline_data.docx")
#                     temp_txt_filename = os.path.join(temp_dir, "str_offline_data.txt")

#                     # Generate CSV
#                     df = pd.DataFrame([flattened_data])
#                     df.to_csv(temp_csv_filename, index=False)
#                     zipf.write(temp_csv_filename, os.path.basename(temp_csv_filename))


#                     generate_pdf_from_csv(temp_csv_filename, temp_pdf_filename, [flattened_data])
#                     zipf.write(temp_pdf_filename, os.path.basename(temp_pdf_filename))

#                     generate_docx_from_dataframe(flattened_data, temp_docx_filename)
#                     zipf.write(temp_docx_filename, os.path.basename(temp_docx_filename))

#                     generate_txt_from_dataframe(flattened_data, temp_txt_filename)
#                     zipf.write(temp_txt_filename, os.path.basename(temp_txt_filename))

#             if os.path.exists(zip_file_path):
#                 return send_from_directory(temp_dir, "offline_data.zip", as_attachment=True)
        
#         connAdminofflineSTRpdf.close()

#         return render_template('FINnet_report.html', message='No approved cases found', type='FINnetReports', role='IT OFFICER')

#     except Exception as e:
#         mysqlofflineSTRpdf.rollback()
#         connAdminofflineSTRpdf.close()
#         return f"Something went wrong {e} , please Re-Login Again",500


@app.route('/download_pdf_offline_str', methods=['POST'])
@secure_route(required_role=['IT OFFICER'])
def download_pdf_offline_str():
    
    accNo = request.form.get('accNumber')
    dateSubmited = request.form.get('date')
    TicketId = request.form.get('TicketId')
    format = request.form.get('format')  # corrected variable name

    # cursor.execute("SELECT o1.[ticket_id],o1.[Created_Date],o1.[Customerno],o1.[casename],o1.[scenario],o1.[Guidance],o1.[RuleScenario],o1.[personname],o1.[SourceofAlert],o1.[alertindicator],o1.[SuspiciousDueToproceedofCrime],o1.[SuspiciousDueToComplexTranscaction],o1.[SuspiciousDueToNoecoRational],o1.[SuspiciousDueToFinancingTerrorism],o1.[AttemptedTranscaction],o1.[LEAInformed],o1.[PriorityRating],o1.[ReportCoverage],o1.[leadetails],o1.[AdditionalDocument],o1.[Aroundofsuspision],o1.[DetailsofInvestigation],o1.[AccountNumber],o1.[AccountType],o1.[holdername],o1.[AccountHolderType],o1.[AccountStatus],o1.[DateofOpening],o1.[RiskRating],o1.[CummulativeCerditTurnover],o1.[CummulativeDebitTurnover],o1.[CummulativeCashDepositTurnover],o1.[CummulativeCashWithdrawalTurnover],o1.[NoOfTransactionsToBeReported],o1.[TransactionDate],o1.[TransactionsID],o1.[TransactionMode],o1.[DebitCredit],o1.[amount],o1.[TransactionsCurrency],o1.[ProductType],o1.[ProductIdentifiers],o1.[TransactionType],o1.[unit],o1.[Date],o1.[DispositionOfFunds],o1.[RelatedAccountNumber],o1.[RelatedInstitutionName],o1.[Remark],o1.[ROS_cmt],o1.[Created_By],o1.[DGM_cmt] FROM [dbo].[offline_collection] o1, [dbo].[offline_tickets] o2 WHERE o1.[ticket_id] = o2.[ticketId] AND o2.[approved] = 1 AND [ROS_cmt] IS NOT NULL AND [DGM_cmt] IS NOT NULL AND o1.[ticket_id] = ?", (TicketId,))

    cursor.execute("SELECT [ticket_id],[Created_Date],[Customerno],[casename],[scenario],[Guidance],[RuleScenario],[personname],[SourceofAlert],[alertindicator],[SuspiciousDueToproceedofCrime],[SuspiciousDueToComplexTranscaction],[SuspiciousDueToNoecoRational],[SuspiciousDueToFinancingTerrorism],[AttemptedTranscaction],[LEAInformed],[PriorityRating],[ReportCoverage],[leadetails],[AdditionalDocument],[Aroundofsuspision],[DetailsofInvestigation],[AccountNumber],[AccountType],[holdername],[AccountHolderType],[AccountStatus],[DateofOpening],[RiskRating],[CummulativeCerditTurnover],[CummulativeDebitTurnover],[CummulativeCashDepositTurnover],[CummulativeCashWithdrawalTurnover],[NoOfTransactionsToBeReported],[TransactionDate],[TransactionsID],[TransactionMode],[DebitCredit],[amount],[TransactionsCurrency],[ProductType],[ProductIdentifiers],[TransactionType],[unit],[Date],[DispositionOfFunds],[RelatedAccountNumber],[RelatedInstitutionName],[Remark],[ROS_cmt],[Created_By],[DGM_cmt] FROM [dbo].[offline_scenarios] WHERE [approved] = 1 AND [ROS_cmt] IS NOT NULL AND [DGM_cmt] IS NOT NULL AND [ticket_id] = ?", (TicketId,))
    rows = cursor.fetchall()

    if rows:
        temp_dir = tempfile.mkdtemp()
        zip_file_path = os.path.join(temp_dir, "offline_data.zip")

        with zipfile.ZipFile(zip_file_path, 'w') as zipf:
            for row in rows:
                flattened_data = flatten_data(row)
                temp_csv_filename = os.path.join(temp_dir, "str_offline_data.csv")
                temp_pdf_filename = os.path.join(temp_dir, "str_offline_data.pdf")
                temp_docx_filename = os.path.join(temp_dir, "str_offline_data.docx")
                temp_txt_filename = os.path.join(temp_dir, "str_offline_data.txt")

                # Generate CSV
                df = pd.DataFrame([flattened_data])
                df.to_csv(temp_csv_filename, index=False)
                zipf.write(temp_csv_filename, os.path.basename(temp_csv_filename))

                # Generate PDF
                generate_pdf_offline_from_csv(temp_csv_filename, temp_pdf_filename, [flattened_data])
                zipf.write(temp_pdf_filename, os.path.basename(temp_pdf_filename))

                # Generate DOCX
                # Assuming you have a function to generate DOCX from data
                generate_docx_from_data(flattened_data, temp_docx_filename)
                zipf.write(temp_docx_filename, os.path.basename(temp_docx_filename))

                # Generate TXT
                # Assuming you have a function to generate TXT from data
                generate_txt_from_data(flattened_data, temp_txt_filename)
                zipf.write(temp_txt_filename, os.path.basename(temp_txt_filename))

        if os.path.exists(zip_file_path):
            return send_from_directory(temp_dir, "offline_data.zip", as_attachment=True)

    return render_template('FINnet_report.html', message='No approved cases found', type='FINnetReports', role='IT OFFICER')








def flatten_data(data, prefix=""):
    flattened_data = {}
    if isinstance(data, dict) or isinstance(data, pyodbc.Row):
        if isinstance(data, pyodbc.Row):
            for column in data.cursor_description:
                column_name = column[0]
                column_value = getattr(data, column_name)
                if isinstance(column_value, (dict, list, tuple, pyodbc.Row)):
                    flattened_data.update(flatten_data(column_value, f"{prefix}_{column_name}" if prefix else column_name))
                else:
                    flattened_data[f"{prefix}_{column_name}" if prefix else column_name] = column_value
        else:
            for key, value in data.items():
                if isinstance(value, (dict, list, tuple, pyodbc.Row)):
                    flattened_data.update(flatten_data(value, f"{prefix}_{key}" if prefix else key))
                else:
                    flattened_data[f"{prefix}_{key}" if prefix else key] = value
    elif isinstance(data, list):
        for index, item in enumerate(data):
            flattened_data.update(flatten_data(item, f"{prefix}_{index}" if prefix else str(index)))
    elif isinstance(data, tuple):
        for index, item in enumerate(data):
            flattened_data.update(flatten_data(item, f"{prefix}_{index}" if prefix else str(index)))
    return flattened_data



# ------------------ """ FUNCTIONS RELATED TO CONVERTING DATA TO DIFF FORMATE FILKES ONLINE STR / OFFLINE STR  """   -------------------------------


def generate_txt_from_dataframe(dataframe, filename):
    with open(filename, 'w') as f:
        f.write(dataframe.to_csv(sep='|', index=False))
def generate_docx_from_dataframe(dataframe, filename):
    doc = Document()

    table = doc.add_table(rows=len(dataframe.columns) + 1, cols=2)  # Two columns for keys and values
    
    # Add column headers
    table.cell(0, 0).text = ''
    table.cell(0, 1).text = ''
    
    # Create a dictionary to store unique values for each key
    unique_values = {}
    
    # Add data rows
    for col_index, column_name in enumerate(dataframe.columns, start=1):
        table.cell(col_index, 0).text = column_name
        values = dataframe[column_name].unique()  # Get unique values for the current column
        unique_values[column_name] = values  # Store unique values in the dictionary
        table.cell(col_index, 1).text = ", ".join(str(value) for value in values)
    
    doc.save(filename)

def generate_pdf_from_csv(csv_filename, pdf_filename, approved_objects):
   
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
  

    
    # Define styles
    if approved_objects[0].get('reportedDate') is not None:
        dateSub = str(approved_objects[0].get('reportedDate').date())
    else:
        dateSub = "None"
    # dateSub = str(approved_objects[0].get('reportedDate').date())


    # Create table data
    table_data = []
    for key, value in approved_objects[0].items():
        if value is not None and key != 'ad_id' and key != 'Largest_Cash_Credit':
                wrapper = textwrap.TextWrapper(width=70)
                wrapped_lines = wrapper.wrap(text=str(value))
                wrapped_text = "\n".join(wrapped_lines)
                table_data.append([key, wrapped_text])

    # Create table
    table = Table(table_data, colWidths=[200, 400],  # Adjust the column widths as needed
                          style=[
                              ('GRID', (0, 0), (-1, -1), 1, colors.black),  # Add grid lines
                              ('BACKGROUND', (0, 0), (-1, 0), colors.white),  # Add background color to header row
                              ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),  # Set text color in header row
                              ('ALIGN', (0, 0), (-1, -1), 'LEFT'),  # Center align all cells
                              ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Middle align all cells
                              ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),  # Inner grid lines
                              ('BOX', (0, 0), (-1, -1), 0.25, colors.black),  # Add box around the table
                          ])

    # Build the story

    all_names = ["MLRO", "CM", "GM/PO \n (Approved)"]
    names_table_data = [[all_names[0], all_names[1], all_names[2]]]
    names_table = Table(names_table_data, colWidths=[180, 180, 180],
                            style=[
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                                ('ALIGN', (2, 0), (-1, -1), 'RIGHT')

                            ])

    spacer = Spacer(1, 70)


        
    # doc.build([ header_content, table,names_table], onFirstPage=_header_footer, onLaterPages=_header_footer)
    doc.build([table,spacer,names_table], onFirstPage=lambda canvas, doc: _header_footer(canvas, doc, dateSub), 
                                          onLaterPages=  lambda canvas, doc: _header_footer(canvas, doc, dateSub))
    buffer.seek(0)
    with open(pdf_filename, 'wb') as f:
        f.write(buffer.read())

def _header_footer(canvas, doc,dateSub):
    # Header
    canvas.saveState()
    image_x = (letter[0] - 4*inch) / 2

    canvas.drawImage("C:/packageoffline/Odisha_Gramya_Bank_App - V1/Odisha_Gramya_Bank_App - V1/Odisha_Gramya_Bank_App - SQLMain Code/static/assets/img/logo.png", image_x, 720, width=4*inch, height=1*inch)
    # canvas.drawImage("./static/assets/img/logo.png", image_x, 720, width=4*inch, height=1*inch)
   
    canvas.restoreState()
    

    # Footer
    canvas.saveState()
    
    all_names_DATE = ['CONFIDENTIAL**']
    names_table_data_DATE = [[all_names_DATE[0]]]
    names_table_DATE = Table(names_table_data_DATE, colWidths=[None], rowHeights=[50],
                             style=[
                                 ('GRID', (0, 0), (-1, -1), 1, colors.black),  # Add grid lines
                                 ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                                 ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # Align text to center
                                 ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                             ])
    # Adjust horizontal alignment
    names_table_DATE.hAlign = 'RIGHT'
    names_table_DATE.wrapOn(canvas, doc.width, doc.topMargin)
    names_table_DATE.drawOn(canvas, doc.width - names_table_DATE._width  + 1.65 * inch, doc.height + 1 * inch)
    
    all_names_DATErr = [f'Date of Report : {dateSub}']
    names_table_data_DATErr = [[all_names_DATErr[0]]]
    names_table_DATErr = Table(names_table_data_DATErr, colWidths=[None], rowHeights=[50])
    # Adjust horizontal alignment
    names_table_DATErr.hAlign = 'LEFT'
    names_table_DATErr.wrapOn(canvas, doc.width, doc.topMargin)
    names_table_DATErr.drawOn(canvas, doc.width - names_table_DATE._width  - 5 * inch, doc.height + 1 * inch)


    canvas.rect(1.5, 1, doc.width + 1.95*inch, doc.height + 1.98*inch)


    page_number = canvas.getPageNumber()
    total_pages = 3
    text = "%d of %d Pages" % (page_number, total_pages)
    canvas.drawRightString(1.25*inch, 0.75*inch, text)


    canvas.restoreState()



# ========================= IT OFFICER / HO ADMIN CODE END'S HERE  ==========================================================


# ----------------------------------------------------------------------------------------------------------------------------


# ========================= MLRO OFFICER CODE START'S HERE ===================================================================



# ----------------------- MLRO OFFICER LANDING / DASHBOARD PAGE -------------------------------------------------------


@app.route('/MLROdashboard', methods=['GET'])
@secure_route(required_role='MLRO')
def MLROdashboard():
    # connMLRODash = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connMLRODash = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlMlroDash = connMLRODash.cursor()

    try:
        if 'email_id' not in session:
            return redirect(url_for('post_login'))
        
        mlro_email = session['email_id']


        query = "SELECT id FROM [user] WHERE EmailId = ?"
        mysqlMlroDash.execute(query, (mlro_email,))
        user = mysqlMlroDash.fetchone()


        count_allocated = 0
        count_submitted = 0
        count_reallocated_mlro = 0
        count_commented = 0

        user_id = user[0]

        try:
            mysqlMlroDash.execute("SELECT COUNT(allocatedTicket) FROM scenarios WHERE allocatedTicket = ?", (user_id,))
            temp_allocated = mysqlMlroDash.fetchone()[0]
            if temp_allocated != 0:
                count_allocated = temp_allocated
            else:
                count_allocated = 0
            

            mysqlMlroDash.execute("SELECT COUNT(mlroCasesTicket) FROM scenarios WHERE mlroCasesTicket = ?", (user_id,))
            temp_mlro_cases = mysqlMlroDash.fetchone()[0]
            if temp_mlro_cases != 0:
                count_submitted = temp_mlro_cases
            else:
                count_submitted = 0

            mysqlMlroDash.execute("SELECT COUNT(unsatisfiedTicket) FROM scenarios WHERE unsatisfiedTicket = ?", (user_id,))
            temp_mlro_reallocated = mysqlMlroDash.fetchone()[0]
            if temp_mlro_reallocated != 0:
                count_reallocated_mlro = temp_mlro_reallocated
            else:
                count_reallocated_mlro = 0
            mysqlMlroDash.execute("SELECT COUNT(mlroClosedTicket) FROM scenarios WHERE mlroClosedTicket = ?", (user_id,))
            temp_mlro_closed = mysqlMlroDash.fetchone()[0]
            if temp_mlro_closed != 0:
                count_commented = temp_mlro_closed
            else:
                count_commented = 0
        except:
            count_allocated = 0
            count_submitted = 0
            count_reallocated_mlro = 0
            count_commented = 0
        
        
        connMLRODash.close()

        return render_template('MLRO Officer.html', count=count_allocated, countSubmited=count_submitted, countcommentedtickets=count_commented, countSentBackCaseAlerts=count_reallocated_mlro, resultlist=[], commentdata=[], caseperday=[], type='MLROdashboard', role='MLRO')

    except Exception as e:
        mysqlMlroDash.rollback()
        connMLRODash.close()
        return f'Something Went Wrong: {e} , Please  Re-Login Again', 500



# ------------------------ MLRO PENDDING CASES END - POINT ----------------------------------------------




@app.route('/MLRONextLevel', methods=['GET', 'POST'])
@secure_route(required_role=['MLRO','DGM/PO'])

def MLRONextLevel():
    mlroMailid = request.form.get('u_mailid')
    session['mlroMailid'] = mlroMailid 
    print(mlroMailid,"dgmmmmmmmmmmm")

    success_message = session.pop('success_message', None)

    # connMLRONext = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connMLRONext = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlMlroNext = connMLRONext.cursor()


    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    try:
    
        mlro_email = session['email_id']
        
        query = "SELECT * FROM [user] WHERE EmailId = ?" 

        mysqlMlroNext.execute(query, (mlro_email,))
        
        rows = mysqlMlroNext.fetchall()

        des = mysqlMlroNext.description

        columns = [col[0] for col in des]

        mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

        mlro = mlro[0]


        
        if mlro is None:
            return "User data not found. Please log in again."
        
        if 'image' in mlro:
                    # Encode the image data as a base64 string
                    mlro['image'] = base64.b64encode(mlro['image']).decode('utf-8')

        role = mlro.get('Role')
        
        if role == "DGM/PO":
            mlro_email = session['mlroMailid']
            
            mysqlMlroNext.execute("SELECT * FROM [user] WHERE EmailId = ?", (mlro_email,))
            rows = mysqlMlroNext.fetchall()
            
            if not rows:
                return "User data not found. Please log in again."

            mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows][0]

        mlroId = mlro.get('id')


        # dataQuery =f"SELECT ACCTNO,CUSTCD,PRIMARY_SOL_ID,alert_created_on,CUST_TYPE_CODE,RIP,scenario_code,TXDATE,TXTYPE,TRNFLOWTYPE,TXAMT,ticketid  FROM scenarios WHERE allocatedTicket = ? " 


        # mysqlMlroNext.execute(dataQuery,(mlroId,))


        # columns = [desc[0] for desc in mysqlMlroNext.description]

        # res = mysqlMlroNext.fetchall()

        # connMLRONext.close()

        # data = []


        # for row in res:
        #     scenario_object = {}
        #     for i, value in enumerate(row):
        #         if columns[i] == 'alert_created_on':
        #             dateConvert = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
        #             scenario_object[columns[i]] = dateConvert
        #         else:
        #             scenario_object[columns[i]] = value
        #     data.append(scenario_object)

        return render_template('alertOperationMLRO.html',  success_message=None,mlrouser=mlro,type='MLRONextLevel',role=role)

    except Exception as e:
            
            mysqlMlroNext.rollback()
            connMLRONext.close()

            return f'Something Went Wrong: {e} Try to Re-Login Again', 500


@app.route('/MLRONextLeveljson', methods=['GET', 'POST'])
@secure_route(required_role=['MLRO', 'DGM/PO'])
def MLRONextLeveljson():
    mlroMailid = request.form.get('u_mailid')
    print(mlroMailid,"dgmdgmdgmdgm")
    success_message = session.pop('success_message', None)

    try:
        connMLRONext = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
        mysqlMlroNext = connMLRONext.cursor()

        mlro_email = session['email_id']
        query = "SELECT * FROM [user] WHERE CAST(EmailId AS nvarchar(max)) = ?"
        mysqlMlroNext.execute(query, (mlro_email,))
        rows = mysqlMlroNext.fetchall()

        if not rows:
            return "User data not found. Please log in again."

        des = mysqlMlroNext.description
        columns = [col[0] for col in des]
        mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows][0]

        if 'image' in mlro:
            mlro['image'] = base64.b64encode(mlro['image']).decode('utf-8')

        role = mlro.get('Role')
        if role == "DGM/PO":
            mlro_email = mlroMailid
            mysqlMlroNext.execute(query, (mlro_email,))
            rows = mysqlMlroNext.fetchall()

            if not rows:
                return "User data not found. Please log in again."

            mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows][0]

        mlroId = mlro.get('id')
        dataQuery = "SELECT ACCTNO, CUSTCD, PRIMARY_SOL_ID, alert_created_on, CUST_TYPE_CODE, RIP, scenario_code, TXDATE, TXTYPE, TRNFLOWTYPE, TXAMT, ticketid FROM scenarios WHERE allocatedTicket = ?"
        mysqlMlroNext.execute(dataQuery, (mlroId,))

        columns = [desc[0] for desc in mysqlMlroNext.description]
        res = mysqlMlroNext.fetchall()

        connMLRONext.close()

        data = []
        for row in res:
            scenario_object = {}
            for i, value in enumerate(row):
                if columns[i] == 'alert_created_on':
                    dateConvert = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
                    scenario_object[columns[i]] = dateConvert
                else:
                    scenario_object[columns[i]] = value
            data.append(scenario_object)

        current_datetime = datetime.now()
        current_date = current_datetime.date()
        curentdate = datetime.combine(current_date, datetime.min.time())

        draw = request.args.get('draw', type=int)
        start = request.args.get('start', type=int)
        length = request.args.get('length', type=int)
        search_value = request.args.get('search[value]', '')

        start = start or 0
        length = length or 10

        filtered_data = [item for item in data if search_value.lower() in str(item).lower()]

        total_records = len(data)
        total_filtered_records = len(filtered_data)

        paginated_data = filtered_data[start:start + length]

        return jsonify({
            'draw': draw,
            'recordsTotal': total_records,
            'recordsFiltered': total_filtered_records,
            'data': paginated_data,
            'currentDate': curentdate
        })

    except Exception as e:
        if connMLRONext:
            connMLRONext.close()
        return f'Something Went Wrong: {e} Try to Re-Login Again', 500


     

# ---------------------- MLRO SUBMITED CASES PAGE END - POINT --------------------------



@app.route('/MLRONextLevelSubmitView', methods=['GET', 'POST'])
@secure_route(required_role=['MLRO','DGM/PO'])
def MLRONextLevelSubmitView(): 
    mlroMailid = request.form.get('u_mailid')
    success_message = session.pop('success_message', None)
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    
    mlro_email = session['email_id']

    # connMLROSub = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connMLROSub = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlMLROSub = connMLROSub.cursor()

    try:

        mysqlMLROSub.execute("SELECT * FROM [user] WHERE EmailId = ?", (mlro_email,))
        rows = mysqlMLROSub.fetchall()
        
        columns = [col[0] for col in mysqlMLROSub.description]
        mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows][0]
        if mlro is None:
            return "User data not found. Please log in again."
        if 'image' in mlro:
            mlro['image'] = base64.b64encode(mlro['image']).decode('utf-8')
        role = mlro.get('Role')
        
        if role == "DGM/PO":
            mlro_email = mlroMailid
            mysqlMLROSub.execute("SELECT * FROM [user] WHERE EmailId = ?", (mlro_email,))
            rows = mysqlMLROSub.fetchall()
            
            if not rows:
                return "User data not found. Please log in again."

            mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows][0]
  

        mlroId = mlro.get('id')

        dataQuery =f"SELECT ACCTNO,CUSTCD,PRIMARY_SOL_ID,alert_created_on,CUST_TYPE_CODE,RIP,scenario_code,TXDATE,TXTYPE,TRNFLOWTYPE,TXAMT,ticketid  FROM scenarios WHERE mlroCasesTicket = ? " 


        mysqlMLROSub.execute(dataQuery,(mlroId,))


        columns = [desc[0] for desc in mysqlMLROSub.description]

        res = mysqlMLROSub.fetchall()
            
        connMLROSub.close()

        data = []

        for row in res:
            scenario_object = {}
            for i, value in enumerate(row):
                if columns[i] == 'alert_created_on':
                    dateConvert = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
                    scenario_object[columns[i]] = dateConvert
                else:
                    scenario_object[columns[i]] = value
            data.append(scenario_object)
    

        return render_template('MLROhigherLevel.html', data=data, success_message=success_message,  type='MLRONextLevelSubmitView',role=role)
    
    except Exception as e:

        mysqlMLROSub.rollback()
        connMLROSub.close()
        
        return f"Something went Wrong {e} , Please Re-Login Again",500    



# ----------------------- MLRO CLOSED CASES PAGE END - POINT ---------------------------



@app.route('/Closed_Mlro_Alerts', methods=['GET', 'POST'])
@secure_route(required_role=['MLRO','DGM/PO'])
def Closed_Mlro_Alerts():
    mlroMailid = request.form.get('u_mailid')
    success_message = session.pop('success_message', None)
    # Check user login
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    mlro_email = session['email_id']

    # connClosed = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connClosed = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlCloasedMlro = connClosed.cursor()

    try:

        query = "SELECT * FROM [user] WHERE EmailId = ?"

        mysqlCloasedMlro.execute(query, (mlro_email,))
        
        rows = mysqlCloasedMlro.fetchall()

        columns = [col[0] for col in mysqlCloasedMlro.description]

        mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

        mlro = mlro[0]


        
        if mlro is None:
            return "User data not found. Please log in again."
        
        if 'image' in mlro:
                    # Encode the image data as a base64 string
                    mlro['image'] = base64.b64encode(mlro['image']).decode('utf-8')
        # if 'ACCTNO' in session and 'ticket_id' in session and 'Customer Number' in session:
        #     session.pop('ACCTNO')
        #     session.pop('ticket_id')
        #     session.pop('Customer Number')
        # ticket_numbers = mlro.get("allocated_tickets", [])
        role = mlro.get('Role')
        
        if role == "DGM/PO":
            mlro_email = mlroMailid
            mysqlCloasedMlro.execute(query, (mlro_email,))
            rows = mysqlCloasedMlro.fetchall()
            
            if not rows:
                return "User data not found. Please log in again."

            mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows][0]


        mlroId = mlro.get('id')


        dataQuery =f"SELECT ACCTNO,CUSTCD,PRIMARY_SOL_ID,alert_created_on,CUST_TYPE_CODE,RIP,scenario_code,TXDATE,TXTYPE,TRNFLOWTYPE,TXAMT,ticketid  FROM scenarios WHERE mlroClosedTicket = ? " 


        mysqlCloasedMlro.execute(dataQuery,(mlroId,))


        columns = [desc[0] for desc in mysqlCloasedMlro.description]

        res = mysqlCloasedMlro.fetchall()
            
        connClosed.close()

        data = []

        for row in res:
            scenario_object = {}
            for i, value in enumerate(row):
                scenario_object[columns[i]] = value
            data.append(scenario_object)
        
        

        return render_template('Closed_Mlro_Alerts.html',data=data,success_message=success_message,mlrouser=mlro,type='Closed_Mlro_Alerts',role=role)
    
    except Exception as e:

        mysqlCloasedMlro.rollback()
        connClosed.close()

        return f"Somthing went Wrong {e} , Please Re-Login Again",500



# ------------------------ MLRO RETURNED CLOSED CASES FROM CM TO MAKE THEM RE SUBMIT AS REASON BY CM END - POINT -----------------------



@app.route('/return_Mlro_Alerts',methods=['POST','GET'])
@secure_route(required_role='MLRO')
def return_Mlro_Alerts():
    success_message = session.pop('success_message', None)
    # Check user login
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    
    mlroEmail = session['email_id']

    # connsnetbackClosedmlro = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connsnetbackClosedmlro = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlsentbackClosedmlro = connsnetbackClosedmlro.cursor()

    try:
        query = "SELECT * FROM [user] WHERE EmailId = ?"

        mysqlsentbackClosedmlro.execute(query, (mlroEmail,))
        
        rows = mysqlsentbackClosedmlro.fetchall()

        columns = [col[0] for col in mysqlsentbackClosedmlro.description]
        # print('columns:',columns)

        mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

        mlro = mlro[0]


        
        if mlro is None:
            return "User data not found. Please log in again."
        
        if 'image' in mlro:
                    # Encode the image data as a base64 string
                    mlro['image'] = base64.b64encode(mlro['image']).decode('utf-8')
        # if 'ACCTNO' in session and 'ticket_id' in session and 'Customer Number' in session:
        #     session.pop('ACCTNO')
        #     session.pop('ticket_id')
        #     session.pop('Customer Number')
        # ticket_numbers = mlro.get("allocated_tickets", [])

        mlroId = mlro.get('id')
        
        dataQuery =f"SELECT ACCTNO,CUSTCD,PRIMARY_SOL_ID,alert_created_on,CUST_TYPE_CODE,RIP,scenario_code,TXDATE,TXTYPE,TRNFLOWTYPE,TXAMT,ticketid  FROM scenarios WHERE sentBackClosedTicket = ? " 


        mysqlsentbackClosedmlro.execute(dataQuery,(mlroId,))


        columns = [desc[0] for desc in mysqlsentbackClosedmlro.description]

        res = mysqlsentbackClosedmlro.fetchall()
            
        connsnetbackClosedmlro.close()

        data = []
        
        for row in res:
            scenario_object = {}
            for i, value in enumerate(row):
                if columns[i] == 'alert_created_on':
                    dateConvert = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
                    scenario_object[columns[i]] = dateConvert
                else:
                    scenario_object[columns[i]] = value
            data.append(scenario_object)

        current_datetime = datetime.now()
        current_date = current_datetime.date()
        curentdate = datetime.combine(current_date, datetime.min.time())
            


        return render_template('returned_Alerts_MLRO.html', data=data,currentDate = curentdate, success_message=success_message,mlrouser=mlro,role='MLRO',type='return_Mlro_Alerts')

    except Exception as e:

        mysqlsentbackClosedmlro.rollback()
        connsnetbackClosedmlro.close()
        
        return f"Somthing Went Wrong {e} Please Re-Login Again.",500

# ======================================EMAIL==================================
import datetime
import time
import pyodbc
import datetime
from datetime import datetime, timedelta
import schedule
import time
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import threading
from flask import Flask



def count_tickets(mlro_id):
    current_date = datetime.date.today().strftime('%Y-%m-%d')
    conn = pyodbc.connect("Driver={SQL Server};SERVER=MSI;Database=ticket_id;Trusted_Connection=yes;MARS_Connection=yes")
    cursor = conn.cursor()
    query = "SELECT COUNT(*) FROM tickets WHERE mlroCasesTicket = ? AND CONVERT(DATE, currentDate) = ?"
    cursor.execute(query, (mlro_id, current_date))
    ticket_count = cursor.fetchone()[0]
    cursor.close()
    conn.close()
    return ticket_count
def send_email(sender_email, sender_password, recipient_email, subject, body):
    smtp_server = 'smtp.office365.com'
    smtp_port = 587
    smtp_username = 'triveni05@outlook.com'
    smtp_password = 'bwgcfnxylxmimoed'
    message = MIMEMultipart()
    message['From'] = smtp_username
    message['To'] = recipient_email
    message['Subject'] = subject
    print("to mail:", recipient_email)
    message.attach(MIMEText(body, 'plain'))
    try:
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(smtp_username, smtp_password)
        server.sendmail(smtp_username, recipient_email, message.as_string())
        server.quit()
        print(f"Email sent to {recipient_email}")
    except Exception as e:
        print(f"Failed to send email to {recipient_email}: {str(e)}")
def get_mlro_details():
    conn = pyodbc.connect("Driver={SQL Server};SERVER=MSI;Database=ticket_id;Trusted_Connection=yes;MARS_Connection=yes")
    cursor = conn.cursor()
    query = "SELECT EmailId, id FROM [user] WHERE Role = 'MLRO'"
    cursor.execute(query)
    mlro_data = cursor.fetchall()
    cursor.close()
    conn.close()
    return mlro_data
def count_completed_tickets(mlro_id):
    current_date = datetime.date.today().strftime('%Y-%m-%d')
    conn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid;Trusted_Connection=yes;MARS_Connection=yes")
    cursor = conn.cursor()
    query = "SELECT COUNT(*) FROM tickets WHERE mlroClosedTicket = ? AND CONVERT(DATE, currentDate) = ?"
    cursor.execute(query, (mlro_id, current_date))
    completed_count = cursor.fetchone()[0]
    cursor.close()
    conn.close()
    return completed_count
def count_pending_tickets(mlro_id):
    current_date = datetime.date.today().strftime('%Y-%m-%d')
    conn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid;Trusted_Connection=yes;MARS_Connection=yes")
    cursor = conn.cursor()
    query = "SELECT COUNT(*) FROM tickets WHERE allocatedTicket = ? AND CONVERT(DATE, currentDate) = ?"
    cursor.execute(query, (mlro_id, current_date))
    pending_count = cursor.fetchone()[0]
    cursor.close()
    conn.close()
    return pending_count
def main_morning():
    sender_email = "triveni05@outlook.com"
    sender_password = "bwgcfnxylxmimoed"
    recipient_email = "poojitha@corundumtechnologies.com"
    current_date = datetime.date.today().strftime("%Y-%m-%d")
    mlro_data = get_mlro_details()
    for mlro_email, mlro_id in mlro_data:
        print("mlro email:", mlro_email)
        print("mlro id:", mlro_id)
        ticket_count = count_tickets(mlro_id)
        print("count:", ticket_count)
        print(f"Sending email to MLRO with ticket count: {ticket_count}")
        send_email(sender_email, sender_password, recipient_email, "Ticket Count Report", f"Your daily ticket count is: {ticket_count}")
def main_evening():
    sender_email = "triveni05@outlook.com"
    sender_password = "bwgcfnxylxmimoed"
    recipient_email = "poojitha@corundumtechnologies.com"
    current_date = datetime.date.today().strftime("%Y-%m-%d")
    mlro_data = get_mlro_details()
    for mlro_email, mlro_id in mlro_data:
        print("mlro email:", mlro_email)
        print("mlro id:", mlro_id)
        completed_count = count_completed_tickets(mlro_id)
        print("completed count:", completed_count)
        pending_count = count_pending_tickets(mlro_id)
        print("pending count:", pending_count)
        subject = "Ticket Status Report"
        body = f"Dear MLRO,\n\nHere is your ticket status report for {current_date}:\n\nCompleted Tickets: {completed_count}\nPending Tickets: {pending_count}\n\nRegards,\nYour Company"
        send_email(sender_email, sender_password, recipient_email, subject, body)
def schedule_emails():
    schedule.every().day.at("16:").do(main_morning)
    schedule.every().day.at("16:00").do(main_evening)
    while True:
        schedule.run_pending()
        time.sleep(1)


#####################









   

kamal_server = 'tcp:Charan\\MSSQLSERVER04,49172'
kamal_database = 'Kamal_upload'
kamal_conn_str = (
    'DRIVER={ODBC Driver 17 for SQL Server};SERVER=' + kamal_server +
    ';DATABASE=' + kamal_database + ';Trusted_Connection=yes;CHARSET=UTF8;'
)
cust_server = 'tcp:Charan\\MSSQLSERVER04,49172'
cust_database = 'ticketid'
cust_conn_str = (
    'DRIVER={ODBC Driver 17 for SQL Server};SERVER=' + cust_server +
    ';DATABASE=' + cust_database + ';Trusted_Connection=yes;CHARSET=UTF8;'
)
def extract_letters(text):
    return re.sub(r'[^a-zA-Z0-9&\-.(),]', '', text)
def get_db_connection(conn_str):
    try:
        conn = pyodbc.connect(conn_str)
        return conn
    except Exception as e:
        print(f"Error connecting to database: {e}")
        return None
def fetch_data(cursor, query, params=None):
    cursor.execute(query, params or ())
    columns = [column[0] for column in cursor.description]
    return [dict(zip(columns, row)) for row in cursor.fetchall()], columns
def column_exists(cursor, table_name, column_name):
    try:
        query = "SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = ? AND COLUMN_NAME = ?"
        cursor.execute(query, (table_name, column_name))
        return cursor.fetchone() is not None
    except Exception as e:
        print(f"Error checking column existence: {e}")
        return False
MONTHS = {
    'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
    'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
}
def normalize_date(date_str):
    if isinstance(date_str, str):
        date_str = re.sub(r'[/\-.]', '-', date_str)
        date_parts = date_str.split('-')
        if len(date_parts) == 1 and len(date_parts[0]) == 4 and date_parts[0].isdigit():
            return date_parts[0]
        elif len(date_parts) == 3 and len(date_parts[0]) == 4 and all(part.isdigit() for part in date_parts):
            return '-'.join(date_parts)
        elif len(date_parts) == 3 and len(date_parts[-1]) == 4 and all(part.isdigit() for part in date_parts):
            return '-'.join(date_parts[::-1])
        elif len(date_parts) == 3 and len(date_parts[-2]) == 4 and all(part.isdigit() for part in date_parts):
            return '-'.join([date_parts[-1], date_parts[0], date_parts[1]])
        elif len(date_parts) == 3 and any(part[:3].lower() in MONTHS for part in date_parts):
            month_number = str(MONTHS[date_parts[1][:3].lower()])
            return '-'.join([date_parts[-1], month_number, date_parts[0]])
        else:
            return None
    elif isinstance(date_str, list):
        if len(date_str) == 3:
            if any(len(part) == 4 for part in date_str):
                return '-'.join(date_str)
            else:
                return None
        else:
            return None
    else:
        return None
def docs_matched_cust(documents, keyword_names_list, pan, dob, rangeInputIndex=80, kamal_address=None, customer_address=None, permanent_address=None):
    docs_list = []
    rangeInputIndex = int(rangeInputIndex)
    for doc in documents:
        print(f"Processing document: {doc}")
        name_match = all(keyword.lower() in doc['name'].lower() for keyword in keyword_names_list)
        pan_match = pan and doc.get('pan') and doc['pan'] == pan
        dob_match = dob and doc.get('dob') and normalize_date(doc['dob']) == normalize_date(dob)
        address_match = (
            (customer_address and doc.get('address') and doc['address'] == customer_address) or
            (permanent_address and doc.get('address') and doc['address'] == permanent_address) or
            (doc.get('address') and doc['address'] == kamal_address)
        )
        similarity_ratio = 0
        if name_match:
            similarity_ratio = fuzz.WRatio(" ".join(keyword_names_list), doc['name'].lower())
            print(f"Calculated similarity ratio: {similarity_ratio}")
            if similarity_ratio >= rangeInputIndex or pan_match or dob_match or address_match:
                doc['fuzzy_value'] = similarity_ratio
                new_list = [name.capitalize() for name in keyword_names_list]
                doc['key_name'] = " ".join(new_list)
            else:
                doc['fuzzy_value'] = None
        else:
            doc['fuzzy_value'] = None
        if pan_match:
            document_pan_str = doc.get("pan")
            if document_pan_str:
                doc['pan'] = document_pan_str
        if dob_match:
            document_dob_str = doc.get("dob")
            normalized_dob_str = normalize_date(document_dob_str)
            if normalized_dob_str:
                doc['dob'] = normalized_dob_str
        if address_match:
            document_address_str = doc.get("address")
            if document_address_str:
                doc['address'] = document_address_str
        if (similarity_ratio >= rangeInputIndex):
            if pan_match and name_match:
                docs_list.append(doc)
            elif name_match:
                docs_list.append(doc)
            elif pan_match and dob_match:
                docs_list.append(doc)
            elif name_match and dob_match:
                docs_list.append(doc)
            elif name_match and address_match:
                docs_list.append(doc)
    return docs_list
def customer_data(fuzzy_threshold=80, max_sanctions_length=500):
    conn_customer = get_db_connection(cust_conn_str)
    conn_kamal_new = get_db_connection(kamal_conn_str)
    if not conn_customer or not conn_kamal_new:
        return "Error connecting to databases.", 500
    skipped_ids = []
    try:
        cursor_customer = conn_customer.cursor()
        customer_data_query = """
            SELECT Top 10 CUSTCD, CustomerName, DOB, PAN, Communication_Address , Permanent_Address
            FROM CUSTOMERS
            WHERE CONVERT(DATE, CUSTID_OPENDATE, 105) = '2024-01-01'
            GROUP BY CUSTCD, CustomerName, DOB, PAN, Communication_Address, Permanent_Address;
        """
        data_customer_data, customer_columns = fetch_data(cursor_customer, customer_data_query)
        cursor_kamal_new = conn_kamal_new.cursor()
        tables_query = "SELECT TABLE_NAME FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_TYPE = 'BASE TABLE'"
        tables, _ = fetch_data(cursor_kamal_new, tables_query)
        table_names = [table['TABLE_NAME'] for table in tables]
        matched_records = defaultdict(list)
        output_table_name = "OutputTable"
        for record in data_customer_data:
            customer_name = record.get('CustomerName')
            customer_dob = normalize_date(record.get('DOB'))
            customer_pan = record.get('PAN')
            customer_address = record.get('Communication_Address')
            permanent_address = record.get('Permanent_Address')
            cust_cd = record.get('CUSTCD')
            print(customer_name)
            if customer_name:
                keywords = re.split(r'[,\s]+', customer_name.strip().lower())
                keyword_conditions = " AND ".join(["name LIKE ?"] * len(keywords))
                params = [f"%{keyword}%" for keyword in keywords]
                for table_name in table_names:
                    if column_exists(cursor_kamal_new, table_name, 'name'):
                        query_table_data = f"SELECT * FROM {table_name} WHERE ({keyword_conditions})"
                        table_data, columns = fetch_data(cursor_kamal_new, query_table_data, params)
                        for rec in table_data:
                            matched_docs = docs_matched_cust([rec], keywords, customer_pan, customer_dob, fuzzy_threshold,
                                                        kamal_address=None, customer_address=customer_address,
                                                        permanent_address=permanent_address)
                            if matched_docs:
                                existing_ids = []
                                if column_exists(cursor_customer, output_table_name, "CUSTCD"):
                                    existing_data_query = f"SELECT SanctionsOutput FROM {output_table_name} WHERE CUSTCD = ?"
                                    existing_data, _ = fetch_data(cursor_customer, existing_data_query, (cust_cd,))
                                    for existing_record in existing_data:
                                        existing_ids.extend([doc['_id'] for doc in eval(existing_record['SanctionsOutput'])])
                                for doc in matched_docs:
                                    if doc['_id'] not in existing_ids:
                                        sanctions_output = {
                                            '_id': doc.get('_id'),
                                            'category': doc.get('category'),
                                            'sub_category': doc.get('sub_category'),
                                            'fuzzy_value': doc.get('fuzzy_value')
                                        }
                                        matched_records[(cust_cd, customer_name)].append(sanctions_output)
                                    else:
                                        skipped_ids.append(doc['_id'])
        if matched_records:
            create_table_query = f"""
                IF OBJECT_ID('{output_table_name}', 'U') IS NULL
                CREATE TABLE {output_table_name} (
                    CUSTCD NVARCHAR(MAX),
                    CustomerName NVARCHAR(MAX),
                    SanctionsOutput NVARCHAR(MAX),
                    Status INT
                )
            """
            cursor_customer.execute(create_table_query)
            insert_query = f"""
                INSERT INTO {output_table_name} (CUSTCD, CustomerName, SanctionsOutput, Status)
                VALUES (?, ?, ?, ?)
            """
            for (cust_cd, customer_name), sanctions_output in matched_records.items():
                status = 0 if len(sanctions_output) > max_sanctions_length else 1
                if len(sanctions_output) > max_sanctions_length:
                    for i in range(0, len(sanctions_output), max_sanctions_length):
                        chunk = sanctions_output[i:i + max_sanctions_length]
                        cursor_customer.execute(insert_query, (cust_cd, customer_name, str(chunk), status))
                else:
                    cursor_customer.execute(insert_query, (cust_cd, customer_name, str(sanctions_output), status))
            conn_customer.commit()
            print(f"{len(matched_records)} customer names inserted into {output_table_name}.")
    except Exception as e:
        print(f"Error during data fetching and processing: {e}")
        return "Error processing data.", 500
    finally:
        cursor_customer.close()
        cursor_kamal_new.close()
    return matched_records
def run_scheduled_job():
    print("Running scheduled job...")
    customer_data()
schedule_time = "17:03"
schedule.every().day.at(schedule_time).do(run_scheduled_job)









# =========================EMAIL=======================================
# ========================= MLRO OFFICER CODE END'S HERE ===================================================================


# ----------------------------------------------------------------------------------------------------------------------------



# ========================= CM OFFICER CODE START's HERE ================================================================= 


# --------------------------- CM OFFICER LANDING PAGE OR DASHBOARD ------------------------------------------------------


@app.route('/CM_SM_dashboard', methods=['GET', 'POST'])
@secure_route(required_role='CM/SM')
def CM_SM_dashboard():

    if 'email_id' not in session:
            return redirect(url_for('post_login'))
        
    cm_email = session['email_id']

    # connCMDash = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connCMDash = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")


    mysqlCmDash = connCMDash.cursor()

    try:

        query = "SELECT id FROM [user] WHERE EmailId = ?"
        mysqlCmDash.execute(query, (cm_email,))
        user = mysqlCmDash.fetchone()

        count_allocated = 0
        count_submitted = 0
        count_reallocated_cm = 0
        count_commented_cm = 0

        user_id = user[0]

        try:
            # Count tickets allocated to the CM/SM user
            mysqlCmDash.execute("SELECT COUNT(*) FROM scenarios WHERE allocatedTicket = ?", (user_id,))
            temp_allocated = mysqlCmDash.fetchone()[0]
            if temp_allocated != 0:
                count_allocated = temp_allocated
            else:
                count_allocated = 0

            # Count tickets where the CM/SM user is assigned as cmSMCasesTicket
            mysqlCmDash.execute("SELECT COUNT(*) FROM scenarios WHERE cmSMCasesTicket = ?", (user_id,))
            temp_cm_sm_cases = mysqlCmDash.fetchone()[0]
            if temp_cm_sm_cases != 0:
                count_submitted = temp_cm_sm_cases
            else:
                count_submitted = 0

            mysqlCmDash.execute("SELECT COUNT(*) FROM scenarios WHERE cmSmClosedTicket = ?", (user_id,))
            temp_cm_sm_closed = mysqlCmDash.fetchone()[0]
            if temp_cm_sm_closed != 0:
                count_commented_cm = temp_cm_sm_closed
            else:
                count_commented_cm = 0
            mysqlCmDash.execute("SELECT COUNT(*) FROM scenarios WHERE unsatisfiedTicket = ?", (user_id,))
            temp_cm_sm_reallocate = mysqlCmDash.fetchone()[0]
            if temp_cm_sm_reallocate != 0:
                count_reallocated_cm = temp_cm_sm_reallocate
            else:
                count_reallocated_cm = 0
        except:
            count_allocated = 0
            count_submitted = 0
            count_reallocated_cm = 0
            count_commented_cm = 0



        # Fetch user with Role 'MLRO'
        query1 = "SELECT * FROM [user] WHERE Role = 'MLRO'"
        mysqlCmDash.execute(query1)
        rows = mysqlCmDash.fetchall()

        columns = [col[0] for col in mysqlCmDash.description]

        mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

        mlro = mlro[0]
        count_allocated_mlro = 0
        count_submitted_mlro = 0
        count_commented_mlro = 0 
        mlroid=mlro.get('id')
        mlroEid=mlro.get('EmpId')
        mlro_mail=mlro.get('EmailId') 
        try:

            mysqlCmDash.execute("SELECT COUNT(*) FROM scenarios WHERE allocatedTicket = ? ", (mlroid,))
            temp_allocated_mlro = mysqlCmDash.fetchone()[0]
            if temp_allocated_mlro != 0:
                count_allocated_mlro = temp_allocated_mlro
            else:
                count_allocated_mlro = 0

            mysqlCmDash.execute("SELECT COUNT(*) FROM scenarios WHERE mlroCasesTicket = ?", (mlroid,))
            temp_mlro_cases = mysqlCmDash.fetchone()[0]
            if temp_mlro_cases != 0:
                count_submitted_mlro = temp_mlro_cases
            else:
                count_submitted_mlro = 0

            mysqlCmDash.execute("SELECT COUNT(*) FROM scenarios WHERE mlroClosedTicket = ?", (mlroid,))

            temp_mlro_closed = mysqlCmDash.fetchone()[0]
            if temp_mlro_closed != 0:
                count_commented_mlro = temp_mlro_closed
            else:
                count_commented_mlro = 0
           
        except:
            count_allocated_mlro = 0
            count_submitted_mlro = 0
            count_commented_mlro = 0

        connCMDash.close()

        return render_template('CM_SM_dashboard.html',mail=mlro_mail,id=mlroEid, count_mlro=count_allocated_mlro, countSubmited_mlro=count_submitted_mlro, countcommentedtickets_mlro=count_commented_mlro,count=count_allocated, countSubmited=count_submitted, countcommentedtickets=count_commented_cm, countSentBackCaseAlerts=count_reallocated_cm, allocateresultlist=[], caseresultlist=[], type='CM_SM_dashboard', role='CM/SM')
                                                                                
    except Exception as e:

        mysqlCmDash.rollback()
        connCMDash.close()

        return f'Something Went Wrong: {e} , Please Re-Login Again', 500



# ------------------------ CM OFFICER PENDDING CASES END - POINT ----------------------------------------------



@app.route('/CM_SM_NextLevel', methods=['GET','POST'])
@secure_route(required_role=['CM/SM','DGM/PO'])
def CM_SM_NextLevel(): 
    cmMailid = request.form.get('u_mailid')
    print(cmMailid,"mmmmmmmmmmmmmm")
    success_message = session.pop('success_message', None)

    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    
    cm_email = session['email_id']

    # connCMNext = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connCMNext = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlCMNext = connCMNext.cursor()

    try:

        query = "SELECT * FROM [user] WHERE EmailId = ?"

        mysqlCMNext.execute(query, (cm_email,))
        
        rows = mysqlCMNext.fetchall()

        columns = [col[0] for col in mysqlCMNext.description]

        cm = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

        cm = cm[0]


        if 'image' in cm:
                    # Encode the image data as a base64 string
                    cm['image'] = base64.b64encode(cm['image']).decode('utf-8')
        # if 'ACCTNO' in session and 'ticket_id' in session and 'Customer Number' in session:
        #     session.pop('ACCTNO')
        #     session.pop('ticket_id')
        #     session.pop('Customer Number')
        # ticket_numbers = mlro.get("allocated_tickets", [])

        role=cm.get('Role')
        if role == "DGM/PO":
            cm_email = cmMailid
            mysqlCMNext.execute(query, (cm_email,))
            rows = mysqlCMNext.fetchall()
            
            if not rows:
                return "User data not found. Please log in again."

            cm = [{columns[i]: row[i] for i in range(len(columns))} for row in rows][0]

        

        cmId = cm.get('id')
        
        dataQuery =f"SELECT ACCTNO,CUSTCD,PRIMARY_SOL_ID,alert_created_on,CUST_TYPE_CODE,RIP,scenario_code,TXDATE,TXTYPE,TRNFLOWTYPE,TXAMT,ticketid  FROM scenarios WHERE allocatedTicket = ? " 


        mysqlCMNext.execute(dataQuery,(cmId,))


        columns = [desc[0] for desc in mysqlCMNext.description]

        res = mysqlCMNext.fetchall()

        connCMNext.close()

        data = []
            
        for row in res:
            scenario_object = {}
            for i, value in enumerate(row):
                if columns[i] == 'alert_created_on':
                    dateConvert = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
                    scenario_object[columns[i]] = dateConvert
                else:
                    scenario_object[columns[i]] = value
            data.append(scenario_object)

        current_datetime = datetime.now()
        current_date = current_datetime.date()
        curentdate = datetime.combine(current_date, datetime.min.time())

        return render_template('alertOperation_CM_SM.html', data=data,currentDate = curentdate,success_message=success_message,cmuser=cm,type='CM_SM_NextLevel',role=role)
    except Exception as e:

        mysqlCMNext.rollback()
        connCMNext.close()

        return f"Somthing went Wrong {e} , Please Re-Login Again",500



# ----------------------- CM OFFICER 10 % CLOSED DATA FROM MLRO OFFICER ---------------------------------------



@app.route('/closed_Data_To_CM',methods=['GET','POST']) 
@secure_route(required_role='CM/SM')
def closed_Data_To_CM():
    success_message = session.pop('success_message', None)

    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    
    cm_email = session['email_id']

    # connClosedCm = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connClosedCm = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlClosedCm = connClosedCm.cursor()

    try:

        query = "SELECT * FROM [user] WHERE EmailId = ?"

        mysqlClosedCm.execute(query, (cm_email,))
        
        rows = mysqlClosedCm.fetchall()

        columns = [col[0] for col in mysqlClosedCm.description]

        cm = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

        cm = cm[0]

        
        if cm is None:
            return "User data not found. Please log in again."
        
        if 'image' in cm:
                    # Encode the image data as a base64 string
                    cm['image'] = base64.b64encode(cm['image']).decode('utf-8')
        # if 'ACCTNO' in session and 'ticket_id' in session and 'Customer Number' in session:
        #     session.pop('ACCTNO')
        #     session.pop('ticket_id')
        #     session.pop('Customer Number')
        # ticket_numbers = mlro.get("allocated_tickets", [])
      

        cmId = cm.get('id')

        dataQuery =f"SELECT ACCTNO,CUSTCD,PRIMARY_SOL_ID,alert_created_on,CUST_TYPE_CODE,RIP,scenario_code,TXDATE,TXTYPE,TRNFLOWTYPE,TXAMT,ticketid  FROM scenarios WHERE ten_percent_ticket = ? " 


        mysqlClosedCm.execute(dataQuery,(cmId,))


        columns = [desc[0] for desc in mysqlClosedCm.description]

        res = mysqlClosedCm.fetchall()

        connClosedCm.close()

        data = []
        for row in res:
            scenario_object = {}
            for i, value in enumerate(row):
                if columns[i] == 'alert_created_on':
                    dateConvert = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
                    scenario_object[columns[i]] = dateConvert
                else:
                    scenario_object[columns[i]] = value
            data.append(scenario_object)

        current_datetime = datetime.now()
        current_date = current_datetime.date()
        curentdate = datetime.combine(current_date, datetime.min.time())

        return render_template('rised_closed_cm.html', data=data, currentDate = curentdate ,success_message=success_message, cmuser=cm, role='CM/SM', type='closed_Data_To_CM')

    except Exception as e:

        mysqlClosedCm.rollback()
        connClosedCm.close()

        return f"Somthing went Wrong {e} , Please Re-Login Again",500



# ----------------------- CM SUBMITED CASES END - POINT ----------------------------------------------



@app.route('/CM_SM_NextLevelSubmitView', methods=['GET', 'POST'])
@secure_route(required_role=['CM/SM','DGM/PO'])
def CM_SM_NextLevelSubmitView():
    cmMailid = request.form.get('u_mailid')
    success_message = session.pop('success_message', None)
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    cm_email = session['email_id']

    # connCMSub = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connCMSub = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlCMSub = connCMSub.cursor()

    try:

        mysqlCMSub.execute("SELECT * FROM [user] WHERE EmailId = ?", (cm_email,))
        rows = mysqlCMSub.fetchall()
        

        columns = [col[0] for col in mysqlCMSub.description]
        cm = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]
        cm = cm[0]

        if cm is None:
            return "User data not found. Please log in again."

        if 'image' in cm:
            cm['image'] = base64.b64encode(cm['image']).decode('utf-8')

        role=cm.get('Role')
        
        if role == "DGM/PO":
            cm_email = cmMailid
            mysqlCMSub.execute("SELECT * FROM [user] WHERE EmailId = ?", (cm_email,))
            rows = mysqlCMSub.fetchall()
            
            if not rows:
                return "User data not found. Please log in again."

            cm = [{columns[i]: row[i] for i in range(len(columns))} for row in rows][0]

        

        cmId = cm.get('id')
        
        dataQuery =f"SELECT ACCTNO,CUSTCD,PRIMARY_SOL_ID,alert_created_on,CUST_TYPE_CODE,RIP,scenario_code,TXDATE,TXTYPE,TRNFLOWTYPE,TXAMT,ticketid  FROM scenarios WHERE cmSMCasesTicket = ? " 


        mysqlCMSub.execute(dataQuery,(cmId,))


        columns = [desc[0] for desc in mysqlCMSub.description]

        res = mysqlCMSub.fetchall()

        connCMSub.close()

        data = []

        for row in res:
            scenario_object = {}
            for i, value in enumerate(row):
                if columns[i] == 'alert_created_on':
                    dateConvert = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
                    scenario_object[columns[i]] = dateConvert
                else:
                    scenario_object[columns[i]] = value
            data.append(scenario_object)
            
        return render_template('CM_SM_higherLevel.html',data=data,success_message=success_message,type='CM_SM_NextLevelSubmitView',role=role)
    
    except Exception as e:
        
        mysqlCMSub.rollback()
        connCMSub.close()
        
        return f"Something went Wrong {e} , please Re-Login Again",500     




# ------------------------ CM OFFICER CLOSED CASES FROM THAT 10 % OF CASES CAME FROM MLRO OFFICER END - POINT ----------------------------------------



@app.route('/closed_alerts_cm',methods=['POST','GET'])
@secure_route(required_role=['CM/SM','DGM/PO'])
def closed_alerts_cm():
    cmMailid = request.form.get('u_mailid')

    success_message = session.pop('success_message', None)

    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    
    cm_email = session['email_id']

    # connClosedCm = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connClosedCm = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlClosedCm = connClosedCm.cursor()

    try:

        query = "SELECT * FROM [user] WHERE EmailId = ?"

        mysqlClosedCm.execute(query, (cm_email,))
        
        rows = mysqlClosedCm.fetchall()

        columns = [col[0] for col in mysqlClosedCm.description]
        # print('columns:',columns)

        cm = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

        cm = cm[0]


        
        if cm is None:
            return "User data not found. Please log in again."
        
        if 'image' in cm:
                    # Encode the image data as a base64 string
                    cm['image'] = base64.b64encode(cm['image']).decode('utf-8')
        # if 'ACCTNO' in session and 'ticket_id' in session and 'Customer Number' in session:
        #     session.pop('ACCTNO')
        #     session.pop('ticket_id')
        #     session.pop('Customer Number')
        # ticket_numbers = mlro.get("allocated_tickets", [])
        role=cm.get('Role')
        
        if role == "DGM/PO":
            cm_email = cmMailid
            mysqlClosedCm.execute(query, (cm_email,))
            rows = mysqlClosedCm.fetchall()
            
            if not rows:
                return "User data not found. Please log in again."

            cm = [{columns[i]: row[i] for i in range(len(columns))} for row in rows][0]
        
        cmId = cm.get('id')


        dataQuery =f"SELECT ACCTNO,CUSTCD,PRIMARY_SOL_ID,alert_created_on,CUST_TYPE_CODE,RIP,scenario_code,TXDATE,TXTYPE,TRNFLOWTYPE,TXAMT,ticketid  FROM scenarios WHERE cmSmClosedTicket = ? " 


        mysqlClosedCm.execute(dataQuery,(cmId,))


        columns = [desc[0] for desc in mysqlClosedCm.description]

        res = mysqlClosedCm.fetchall()

        connClosedCm.close()

        data = []
            
        for row in res:
            scenario_object = {}
            for i, value in enumerate(row):
                if columns[i] == 'alert_created_on':
                    dateConvert = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
                    scenario_object[columns[i]] = dateConvert
                else:
                    scenario_object[columns[i]] = value
            data.append(scenario_object)

        return render_template('CM_SM_Closed.html',data=data,success_message=success_message,type='closed_alerts_cm',role=role)

    except Exception as e:

        mysqlClosedCm.rollback()
        connClosedCm.close()

        return f"Something Went Wrong {e} , Please Re-Login Again.",500

        
   

# ========================= CM OFFICER CODE END'S HERE ===================================================================



# ----------------------------------------------------------------------------------------------------------------------------



# ========================= DGM / GM OFFICER CODE START'S HERE ===================================================================



# -------------------------- DGM OFFICER LANDING PAGE OR DASHBOARD -------------------------------------------------------------


@app.route('/DGMdashboard', methods=['GET', 'POST'])
@secure_route(required_role='DGM/PO')
def DGMdashboard():
        if 'email_id' not in session:
                return redirect(url_for('post_login'))
        
        dgm_email = session['email_id']

        # conngmDash = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
        conngmDash = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

        mysqlgmDash = conngmDash.cursor()

        
        try:

            try:

                query = "SELECT id FROM [user] WHERE EmailId = ?"
                mysqlgmDash.execute(query, (dgm_email,))
                user = mysqlgmDash.fetchone()

                user_id = user[0]
                
                count = 0
                Approved = 0
                Rejected = 0 
            
                mysqlgmDash.execute("SELECT COUNT(*) FROM scenarios WHERE allocatedTicket = ?", (user_id,))
                
                tempcount = mysqlgmDash.fetchone()[0]
                if tempcount != 0:
                        count = tempcount
                else:

                    count= 0
                    
                mysqlgmDash.execute("SELECT COUNT(*) FROM scenarios WHERE approved = ?", (user_id,))
                
                tempapproved = mysqlgmDash.fetchone()[0]
                if tempapproved != 0:
                        Approved = tempapproved
                else:

                    Approved= 0
                mysqlgmDash.execute("SELECT COUNT(*) FROM scenarios WHERE rejected = ?", (user_id,))
                
                temprejected = mysqlgmDash.fetchone()[0]
                if temprejected != 0:
                        Rejected = temprejected
                else:

                    Rejected= 0


            except:
                count = 0
                Approved = 0
                Rejected = 0
            
        

            query1 = "SELECT * FROM [user] WHERE Role = 'MLRO'"
            mysqlgmDash.execute(query1)
            rows = mysqlgmDash.fetchall()

            columns = [col[0] for col in mysqlgmDash.description]

            mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

            mlro = mlro[0]

            count_allocated = 0
            count_submitted = 0
            count_reallocated_cm = 0
            count_commented_cm = 0 

            mlroid=mlro.get('id')
            mlroEid=mlro.get('EmpId')
            mlro_mail=mlro.get('EmailId') 

            try:

                # Count tickets allocated to the CM/SM user
                mysqlgmDash.execute("SELECT COUNT(*) FROM scenarios WHERE allocatedTicket = ?  ", (mlroid,))
                temp_allocated = mysqlgmDash.fetchone()[0]
                if temp_allocated != 0:
                    count_allocated = temp_allocated
                else:
                    count_allocated = 0

                # Count tickets where the CM/SM user is assigned as cmSMCasesTicket
                mysqlgmDash.execute("SELECT COUNT(*) FROM scenarios WHERE mlroCasesTicket = ?", (mlroid,))
                temp_cm_sm_cases = mysqlgmDash.fetchone()[0]
                if temp_cm_sm_cases != 0:
                    count_submitted = temp_cm_sm_cases
                else:
                    count_submitted = 0

                mysqlgmDash.execute("SELECT COUNT(*) FROM scenarios WHERE mlroClosedTicket = ?", (mlroid,))
                temp_cm_sm_closed = mysqlgmDash.fetchone()[0]
                if temp_cm_sm_closed != 0:
                    count_commented_cm = temp_cm_sm_closed
                else:
                    count_commented_cm = 0
                mysqlgmDash.execute("SELECT COUNT(*) FROM scenarios WHERE unsatisfiedTicket = ?", (mlroid,))
                temp_cm_sm_reallocate = mysqlgmDash.fetchone()[0]
                if temp_cm_sm_reallocate != 0:
                    count_reallocated_cm = temp_cm_sm_reallocate
                else:
                    count_reallocated_cm = 0
            except:
                count_allocated = 0
                count_submitted = 0
                count_reallocated_cm = 0
                count_commented_cm = 0

    
            query1 = "SELECT * FROM [user] WHERE Role = 'CM/SM'"

            mysqlgmDash.execute(query1)
            rows = mysqlgmDash.fetchall()

            columns = [col[0] for col in mysqlgmDash.description]

            mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

            mlro = mlro[0]
            count_allocated1= 0
            count_submitted1= 0
            count_reallocated1 = 0
            count_commented1 = 0 

            cmid=mlro.get('id')
            cmEid=mlro.get('EmpId')
            cm_mail=mlro.get('EmailId') 

            try:

                mysqlgmDash.execute("SELECT COUNT(*) FROM scenarios WHERE allocatedTicket = ? AND deletedTicket IS NULL ", (cmid,))

                temp_allocated = mysqlgmDash.fetchone()[0]

                if temp_allocated != 0:
                    count_allocated1 = temp_allocated
                else:
                    count_allocated1 = 0

                mysqlgmDash.execute("SELECT COUNT(*) FROM scenarios WHERE cmSMCasesTicket = ?", (cmid,))

                temp_cm_sm_cases = mysqlgmDash.fetchone()[0]

                if temp_cm_sm_cases != 0:
                    count_submitted1 = temp_cm_sm_cases
                else:
                    count_submitted1 = 0

                mysqlgmDash.execute("SELECT COUNT(*) FROM scenarios WHERE cmSmClosedTicket = ?", (cmid,))

                temp_cm_sm_closed = mysqlgmDash.fetchone()[0]

                if temp_cm_sm_closed != 0:
                    count_commented1 = temp_cm_sm_closed
                else:
                    count_commented1 = 0

                mysqlgmDash.execute("SELECT COUNT(*) FROM scenarios WHERE unsatisfiedTicket = ?", (cmid,))

                temp_cm_sm_reallocate = mysqlgmDash.fetchone()[0]

                if temp_cm_sm_reallocate != 0:
                    count_reallocated1 = temp_cm_sm_reallocate
                else:
                    count_reallocated1 = 0
            except:
                count_allocated1 = 0
                count_submitted1 = 0
                count_reallocated1 = 0
                count_commented1 = 0
            
            conngmDash.close()

            return render_template('DGMdashboard.html',count=count,approved_count=Approved,rejected_count=Rejected,DGMRejectedperDay=[],mail=mlro_mail,id=mlroEid,cm_mailid=cm_mail,cm_id=cmEid, count_mlro=count_allocated, countSubmited=count_submitted, countcommentedtickets=count_commented_cm, countSentBackCaseAlerts=count_reallocated_cm,sentbackcount=count_reallocated1, allocateresultlist=[], caseresultlist=[],cm_count=count_allocated1,created_count=count_submitted1,closed_cm_count=count_commented1,allocatedperday=[],DGMApprovedperDay=[],type='DGMdashboard',role='DGM/PO')
        except Exception as e:

            mysqlgmDash.rollback()
            conngmDash.close()
            
            return f'Something Went Wrong: {e} , Please Re-Login Again', 500


# -------------------------- DGM OFFICER PENDDING CASES END - POINT -----------------------------------------------------------



@app.route('/DGMNextLevel', methods=['GET', 'POST'])
@secure_route(required_role='DGM/PO')
def DGMNextLevel():
    success_message = session.pop('success_message', None)
    # Check user login
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    email = session['email_id']

    # connDGMNext = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connDGMNext = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlDGMNext = connDGMNext.cursor()

    try:

        query = "SELECT * FROM [user] WHERE EmailId = ?"

        mysqlDGMNext.execute(query, (email,))
        
        rows = mysqlDGMNext.fetchall()

        columns = [col[0] for col in mysqlDGMNext.description]

        gm = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

        gm = gm[0]

        
        if 'image' in gm:
                    # Encode the image data as a base64 string
                    gm['image'] = base64.b64encode(gm['image']).decode('utf-8')
        # if 'ACCTNO' in session and 'ticket_id' in session and 'Customer Number' in session:
        #     session.pop('ACCTNO')
        #     session.pop('ticket_id')
        #     session.pop('Customer Number')
        # ticket_numbers = mlro.get("allocated_tickets", [])
        

        gmId = gm.get('id')
        
        
        dataQuery =f"SELECT ACCTNO,CUSTCD,PRIMARY_SOL_ID,alert_created_on,CUST_TYPE_CODE,RIP,scenario_code,TXDATE,TXTYPE,TRNFLOWTYPE,TXAMT,ticketid  FROM scenarios WHERE allocatedTicket = ? " 


        mysqlDGMNext.execute(dataQuery,(gmId,))


        columns = [desc[0] for desc in mysqlDGMNext.description]

        res = mysqlDGMNext.fetchall()

        connDGMNext.close()

        data = []
            

        for row in res:
            scenario_object = {}
            for i, value in enumerate(row):
                if columns[i] == 'alert_created_on':
                    dateConvert = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
                    scenario_object[columns[i]] = dateConvert
                else:
                    scenario_object[columns[i]] = value
            data.append(scenario_object)

        current_datetime = datetime.now()
        current_date = current_datetime.date()
        curentdate = datetime.combine(current_date, datetime.min.time())

        return render_template('alertOperationDGM.html',data=data,currentDate=curentdate,success_message=success_message,dgmuser=gm,role='DGM/PO',type='DGMNextLevel')

    except Exception as e:

        mysqlDGMNext.rollback()
        connDGMNext.close()

        return f"Something went Wrong {e} , Please Re-Login Again",500



# ---------------------------- DGM OFFICER APPROVED CASES END - POINT ----------------------------------------------------------


@app.route('/DGMNextLevelSubmitView', methods=['GET', 'POST'])
@secure_route(required_role='DGM/PO')
def DGMNextLevelSubmitView():
        success_message = session.pop('success_message', None)

        if 'email_id' not in session:
            return redirect(url_for('post_login'))
        
        dgm_email = session['email_id']

        # connGMSub = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
        connGMSub = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

        mysqlGMSub = connGMSub.cursor()

        try:

            mysqlGMSub.execute("SELECT * FROM [user] WHERE EmailId = ?", (dgm_email,))

            rows = mysqlGMSub.fetchall()
            
            columns = [col[0] for col in mysqlGMSub.description]

            gm = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

            gm = gm[0]

            if 'image' in gm:
                        # Encode the image data as a base64 string
                        gm['image'] = base64.b64encode(gm['image']).decode('utf-8')
            # if 'ACCTNO' in session and 'ticket_id' in session and 'Customer Number' in session:
            #     session.pop('ACCTNO')
            #     session.pop('ticket_id')
            #     session.pop('Customer Number')
            # ticket_numbers = mlro.get("allocated_tickets", [])
         
            gmId = gm.get('id')

            dataQuery =f"SELECT ACCTNO,CUSTCD,PRIMARY_SOL_ID,alert_created_on,CUST_TYPE_CODE,RIP,scenario_code,TXDATE,TXTYPE,TRNFLOWTYPE,TXAMT,ticketid  FROM scenarios WHERE approved = ? " 


            mysqlGMSub.execute(dataQuery,(gmId,))


            columns = [desc[0] for desc in mysqlGMSub.description]

            res = mysqlGMSub.fetchall()

            connGMSub.close()

            data = []

            for row in res:
                scenario_object = {}
                for i, value in enumerate(row):
                    if columns[i] == 'alert_created_on':
                        dateConvert = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
                        scenario_object[columns[i]] = dateConvert
                    else:
                        scenario_object[columns[i]] = value
                data.append(scenario_object)

            return render_template('DGMhigherLevel.html',data=data,success_message=success_message,role='DGM/PO',type='DGMNextLevelSubmitView')
            
        except Exception as e:

            mysqlGMSub.rollback()
            connGMSub.close()

            return f"Something went Wrong {e} , Please Re-Login Again",500 



# ---------------------------- DGM OFFICER REJECTED CASES END - POINT ----------------------------------------------------------



@app.route('/DGMNextLevelSubmitViewRejected', methods=['GET', 'POST'])
@secure_route(required_role='DGM/PO')
def DGMNextLevelSubmitViewRejected():
        
        success_message = session.pop('success_message', None)

        if 'email_id' not in session:
            return redirect(url_for('post_login'))
        
        dgm_email = session['email_id']

        # connGMRejSub = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
        connGMRejSub = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

        mysqlGMRejSub = connGMRejSub.cursor()

        try:

            mysqlGMRejSub.execute("SELECT * FROM [user] WHERE EmailId = ?", (dgm_email,))

            rows = mysqlGMRejSub.fetchall()

            columns = [col[0] for col in mysqlGMRejSub.description]

            gm = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

            gm = gm[0]
          

            if 'image' in gm:
                        # Encode the image data as a base64 string
                        gm['image'] = base64.b64encode(gm['image']).decode('utf-8')
            # if 'ACCTNO' in session and 'ticket_id' in session and 'Customer Number' in session:
            #     session.pop('ACCTNO')
            #     session.pop('ticket_id')
            #     session.pop('Customer Number')
            # ticket_numbers = mlro.get("allocated_tickets", [])
         
            gmId = gm.get('id')

            dataQuery =f"SELECT ACCTNO,CUSTCD,PRIMARY_SOL_ID,alert_created_on,CUST_TYPE_CODE,RIP,scenario_code,TXDATE,TXTYPE,TRNFLOWTYPE,TXAMT,ticketid  FROM scenarios WHERE rejected = ? " 


            mysqlGMRejSub.execute(dataQuery,(gmId,))


            columns = [desc[0] for desc in mysqlGMRejSub.description]

            res = mysqlGMRejSub.fetchall()

            connGMRejSub.close()
                
            data = []

            for row in res:
                scenario_object = {}
                for i, value in enumerate(row):
                    if columns[i] == 'alert_created_on':
                        dateConvert = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
                        scenario_object[columns[i]] = dateConvert
                    else:
                        scenario_object[columns[i]] = value
                data.append(scenario_object)

            return render_template('DGMhigherLevelRejected.html',data=data,success_message=success_message,role='DGM/PO',type='DGMNextLevelSubmitViewRejected')
        
        except Exception as e:

            mysqlGMRejSub.rollback()
            connGMRejSub.close()

            return f"Something went Wrong {e} , Please Re-Login Again",500 



# ----------------------------- DGM OFFICER OFFLINE DASHBOARD ----------------------------------------------------------



# @app.route('/DGMofflinedashboard', methods=['GET', 'POST'])
# @secure_route(required_role='DGM/PO')
# def DGMofflinedashboard():
#     # Redirect if the user is not logged in
#     if 'email_id' not in session:
#         return redirect(url_for('post_login'))
    
#     dgm_email = session['email_id']


#     # connGMOfflineDash = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
#     connGMOfflineDash = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

#     mysqlGMOfflineDash = connGMOfflineDash.cursor()


#     try:
#         select_query = "SELECT * FROM [user] WHERE EmailId = ?"

#         mysqlGMOfflineDash.execute(select_query,(dgm_email,))

#         dgmuser = mysqlGMOfflineDash.fetchone()

#         if 'image' in dgmuser:
#             # Encode the image data as a base64 string
#             dgmuser['image'] = base64.b64encode(dgmuser['image']).decode('utf-8')


#         select_get_user_id_query = "SELECT id FROM [user] WHERE Role = 'ROS' "

#         mysqlGMOfflineDash.execute(select_get_user_id_query)

#         User_ids = mysqlGMOfflineDash.fetchall()

#         pending_cnt_dgm = 0
#         submited_cnt_dgm = 0
#         approved_count_dgm = 0
#         rejected_count_dgm = 0

#         info = []

#         try:
#             for user_id_tuple in User_ids:
#                 user_id = user_id_tuple[0]

#                 select_get_all_pending_data_query = "SELECT * FROM [dbo].[offline_tickets] o1, [dbo].[offline_collection] o2 WHERE o1.[ticketId] = o2.[ticket_id] AND o1.[allocatedTicket] = ? AND o1.[DGMCasesTicket] IS NULL AND o2.[ROS_cmt] IS NOT NULL AND o2.[DGM_cmt] IS NULL"
                
#                 mysqlGMOfflineDash.execute(select_get_all_pending_data_query, (user_id,))
                
#                 info.append(mysqlGMOfflineDash.fetchall())
            
#             total_pending_cnt = sum(len(sublist) for sublist in info)

#         except:

#             total_pending_cnt= 0 

        
#         if total_pending_cnt != 0:
#             pending_cnt_dgm = total_pending_cnt
#         else:
#             pending_cnt_dgm = 0

#         try:

#             select_total_submited_data_query = "SELECT COUNT(*) FROM [dbo].[offline_tickets] WHERE [allocatedTicket] = ? AND [DGMCasesTicket] = ? AND [approved] = 1"
#             mysqlGMOfflineDash.execute(select_total_submited_data_query, (dgmuser[0],dgmuser[0]))
#             total_submited_cnt = mysqlGMOfflineDash.fetchone()[0]

#         except:
#             total_submited_cnt = 0

#         if total_submited_cnt != 0:
#             submited_cnt_dgm = total_submited_cnt
#         else:
#             submited_cnt_dgm = 0

#         try:

#             select_total_approved_data_query = "SELECT COUNT(*) FROM [dbo].[offline_tickets] WHERE [approved] = 1"
#             mysqlGMOfflineDash.execute(select_total_approved_data_query)
#             approved_count = mysqlGMOfflineDash.fetchone()[0]

#         except:

#             approved_count = 0

#         if approved_count != 0:
#             approved_count_dgm = approved_count
#         else:
#             approved_count_dgm = 0

#         try:

#             select_total_rejected_data_query = "SELECT COUNT(*) FROM [dbo].[offline_tickets] WHERE [rejected] = 1"
#             mysql2.execute(select_total_rejected_data_query)
#             rejected_count = mysqlGMOfflineDash.fetchone()[0]

#         except:

#             rejected_count = 0

#         if rejected_count != 0:
#             rejected_count_dgm = rejected_count
#         else:
#             rejected_count_dgm = 0

#         # ROS Datas
#         ros_data = []
#         ros_pending_cnt = 0
#         ros_submited_cnt = 0
#         for user_id_tuple in User_ids:
#             user_id = user_id_tuple[0]

#             # Get User Details
#             select_get_user_details_query = "SELECT EmpId,EmailId FROM [user] WHERE id = ?"

#             mysqlGMOfflineDash.execute(select_get_user_details_query,(user_id,))

#             user_details = mysqlGMOfflineDash.fetchone()

#             EmpID = user_details[0]
#             Email = user_details[1]

#             # Total Pending Ticket Count ROS
#             select_get_brach_code_query = "SELECT BranchCode FROM [user] WHERE [id] = ?"

#             mysqlGMOfflineDash.execute(select_get_brach_code_query,(user_id,))

#             B_code = mysqlGMOfflineDash.fetchone()[0]
            
#             select_get_b_user_id_query = "SELECT id FROM [user] WHERE BranchCode = ? AND Role = 'BranchMakers' "

#             mysqlGMOfflineDash.execute(select_get_b_user_id_query,(B_code,))

#             B_User_id = mysqlGMOfflineDash.fetchone()[0]

#             try:
#                 select_total_pending_cnt_query = "SELECT COUNT(*) FROM [dbo].[offline_tickets] o1, [dbo].[offline_collection] o2 WHERE o1.[ticketId] = o2.[ticket_id] AND o1.[allocatedTicket] = ? AND o1.[ROSCasesTicket] IS NULL AND o2.[ROS_cmt] IS NULL AND o2.[DGM_cmt] IS NULL"
                
#                 mysqlGMOfflineDash.execute(select_total_pending_cnt_query,(B_User_id,)) 
                
#                 total_pending_cnt = mysqlGMOfflineDash.fetchone()[0]

#             except:

#                 total_pending_cnt = 0

#             if total_pending_cnt != 0:
#                 ros_pending_cnt = total_pending_cnt
#             else:
#                 ros_pending_cnt = 0
            
#             try:

#                 select_total_submited_cnt_query = "SELECT COUNT(*) FROM [dbo].[offline_tickets] WHERE [allocatedTicket] IN(?,5) AND [ROSCasesTicket] = ?"

#                 mysqlGMOfflineDash.execute(select_total_submited_cnt_query,(user_id,user_id)) 

#                 total_submited_cnt = mysqlGMOfflineDash.fetchone()[0]
                
#             except:

#                 total_submited_cnt = 0

            
#             if total_submited_cnt != 0:

#                 ros_submited_cnt = total_submited_cnt

#             else:

#                 ros_submited_cnt = 0

#             ros_data.append(
#                 {
#                     "ros_email_id": Email,
#                     "empid": EmpID,
#                     "offline_ticket_count": ros_pending_cnt,
#                     "offline_submited_ticket_count": ros_submited_cnt
#                 }
#             )
#         connGMOfflineDash.close()

#         return render_template('DGM_offline_dashboard.html', ros_data=ros_data, count=pending_cnt_dgm,assignedperdaydgm=[],DGMApprovedperDayoffline=[],DGMRejectedperDayoffline=[], count_submitted=submited_cnt_dgm, approved_count=approved_count_dgm, rejected_count=rejected_count_dgm, dgmuser=dgmuser, role='DGM/PO', type='DGMofflinedashboard')
    
#     except Exception as e:

#         mysqlGMOfflineDash.rollback()
#         connGMOfflineDash.close()

#         return f"Something went worng {e} , Please Re-Login ",500


@app.route('/DGMofflinedashboard', methods=['GET', 'POST'])
@secure_route(required_role='DGM/PO')
def DGMofflinedashboard():
    conn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
    cursor = conn.cursor()
    # Redirect if the user is not logged in
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    dgm_email = session['email_id']
    try:
        select_query = "SELECT * FROM [dbo].[user] WHERE [EmailId] = ?"
        cursor.execute(select_query,(dgm_email,))
        dgmuser = cursor.fetchone()
        if 'image' in dgmuser:
            # Encode the image data as a base64 string
            dgmuser['image'] = base64.b64encode(dgmuser['image']).decode('utf-8')

        select_get_user_id_query = "SELECT [id] FROM [dbo].[user] WHERE [Role] = 'ROS' "
        cursor.execute(select_get_user_id_query)
        User_ids = cursor.fetchall()

        pending_cnt_dgm = 0
        submited_cnt_dgm = 0
        approved_count_dgm = 0
        rejected_count_dgm = 0

        # Total Pending Ticket Count
        info = []
        for user_id_tuple in User_ids:
            user_id = user_id_tuple[0]
            # select_get_all_pending_data_query = "SELECT * FROM [dbo].[offline_tickets] WHERE [allocatedTicket] = ? AND [DGMCasesTicket] IS NULL"
            #select_get_all_pending_data_query = "SELECT * FROM [dbo].[offline_tickets] o1, [dbo].[offline_collection] o2 WHERE o1.[ticketId] = o2.[ticket_id] AND o1.[allocatedTicket] = ? AND o1.[DGMCasesTicket] IS NULL AND o2.[ROS_cmt] IS NOT NULL AND o2.[DGM_cmt] IS NULL"
            #MY QUR
            select_get_all_pending_data_query = """
            SELECT * FROM [ticketid].[dbo].[offline_scenarios] 
            WHERE [allocatedTicket] = ? AND 
            [DGMCasesTicket] IS NULL AND 
            [ROS_cmt] IS NOT NULL AND 
            [DGM_cmt] IS NULL"""
            cursor.execute(select_get_all_pending_data_query, (user_id,))
            info.append(cursor.fetchall())
        print(info)

        total_pending_cnt = sum(len(sublist) for sublist in info)
        print("Total number of tuples:", total_pending_cnt)

        if total_pending_cnt != 0:
            pending_cnt_dgm = total_pending_cnt
        else:
            pending_cnt_dgm = 0
        print('total:',pending_cnt_dgm)

        # Total Submited Ticket Count
        select_total_submited_data_query = "SELECT COUNT(*) FROM [dbo].[offline_scenarios] WHERE [allocatedTicket] = ? AND [DGMCasesTicket] = ? AND [approved] = 1"
        cursor.execute(select_total_submited_data_query, (dgmuser[0],dgmuser[0]))
        total_submited_cnt = cursor.fetchone()[0]
        if total_submited_cnt != 0:
            submited_cnt_dgm = total_submited_cnt
        else:
            submited_cnt_dgm = 0

        # Total  Approved Count Ticket Count
        select_total_approved_data_query = "SELECT COUNT(*) FROM [dbo].[offline_scenarios] WHERE [approved] = 1"
        cursor.execute(select_total_approved_data_query)
        approved_count = cursor.fetchone()[0]
        if approved_count != 0:
            approved_count_dgm = approved_count
        else:
            approved_count_dgm = 0

        # Total Rejected Count Ticket Count
        select_total_rejected_data_query = "SELECT COUNT(*) FROM [dbo].[offline_scenarios] WHERE [rejected] = 1"
        cursor.execute(select_total_rejected_data_query)
        rejected_count = cursor.fetchone()[0]
        if rejected_count != 0:
            rejected_count_dgm = rejected_count
        else:
            rejected_count_dgm = 0

        # ROS Datas
        ros_data = []
        ros_pending_cnt = 0
        ros_submited_cnt = 0
        for user_id_tuple in User_ids:
            user_id = user_id_tuple[0]

            # Get User Details
            select_get_user_details_query = "SELECT [EmpId],[EmailId] FROM [dbo].[user] WHERE [id] = ?"
            cursor.execute(select_get_user_details_query,(user_id,))
            user_details = cursor.fetchone()
            EmpID = user_details[0]
            Email = user_details[1]

            # Total Pending Ticket Count ROS
            select_get_brach_code_query = "SELECT [BranchCode] FROM [dbo].[user] WHERE [id] = ?"
            cursor.execute(select_get_brach_code_query,(user_id,))
            B_code = cursor.fetchone()[0]
            
            select_get_b_user_id_query = "SELECT [id] FROM [dbo].[user] WHERE [BranchCode] = ? AND [Role] = 'BranchMakers' "
            cursor.execute(select_get_b_user_id_query,(B_code,))
            B_User_id = cursor.fetchone()[0]

            # select_total_pending_cnt_query = " SELECT COUNT(*) FROM [dbo].[offline_tickets] WHERE [allocatedTicket] = ? AND [ROSCasesTicket] IS NULL "
            # select_total_pending_cnt_query = "SELECT COUNT(*) FROM [dbo].[offline_tickets] o1, [dbo].[offline_collection] o2 WHERE o1.[ticketId] = o2.[ticket_id] AND o1.[allocatedTicket] = ? AND o1.[ROSCasesTicket] IS NULL AND o2.[ROS_cmt] IS NULL AND o2.[DGM_cmt] IS NULL"
            #MY QRY
            select_total_pending_cnt_query = """
            SELECT COUNT(*) FROM [ticketid].[dbo].[offline_scenarios] 
            WHERE [allocatedTicket] = ? AND 
            [ROSCasesTicket] IS NULL AND 
            [ROS_cmt] IS NULL AND 
            [DGM_cmt] IS NULL """
            
            cursor.execute(select_total_pending_cnt_query,(B_User_id,)) 
            total_pending_cnt = cursor.fetchone()[0]
            if total_pending_cnt != 0:
                ros_pending_cnt = total_pending_cnt
            else:
                ros_pending_cnt = 0
            
            # Total Submited Ticket Count ROS
            select_total_submited_cnt_query = "SELECT COUNT(*) FROM [dbo].[offline_scenarios] WHERE [allocatedTicket] IN(?,5) AND [ROSCasesTicket] = ?"
            cursor.execute(select_total_submited_cnt_query,(user_id,user_id)) 
            total_submited_cnt = cursor.fetchone()[0]
            if total_submited_cnt != 0:
                ros_submited_cnt = total_submited_cnt
            else:
                ros_submited_cnt = 0

            ros_data.append(
                {
                    "ros_email_id": Email,
                    "empid": EmpID,
                    "offline_ticket_count": ros_pending_cnt,
                    "offline_submited_ticket_count": ros_submited_cnt
                }
            )
        print("ros:",ros_data)
        return render_template('DGM_offline_dashboard.html', ros_data=ros_data, count=pending_cnt_dgm, count_submitted=submited_cnt_dgm, approved_count=approved_count_dgm, rejected_count=rejected_count_dgm, dgmuser=dgmuser, role='DGM/PO', type='DGMofflinedashboard')
     
    except pyodbc.Error as e:
        print(f"Database error occurred: {e}")
        return render_template('DGM_offline_dashboard.html', ros_data=[], count=0, count_submitted=0, approved_count=0, rejected_count=0,  dgmuser={}, role='DGM/PO', type='DGMofflinedashboard')


@app.route('/acc_holder_offline_history', methods=['POST','GET'])
@secure_route(required_role=['ROS','DGM/PO','BranchMakers'])
def acc_holder_offline_history():
    conn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
    cursor = conn.cursor()
    accnum = request.form.get('accnum')
    print('accno:',accnum)
    custid = request.form.get('custid')
    print('custid:',custid)
    txdate = request.form.get('txdate')
  
    success_message = session.pop('success_message', None)
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    mlro_email = session['email_id']
    cur = mysql2.connection.cursor()

    query = "SELECT * FROM [user] WHERE EmailId = ?"

    cur.execute(query, (mlro_email,))
    
    rows = cur.fetchall()

    columns = [col[0] for col in cur.description]

    mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

    mlro = mlro[0]
    if mlro is None:
        return "User data not found. Please log in again."
    data = []
    data_scenarios = []
    scenariotransactiondata=[]
    accounts=[]
    if 'image' in mlro:
                # Encode the image data as a base64 string
                mlro['image'] = base64.b64encode(mlro['image']).decode('utf-8')
   
    role=mlro.get('Role')
    if role in ['DGM/PO','ROS','BranchMakers']:
        cur = mysql2.connection.cursor()
        query_scenarios = "  SELECT * FROM [dbo].[scenarios] WHERE CAST([TXNNO] AS NVARCHAR(20)) IN (SELECT CAST([TXNNO] AS NVARCHAR(20)) FROM merged_collection WHERE [CUSTCD] = ? AND [ACCTNO] = ?)"

        cur.execute(query_scenarios, (str(custid), str(accnum)))
        columns_scenarios = [desc[0] for desc in cur.description]
        unique_txnno1 = set()

        for row in cur.fetchall():
            scenario_object = {}
            for i, value in enumerate(row):
                scenario_object[columns_scenarios[i]] = value
            if scenario_object['TXNNO'] not in unique_txnno1:
                data_scenarios.append(scenario_object)
        # Add the TXNNO to the set to track uniqueness
                unique_txnno1.add(scenario_object['TXNNO'])
            # data_scenarios.append(scenario_object)
        
        query = "SELECT * FROM merged_collection WHERE CUSTCD = ? AND ACCTNO = ?"

        cur.execute(query, (str(custid),str(accnum)))
        
        columns = [desc[0] for desc in cur.description]
        # print(columns)
        unique_txnno = set()
            # Fetch rows and create objects dynamically
        for row in cur.fetchall():
            scenario_object = {}
            for i, value in enumerate(row):
                scenario_object[columns[i]] = value 
            if scenario_object['TXNNO'] not in unique_txnno:
                data.append(scenario_object)
        # Add the TXNNO to the set to track uniqueness
                unique_txnno.add(scenario_object['TXNNO'])
        print('data:',unique_txnno)

        formatted_date = ''
        if txdate == None or txdate == '':
            formatted_date = ''
        else:
            # Convert string to datetime object
            date_obj = datetime.strptime(txdate, "%Y-%m-%d")

            # Format datetime object to dd-mm-yyyy
            formatted_date = date_obj.strftime("%d-%m-%Y")
        print("datadatadatadata: ",type(str(formatted_date)))

        query1 = "SELECT * FROM merged_collection WHERE CUSTCD = ? AND ACCTNO = ? AND TXDATE = ? "

        cur.execute(query1, (str(custid), str(accnum), str(formatted_date)))

        # Fetch rows and create objects dynamically
        columns = [desc[0] for desc in cur.description]
        for row in cur.fetchall():
            scenario_object = {}
            for i, value in enumerate(row):
                scenario_object[columns[i]] = value
            scenariotransactiondata.append(scenario_object)

        # print('data:',unique_txnno2)
        # print('datas:',data)
        # print('scenariotransactiondata:',scenariotransactiondata)
        # print('data_scenarios:',data_scenarios)

    return render_template('account_holder_details.html', scenariotransactiondata=scenariotransactiondata,data=data,data_scenarios=data_scenarios,role=role,allImages=mlro)

# --------------------------------- DGM OFFICER OFFLINE PENDDING CASES END - POINT ------------------------------------------------



# @app.route('/offline_dgm_Str', methods=['GET', 'POST'])
# @secure_route(required_role='DGM/PO')
# def offline_dgm_Str():
#     success_message = session.pop('success_message', None)

#     if 'email_id' not in session:
#         return redirect(url_for('post_login'))
    
#     # connGMofflineStr = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
#     connGMofflineStr = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

#     mysqlGMofflineStr = connGMofflineStr.cursor()
    
#     try:

#         mysqlGMofflineStr.execute("SELECT * FROM [offline_collection] WHERE DGM_cmt IS NULL AND ROS_cmt IS NOT NULL")
#         info = mysqlGMofflineStr.fetchall()

#         connGMofflineStr.close()
#     except:
#         info = []
#     return render_template('offline_str_dgm.html',success_message=success_message,dgmuser=info,role='DGM/PO',type='offline_dgm_Str')

@app.route('/offline_dgm_Str', methods=['GET', 'POST'])
@secure_route(required_role='DGM/PO')
def offline_dgm_Str():
    conn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
    cursor = conn.cursor()
    success_message = session.pop('success_message', None)
    # Check user login
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    try:
        cursor.execute("SELECT * FROM [dbo].[offline_scenarios] WHERE DGM_cmt IS NULL AND ROS_cmt IS NOT NULL")
        info = cursor.fetchall()
        return render_template('offline_str_dgm.html',success_message=success_message,dgmuser=info,role='DGM/PO',type='offline_dgm_Str')
    except pyodbc.Error as e:
        print(f"Database error occurred: {e}")
        return render_template('offline_str_dgm.html',success_message=success_message,dgmuser=[],role='DGM/PO',type='offline_dgm_Str')


# --------------------------------- DGM OFFICER OFFLINE SUBMITED CASES END - POINT ------------------------------------------------



# @app.route('/offline_dgm_submited_Str', methods=['GET', 'POST'])
# @secure_route(required_role='DGM/PO')
# def offline_dgm_submited_Str():
#     success_message = session.pop('success_message', None)

#     if 'email_id' not in session:
#         return redirect(url_for('post_login'))
    
#     ros_email = session['email_id']

#     # connGMofflineStrSub = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
#     connGMofflineStrSub = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

#     mysqlGMofflineStrSub = connGMofflineStrSub.cursor()

#     try:
#         # Here we are try to the data arrengment for DGM 
#         select_query1 = "SELECT id from [user] where [EmailId] = ?"
#         mysqlGMofflineStrSub.execute(select_query1, (ros_email,))
#         branch_code = mysqlGMofflineStrSub.fetchone()
#         B_id = branch_code[0]
        
#         try:

#             select_query2 = "SELECT ticketid FROM [offline_tickets] where DGMCasesTicket = ?"
#             mysqlGMofflineStrSub.execute(select_query2, (B_id,))
#             ros_email_id = mysqlGMofflineStrSub.fetchall()
#         except:
#             ros_email_id = []

#         info = []
#         T_id = []
#         if ros_email_id:
#             for tup in ros_email_id:
#                 B_email = tup
#                 T_id.append(B_email[0])

#         if T_id:
#             for i in T_id:
#                 select_query3 = "SELECT o1.[ticket_id],o1.[Created_Date],o1.[Customerno],o1.[casename],o1.[scenario],o1.[Guidance],o1.[RuleScenario],o1.[personname],o1.[SourceofAlert],o1.[alertindicator],o1.[SuspiciousDueToproceedofCrime],o1.[SuspiciousDueToComplexTranscaction],o1.[SuspiciousDueToNoecoRational],o1.[SuspiciousDueToFinancingTerrorism],o1.[AttemptedTranscaction],o1.[LEAInformed],o1.[PriorityRating],o1.[ReportCoverage],o1.[leadetails],o1.[AdditionalDocument],o1.[Aroundofsuspision],o1.[DetailsofInvestigation],o1.[AccountNumber],o1.[AccountType],o1.[holdername],o1.[AccountHolderType],o1.[AccountStatus],o1.[DateofOpening],o1.[RiskRating],o1.[CummulativeCerditTurnover],o1.[CummulativeDebitTurnover],o1.[CummulativeCashDepositTurnover],o1.[CummulativeCashWithdrawalTurnover],o1.[NoOfTransactionsToBeReported],o1.[TransactionDate],o1.[TransactionsID],o1.[TransactionMode],o1.[DebitCredit],o1.[amount],o1.[TransactionsCurrency],o1.[ProductType],o1.[ProductIdentifiers],o1.[TransactionType],o1.[unit],o1.[Date],o1.[DispositionOfFunds],o1.[RelatedAccountNumber],o1.[RelatedInstitutionName],o1.[Remark],o1.[ROS_cmt],o1.[Created_By],o1.[DGM_cmt] FROM [offline_collection] o1, [offline_tickets] o2 WHERE o1.[ticket_id] = o2.[ticketId] AND o2.[approved] = 1 AND [ROS_cmt] IS NOT NULL AND [DGM_cmt] IS NOT NULL AND o1.[ticket_id] = ?"
#                 mysqlGMofflineStrSub.execute(select_query3, (i,))
#                 info.append(mysqlGMofflineStrSub.fetchall())

#         connGMofflineStrSub.close()

#         return render_template("offline_submited_data.html", success_message=success_message, data=info, role='DGM/PO',type='offline_dgm_submited_Str')
#     except Exception as e:
#         mysqlGMofflineStrSub.rollback()
#         connGMofflineStrSub.close()
#         return f"Something went Wrong {e} please Re-Login Again",500 


@app.route('/offline_dgm_submited_Str', methods=['GET', 'POST'])
@secure_route(required_role='DGM/PO')
def offline_dgm_submited_Str():
    conn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
    cursor = conn.cursor()
    success_message = session.pop('success_message', None)
    ros_email = session['email_id']
    # Check user login
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    
    # Here we are try to the data arrengment for DGM 
    # Getting the user id based on Email Id(Email Id stored in Session)
    select_query1 = "SELECT id from [dbo].[user] where [EmailId] = ?"
    try:
        cursor.execute(select_query1, (ros_email,))
        branch_code = cursor.fetchone()
        B_id = branch_code[0]
        # print(branch_code[0])

        # Getting the ticket id in the offline_tickets table based on the user id
        select_query2 = "SELECT ticket_id FROM [dbo].[offline_scenarios] where DGMCasesTicket = ?"
        cursor.execute(select_query2, (B_id,))
        ros_email_id = cursor.fetchall()
        # print(ros_email_id)
        T_id = []
        # checking the data like if data is present or not
        if ros_email_id:
            # In the above method will give us tuple of Email Id so we have to convert it into list
            for tup in ros_email_id:
                B_email = tup
                T_id.append(B_email[0])
        info = []
        for i in T_id:
            # select_query3 = "SELECT o1.[ticket_id],o1.[Created_Date],o1.[Customerno],o1.[casename],o1.[scenario],o1.[Guidance],o1.[RuleScenario],o1.[personname],o1.[SourceofAlert],o1.[alertindicator],o1.[SuspiciousDueToproceedofCrime],o1.[SuspiciousDueToComplexTranscaction],o1.[SuspiciousDueToNoecoRational],o1.[SuspiciousDueToFinancingTerrorism],o1.[AttemptedTranscaction],o1.[LEAInformed],o1.[PriorityRating],o1.[ReportCoverage],o1.[leadetails],o1.[AdditionalDocument],o1.[Aroundofsuspision],o1.[DetailsofInvestigation],o1.[AccountNumber],o1.[AccountType],o1.[holdername],o1.[AccountHolderType],o1.[AccountStatus],o1.[DateofOpening],o1.[RiskRating],o1.[CummulativeCerditTurnover],o1.[CummulativeDebitTurnover],o1.[CummulativeCashDepositTurnover],o1.[CummulativeCashWithdrawalTurnover],o1.[NoOfTransactionsToBeReported],o1.[TransactionDate],o1.[TransactionsID],o1.[TransactionMode],o1.[DebitCredit],o1.[amount],o1.[TransactionsCurrency],o1.[ProductType],o1.[ProductIdentifiers],o1.[TransactionType],o1.[unit],o1.[Date],o1.[DispositionOfFunds],o1.[RelatedAccountNumber],o1.[RelatedInstitutionName],o1.[Remark],o1.[ROS_cmt],o1.[Created_By],o1.[DGM_cmt] FROM [dbo].[offline_collection] o1, [dbo].[offline_tickets] o2 WHERE o1.[ticket_id] = o2.[ticketId] AND o2.[approved] = 1 AND [ROS_cmt] IS NOT NULL AND [DGM_cmt] IS NOT NULL AND o1.[ticket_id] = ?"
            #MY QRY
            select_query3 = "SELECT [ticket_id],[Created_Date],[Customerno],[casename],[scenario],[Guidance],[RuleScenario],[personname],[SourceofAlert],[alertindicator],[SuspiciousDueToproceedofCrime],[SuspiciousDueToComplexTranscaction],[SuspiciousDueToNoecoRational],[SuspiciousDueToFinancingTerrorism],[AttemptedTranscaction],[LEAInformed],[PriorityRating],[ReportCoverage],[leadetails],[AdditionalDocument],[Aroundofsuspision],[DetailsofInvestigation],[AccountNumber],[AccountType],[holdername],[AccountHolderType],[AccountStatus],[DateofOpening],[RiskRating],[CummulativeCerditTurnover],[CummulativeDebitTurnover],[CummulativeCashDepositTurnover],[CummulativeCashWithdrawalTurnover],[NoOfTransactionsToBeReported],[TransactionDate],[TransactionsID],[TransactionMode],[DebitCredit],[amount],[TransactionsCurrency],[ProductType],[ProductIdentifiers],[TransactionType],[unit],[Date],[DispositionOfFunds],[RelatedAccountNumber],[RelatedInstitutionName],[Remark],[ROS_cmt],[Created_By],[DGM_cmt] FROM [dbo].[offline_scenarios] WHERE [approved] = 1 AND [ROS_cmt] IS NOT NULL AND [DGM_cmt] IS NOT NULL AND [ticket_id] = ?"

            cursor.execute(select_query3, (i,))
            info.append(cursor.fetchall())
            print("InfoData----",info)

        return render_template("offline_submited_data.html", success_message=success_message, data=info, role='DGM/PO',type='offline_dgm_submited_Str')
    except pyodbc.Error as e:
        print(f"Database error occurred: {e}")
        return render_template("offline_submited_data.html", success_message=success_message, data=[], role='DGM/PO',type='offline_dgm_submited_Str')




# ---------------------------------- DGM OFFICER OFFLINE REJECTED CASES END - POINT -----------------------------------------------



# @app.route('/offline_dgm_rejected_submited_Str', methods=['GET', 'POST'])
# @secure_route(required_role='DGM/PO')
# def offline_dgm_rejected_submited_Str():
#     success_message = session.pop('success_message', None)
#     # Check user login
#     if 'email_id' not in session:
#         return redirect(url_for('post_login'))
    
#     # connOfflineRej = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
#     connOfflineRej = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

#     mysqlOfflineRej = connOfflineRej.cursor()

#     try:
        
#         select_reject_query = "SELECT o1.[ticket_id],o1.[AccountNumber],o1.[Customerno],o1.[RiskRating],o1.[DateofOpening],o1.[amount],o1.[ROS_cmt],o1.[DGM_cmt] FROM [offline_collection] o1, [offline_tickets] o2 WHERE o1.[ticket_id] = o2.[ticketId] AND o2.[rejected] = 1"
#         mysqlOfflineRej.execute(select_reject_query)
#         info = mysqlOfflineRej.fetchall()

#         connOfflineRej.close()

#     except:
#         info = []
#     return render_template('offline_dgm_rejected.html',success_message=success_message,data=info,role='DGM/PO',type='offline_dgm_rejected_submited_Str') 


@app.route('/offline_dgm_rejected_submited_Str', methods=['GET', 'POST'])
@secure_route(required_role='DGM/PO')
def offline_dgm_rejected_submited_Str():
    conn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
    cursor = conn.cursor()
    success_message = session.pop('success_message', None)
    # Check user login
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    # select_reject_query = "SELECT o1.[ticket_id],o1.[AccountNumber],o1.[Customerno],o1.[RiskRating],o1.[DateofOpening],o1.[amount],o1.[ROS_cmt],o1.[DGM_cmt] FROM [dbo].[offline_collection] o1, [dbo].[offline_tickets] o2 WHERE o1.[ticket_id] = o2.[ticketId] AND o2.[rejected] = 1"
    #MY QRY
    select_reject_query = "SELECT [ticket_id],[AccountNumber],[Customerno],[RiskRating],[DateofOpening],[amount],[ROS_cmt],[DGM_cmt] FROM [dbo].[offline_scenarios] WHERE [ticket_id] = [ticket_id] AND [rejected] = 1"
    try:
        cursor.execute(select_reject_query)
        info = cursor.fetchall()
        return render_template('offline_dgm_rejected.html',success_message=success_message,data=info,role='DGM/PO',type='offline_dgm_rejected_submited_Str') 
    except pyodbc.Error as e:
        print(f"Database error occurred: {e}")
        return render_template('offline_dgm_rejected.html',success_message=success_message,data=[],role='DGM/PO',type='offline_dgm_rejected_submited_Str') 



# ----------------------------------- DGM OFFICER THRESHOLD PAGE SETTING END - POINT ----------------------------------------------




@app.route('/update_threshold_values')
@secure_route(required_role=['DGM/PO'])
def update_threshold_values():

    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    
    dgm_email = session['email_id']

    # connThreshold = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connThreshold = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlThreshold = connThreshold.cursor()

    try:

        select_query = "SELECT * FROM [user] WHERE EmailId = ?"

        mysqlThreshold.execute(select_query,(dgm_email,))

        dgmuser = mysqlThreshold.fetchone()

        connThreshold.close()

        if 'image' in dgmuser:
            dgmuser['image'] = base64.b64encode(dgmuser['image']).decode('utf-8')
        
        codes=fetch_threshold_codes()

        data = create_group_dictionary(codes)


        return render_template("threshold_page.html",data=data, dgmuser=dgmuser, role='DGM/PO', type='update_threshold_values')
    
    except Exception as e:

        mysqlThreshold.rollback()
        connThreshold.close()

        return f"Something went wrong {e} , Please Re-Login Again",500
     


def fetch_threshold_codes():

    # connFetchThreshold = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connFetchThreshold = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlFetchThreshold = connFetchThreshold.cursor()

    try:

                
        query = "SELECT code FROM Thresholds"
        
        mysqlFetchThreshold.execute(query)

        codes = [code[0] for code in mysqlFetchThreshold.fetchall()]

        connFetchThreshold.close()        

        return codes
    
    except Exception as e:
        mysqlFetchThreshold.rollback()
        connFetchThreshold.close()
        print(f"An unexpected error occurred: {e}")
        return None

def create_group_dictionary(codes):
    group_dict = {}
    for code in codes:
        group_name = code.split('_')[0]
        if group_name not in group_dict:
            group_dict[group_name] = []
        group_dict[group_name].append(code)
    return group_dict






@app.route('/DV_threshold_values', methods=['POST','GET'])
@secure_route(required_role=['DGM/PO'])
def DV_threshold_values():

    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    

    # connThresholdValues = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connThresholdValues = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlThresholdValues = connThresholdValues.cursor()
    
    dgm_email = session['email_id']

    try:

        select_query = "SELECT * FROM [user] WHERE EmailId = ?"

        mysqlThresholdValues.execute(select_query,(dgm_email,))

        dgmuser = mysqlThresholdValues.fetchone()

        if 'image' in dgmuser:
            dgmuser['image'] = base64.b64encode(dgmuser['image']).decode('utf-8')


        if request.method == 'POST':
            code_or_prefix = request.form.get('route_code')
        else:
            code_or_prefix=session['route_code']

        mysqlThresholdValues.execute("SELECT Alert_title, current_values, previous_values FROM Thresholds WHERE code = ?", (code_or_prefix))
        row = mysqlThresholdValues.fetchone()

        connThresholdValues.close()


        if row:
            alert_title, current_values_json, previous_values_json = row
            
            current_values = json.loads(current_values_json)


            previous_values = json.loads(previous_values_json) if previous_values_json else {}

            return render_template("DV_TM.html", tm_sub_heading=code_or_prefix, tm_heading=alert_title, current_values=current_values, previous_values=previous_values, dgmuser=dgmuser, role='DGM/PO', type='update_threshold_values')
        else:
            return "No matching data found"
    except Exception as e:

        mysqlThresholdValues.rollback()
        connThresholdValues.close()

        return f"Something went Wrong {e} , Please Re-Login Again",500



@app.route('/update_threshold', methods=['POST'])
@secure_route(required_role=['DGM/PO'])
def update_threshold():
    if request.method == 'POST':

        if 'email_id' not in session:
            return redirect(url_for('post_login'))
        
        # connThresholdUpadte = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
        connThresholdUpadte = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

        mysqlThresholdUpadte = connThresholdUpadte.cursor()

        code = request.form.get('route_code')

        session['route_code'] = code

        try:
            

            mysqlThresholdUpadte.execute("SELECT current_values, previous_values FROM Thresholds WHERE code = ?", (code,))
            row = mysqlThresholdUpadte.fetchone()

            if row:
                current_values_json, previous_values_json = row
                current_values = json.loads(current_values_json)
                previous_values = json.loads(previous_values_json) if previous_values_json else []

                current_values['timestamp'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                previous_values.append(current_values)

                updated_previous_values_json = json.dumps(previous_values)

                new_values = {}
                for key in request.form.keys():
                    if key != 'route_code':
                        new_values[key] = int(request.form[key])

                new_values_json = json.dumps(new_values)

                mysqlThresholdUpadte.execute("UPDATE Thresholds SET current_values = ?, previous_values = ? WHERE code = ?", 
                            (new_values_json, updated_previous_values_json, code))
                
                mysqlThresholdUpadte.commit()

                connThresholdUpadte.close()

                return redirect(url_for('DV_threshold_values'))
            
        except Exception as e:

            mysqlThresholdUpadte.rollback()
            connThresholdUpadte.close()

            return f"Somethong went wrong {e} While updating Threshold Values , Please Re-Login Again.", 500
    else:

        return "Invalid request"




# ========================= DGM / GM OFFICER CODE END'S HERE ===================================================================


# ------------------------------------------------------------------------------------------------------------------------------------------------------------



# ========================================== COMMOUN END - POINTS FOR ONLINE START's HERE =====================================================


# ---------------------------- DISPLAY ALL THE UNSATISFIED CASES CAME FROM HIGHER LEVEL OFFICER -------------------------------


@app.route('/display_Sent_Back_Alerts',methods=['POST','GET'])
@secure_route(required_role=['MLRO','CM/SM','DGM/PO'])
def display_Sent_Back_Alerts():
    user_mail = request.form.get('u_mailid')
    success_message = session.pop('success_message', None)
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
  
    email = session['email_id']

    # connDispSentBackCase = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connDispSentBackCase = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlDispSentBackCase = connDispSentBackCase.cursor()

    try:

        query = "SELECT * FROM [user] WHERE EmailId = ?"

        mysqlDispSentBackCase.execute(query, (email,))
        
        rows = mysqlDispSentBackCase.fetchall()

        columns = [col[0] for col in mysqlDispSentBackCase.description]

        cm = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

        cm = cm[0]

        
        if cm is None:
            return "User data not found. Please log in again."
        
        if 'image' in cm:
                    # Encode the image data as a base64 string
                    cm['image'] = base64.b64encode(cm['image']).decode('utf-8')
        # if 'ACCTNO' in session and 'ticket_id' in session and 'Customer Number' in session:
        #     session.pop('ACCTNO')
        #     session.pop('ticket_id')
        #     session.pop('Customer Number')
        # ticket_numbers = mlro.get("allocated_tickets", [])
        role = cm.get('Role')

        if role == "DGM/PO":
            email = user_mail 
            mysqlDispSentBackCase.execute(query, (email,))
            rows = mysqlDispSentBackCase.fetchall()
            
            if not rows:
                return "User data not found. Please log in again."

            cm = [{columns[i]: row[i] for i in range(len(columns))} for row in rows][0]

        cmId = cm.get('id')

        dataQuery =f"SELECT ACCTNO,CUSTCD,PRIMARY_SOL_ID,alert_created_on,CUST_TYPE_CODE,RIP,scenario_code,TXDATE,TXTYPE,TRNFLOWTYPE,TXAMT,ticketid  FROM scenarios WHERE unsatisfiedTicket = ? " 


        mysqlDispSentBackCase.execute(dataQuery,(cmId,))


        columns = [desc[0] for desc in mysqlDispSentBackCase.description]

        res = mysqlDispSentBackCase.fetchall()

        connDispSentBackCase.close()

        data = []
            

        for row in res:
            scenario_object = {}
            for i, value in enumerate(row):
                if columns[i] == 'alert_created_on':
                    dateConvert = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
                    scenario_object[columns[i]] = dateConvert
                else:
                    scenario_object[columns[i]] = value
            data.append(scenario_object)

        current_datetime = datetime.now()
        current_date = current_datetime.date()
        curentdate = datetime.combine(current_date, datetime.min.time())

        response = make_response(render_template('case_sent_Back_Alerts.html',data=data,currentDate = curentdate,success_message=success_message,allImages=cm,type='display_Sent_Back_Alerts',role=role ))
        response.set_cookie('display_Sent_Back_Alerts','I am cookieee',secure=True,samesite='Lax',httponly=True)
        return response
    
    except Exception as e:

        mysqlDispSentBackCase.rollback()
        connDispSentBackCase.close()
        
        return f"Something went Wrong {e} please Re-Login Again",500 




# -------------------------- CM OFFICER REJECTED CASES END - POINT ------------------------------------------------



@app.route('/rejectAlert',methods=['POST'])
@secure_route(required_role=['MLRO','CM/SM'])
def rejectAlert():
    ticketId = request.form.get('tickId')
    comment = request.form.get('reason')
    
    allocateTo = request.form.get('allocateTo')

    cm_email = session['email_id']
    
    # connsentbackAlertRej = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connsentbackAlertRej = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlSentbackAlertRej = connsentbackAlertRej.cursor()
        

    try:
        query = "SELECT id FROM [user] WHERE EmailId = ?"
        mysqlSentbackAlertRej.execute(query, (cm_email,))
        cmId_row = mysqlSentbackAlertRej.fetchone()
        
        if cmId_row:
            cmId = cmId_row[0]
        else:
            return "User data not found. Please log in again."
    
        update_query = "UPDATE scenarios SET mlroClosedTicket = NULL, ten_percent_ticket = NULL,cmSmClosedTicket = ? WHERE ticketid = ?"
        mysqlSentbackAlertRej.execute(update_query, (cmId,ticketId))
        mysqlSentbackAlertRej.commit()


        mysqlSentbackAlertRej.execute("SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = 'scenarios' AND COLUMN_NAME = 'cm_level'")
        mlro_column_exists = mysqlSentbackAlertRej.fetchone()

        if not mlro_column_exists:
            mysqlSentbackAlertRej.execute("ALTER TABLE scenarios ADD cm_level NVARCHAR(MAX)")
            mysqlSentbackAlertRej.commit()

        mysqlSentbackAlertRej.execute("UPDATE scenarios SET cm_level = ? WHERE ticketid = ?", (cm_email, ticketId))

        mysqlSentbackAlertRej.commit()

        mysqlSentbackAlertRej.execute("SELECT COUNT(*) FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = 'scenarios' AND COLUMN_NAME = 'cm_sm_closed_comment'")
        column_exists = mysqlSentbackAlertRej.fetchone()[0]

        if not column_exists:
            mysqlSentbackAlertRej.execute("ALTER TABLE scenarios ADD cm_sm_closed_comment NVARCHAR(MAX)")

        if comment:
            update_comment_query = "UPDATE scenarios SET cm_sm_closed_comment = ? WHERE ticketid = ? ;"
            mysqlSentbackAlertRej.execute(update_comment_query, (str(comment), str(ticketId)))
            mysqlSentbackAlertRej.commit()            

        connsentbackAlertRej.close()
        return redirect(url_for("closed_Data_To_CM"))
    except:
        mysqlSentbackAlertRej.rollback()
        connsentbackAlertRej.close()
        return redirect(url_for("closed_Data_To_CM"))

    

# ------------------------ DELETING ONLINE CASES END - POINT ------------------------------------------------



@app.route('/deletedcases', methods=['GET','POST'])
@secure_route(required_role=['MLRO','CM/SM','ROS','DGM/PO'])
def deletedcases():
    if request.method == 'POST':

   
        if 'email_id' not in session:
            return redirect(url_for('post_login'))
        
        # connsentdDeletedCases = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
        connsentdDeletedCases = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

        mysqlDeletedCases = connsentdDeletedCases.cursor()
        
        mlro_email = session['email_id']

        try:

            ticket_id = request.form.get('ticket_id')
            rediretingEndPoint = request.form.get('type')

            query = "SELECT * FROM [user] WHERE EmailId = ?" 

            mysqlDeletedCases.execute(query, (mlro_email,))
            
            rows = mysqlDeletedCases.fetchall()

            columns = [col[0] for col in mysqlDeletedCases.description]
            # print('columns:',columns)

            mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

            mlro = mlro[0]

            if mlro is None:
                return "User data not found. Please log in again."
            
            if 'image' in mlro:
                mlro['image'] = base64.b64encode(mlro['image']).decode('utf-8')
            
            mlroId = mlro.get('id')

            updateQuery = "UPDATE scenarios SET deletedTicket = ?,allocatedTicket = NULL  WHERE ticketid = ?"

            mysqlDeletedCases.execute(updateQuery, (mlroId, ticket_id))

            mysqlDeletedCases.commit()
            
            connsentdDeletedCases.close()
            
            return redirect(url_for(rediretingEndPoint))

        except Exception as e:
            
            mysqlDeletedCases.rollback()
            connsentdDeletedCases.close()

            return f"Something went wrong {e} , Please Re-Login Again",500
        
    else:
        return "Bad Request Please Re-Login Again",400





# ------------------------ DISPLAY DELETED ONLINE CASES END - POINT ------------------------------------------------




@app.route('/Display_Deleted_Mlro_Alerts', methods=['GET', 'POST'])
@secure_route(required_role=['MLRO','CM/SM','ROS','DGM/PO'])
def Display_Deleted_Mlro_Alerts():
    success_message = session.pop('success_message', None)

    if 'email_id' not in session:
        return redirect(url_for('post_login'))

    # connDeleteCases = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connDeleteCases = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlDeleteCases = connDeleteCases.cursor()
 
    mlro_email = session['email_id']

    try:


        query = "SELECT * FROM [user] WHERE EmailId = ?" 

        mysqlDeleteCases.execute(query, (mlro_email,))
        
        rows = mysqlDeleteCases.fetchall()

        columns = [col[0] for col in mysqlDeleteCases.description]

        mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

        mlro = mlro[0]

        


        # return mlro[0].get('EmailId')
        
        if mlro is None:
            return "User data not found. Please log in again."
        
        if 'image' in mlro:
                    mlro['image'] = base64.b64encode(mlro['image']).decode('utf-8')
        # if 'ACCTNO' in session and 'ticket_id' in session and 'Customer Number' in session:
        #     session.pop('ACCTNO')
        #     session.pop('ticket_id')
        #     session.pop('Customer Number')
        # ticket_numbers = mlro.get("allocated_tickets", [])
        
        data = []

        mlroId = mlro.get('id')

    
        dataQuery =f"SELECT ACCTNO,CUSTCD,PRIMARY_SOL_ID,alert_created_on,CUST_TYPE_CODE,RIP,scenario_code,TXDATE,TXTYPE,TRNFLOWTYPE,TXAMT,ticketid  FROM scenarios WHERE deletedTicket = ? " 


        mysqlDeleteCases.execute(dataQuery,(mlroId,))


        columns = [desc[0] for desc in mysqlDeleteCases.description]

        res = mysqlDeleteCases.fetchall()

        connDeleteCases.close()

        for row in res:
            scenario_object = {}
            for i, value in enumerate(row):
                if columns[i] == 'alert_created_on':
                    dateConvert = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
                    scenario_object[columns[i]] = dateConvert
                else:
                    scenario_object[columns[i]] = value
            data.append(scenario_object)


        return render_template('deletedalerts.html', data=data, success_message=success_message,type='Display_Deleted_Mlro_Alerts',role=mlro.get('Role'))

    except Exception as e:

        mysqlDeleteCases.rollback()
        connDeleteCases.close()

        return f"Something went wrong , {e} , Please Re-Login Again",500


# ------------------------ ENABLING DELETED ONLINE CASES END - POINT ------------------------------------------------





@app.route('/enableCases', methods=['GET', 'POST'])
@secure_route(required_role=['MLRO','CM/SM','ROS','DGM/PO'])
def enableCases():
    if request.method == 'POST':
        ticket_id = request.form.get('ticket_id')

        if 'email_id' not in session:
            return redirect(url_for('post_login'))
        mlro_email = session['email_id']


        # connEnablCases = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
        connEnablCases = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

        mysqlEnableCases = connEnablCases.cursor()

        try:
            query = "SELECT * FROM [user] WHERE EmailId = ?"

            mysqlEnableCases.execute(query, (mlro_email,))

            rows = mysqlEnableCases.fetchall()
            
            columns = [col[0] for col in mysqlEnableCases.description]

            mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

            mlro = mlro[0]

            if mlro is None:
                return "User data not found. Please log in again."
            if 'image' in mlro:
                    mlro['image'] = base64.b64encode(mlro['image']).decode('utf-8')

            mlroId = mlro.get('id')

            updateQuery = "UPDATE scenarios SET deletedTicket = NULL,allocatedTicket = ? WHERE  ticketid = ?"

            mysqlEnableCases.execute(updateQuery, (mlroId, ticket_id))
            mysqlEnableCases.commit()
                
            connEnablCases.close()

            return redirect('Display_Deleted_Mlro_Alerts')

        except Exception as e:

            mysqlEnableCases.rollback()
            connEnablCases.close()

            return f"Something went wrong {e} , Please Re-Login Again",500
    else:
        return "Bad Request Please Re-Login Again",400


    
   




# ========================= COMMOUN END - POINTS FOR ONLINE END's HERE =====================================================



# --------------------------------------------------------------------------------------------------------------------------------------------------------------




# ========================= MAKER / BRANCHMAKER OFFICER CODE START'S HERE ===================================================================


# ------------------------ BRANCHMAKER LANDING PAGE OR FORM PAGE -------------------------------------------------


# @app.route('/branchmakers')
# @secure_route(required_role='BranchMakers')
# def branchmakers():
#     if 'email_id' not in session:
#         return redirect(url_for('post_login'))
    
#     email_id = session['email_id']

#     # connBranchland = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
#     connBranchland = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

#     mysqlBranchland = connBranchland.cursor()
    
#     try:
#         mysqlBranchland.execute("SELECT * FROM [user] WHERE EmailId = ?", (email_id,))

#         user = mysqlBranchland.fetchone()
        
#         connBranchland.close()

#         branchmakeruser = {}
#         if user and 'image' in user:
#             branchmakeruser['image'] = base64.b64encode(user['image']).decode('utf-8')
        
#         success_message = session.pop('success_message', None)
        
#         return render_template('Branch_makers.html',success_message=success_message,branchmakeruser=branchmakeruser,role='BranchMakers',type='branchmakers')
    
#     except Exception as e:

#         mysqlBranchland.rollback()
#         connBranchland.close()

#         return f"Somethong went wrong {e} , Please Re-Login Again ",500


# Calling the Branchmaker Dashboard
@app.route('/branchmakers')
@secure_route(required_role='BranchMakers')
def branchmakers():
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    email_id = session['email_id']
    
    # Fetch user info from MS SQL Server
    cursor = mysql2.connection.cursor()
    cursor.execute("SELECT * FROM [dbo].[user] WHERE [EmailId]=?", (email_id,))
    user = cursor.fetchone()

    branchmakeruser = {}
    if user and 'image' in user:
        # Encode the image data as a base64 string
        branchmakeruser['image'] = base64.b64encode(user['image']).decode('utf-8')
    
    success_message = session.pop('success_message', None)
    # mysql2.connection.close()
    
    return render_template('Branch_makers.html',success_message=success_message,branchmakeruser=branchmakeruser,role='BranchMakers',type='branchmakers')




# ------------------------------ AUTO FILLING CUSTOMER DETAILS IN BRANCHMAKER FORM END - POINT --------------------------------



@app.route('/searchUserOffline', methods=['POST'])
@secure_route(required_role='BranchMakers')
def searchUserOffline():
    userInfo = request.get_json()
    user = userInfo["user"]
    field = userInfo["field"]
    # connBranchSearchCust = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connBranchSearchCust = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
    mysqlBranchSearchCust = connBranchSearchCust.cursor()
    if field == "Account Number":
        # SQL query to retrieve user details by account number
        mysqlBranchSearchCust.execute("SELECT * FROM CUSTOMERS WHERE ACCTNO = ?", user)
    else:
        # SQL query to retrieve user details by customer ID
        mysqlBranchSearchCust.execute("SELECT * FROM CUSTOMERS WHERE CUSTCD = ?", user)
    userdetails = mysqlBranchSearchCust.fetchone()
    if userdetails:
        userAllInfo = {
            "personname": f"{userdetails.__getattribute__('Title')} {userdetails.__getattribute__('CustomerName')}",
            "custCode": userdetails.__getattribute__('CUSTCD'),
            "AccNo": userdetails.__getattribute__('ACCTNO'),
            "AccType": userdetails.__getattribute__('CUST_TYPE_CODE'),
            "opingDate": userdetails.__getattribute__('ACCT_OPENDATE'),
            "AccStatus": userdetails.__getattribute__('ACCT_STATUS'),
            "holderType": userdetails.__getattribute__('CUST_TYPE_CODE'),
            "holderName": f"{userdetails.__getattribute__('Title')},{userdetails.__getattribute__('CustomerName')}",
            "debitCard": "Null",          # Assuming this value is constant
            "Address": userdetails.__getattribute__('Permanent_Address'),     # Assuming this field represents address
            "Pincode": userdetails.__getattribute__('PINCODE'),     # Assuming this field represents pincode
            "City": userdetails.__getattribute__('STATE'),        # Assuming this field represents city
            "DOB": userdetails.__getattribute__('DOB'),         # Assuming this field represents DOB
            "MobileNumber": 'Null',# Assuming this field represents mobile number
            "PAN": userdetails.__getattribute__('PAN'),
            "TransactionAmount": None,
            "TransactionType":  None,
            "TransactionCategory": None,
            "TransactionCurrency": None,
            "BankName": 'Odisha Gramya Bank',
            "BankState": userdetails.__getattribute__('PRIMARY_SOL_ID')
        }
        return userAllInfo
    else:
        return {"none": "no data"}



# ------------------------ BRANCHMAKER OFFLINE SUBMIT PAGE END - POINT -------------------------------------------------


# @app.route('/offline_sumited_branchmaker', methods=['GET'])
# @secure_route(required_role='BranchMakers')
# def offline_sumited_branchmaker():

#     success_message = session.pop('success_message', None)

#     if 'email_id' not in session:
#         return redirect(url_for('post_login'))
    
#     email = session['email_id']

#     # connBranchSub = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
#     connBranchSub = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

#     mysqlBranchSub = connBranchSub.cursor()

#     try:
#         try:
#             select_query = "SELECT * FROM [offline_collection] o1 JOIN [offline_tickets] o2 ON o1.ticket_id = o2.ticketId WHERE (o1.ticket_id = o2.ticketId AND o1.[Created_By] = ?) AND ( ((o1.[ROS_cmt] IS NULL OR o1.[DGM_cmt] IS NULL) OR (o1.[ROS_cmt] != 'archive' AND o1.[DGM_cmt] != 'archive')) )"
        
#             mysqlBranchSub.execute(select_query, email)
        
#             info = mysqlBranchSub.fetchall()

#             connBranchSub.close()

#         except:
      
#             info = []

#         return render_template("offline_sumited_branchmaker.html", success_message=success_message, data=info, role='BranchMakers',type='offline_sumited_branchmaker')
    
#     except Exception as e:

#         mysqlBranchSub.rollback()
#         connBranchSub.close()

#         return f"Something went wrong {e} , Please Re - Login Again ",500
        
@app.route('/offline_sumited_branchmaker', methods=['GET', 'POST'])
@secure_route(required_role='BranchMakers')
def offline_sumited_branchmaker():
    conn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
    cursor = conn.cursor()
    success_message = session.pop('success_message', None)
    
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    
    email = session['email_id']
    
    select_query = """
    SELECT * 
    FROM [dbo].[offline_scenarios]
    WHERE [Created_By] = ? AND
    (([ROS_cmt] IS NULL OR [DGM_cmt] IS NULL) 
    OR ([ROS_cmt] != 'archive' AND [DGM_cmt] != 'archive'))
    """
    
    try:    
        # Execute the query
        cursor.execute(select_query, email)
        info = cursor.fetchall()
        
        # Close the cursor and connection
        cursor.close()
        # conn.close()
    except pyodbc.Error as e:
        print(f"Database error occurred: {e}")
        return render_template("offline_sumited_branchmaker.html", success_message=None, data=[], role='BranchMakers', type='offline_sumited_branchmaker', error_message="An error occurred while fetching data. Please try again later.")
    
    return render_template("offline_sumited_branchmaker.html", success_message=success_message, data=info, role='BranchMakers', type='offline_sumited_branchmaker')


# ========================= MAKER / BRANCHMAKER OFFICER CODE END'S HERE ===================================================================


# ------------------------------------------------------------------------------------------------------------------------------------------------------


# ========================= CHECKER / ROS OFFICER CODE START'S HERE ===================================================================


# --------------------------------- ROS OFFICER LANDING PAGE OR DASHBOARD -------------------------------------------------------



# @app.route('/ROSDashboard', methods=['GET'])
# @secure_route(required_role='ROS')
# def ROSDashboard():
    
#     if 'email_id' not in session:
#         return redirect(url_for('post_login'))
    
#     ros_email = session['email_id']

#     # connROSDash = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
#     connROSDash = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

#     mysqlROSDash = connROSDash.cursor()

#     try:

#         mysqlROSDash.execute("SELECT * FROM [user] WHERE EmailId=?", (ros_email,))

#         rosuser = mysqlROSDash.fetchone()

#         if 'image' in rosuser:
#                 rosuser['image'] = base64.b64encode(rosuser['image']).decode('utf-8')
            
        
#         select_get_brach_code_query = "SELECT BranchCode FROM [user] WHERE EmailId = ?"

#         mysqlROSDash.execute(select_get_brach_code_query,(ros_email,))

#         B_code = mysqlROSDash.fetchone()[0]
        
#         select_get_user_id_query = "SELECT id FROM [user] WHERE BranchCode = ? AND Role = 'BranchMakers' "

#         mysqlROSDash.execute(select_get_user_id_query,(B_code,))

#         User_id = mysqlROSDash.fetchone()[0]
        
#         pending_cnt = 0
#         submited_cnt = 0
#         perday_pending_cnt = []
#         perday_submited_cnt = []

#         try:

#             select_total_pending_cnt_query = "SELECT COUNT(*) FROM [dbo].[offline_tickets] o1, [dbo].[offline_collection] o2 WHERE o1.[ticketId] = o2.[ticket_id] AND o1.[allocatedTicket] = ? AND o1.[ROSCasesTicket] IS NULL AND o2.[ROS_cmt] IS NULL AND o2.[DGM_cmt] IS NULL"
            
#             mysqlROSDash.execute(select_total_pending_cnt_query,(User_id,)) 
            
#             total_pending_cnt = mysqlROSDash.fetchone()[0]
            
#             if total_pending_cnt != 0:
#                 pending_cnt = total_pending_cnt
#             else:
#                 pending_cnt = 0


#             select_total_submited_cnt_query = "SELECT COUNT(*) FROM [dbo].[offline_tickets] WHERE [allocatedTicket] IN(?,5) AND [ROSCasesTicket] = ?"
            
#             mysqlROSDash.execute(select_total_submited_cnt_query,(rosuser[0],rosuser[0])) 
        
#             total_submited_cnt = mysqlROSDash.fetchone()[0]
        
#             if total_submited_cnt != 0:
#                 submited_cnt = total_submited_cnt
        
#             else:
#                 submited_cnt = 0

#             select_perday_total_pending_cnt_query = "SELECT O2.[Created_Date], COUNT(*) as cnt FROM [dbo].[offline_tickets] O1, [dbo].[offline_collection] O2 WHERE O1.[ticketId] = O2.[ticket_id] AND O1.[allocatedTicket] = ? AND O1.[ROSCasesTicket] IS NULL GROUP BY O2.[Created_Date]"
            
#             mysqlROSDash.execute(select_perday_total_pending_cnt_query,(User_id,))
            
#             perday_total_pending_cnt = mysqlROSDash.fetchall()
            
#             if perday_total_pending_cnt:
#                 for row in perday_total_pending_cnt:
#                     date_str = row[0] 
#                     count = row[1]
#                     perday_pending_cnt.append({"date": date_str, "count": count})
#             else:
#                 perday_pending_cnt = []


#             select_perday_total_submited_cnt_query = "SELECT O2.[Created_Date], COUNT(*) FROM [dbo].[offline_tickets] O1, [dbo].[offline_collection] O2 WHERE O1.[ticketId] = O2.[ticket_id] AND O1.[allocatedTicket] IN (?,5) AND O1.[ROSCasesTicket] = ? GROUP BY O2.[Created_Date]"
            
#             mysqlROSDash.execute(select_perday_total_submited_cnt_query,(rosuser[0], rosuser[0]))
            
#             perday_total_submited_cnt = mysqlROSDash.fetchall()
            
#             if perday_total_submited_cnt:
#                 for row in perday_total_submited_cnt:
#                     date_str = row[0]
#                     count = row[1]
#                     perday_submited_cnt.append({"date": date_str, "count": count})
#             else:
#                 perday_submited_cnt = []

#         except:
#             pending_cnt = 0
#             submited_cnt = 0
#             perday_pending_cnt = []
#             perday_submited_cnt = []

#         connROSDash.close()

#         return render_template('ros_dashboard.html', count=pending_cnt, countSubmited=submited_cnt, assignedperday=perday_pending_cnt, submittedperday=perday_submited_cnt, rosuser=rosuser, role='ROS', type='ROSDashboard',notify="1")
    
#     except Exception as e:

#         connROSDash.rollback()
#         connROSDash.close()

#         return f"Something went Wrong {e} , Please Re-Login Again",500



# @app.route('/ROSDashboard', methods=['GET', 'POST'])
# @secure_route(required_role='ROS')
# def ROSDashboard():
    
#     if 'email_id' not in session:
#         return redirect(url_for('post_login'))
#     ros_email = session['email_id']
#     cursor = mysql2.connection.cursor()

#     select_query = "SELECT * FROM [dbo].[user] WHERE [EmailId] = ?"
#     cursor.execute(select_query, ros_email)
#     rosuser = cursor.fetchone()
#     if 'image' in rosuser:
#         rosuser['image'] = base64.b64encode(rosuser['image']).decode('utf-8')
    
    
#     select_get_brach_code_query = "SELECT [BranchCode] FROM [dbo].[user] WHERE [EmailId] = ?"
#     cursor.execute(select_get_brach_code_query,(ros_email,))
#     B_code = cursor.fetchone()[0]
#     print('B_code :', B_code)
    
#     select_get_user_id_query = "SELECT [id] FROM [dbo].[user] WHERE [BranchCode] = ? AND [Role] = 'BranchMakers' "
#     cursor.execute(select_get_user_id_query,(B_code,))
#     User_id = cursor.fetchone()[0]
#     print('User_id :', User_id)
#     # notify = notification(ros_email)
    
#     pending_cnt = 0
#     submited_cnt = 0
#     perday_pending_cnt = []
#     perday_submited_cnt = []

#     # Total Pending Ticket Count
#     # select_total_pending_cnt_query = "SELECT COUNT(*) FROM [dbo].[offline_tickets] o1, [dbo].[offline_collection] o2 WHERE o1.[ticketId] = o2.[ticket_id] AND o1.[allocatedTicket] = ? AND o1.[ROSCasesTicket] IS NULL AND o2.[ROS_cmt] IS NULL AND o2.[DGM_cmt] IS NULL"
#     #MY QUR
#     select_total_pending_cnt_query ="""
#     SELECT COUNT(*) FROM [offline_scenarios]
#     WHERE [allocatedTicket] = ? AND
#     [ROSCasesTicket] IS NULL AND [ROS_cmt] IS NULL AND [DGM_cmt] IS NULL 
#     """
    
#     cursor.execute(select_total_pending_cnt_query,(User_id,)) 
#     total_pending_cnt = cursor.fetchone()[0]
#     if total_pending_cnt != 0:
#         pending_cnt = total_pending_cnt
#     else:
#         pending_cnt = 0
#     print('total_pending_cnt :', pending_cnt)

#     # Total Submited Ticket Count	
#     select_total_submited_cnt_query = "SELECT COUNT(*) FROM [dbo].[offline_scenarios] WHERE [allocatedTicket] IN(?,5) AND [ROSCasesTicket] = ?"
#     cursor.execute(select_total_submited_cnt_query,(rosuser[0],rosuser[0])) 
#     total_submited_cnt = cursor.fetchone()[0]
#     if total_submited_cnt != 0:
#         submited_cnt = total_submited_cnt
#     else:
#         submited_cnt = 0
#     print('total_submited_cnt :', submited_cnt)

#     # Perday Pending Ticket Count
#     # select_perday_total_pending_cnt_query = "SELECT O2.[Created_Date], COUNT(*) as cnt FROM [dbo].[offline_tickets] O1, [dbo].[offline_collection] O2 WHERE O1.[ticketId] = O2.[ticket_id] AND O1.[allocatedTicket] = ? AND O1.[ROSCasesTicket] IS NULL GROUP BY O2.[Created_Date]"
#     #MY QUR
#     select_perday_total_pending_cnt_query ="""
#     SELECT [Created_Date], COUNT(*) as cnt 
#     FROM [ticketid].[dbo].[offline_scenarios]
#     WHERE [allocatedTicket] = ? AND 
#     [ROSCasesTicket] IS NULL GROUP BY [Created_Date]
#     """
    
#     cursor.execute(select_perday_total_pending_cnt_query,(User_id,))
#     perday_total_pending_cnt = cursor.fetchall()
#     if perday_total_pending_cnt:
#         for row in perday_total_pending_cnt:
#             date_str = row[0] # Assuming the date is returned as string
#             count = row[1]
#             perday_pending_cnt.append({"date": date_str, "count": count})
#     else:
#         perday_pending_cnt = []
#     print('perday_pending_cnt :', perday_pending_cnt)

#     # Perday Submited Ticket Count
#     #select_perday_total_submited_cnt_query = "SELECT O2.[Created_Date], COUNT(*) FROM [dbo].[offline_tickets] O1, [dbo].[offline_collection] O2 WHERE O1.[ticketId] = O2.[ticket_id] AND O1.[allocatedTicket] IN (?,5) AND O1.[ROSCasesTicket] = ? GROUP BY O2.[Created_Date]"
#     select_perday_total_submited_cnt_query = """
#     SELECT [Created_Date], COUNT(*) 
#     FROM [ticketid].[dbo].[offline_scenarios]
#     WHERE [allocatedTicket] IN (?,5) AND
#     [ROSCasesTicket] = ? GROUP BY [Created_Date]
#     """
#     cursor.execute(select_perday_total_submited_cnt_query,(rosuser[0], rosuser[0]))
#     perday_total_submited_cnt = cursor.fetchall()
#     if perday_total_submited_cnt:
#         for row in perday_total_submited_cnt:
#             date_str = row[0] # Assuming the date is returned as string
#             count = row[1]
#             perday_submited_cnt.append({"date": date_str, "count": count})
#     else:
#         perday_submited_cnt = []
#     print('perday_submited_cnt :', perday_submited_cnt)

#     return render_template('ros_dashboard.html', count=pending_cnt, countSubmited=submited_cnt, assignedperday=perday_pending_cnt, submittedperday=perday_submited_cnt, rosuser=rosuser, role='ROS', type='ROSDashboard',notify="1")


@app.route('/ROSDashboard', methods=['GET', 'POST'])
@secure_route(required_role='ROS')
def ROSDashboard():
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    ros_email = session['email_id']
    cursor = mysql2.connection.cursor()

    try:
        select_query = "SELECT * FROM [dbo].[user] WHERE [EmailId] = ?"
        cursor.execute(select_query, (ros_email,))
        rosuser = cursor.fetchone()
        if 'image' in rosuser:
            rosuser['image'] = base64.b64encode(rosuser['image']).decode('utf-8')

        select_get_branch_code_query = "SELECT [BranchCode] FROM [dbo].[user] WHERE [EmailId] = ?"
        cursor.execute(select_get_branch_code_query, (ros_email,))
        B_code = cursor.fetchone()[0]
        print('B_code :', B_code)

        select_get_user_id_query = "SELECT [id] FROM [dbo].[user] WHERE [BranchCode] = ? AND [Role] = 'BranchMakers' "
        cursor.execute(select_get_user_id_query, (B_code,))
        User_id = cursor.fetchone()[0]
        print('User_id :', User_id)

        pending_cnt = 0
        submited_cnt = 0
        perday_pending_cnt = []
        perday_submited_cnt = []

        # Total Pending Ticket Count
        select_total_pending_cnt_query = """
        SELECT COUNT(*) FROM [dbo].[offline_scenarios]
        WHERE [allocatedTicket] = ? AND
        [ROSCasesTicket] IS NULL AND [ROS_cmt] IS NULL AND [DGM_cmt] IS NULL 
        """
        cursor.execute(select_total_pending_cnt_query, (User_id,))
        total_pending_cnt = cursor.fetchone()[0]
        if total_pending_cnt:
            pending_cnt = total_pending_cnt
        else:
            pending_cnt = 0
        print('total_pending_cnt :', pending_cnt)

        # Total Submited Ticket Count
        select_total_submited_cnt_query = """
        SELECT COUNT(*) FROM [dbo].[offline_scenarios] 
        WHERE [allocatedTicket] IN (?, 5) AND [ROSCasesTicket] = ?
        """
        cursor.execute(select_total_submited_cnt_query, (rosuser[0], rosuser[0]))
        total_submited_cnt = cursor.fetchone()[0]
        if total_submited_cnt:
            submited_cnt = total_submited_cnt
        else:
            submited_cnt = 0
        print('total_submited_cnt :', submited_cnt)

        # Perday Pending Ticket Count
        select_perday_total_pending_cnt_query = """
        SELECT [Created_Date], COUNT(*) as cnt 
        FROM [dbo].[offline_scenarios]
        WHERE [allocatedTicket] = ? AND 
        [ROSCasesTicket] IS NULL GROUP BY [Created_Date]
        """
        cursor.execute(select_perday_total_pending_cnt_query, (User_id,))
        perday_total_pending_cnt = cursor.fetchall()
        if perday_total_pending_cnt:
            for row in perday_total_pending_cnt:
                date_str = row[0]  # Assuming the date is returned as string
                count = row[1]
                perday_pending_cnt.append({"date": date_str, "count": count})
        else:
            perday_pending_cnt = []
        print('perday_pending_cnt :', perday_pending_cnt)

        # Perday Submited Ticket Count
        select_perday_total_submited_cnt_query = """
        SELECT [Created_Date], COUNT(*) 
        FROM [dbo].[offline_scenarios]
        WHERE [allocatedTicket] IN (?, 5) AND
        [ROSCasesTicket] = ? GROUP BY [Created_Date]
        """
        cursor.execute(select_perday_total_submited_cnt_query, (rosuser[0], rosuser[0]))
        perday_total_submited_cnt = cursor.fetchall()
        if perday_total_submited_cnt:
            for row in perday_total_submited_cnt:
                date_str = row[0]  # Assuming the date is returned as string
                count = row[1]
                perday_submited_cnt.append({"date": date_str, "count": count})
        else:
            perday_submited_cnt = []
        print('perday_submited_cnt :', perday_submited_cnt)

        return render_template('ros_dashboard.html', count=pending_cnt, countSubmited=submited_cnt, assignedperday=perday_pending_cnt, submittedperday=perday_submited_cnt, rosuser=rosuser, role='ROS', type='ROSDashboard', notify="1")
    
    except pyodbc.Error as e:
        print(f"Database error occurred: {e}")
        return render_template('ros_dashboard.html', count=0, countSubmited=0, assignedperday=[], submittedperday=[], rosuser={}, role='ROS', type='ROSDashboard', notify="1")




# ----------------------------------- ROS OFFICER PENDING CASES END - POINT ----------------------------------------------------



# @app.route('/ros', methods=['GET', 'POST'])
# @secure_route(required_role='ROS')
# def ros():

#     success_message = session.pop('success_message', None)

#     if 'email_id' not in session:
#         return redirect(url_for('post_login'))
    

#     # connROSPending = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
#     connROSPending = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

#     mysqlROSPending = connROSPending.cursor()
    
#     ros_email = session['email_id']

#     try:
#         try:    
    
#             select_query1 = "SELECT BranchCode FROM [user] WHERE EmailId = ?"

#             mysqlROSPending.execute(select_query1, (ros_email,))

#             branch_code = mysqlROSPending.fetchone()

#             B_code = branch_code[0]

#             select_query2 = "SELECT EmailId FROM user WHERE BranchCode = ? AND Role = 'BranchMakers'"

#             mysqlROSPending.execute(select_query2, (B_code,))

#             ros_email_id = mysqlROSPending.fetchall()
            
#             for tup in ros_email_id:    
#                 B_email = tup

#             createdby = B_email[0]

#             select_query3 = "SELECT * FROM [offline_collection] WHERE [Created_By] = ?  AND [ROS_cmt] IS NULL AND [DGM_cmt] IS NULL"

#             mysqlROSPending.execute(select_query3, (createdby,))
            
#             info = mysqlROSPending.fetchall()

#             connROSPending.close()

#         except:
#             info = []

#         return render_template("ros.html", success_message=success_message, data=info, role='ROS',type='ros')
    
#     except Exception as e:
        
#         mysqlROSPending.rollback()
#         connROSPending.close()

#         return f"Something went wrong {e} , Please Re-Login Again ",500


@app.route('/ros', methods=['GET', 'POST'])
@secure_route(required_role='ROS')
def ros():
    conn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
    cursor = conn.cursor()
    ros_email = session['email_id']
    print('RoS Email id:',ros_email)
    success_message = session.pop('success_message', None)
    # Check user login
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    
    # Here we are try to the data arrengment for ROS because now we are dealing with multiple ROS
    # Getting the Branch Code based on Email Id(Email Id stored in Session)
    select_query1 = "SELECT [BranchCode] FROM [user] WHERE [EmailId] = ?"
    cursor.execute(select_query1, (ros_email,))
    branch_code = cursor.fetchone()
    B_code = branch_code[0]
    # print(branch_code[0])

    # Getting the Email Id of Branchmaker based on Branch Code and Role
    select_query2 = "SELECT [EmailId] FROM [user] WHERE [BranchCode] = ? AND [Role] = 'BranchMakers'"
    cursor.execute(select_query2, (B_code,))
    ros_email_id = cursor.fetchall()
    #print(ros_email_id)
    
    # In the above method will give us tuple of Email Id so we have to convert it into list
    for tup in ros_email_id:    
        B_email = tup
    print(B_email[0])
    createdby = B_email[0]
    
    # Getting the data from offline_collection table based on Created By
    # cursor.execute("SELECT * FROM [ticketid].[dbo].[offline_collection] WHERE [ROS_cmt] IS NULL")
    select_query3 = "SELECT * FROM [dbo].[offline_scenarios] WHERE [Created_By] = ?  AND [ROS_cmt] IS NULL AND [DGM_cmt] IS NULL"
    try:
        cursor.execute(select_query3, (createdby,))
        info = cursor.fetchall()
        #print(info)
        return render_template("ros.html", success_message=success_message, data=info, role='ROS',type='ros')
    except pyodbc.Error as e:
        print(f"Database error occurred: {e}")
        return render_template("ros.html", success_message=success_message, data=[], role='ROS',type='ros')
    


# ---------------------------------- ROS OFFICER SUBMITED PAGE END - POINT -------------------------------------------------



# @app.route('/offline_submited_Str',methods=['POST','GET'])
# @secure_route(required_role='ROS')
# def offline_submited_Str():
#     success_message = session.pop('success_message', None)

#     if 'email_id' not in session:
#         return redirect(url_for('post_login'))
    

#     # connROSSub = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
#     connROSSub = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

#     mysqlROSSub = connROSSub.cursor()
    
#     ros_email = session['email_id']

#     try:

#         select_query1 = "SELECT id from [user] where [EmailId] = ?"

#         mysqlROSSub.execute(select_query1, (ros_email,))

#         branch_code = mysqlROSSub.fetchone()

#         B_id = branch_code[0]

#         try:

#             select_query2 = "SELECT ticketid FROM [offline_tickets] where ROSCasesTicket = ?"

#             mysqlROSSub.execute(select_query2, (B_id,))

#             ros_email_id = mysqlROSSub.fetchall()

#         except:

#             ros_email_id = []

#         info = []

#         T_id = []

#         if ros_email_id:

#             for tup in ros_email_id:
#                 B_email = tup
#                 T_id.append(B_email[0])

#         if T_id:
#             for i in T_id:
#                 select_query3 = "SELECT * FROM [offline_collection] where ticket_id = ?"
            
#                 mysqlROSSub.execute(select_query3, (i,))
            
#                 info.append(mysqlROSSub.fetchall())

#         connROSSub.close()

#         return render_template("offline_submited_data.html", success_message=success_message, data=info, role='ROS',type='offline_submited_Str')

#     except Exception as e:

#         mysqlROSSub.rollback()
#         connROSSub.close()

#         return f"Something went wrong {e} , Please Re-Login Again ",500


@app.route('/offline_submited_Str',methods=['POST','GET'])
@secure_route(required_role=['AGM','BranchMakers','ROS','DGM/PO'])
def offline_submited_Str():
    conn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
    cursor = conn.cursor()
    success_message = session.pop('success_message', None)
    ros_email = session['email_id']
    # Check user login
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    
    # Here we are try to the data arrengment for ROS because now we are dealing with multiple ROS
    # Getting the user id based on Email Id(Email Id stored in Session)
    select_query1 = "SELECT id from [dbo].[user] where [EmailId] = ?"
    cursor.execute(select_query1, (ros_email,))
    branch_code = cursor.fetchone()
    
    if not branch_code:
        return "User not found", 404
    
    B_id = branch_code[0]
    # print(branch_code[0])

    # Getting the ticket ids from the offline_scenarios table based on the user id
    # select_query2 = "SELECT ticket_id FROM [dbo].[offline_scenarios] WHERE ROSCasesTicket = ?"
    # cursor.execute(select_query2, (B_id,))
    # ticket_ids = cursor.fetchall()

    # print(ros_email_id)

    # checking the data like if data is present or not
    # if ticket_ids:
    #     T_id = []
        # In the above method will give us tuple of Ticket Id so we have to convert it into list
        # for tup in ticket_ids:
        #     B_email = tup
        #     T_id.append(B_email[0])
            
    # print("T_id.......", T_id)
    # info = []
    # for i in T_id:
    #     select_query3 = "SELECT * FROM [dbo].[offline_scenarios] where ticket_id = ?"
    #     cursor.execute(select_query3, (i,))
    #     info.append(cursor.fetchall())
        # print("iiiiiiiiinnnnnfffffffff",info)

    # Getting the tickets from the offline_scenarios table based on the user id
    select_query2 = "SELECT * FROM [dbo].[offline_scenarios] WHERE ROSCasesTicket = ?"
    try:
        cursor.execute(select_query2, (B_id,))
        scenarios = cursor.fetchall()
        print("Fetched scenarios:", scenarios)
        if not scenarios:
            scenarios = []
        print("scenarios-------",scenarios)
        # Render the template with the fetched data
        print("Data passed to template:", scenarios)
        return render_template("offline_submited_data.html", success_message=success_message, data=scenarios, role='ROS',type='offline_submited_Str')
    except pyodbc.Error as e:
        print(f"Database error occurred: {e}")
        return render_template("offline_submited_data.html", success_message=success_message, data=[], role='ROS',type='offline_submited_Str')
            

# ========================= CHECKER / RO OFFICER CODE ENDS'S HERE ===================================================================


# -----------------------------------------------------------------------------------------------------------------------------------------------------------



# ============================== ALL FORM PAGES FOR DISPLAYING THE DATA END - POINT CODE START's HERE ===================================================



# ------------------------------ MLRO INITIAL CASE CREATION FORM END - POINT ------------------------------------------


@app.route('/caseFormPage', methods=['GET','POST'])
@secure_route(required_role='MLRO')
def caseFormPage():
    if request.method == 'POST':

        # connMlroForm = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
        connMlroForm = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

        mysqlMlroForm = connMlroForm.cursor()
        
        try:
            ticket_id_from_form = request.form['ticket_id']
            
            mlro_email = session['email_id']

            query = "SELECT * FROM [user] WHERE EmailId = ?"

            mysqlMlroForm.execute(query, (mlro_email,))
            
            rows = mysqlMlroForm.fetchall()

            columns = [col[0] for col in mysqlMlroForm.description]

            mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

            mlro = mlro[0]


            data = []

            
            if 'image' in mlro:
                    mlro['image'] = base64.b64encode(mlro['image']).decode('utf-8')

            mysqlMlroForm.execute("SELECT EmailId FROM [user] WHERE Role = ?",('CM/SM',))
            allocatedCM = mysqlMlroForm.fetchone()[0]


            dataQuery = "SELECT * FROM scenarios WHERE ticketid = ?"  

            mysqlMlroForm.execute(dataQuery, (ticket_id_from_form,))  

            columns = [desc[0] for desc in mysqlMlroForm.description]

            res = mysqlMlroForm.fetchall()

            connMlroForm.close()

            for row in res:
                scenario_object = {}
                for i, value in enumerate(row):
                    scenario_object[columns[i]] = value
                data.append(scenario_object)


            accNum = scenario_object['ACCTNO']
            CustomerId = scenario_object['CUSTCD']
            CustomerName = scenario_object['CustomerName']

            txdate = scenario_object['TXDATE']

            debit_details,credit_details,cash_deposits,cash_withdrawals = cummilativeFunction(accNum,None,txdate)


            debitTurnOverOver = debit_details[0]
            creditTurnOverOver = credit_details[0]
            cash_deposits = cash_deposits[0]
            cash_withdrawals = cash_withdrawals[0]
            
            return render_template('Form_Page.html',ticket_id=ticket_id_from_form,accNum=accNum,CustomerId=CustomerId,allocatedCM=allocatedCM,data=data,role='MLRO',type='MLRONextLevel',mlrouser=mlro,CustomerName=CustomerName,debitTurnOverOver=debitTurnOverOver,creditTurnOverOver=creditTurnOverOver,cash_deposits=cash_deposits,cash_withdrawals=cash_withdrawals)
        
        except Exception as e:

            mysqlMlroForm.rollback()
            connMlroForm.close()

            return f'Something Went Wrong {e} , Please Re-Login Again', 500
    else:
        return 'Bad Request Try to Login Again',400
    



def cummilativeFunction(ACCNO,DATE_of_SUBMITION_GM,txdate):
    quiryForAccTRD = """SELECT 
                    SUM(CONVERT(FLOAT, TXAMT)) AS TotalDebitAmount
                FROM 
                    TRANSACTIONS
                WHERE 
                    TRNFLOWTYPE = 'D' AND ACCTNO = ?
                    AND TRY_CONVERT(datetime, TXDATE, 105) >= DATEADD(MONTH, -12, ?)
                GROUP BY 
                    ACCTNO;
                """
    
    quiryForAccTRC = """SELECT 
                    SUM(CONVERT(FLOAT, TXAMT)) AS TotalCreditAmount
                FROM 
                    TRANSACTIONS
                WHERE 
                    TRNFLOWTYPE = 'C' AND ACCTNO = ? 
                    AND TRY_CONVERT(datetime, TXDATE, 105) >= DATEADD(MONTH, -12, ?)
                GROUP BY 
                    ACCTNO;
                """
    
    quiryForAccTRCdeposits= """SELECT 
                    SUM(CONVERT(FLOAT, TXAMT)) AS TotalCreditAmount
                FROM 
                    TRANSACTIONS
                WHERE 
                    TRNFLOWTYPE = 'C' AND ACCTNO = ? AND TXTYPE = 'C'
                    AND TRY_CONVERT(datetime, TXDATE, 105) >= DATEADD(MONTH, -12, ?)
                GROUP BY 
                    ACCTNO;
                """
    

    quiryForAccTRCWith = """SELECT 
                    SUM(CONVERT(FLOAT, TXAMT)) AS TotalCashDeposits
                FROM 
                    TRANSACTIONS
                WHERE 
                    TXTYPE = 'C' AND TRNFLOWTYPE = 'D' AND ACCTNO = ?
                    AND TRY_CONVERT(datetime, TXDATE, 105) >= DATEADD(MONTH, -12, ?)
                GROUP BY 
                    ACCTNO;
                """

    try:
        divDate = txdate.split(' to ')
        presentDate1 = datetime.strptime(divDate[1], '%Y-%m-%d').strftime('%Y-%m-%d')
        pastDate1 = datetime.strptime(divDate[0], '%Y-%m-%d').strftime('%Y-%m-%d')
    except:
        
        parsed_date = datetime.strptime(txdate, '%d-%m-%Y')
        presentDate1 = parsed_date.strftime('%Y-%m-%d')


    debit_details = fetch_single_result(quiryForAccTRD, (ACCNO,presentDate1 ), (0, 0.0))
    credit_details = fetch_single_result(quiryForAccTRC, (ACCNO,presentDate1 ), (0, 0.0))
    cash_deposits = fetch_single_result(quiryForAccTRCdeposits, (ACCNO,presentDate1 ), (0.0,))
    cash_withdrawals = fetch_single_result(quiryForAccTRCWith, (ACCNO, presentDate1), (0.0,))

    return debit_details,credit_details,cash_deposits,cash_withdrawals



# ------------------------------ AFTER CASE CREATION EDITING THE FORM BY NEXT LEAVEL OFFICER END - POINT ------------------------------------------


@app.route('/caseFormPageEdit',methods=['GET','POST'])
@secure_route(required_role=['CM/SM','DGM/PO','MLRO'])
def caseFormPageEdit():
    if request.method == 'POST':

        # connCaseFormEdit = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
        connCaseFormEdit = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

        mysqlCaseFormEdit = connCaseFormEdit.cursor()

        try:
            ticket_id_from_form = request.form['ticket_id']
        
            email = session['email_id']

            query = "SELECT * FROM [user] WHERE EmailId = ?"

            mysqlCaseFormEdit.execute(query, (email,))
            
            rows = mysqlCaseFormEdit.fetchall()

            columns = [col[0] for col in mysqlCaseFormEdit.description]

            info = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

            info = info[0]


            data = []

            
            if 'image' in info:
                        # Encode the image data as a base64 string
                        info['image'] = base64.b64encode(info['image']).decode('utf-8')

    
            dataQuery = "SELECT * FROM scenarios WHERE ticketid = ?"  

            mysqlCaseFormEdit.execute(dataQuery, (ticket_id_from_form,))  

            columns = [desc[0] for desc in mysqlCaseFormEdit.description]

            for row in mysqlCaseFormEdit.fetchall():
                scenario_object = {}
                for i, value in enumerate(row):
                    scenario_object[columns[i]] = value
                data.append(scenario_object)


            accNum = scenario_object['ACCTNO']
            CustomerId = scenario_object['CUSTCD']
            CustomerName = scenario_object['CustomerName']
            


            if info["Role"] == "CM/SM":
                mysqlCaseFormEdit.execute("SELECT EmailId FROM [user] WHERE Role = ?",('DGM/PO',))
                allocated = mysqlCaseFormEdit.fetchone()[0]
                endpo = 'CM_SM_NextLevel'

            if info["Role"] == "DGM/PO":
                allocated = None
                endpo = 'DGMNextLevel'

            connCaseFormEdit.close()

            return render_template('Form_PageEdit.html',ticket_id=ticket_id_from_form,accNum=accNum,CustomerId=CustomerId,role=info['Role'],allocated=allocated,data=data,type=endpo,allImages=info,CustomerName=CustomerName)
        except:
            mysqlCaseFormEdit.rollback()
            connCaseFormEdit.close()
            return 'Something Went Wrong please Re-Login Again',500
    else:
        return 'Bad Request Try to Submit Again',400




# ----------------------------- SEND BACK CASES FORM EDITABLE WITH REASON FOR SEND BACK --------------------------------------



@app.route('/sent_back_case_FormPage',methods=['POST','GET'])
@secure_route(required_role=['MLRO','AGM','CM/SM'])
def sent_back_case_FormPage():
    if request.method == 'POST':

        ticket_id_from_form = request.form['ticket_id']
        CustomerId = request.form['custId']
       
        email = session['email_id']

        # connDispSentBackCaseForm = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
        connDispSentBackCaseForm = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

        mysqlDispSentBackCaseForm = connDispSentBackCaseForm.cursor()
        
        try:


            query = "SELECT * FROM [user] WHERE EmailId = ?"

            mysqlDispSentBackCaseForm.execute(query, (email,))
            
            rows = mysqlDispSentBackCaseForm.fetchall()

            columns = [col[0] for col in mysqlDispSentBackCaseForm.description]

            info = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

            info = info[0]


            data = []

            
            if 'image' in info:
                        # Encode the image data as a base64 string
                        info['image'] = base64.b64encode(info['image']).decode('utf-8')

    
            dataQuery = "SELECT * FROM scenarios WHERE ticketid = ?"  

            mysqlDispSentBackCaseForm.execute(dataQuery, (ticket_id_from_form,))  

            columns = [desc[0] for desc in mysqlDispSentBackCaseForm.description]

            for row in mysqlDispSentBackCaseForm.fetchall():
                scenario_object = {}
                for i, value in enumerate(row):
                    scenario_object[columns[i]] = value
                data.append(scenario_object)


            accNum = scenario_object['ACCTNO']
            CustomerId = scenario_object['CUSTCD']
            CustomerName = scenario_object['CustomerName']
            


            if info["Role"] == "CM/SM":
                mysqlDispSentBackCaseForm.execute("SELECT EmailId FROM [user] WHERE Role = ?",('DGM/PO',))
                allocated = mysqlDispSentBackCaseForm.fetchone()[0]

            if info["Role"] == "MLRO":
                mysqlDispSentBackCaseForm.execute("SELECT EmailId FROM [user] WHERE Role = ?",('CM/SM',))
                allocated = mysqlDispSentBackCaseForm.fetchone()[0]
        
            connDispSentBackCaseForm.close()

            return render_template('sent_back_case_form.html',ticket_id=ticket_id_from_form,CustomerId=CustomerId,role=info["Role"],allocated=allocated,data=data,type='display_Sent_Back_Alerts',allImages=info,CustomerName=CustomerName,accNum=accNum)

        except Exception as e:

            mysqlDispSentBackCaseForm.rollback()
            connDispSentBackCaseForm.close()

            return f"Something went Wrong {e} , Please Re-Login Again",500
    else:
        return f"Bad Request Try to Login Again",400



# ------------------------------- 10 % SEND BACK CASE FORMPAGE IN THE MLRO WITH REASON TO MAKE IT AS STR END - POINT -------------------------------


@app.route('/returnCaseFormEdit',methods=['POST','GET'])
@secure_route(required_role=['MLRO','CM/SM'])
def returnCaseFormEdit():
    if request.method == 'POST':

        ticket_id_from_form = request.form['ticket_id']

        CustomerId = request.form['custId']

        mlro_email = session['email_id']

        # connsnetbackClosededit = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
        connsnetbackClosededit = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

        mysqlsentbackClosededit = connsnetbackClosededit.cursor()

        try:
    
            query = "SELECT * FROM [user] WHERE EmailId = ?"
            mysqlsentbackClosededit.execute(query, (mlro_email,))
            mlro = mysqlsentbackClosededit.fetchone()
    
            if mlro is None:
                return "User data not found. Please log in again."

            if 'image' in mlro:
                mlro['image'] = base64.b64encode(mlro['image']).decode('utf-8')

            data = []
            

            

            dataQuery = f"SELECT * FROM scenarios WHERE ticketid = ? "
    

            mysqlsentbackClosededit.execute(dataQuery, (ticket_id_from_form,))
            
            columns = [desc[0] for desc in mysqlsentbackClosededit.description]

            res = mysqlsentbackClosededit.fetchall()
            
            connsnetbackClosededit.close()
        
            for row in res:
                scenario_object = {}
                for i, value in enumerate(row):
                    scenario_object[columns[i]] = value
                data.append(scenario_object)
                CustomerId = scenario_object['CUSTCD']
                CustomerName = scenario_object['CustomerName']
                ticket_id = scenario_object['ticketid']
                accNo = scenario_object['ACCTNO']
                allocatedCM = scenario_object['cm_level']
            
            return render_template('closed_Form_Page.html',ticket_id=ticket_id,CustomerName=CustomerName,accNum=accNo,CustomerId=CustomerId,mlrouser=mlro,role='MLRO',allocatedCM=allocatedCM,data=data,type='return_Mlro_Alerts',allImages=mlro)
        
        except Exception as e:

            mysqlsentbackClosededit.rollback()
            connsnetbackClosededit.close()

            return f"Somthing went wrong {e} Please Re-login Again",500



# ------------------ INSIDE FORM PAGE SEDN BACK BUTTON TO SEND THE STR CASE TO PREV OFFICER "" UNSATISFIED CASES "" -------------------------



@app.route('/sendBackCaseCreated',methods=['POST','GET'])
@secure_route(required_role=['CM/SM','DGM/PO'])
def sendBackCaseCreated():
    if request.method == 'POST':


        # connSentBackCase = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
        connSentBackCase = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

        mysqlSentBackCase = connSentBackCase.cursor()

        try:
            ticketId = request.form.get('tickId')

            comment = request.form.get('sendBackComment')

            reverseTo = request.form.get('allocateTo')

            typee = request.form.get('typee')
            

            mailid = session['email_id']


            mysqlSentBackCase.execute("SELECT Role FROM [user] WHERE EmailId = ?", (mailid,))
            role = mysqlSentBackCase.fetchone()[0]


            mysqlSentBackCase.execute("SELECT id FROM [user] WHERE EmailId = ?", (reverseTo,))
            reverseToId = mysqlSentBackCase.fetchone()[0]

            if role == 'DGM/PO':
                mysqlSentBackCase.execute("SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = 'scenarios' AND COLUMN_NAME = 'sentBackReason'")
                column_exists = mysqlSentBackCase.fetchone()

                if not column_exists:
                    mysqlSentBackCase.execute("ALTER TABLE scenarios ADD sentBackReason TEXT")

                query = "UPDATE scenarios SET sentBackReason = ? WHERE ticketid = ?"

                mysqlSentBackCase.execute(query, (comment, ticketId))

                mysqlSentBackCase.commit()

                mysqlSentBackCase.execute("UPDATE scenarios SET unsatisfiedTicket = ?, allocatedTicket = Null ,cmSMCasesTicket = NULL WHERE ticketid = ?", (reverseToId, ticketId))
                mysqlSentBackCase.commit()

            elif role == 'CM/SM':
                mysqlSentBackCase.execute("SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = 'scenarios' AND COLUMN_NAME = 'sentBackReason'")
                column_exists = mysqlSentBackCase.fetchone()

                if not column_exists:
                    mysqlSentBackCase.execute("ALTER TABLE scenarios ADD sentBackReason TEXT")

                query = "UPDATE scenarios SET sentBackReason = ? WHERE ticketid = ?"
                mysqlSentBackCase.execute(query, (comment, ticketId))
                mysqlSentBackCase.commit()

                mysqlSentBackCase.execute("UPDATE scenarios SET unsatisfiedTicket = ?, allocatedTicket = Null ,mlroCasesTicket = NULL WHERE ticketid = ?", (reverseToId, ticketId))
                mysqlSentBackCase.commit()



            if typee == 'CM_SM_NextLevel':
                redirect_to_endpoint = typee
            else:
                redirect_to_endpoint = 'display_Sent_Back_Alerts'

            if role != 'DGM/PO':
                return redirect(url_for(redirect_to_endpoint))
            else:
                return redirect(url_for("DGMNextLevel"))
        except Exception as e:

            mysqlSentBackCase.rollback()
            connSentBackCase.close()

            return f"Somthing went wrong {e} , Please Re-login Again ",500
    else:
        return f"Bad Request Try to Login Again",400



# ------------------------------- BRANCHMAKER / ROS / DGM OFFLINE EDITABLE FORM PAGE --------------------------------------------------



# @app.route('/offline_Form_Edit',methods=['POST','GET'])
# @secure_route(required_role=['ROS','DGM/PO','BranchMakers'])
# def offline_Form_Edit():
#     success_message = session.pop('success_message', None)
#     # Check user login
#     if 'email_id' not in session:
#         return redirect(url_for('post_login'))
#     account_number = request.args.get('account_number')
    
#     emailId = session['email_id']

#     cursor = mysql2.connection.cursor()

#     select_query = "SELECT * FROM [user] WHERE [EmailId] = ?"
#     cursor.execute(select_query, (emailId,))
#     user = cursor.fetchone()
  

#     cursor.execute("SELECT * FROM [offline_collection] WHERE CAST([ticket_id] AS VARCHAR(MAX)) = ?",(str(account_number),))
#     info = cursor.fetchall()
#     columns = [column[0] for column in cursor.description]  # Get column names
#     print(columns)
#     role = ''
#     typeEnd = ''
#     for col in user:
#         print(col)
#         if col == 'DGM/PO':
#             typeEnd = 'offline_dgm_Str'
#             role = 'DGM/PO'  
#         elif col == 'ROS':
#             typeEnd = 'ros'
#             role = 'ROS'
#         elif col == 'BranchMakers':
#             typeEnd = 'archived'
#             role = 'BranchMakers'

#     return render_template("offline_Form_edit.html", success_message=success_message, data=info, columns=columns, role=role,type=typeEnd)


@app.route('/offline_Form_Edit',methods=['POST','GET'])
@secure_route(required_role=['AGM','ROS','DGM/PO','BranchMakers'])
def offline_Form_Edit():
    conn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
    cursor = conn.cursor()
    success_message = session.pop('success_message', None)
    # Check user login
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    account_number = request.args.get('account_number')
    
    emailId = session['email_id']

    select_query = "SELECT * FROM [dbo].[user] WHERE [EmailId] = ?"
    cursor.execute(select_query, (emailId,))
    user = cursor.fetchone()
    print(user)
    print(user[6])

    cursor.execute("SELECT * FROM [dbo].[offline_scenarios] WHERE CAST([ticket_id] AS VARCHAR(MAX)) = ?",(str(account_number),))
    info = cursor.fetchall()
    print(info,"info")
    print(info[:1])
    columns = [column[0] for column in cursor.description]  # Get column names
    print(columns)
    role = ''
    typeEnd = ''
    for col in user:
        print(col)
        if col == 'DGM/PO':
            typeEnd = 'offline_dgm_Str'
            role = 'DGM/PO'  
            # return render_template("offline_Form_edit.html", success_message=success_message, data=info, columns=columns, role='DGM/PO',type=typeEnd)
        elif col == 'ROS':
            typeEnd = 'ros'
            role = 'ROS'
            # return render_template("offline_Form_edit.html", success_message=success_message, data=info, columns=columns, role='ROS',type=typeEnd) 
        elif col == 'BranchMakers':
            typeEnd = 'archived'
            role = 'BranchMakers'
    print('role:',role)
    print('typeend:',typeEnd)
    return render_template("offline_Form_edit.html", success_message=success_message, data=info, columns=columns, role=role,type=typeEnd)




# ============================== ALL FORM PAGES FOR DISPLAYING THE DATA END - POINT CODE END's HERE ===================================================




# -----------------------------------------------------------------------------------------------------------------------------------------------------------------





# ============================== ALL FORM PAGES SUBMITE THE DATA END - POINT CODE START's HERE ===================================================



# ------- ONLINE CASE SUBMIT TO THE NEXT LEVEL OFFICER's AND SUBMIT THE SENT BACK FORM FROM LOW LEVEL OFFICER AND CLOSED CASES SUBMITION's ----------------------------


@app.route('/submetCaseForm', methods=['GET','POST'])
@secure_route(required_role=['MLRO','AGM','CM/SM','DGM/PO'])
def submetCaseForm():
    if request.method == 'POST':

            # connCaseForm = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
            connCaseForm = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

            mysqlCaseForm = connCaseForm.cursor()

            ticket_id_from_form = request.form['ticket_id'] 
            sentTo = request.form['sentTo']
            crime = request.form['crime']
            Classification = request.form['caseAllocation']
            SuspeciousDuo = request.form['SuspeciousDuo']
            SuspeciousDuoeco = request.form['SuspeciousDuoeco']
            SuspeciousDuofinancing = request.form['SuspeciousDuofinancing']
            Investigation = request.form['Investigation']
            LEAinformed = request.form['LEAinformed']
            LEADetails = request.form['LEADetails']
            reportCase = request.form['reportCase']
            additionalInfo = request.form['additionalInfo']
            priorityRAG = request.form['priorityRAG']
            cumulativeCrtTurnover = request.form['cumulativeCrtTurnover']
            cumulativeDTTurnover = request.form['cumulativeDTTurnover']
            cumulativeDepositTurnover = request.form['cumulativeDepositTurnover']
            cumulativeWithdrawalTurnover = request.form['cumulativeWithdrawalTurnover']
            NumTransactionReport = request.form['NumTransactionReport']
            remark = request.form['remark']
            Raise_STR = request.form['strraising']
            comment = request.form['comments']
            finalReport = request.form['finalReport']
            prefilled = request.form['prefilled']

            typee = request.form.get('typee', None)
            typeee = request.form.get('typeee', None)
            
            try:



                if 'file' not in request.files:
                    print('No file part')

                file = request.files['file']

                if file.filename == '':
                    print('No selected file')

        
                user = session['email_id']
        
                query = "SELECT * FROM [user] WHERE EmailId = ?"

                mysqlCaseForm.execute(query, (user,))
                
                rows = mysqlCaseForm.fetchall()

                columns = [col[0] for col in mysqlCaseForm.description]

                mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]

                info = mlro[0]
                
                role = info["Role"]

                if role == 'CM/SM':
                    role = "cmsm"
                if role == 'DGM/PO':
                    role = "dgm"

                sentById = info['id']



                formData = {
                    "sentTo":sentTo,
                    "SuspeciousDuoCrime":crime,
                    "Classification":Classification,
                    "SuspeciousDuoComplexTr":SuspeciousDuo,
                    "SuspeciousDuoNoeco":SuspeciousDuoeco,
                    "terrorisumFunding":SuspeciousDuofinancing,
                    "Investigation":Investigation,
                    "LEAinformed":LEAinformed,
                    "LEADetails":LEADetails,
                    "reportCase":reportCase,
                    "additionalInfo":additionalInfo,
                    "priorityRAG":priorityRAG,
                    "cumulativeCrtTurnover":cumulativeCrtTurnover,
                    "cumulativeDTTurnover":cumulativeDTTurnover,
                    "cumulativeDepositTurnover":cumulativeDepositTurnover,
                    "cumulativeWithdrawalTurnover":cumulativeWithdrawalTurnover,
                    "NumTransactionReport":NumTransactionReport,
                    "remark":remark,
                    "Raise_STR":Raise_STR,
                    f"{role}comment":comment,
                    "finalReport":finalReport,
                    f"{role}prefilled":prefilled,
                    "fileName":file.filename
                    }

                allKeys = list(formData.keys())


                for key in allKeys:
                    
                    mysqlCaseForm.execute(f"SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = 'scenarios' AND COLUMN_NAME = ?", (key,))
                    column_exists = mysqlCaseForm.fetchone()

                    if not column_exists:
                        
                        mysqlCaseForm.execute(f"ALTER TABLE scenarios ADD {key} TEXT")



                for key, value in formData.items():
                            query = f"UPDATE scenarios SET {key} = ? WHERE ticketid = ?"
                            mysqlCaseForm.execute(query, (value, ticket_id_from_form))
                            mysqlCaseForm.commit()


                if info["Role"] != 'DGM/PO':
                    mysqlCaseForm.execute("SELECT * FROM [user] WHERE EmailId = ?",(sentTo,))

                    sentToId = mysqlCaseForm.fetchone()[0]

                current_datetime = datetime.now()
                current_date = current_datetime.date()
                midnight_datetime = datetime.combine(current_date, datetime.min.time())
                
                
                if info["Role"] == "MLRO":

                    
                    mysqlCaseForm.execute("SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = 'scenarios' AND COLUMN_NAME = 'mlro_level'")
                    mlro_column_exists = mysqlCaseForm.fetchone()

                    if not mlro_column_exists:
                        mysqlCaseForm.execute("ALTER TABLE scenarios ADD mlro_level TEXT")

                    mysqlCaseForm.execute("UPDATE scenarios SET mlro_level = ? WHERE ticketid = ?", (user, ticket_id_from_form))
                    mysqlCaseForm.commit()

                    if Raise_STR == 'STR':

                        mysqlCaseForm.execute("UPDATE scenarios SET allocatedTicket = ?,mlroCasesTicket = ?,currentDate = ? WHERE ticketid = ?",(sentToId,sentById,str(midnight_datetime),ticket_id_from_form))
                        mysqlCaseForm.commit()

                        if typeee:
                            mysqlCaseForm.execute("UPDATE scenarios SET cm_sm_sentBack_closed_comment = NULL WHERE ticketid = ?", (ticket_id_from_form,))
                            mysqlCaseForm.commit()
                    
                    if Raise_STR == 'NON-STR':
                    # add time for taking 10 % if the time was not there aleready....
                        conformAleredyTen = mysqlCaseForm.execute("SELECT currentDate FROM scenarios WHERE ticketid = ?",(ticket_id_from_form,)).fetchone()

                        if conformAleredyTen[0]:

                            mysqlCaseForm.execute("UPDATE scenarios SET allocatedTicket = ?,mlroClosedTicket = ?, WHERE ticketid = ?",(None,sentById,ticket_id_from_form))
                            mysqlCaseForm.commit()

                        else:
                            mysqlCaseForm.execute("UPDATE scenarios SET allocatedTicket = ?,mlroClosedTicket = ?,currentDate = ? WHERE ticketid = ?",(None,sentById,str(midnight_datetime),ticket_id_from_form))
                            mysqlCaseForm.commit()


                    if typee:

                        mysqlCaseForm.execute("UPDATE scenarios SET unsatisfiedTicket = NULL WHERE ticketid = ?", (ticket_id_from_form,))
                        mysqlCaseForm.commit()

                        try:
                            mysqlCaseForm.execute("UPDATE scenarios SET cmsmcomment = NULL WHERE ticketid = ?", (ticket_id_from_form,))
                            mysqlCaseForm.commit()
                        except:
                            print("No Cm Comment field to clear.")

                        return redirect(url_for('display_Sent_Back_Alerts'))
                    
                    elif typeee:

                        mysqlCaseForm.execute("UPDATE scenarios SET sentBackClosedTicket = NULL WHERE ticketid = ?", (ticket_id_from_form,))
                        mysqlCaseForm.commit()

                    
                        return redirect(url_for('return_Mlro_Alerts'))

                    else:
                        return redirect(url_for('MLRONextLevel'))
                
                if info["Role"] == "CM/SM":
                        
                        mysqlCaseForm.execute("SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = 'scenarios' AND COLUMN_NAME = 'cm_level'")
                        mlro_column_exists = mysqlCaseForm.fetchone()

                        if not mlro_column_exists:
                            mysqlCaseForm.execute("ALTER TABLE scenarios ADD cm_level TEXT")

                        mysqlCaseForm.execute("UPDATE scenarios SET cm_level = ? WHERE ticketid = ?", (user, ticket_id_from_form))
                        mysqlCaseForm.commit()
                
                        mysqlCaseForm.execute("UPDATE scenarios SET allocatedTicket = ?,cmSMCasesTicket = ?,currentDate = ? WHERE ticketid = ?",(sentToId,sentById,str(midnight_datetime),ticket_id_from_form))
                        mysqlCaseForm.commit()


                        if typee:
                            mysqlCaseForm.execute("UPDATE scenarios SET unsatisfiedTicket = NULL WHERE ticketid = ?", (ticket_id_from_form,))

                            mysqlCaseForm.commit()

                            connCaseForm.close()

                            return redirect(url_for('display_Sent_Back_Alerts'))
                        else:
                            
                            connCaseForm.close()

                            return redirect(url_for('CM_SM_NextLevel'))

                if info["Role"] == "DGM/PO":
                        
                
                        if finalReport == 'Approved':
                            mysqlCaseForm.execute("SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = 'scenarios' AND COLUMN_NAME = 'reportedDate'")
                            mlro_column_exists = mysqlCaseForm.fetchone()

                            # If reportedDate column does not exist, add it with datetime data type
                            if not mlro_column_exists:
                                mysqlCaseForm.execute("ALTER TABLE scenarios ADD reportedDate DATETIME")

                            mysqlCaseForm.execute("UPDATE scenarios SET reportedDate = GETDATE() WHERE ticketid = ?", ( ticket_id_from_form,))
                            mysqlCaseForm.commit()

                            mysqlCaseForm.execute("UPDATE scenarios SET allocatedTicket = ?, gmCasesTicket = ?, approved = ?,currentDate = ? WHERE ticketid = ?", (None, sentById, sentById,str(midnight_datetime),ticket_id_from_form))
                            mysqlCaseForm.commit()

                        if finalReport == 'Rejected':

                            mysqlCaseForm.execute("SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = 'scenarios' AND COLUMN_NAME = 'reportedDate'")
                            mlro_column_exists = mysqlCaseForm.fetchone()

                            # If mlro_level column does not exist, add it
                            if not mlro_column_exists:
                                mysqlCaseForm.execute("ALTER TABLE scenarios ADD reportedDate DATETIME")


                            mysqlCaseForm.execute("UPDATE scenarios SET reportedDate = GETDATE() WHERE ticketid = ?", ( ticket_id_from_form,))
                            mysqlCaseForm.commit()

                            mysqlCaseForm.execute("UPDATE scenarios SET allocatedTicket = ?,gmCasesTicket = ?,rejected = ?,currentDate = ? WHERE ticketid = ?",(None,sentById,sentById,str(midnight_datetime),ticket_id_from_form))
                            mysqlCaseForm.commit()

                        connCaseForm.close()
                        
                        return redirect(url_for('DGMNextLevel'))
                
            except Exception as e:

                mysqlCaseForm.rollback();
                connCaseForm.close()

                return f"Somthing Went Wrong {e} , Please Re-Login Again ",500
     
    else:
        return 'Bad Request Try to Submit Again',400



# --------------------------------- OFFLINE CASES SUBMIT TO NEXT LEVEL OFFICERS END - POINT --------------------------------------------


# @app.route('/post_manual_str', methods=['POST'])
# @secure_route(required_role=['BranchMakers'])
# def post_manual_str():


#     # connCaseOfflineFormSub = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
#     connCaseOfflineFormSub = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

#     conn = connCaseOfflineFormSub.cursor()

#     Tick_id = request.args.get('account_number')
#     print('offid:',Tick_id)
#     try:
#         # cursor = conn.cursor()
#         branchmaker_email = session.get('email_id')
#         if request.method == 'POST':
#             current_datetime = datetime.now()
#             current_date = current_datetime.date()
#             current_date_str = current_date.isoformat()

#             Created_Date = current_date_str
#             customerno = request.form.get('Customerno')
#             casename = request.form.get('casename')
#             personname = request.form.get('personname')
#             RuleScenario = request.form.get('RuleScenario')
#             Guidance = request.form.get('Guidance')
#             scenario = request.form.get('scenario')
#             SourceofAlert = request.form.get('SourceofAlert')
#             alertindicator = request.form.get('alertindicator')
#             SuspiciousDueToproceedofCrime = request.form.get('SuspiciousDueToproceedofCrime')
#             SuspiciousDueToComplexTranscaction = request.form.get('SuspiciousDueToComplexTranscaction')
#             SuspiciousDueToNoecoRational = request.form.get('SuspiciousDueToNoecoRational')
#             SuspiciousDueToFinancingTerrorism = request.form.get('SuspiciousDueToFinancingTerrorism')
#             AttemptedTranscaction = request.form.get('AttemptedTranscaction')
#             LEAInformed = request.form.get('LEAInformed')
#             PriorityRating = request.form.get('PriorityRating')
#             ReportCoverage = request.form.get('ReportCoverage')
#             leadetails = request.form.get('leadetails')
#             AdditionalDocument = request.form.get('AdditionalDocument')
#             Aroundofsuspision = request.form.get('Aroundofsuspision')
#             DetailsofInvestigation = request.form.get('DetailsofInvestigation')
#             AccountNumber = request.form.get('AccountNumber')
#             AccountType = request.form.get('AccountType')
#             holdername = request.form.get('holdername')
#             AccountHolderType = request.form.get('AccountHolderType')
#             AccountStatus = request.form.get('AccountStatus')
#             DateofOpening = request.form.get('DateofOpening')
#             RiskRating = request.form.get('RiskRating')
#             CummulativeCerditTurnover = request.form.get('CummulativeCerditTurnover')
#             CummulativeDebitTurnover = request.form.get('CummulativeDebitTurnover')
#             CummulativeCashDepositTurnover = request.form.get('CummulativeCashDepositTurnover')
#             CummulativeCashWithdrawalTurnover = request.form.get('CummulativeCashWithdrawalTurnover')
#             NoOfTransactionsToBeReported = request.form.get('NoOfTransactionsToBeReported')
#             TransactionDate = request.form.get('TransactionDate')
#             TransactionsID = request.form.get('TransactionsID')
#             TransactionMode = request.form.get('TransactionMode')
#             DebitCredit = request.form.get('DebitCredit')
#             amount = request.form.get('amount')
#             TransactionsCurrency = request.form.get('TransactionsCurrency')
#             ProductType = request.form.get('ProductType')
#             ProductIdentifiers = request.form.get('ProductIdentifiers')
#             TransactionType = request.form.get('TransactionType')
#             unit = request.form.get('unit')
#             Date = request.form.get('Date')
#             DispositionOfFunds = request.form.get('DispositionOfFunds')
#             RelatedAccountNumber = request.form.get('RelatedAccountNumber')
#             RelatedInstitutionName = request.form.get('RelatedInstitutionName')
#             RelatedInstitutionRefNum = request.form.get('RelatedInstitutionRefNum')
#             Remark = request.form.get('Remark')
#             offlineArc = request.form.get('offlineArc')

#             ticket_id = request.form.get('ticket_id')

#             # Check if any of the fields are empty and set them to None
#             if (leadetails == ''):
#                 leadetails = None
#             if( Aroundofsuspision == ''):
#                 Aroundofsuspision = None
#             if(DetailsofInvestigation == ''):
#                 DetailsofInvestigation = None
#             if(holdername == ''):
#                 holdername = None
#             if(DateofOpening == ''):
#                 DateofOpening = None
#             if(CummulativeCerditTurnover == ''):
#                 CummulativeCerditTurnover = None 
#             if(CummulativeDebitTurnover == ''):
#                 CummulativeDebitTurnover = None
#             if(CummulativeCashDepositTurnover == ''):
#                 CummulativeCashDepositTurnover = None
#             if(CummulativeCashWithdrawalTurnover == ''):
#                 CummulativeCashWithdrawalTurnover = None
#             if(NoOfTransactionsToBeReported == ''):
#                 NoOfTransactionsToBeReported = None
#             if(TransactionsID == ''):
#                 TransactionsID = None
#             if(amount == ''):
#                 amount = None
#             if(ProductIdentifiers == ''):
#                 ProductIdentifiers = None
#             if(unit == ''):
#                 unit = None 
#             if(DispositionOfFunds == ''):
#                 DispositionOfFunds = None
#             if(RelatedInstitutionName == ''):
#                 RelatedInstitutionName = None
#             if(RelatedInstitutionRefNum == ''):
#                 RelatedInstitutionRefNum = None
#             if(Remark == ''):
#                 Remark = None
                
#             obj = {
#                 'ticket_id': request.form.get('ticket_id') or f"ARM-VRV-{str(uuid.uuid4())}",
#                 'Created_Date': current_date_str,
#                 'Customerno': customerno,
#                 'casename': casename,
#                 'scenario': scenario,
#                 'Guidance': Guidance,
#                 'RuleScenario': RuleScenario,
#                 'personname': personname,
#                 'SourceofAlert': SourceofAlert,
#                 'alertindicator': alertindicator,
#                 'SuspiciousDueToproceedofCrime': SuspiciousDueToproceedofCrime,
#                 'SuspiciousDueToComplexTranscaction': SuspiciousDueToComplexTranscaction,
#                 'SuspiciousDueToNoecoRational': SuspiciousDueToNoecoRational,
#                 'SuspiciousDueToFinancingTerrorism': SuspiciousDueToFinancingTerrorism,
#                 'AttemptedTranscaction': AttemptedTranscaction,
#                 'LEAInformed': LEAInformed,
#                 'PriorityRating': PriorityRating,
#                 'ReportCoverage': ReportCoverage,
#                 'leadetails': leadetails,
#                 'AdditionalDocument': AdditionalDocument,
#                 'Aroundofsuspision': Aroundofsuspision,
#                 'DetailsofInvestigation': DetailsofInvestigation,
#                 'AccountNumber': AccountNumber,
#                 'AccountType': AccountType,
#                 'holdername': holdername,
#                 'AccountHolderType': AccountHolderType,
#                 'AccountStatus': AccountStatus,
#                 'DateofOpening': DateofOpening,
#                 'RiskRating': RiskRating,
#                 'CummulativeCerditTurnover': CummulativeCerditTurnover,
#                 'CummulativeDebitTurnover': CummulativeDebitTurnover,
#                 'CummulativeCashDepositTurnover': CummulativeCashDepositTurnover,
#                 'CummulativeCashWithdrawalTurnover': CummulativeCashWithdrawalTurnover,
#                 'NoOfTransactionsToBeReported': NoOfTransactionsToBeReported,
#                 'TransactionDate': TransactionDate,
#                 'TransactionsID': TransactionsID,
#                 'TransactionMode': TransactionMode,
#                 'DebitCredit': DebitCredit,
#                 'amount': amount,
#                 'TransactionsCurrency': TransactionsCurrency,
#                 'ProductType': ProductType,
#                 'ProductIdentifiers': ProductIdentifiers,
#                 'TransactionType': TransactionType,
#                 'unit': unit,
#                 'Date': Date,
#                 'DispositionOfFunds': DispositionOfFunds,
#                 'RelatedAccountNumber': RelatedAccountNumber,
#                 'RelatedInstitutionName': RelatedInstitutionName,
#                 'RelatedInstitutionRefNum': RelatedInstitutionRefNum,
#                 'Remark': Remark,
#                 'offlineArc': offlineArc,
#                 'Created_By': branchmaker_email
#             }
#             if obj['offlineArc'] != 'archive' and not obj['ticket_id']:
#                 obj['ticket_id'] = f"ARM-VRV-{str(uuid.uuid4())}"
            
#             if Tick_id == None:
#                 # Ensure database connection is open
#                 # cursor = mysql2.connection.cursor()
#                 # Insert the new ticket into the offline_collection database
#                 conn.execute("SELECT * FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'offline_scenarios'")
#                 table_exists = conn.fetchone()
                
#                 # Getting the user id of the branchmaker
#                 select_query = "SELECT [id] FROM [dbo].[user] WHERE [EmailId] = ?"
#                 conn.execute(select_query, (branchmaker_email))
#                 result = conn.fetchone()
#                 obj['id'] = result[0]
#                 print(' ID: ',obj['id'])
                
#                 # If table does not exist, create it
#                 if not table_exists:
#                     create_table_query = ("CREATE TABLE offline_scenarios("
#                                         "[ticket_id] [varchar](max) NULL, "
#                                         "allocatedTicket int NULL, "
#                                         "BranchmakerCasesTicket int NULL, "
#                                         "ROSCasesTicket int NULL, "
#                                         "BranchmakerClosedTicket int NULL, "
#                                         "ROSClosedTicket int NULL, "
#                                         "sentBackClosedTicket int NULL, "
#                                         "unsatisfiedTicket int NULL, "
#                                         "approved int NULL, "
#                                         "rejected int NULL, "
#                                         "DGMCasesTicket int NULL, "
#                                         "[DiffCount] [varchar](max) NULL, "
#                                         "[offlineCasesCount] [varchar](max) NULL, "
#                                         "[PendingCount] [varchar](max) NULL, "
#                                         "[Created_Date] [varchar](max) NULL, "
#                                         "[Customerno] [varchar](max) NULL, "
#                                         "[casename] [varchar](max) NULL, "
#                                         "[scenario] [varchar](max) NULL, "
#                                         "[Guidance] [varchar](max) NULL, "
#                                         "[RuleScenario] [varchar](max) NULL, "
#                                         "[personname] [varchar](max) NULL, "
#                                         "[SourceofAlert] [varchar](max) NULL, "
#                                         "[alertindicator] [varchar](max) NULL, "
#                                         "[SuspiciousDueToproceedofCrime] [varchar](max) NULL, "
#                                         "[SuspiciousDueToComplexTranscaction] [varchar](max) NULL, "
#                                         "[SuspiciousDueToNoecoRational] [varchar](max) NULL, "
#                                         "[SuspiciousDueToFinancingTerrorism] [varchar](max) NULL, "
#                                         "[AttemptedTranscaction] [varchar](max) NULL, "
#                                         "[LEAInformed] [varchar](max) NULL, "
#                                         "[PriorityRating] [varchar](max) NULL, "
#                                         "[ReportCoverage] [varchar](max) NULL, "
#                                         "[leadetails] [varchar](max) NULL, "
#                                         "[AdditionalDocument] [varchar](max) NULL, "
#                                         "[Aroundofsuspision] [varchar](max) NULL, "
#                                         "[DetailsofInvestigation] [varchar](max) NULL, "
#                                         "[AccountNumber] [varchar](max) NULL, "
#                                         "[AccountType] [varchar](max) NULL, "
#                                         "[holdername] [varchar](max) NULL, "
#                                         "[AccountHolderType] [varchar](max) NULL, "
#                                         "[AccountStatus] [varchar](max) NULL, "
#                                         "[DateofOpening] [varchar](max) NULL, "
#                                         "[RiskRating] [varchar](max) NULL, "
#                                         "[CummulativeCerditTurnover] [varchar](max) NULL, "
#                                         "[CummulativeDebitTurnover] [varchar](max) NULL, "
#                                         "[CummulativeCashDepositTurnover] [varchar](max) NULL, "
#                                         "[CummulativeCashWithdrawalTurnover] [varchar](max) NULL, "
#                                         "[NoOfTransactionsToBeReported] [varchar](max) NULL, "
#                                         "[TransactionDate] [varchar](max) NULL, "
#                                         "[TransactionsID] [varchar](max) NULL, "
#                                         "[TransactionMode] [varchar](max) NULL, "
#                                         "[DebitCredit] [varchar](max) NULL, "
#                                         "[amount] [varchar](max) NULL, "
#                                         "[TransactionsCurrency] [varchar](max) NULL, "
#                                         "[ProductType] [varchar](max) NULL, "
#                                         "[ProductIdentifiers] [varchar](max) NULL, "
#                                         "[TransactionType] [varchar](max) NULL, "
#                                         "[unit] [varchar](max) NULL, "
#                                         "[Date] [varchar](max) NULL, "
#                                         "[DispositionOfFunds] [varchar](max) NULL, "
#                                         "[RelatedAccountNumber] [varchar](max) NULL, "
#                                         "[RelatedInstitutionName] [varchar](max) NULL, "
#                                         "[RelatedInstitutionRefNum] [varchar](max) NULL, "
#                                         "[Remark] [varchar](max) NULL, "
#                                         "[ROS_cmt] [varchar](max) NULL, "
#                                         "[Created_By] [varchar](max) NULL, "
#                                         "[DGM_cmt] [varchar](max) NULL))"
#                                         )
#                     # conn = mysql2.connection.cursor()
#                     conn.execute(create_table_query)
#                     conn.commit()


#                     insert_query = "INSERT INTO [dbo].[offline_scenarios]([ticket_id],[allocatedTicket],[BranchmakerCasesTicket],[ROSCasesTicket],[BranchmakerClosedTicket],[ROSClosedTicket],[sentBackClosedTicket],[unsatisfiedTicket],[approved],[rejected],[Created_Date],[Customerno],[casename],[scenario],[Guidance],[RuleScenario],[personname],[SourceofAlert],[alertindicator],[SuspiciousDueToproceedofCrime],[SuspiciousDueToComplexTranscaction],[SuspiciousDueToNoecoRational],[SuspiciousDueToFinancingTerrorism],[AttemptedTranscaction],[LEAInformed],[PriorityRating],[ReportCoverage],[leadetails],[AdditionalDocument],[Aroundofsuspision],[DetailsofInvestigation],[AccountNumber],[AccountType],[holdername],[AccountHolderType],[AccountStatus],[DateofOpening],[RiskRating],[CummulativeCerditTurnover],[CummulativeDebitTurnover],[CummulativeCashDepositTurnover],[CummulativeCashWithdrawalTurnover],[NoOfTransactionsToBeReported],[TransactionDate],[TransactionsID],[TransactionMode],[DebitCredit],[amount],[TransactionsCurrency],[ProductType],[ProductIdentifiers],[TransactionType],[unit],[Date],[DispositionOfFunds],[RelatedAccountNumber],[RelatedInstitutionName],[RelatedInstitutionRefNum],[Remark],[ROS_cmt],[Created_By],[DGM_cmt]) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)"
#                     cursor.execute(insert_query, (obj['ticket_id'],obj['id'], obj['id'], None, None, None, None, None, None, None, obj['Created_Date'], obj['Customerno'], obj['casename'], obj['scenario'],obj['Guidance'], obj['RuleScenario'], obj['personname'], obj['SourceofAlert'], obj['alertindicator'], obj['SuspiciousDueToproceedofCrime'], obj['SuspiciousDueToComplexTranscaction'],obj['SuspiciousDueToNoecoRational'], obj['SuspiciousDueToFinancingTerrorism'], obj['AttemptedTranscaction'], obj['LEAInformed'], obj['PriorityRating'], obj['ReportCoverage'], obj['leadetails'], obj['AdditionalDocument'], obj['Aroundofsuspision'], obj['DetailsofInvestigation'], obj['AccountNumber'], obj['AccountType'], obj['holdername'], obj['AccountHolderType'], obj['AccountStatus'], obj['DateofOpening'], obj['RiskRating'], obj['CummulativeCerditTurnover'], obj['CummulativeDebitTurnover'], obj['CummulativeCashDepositTurnover'], obj['CummulativeCashWithdrawalTurnover'], obj['NoOfTransactionsToBeReported'], obj['TransactionDate'], obj['TransactionsID'], obj['TransactionMode'], obj['DebitCredit'], obj['amount'], obj['TransactionsCurrency'], obj['ProductType'], obj['ProductIdentifiers'],obj['TransactionType'], obj['unit'], obj['Date'], obj['DispositionOfFunds'], obj['RelatedAccountNumber'], obj['RelatedInstitutionName'], obj['RelatedInstitutionRefNum'], obj['Remark'], None, obj['Created_By'], None))
#                     conn.commit()
#                 else:
#                     insert_query = "INSERT INTO [dbo].[offline_scenarios]([ticket_id],[allocatedTicket],[BranchmakerCasesTicket],[ROSCasesTicket],[BranchmakerClosedTicket],[ROSClosedTicket],[sentBackClosedTicket],[unsatisfiedTicket],[approved],[rejected],[Created_Date],[Customerno],[casename],[scenario],[Guidance],[RuleScenario],[personname],[SourceofAlert],[alertindicator],[SuspiciousDueToproceedofCrime],[SuspiciousDueToComplexTranscaction],[SuspiciousDueToNoecoRational],[SuspiciousDueToFinancingTerrorism],[AttemptedTranscaction],[LEAInformed],[PriorityRating],[ReportCoverage],[leadetails],[AdditionalDocument],[Aroundofsuspision],[DetailsofInvestigation],[AccountNumber],[AccountType],[holdername],[AccountHolderType],[AccountStatus],[DateofOpening],[RiskRating],[CummulativeCerditTurnover],[CummulativeDebitTurnover],[CummulativeCashDepositTurnover],[CummulativeCashWithdrawalTurnover],[NoOfTransactionsToBeReported],[TransactionDate],[TransactionsID],[TransactionMode],[DebitCredit],[amount],[TransactionsCurrency],[ProductType],[ProductIdentifiers],[TransactionType],[unit],[Date],[DispositionOfFunds],[RelatedAccountNumber],[RelatedInstitutionName],[RelatedInstitutionRefNum],[Remark],[ROS_cmt],[Created_By],[DGM_cmt]) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)"
#                     cursor.execute(insert_query, (obj['ticket_id'],obj['id'], obj['id'], None, None, None, None, None, None, None, obj['Created_Date'], obj['Customerno'], obj['casename'], obj['scenario'],obj['Guidance'], obj['RuleScenario'], obj['personname'], obj['SourceofAlert'], obj['alertindicator'], obj['SuspiciousDueToproceedofCrime'], obj['SuspiciousDueToComplexTranscaction'],obj['SuspiciousDueToNoecoRational'], obj['SuspiciousDueToFinancingTerrorism'], obj['AttemptedTranscaction'], obj['LEAInformed'], obj['PriorityRating'], obj['ReportCoverage'], obj['leadetails'], obj['AdditionalDocument'], obj['Aroundofsuspision'], obj['DetailsofInvestigation'], obj['AccountNumber'], obj['AccountType'], obj['holdername'], obj['AccountHolderType'], obj['AccountStatus'], obj['DateofOpening'], obj['RiskRating'], obj['CummulativeCerditTurnover'], obj['CummulativeDebitTurnover'], obj['CummulativeCashDepositTurnover'], obj['CummulativeCashWithdrawalTurnover'], obj['NoOfTransactionsToBeReported'], obj['TransactionDate'], obj['TransactionsID'], obj['TransactionMode'], obj['DebitCredit'], obj['amount'], obj['TransactionsCurrency'], obj['ProductType'], obj['ProductIdentifiers'],obj['TransactionType'], obj['unit'], obj['Date'], obj['DispositionOfFunds'], obj['RelatedAccountNumber'], obj['RelatedInstitutionName'], obj['RelatedInstitutionRefNum'], obj['Remark'], None, obj['Created_By'], None))
#                     conn.commit()

                
                        
#                 # Insert the new ticket into the offline_tickets database
#                 # cursor.execute("SELECT * FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'offline_tickets'")
#                 # ticket_table_exists = cursor.fetchone()

#                 # if not ticket_table_exists:
#                     # create_ticket_table_query = "CREATE TABLE offline_tickets (ticketId varchar(max) NULL, allocatedTicket int NULL, BranchmakerCasesTicket int NULL, ROSCasesTicket int NULL, BranchmakerClosedTicket int NULL, ROSClosedTicket int NULL, sentBackClosedTicket int NULL, unsatisfiedTicket int NULL, approved int NULL, rejected int NULL, DGMCasesTicket int NULL)"
#                     # cursor = mysql2.connection.cursor()
#                     # cursor.execute(create_ticket_table_query)
#                     # conn.commit()
                    
#                 #     insert_ticket_query = "INSERT INTO offline_tickets([ticketId],[allocatedTicket],[BranchmakerCasesTicket],[ROSCasesTicket],[BranchmakerClosedTicket],[ROSClosedTicket],[sentBackClosedTicket],[unsatisfiedTicket],[approved],[rejected]) VALUES(?,?,?,?,?,?,?,?,?,?)"
#                 #     cursor.execute(insert_ticket_query, (obj['ticket_id'], obj['id'], obj['id'], None, None, None, None, None, None, None))
#                 #     conn.commit()
#                 # else:
#                 #     insert_ticket_query = "INSERT INTO [dbo].[offline_tickets]([ticketId],[allocatedTicket],[BranchmakerCasesTicket],[ROSCasesTicket],[BranchmakerClosedTicket],[ROSClosedTicket],[sentBackClosedTicket],[unsatisfiedTicket],[approved],[rejected]) VALUES(?,?,?,?,?,?,?,?,?,?)"
#                 #     cursor.execute(insert_ticket_query, (obj['ticket_id'], obj['id'], obj['id'], None, None, None, None, None, None, None))
#                 #     conn.commit()

#                 if obj['offlineArc'] == 'archive':
#                     select_get_doj_data = "SELECT [DateofOpening] FROM [dbo].[offline_scenarios] WHERE [ticket_id] = ?"
#                     conn.execute(select_get_doj_data,(obj['ticket_id'],))
#                     doj_data = conn.fetchone()
#                     doj_data = doj_data[0]
#                     print('doaadata:',doj_data)

#                     if obj['DateofOpening'] is None or obj['DateofOpening'] == '':
#                         obj['DateofOpening'] = doj_data
#                     print('DateofOpening:', obj['DateofOpening'])

#                     update_archive_query = "UPDATE [dbo].[offline_scenarios] SET [Created_Date] = ?,[Customerno] = ?,[casename] = ?,[scenario] = ?,[Guidance] = ?,[RuleScenario] = ?,[personname] = ?,[SourceofAlert] = ?,[alertindicator] = ?,[SuspiciousDueToproceedofCrime] = ?,[SuspiciousDueToComplexTranscaction] = ?,[SuspiciousDueToNoecoRational] = ?,[SuspiciousDueToFinancingTerrorism] = ?,[AttemptedTranscaction] = ?,[LEAInformed] = ?,[PriorityRating] = ?,[ReportCoverage] = ?,[leadetails] = ?,[AdditionalDocument] = ?,[Aroundofsuspision] = ?,[DetailsofInvestigation] = ?,[AccountNumber] = ?,[AccountType] = ?,[holdername] = ?,[AccountHolderType] = ?,[AccountStatus] = ?,[DateofOpening] = ?,[RiskRating] = ?,[CummulativeCerditTurnover] = ?,[CummulativeDebitTurnover] = ?,[CummulativeCashDepositTurnover] = ?,[CummulativeCashWithdrawalTurnover] = ?,[NoOfTransactionsToBeReported] = ?,[TransactionDate] = ?,[TransactionsID] = ?,[TransactionMode] = ?,[DebitCredit] = ?,[amount] = ?,[TransactionsCurrency] = ?,[ProductType] = ?,[ProductIdentifiers] = ?,[TransactionType] = ?,[unit] = ?,[Date] = ?,[DispositionOfFunds] = ?,[RelatedAccountNumber] = ?,[RelatedInstitutionName] = ?,[RelatedInstitutionRefNum] = ?,[Remark] = ?,[ROS_cmt] = ?,[Created_By] = ?,[DGM_cmt] = ? WHERE [ticket_id] = ?"
#                     conn.execute(update_archive_query, (obj['Created_Date'], obj['Customerno'], obj['casename'], obj['scenario'],obj['Guidance'], obj['RuleScenario'], obj['personname'], obj['SourceofAlert'], obj['alertindicator'], obj['SuspiciousDueToproceedofCrime'], obj['SuspiciousDueToComplexTranscaction'],obj['SuspiciousDueToNoecoRational'], obj['SuspiciousDueToFinancingTerrorism'], obj['AttemptedTranscaction'], obj['LEAInformed'], obj['PriorityRating'], obj['ReportCoverage'], obj['leadetails'], obj['AdditionalDocument'], obj['Aroundofsuspision'], obj['DetailsofInvestigation'], obj['AccountNumber'], obj['AccountType'], obj['holdername'], obj['AccountHolderType'], obj['AccountStatus'], obj['DateofOpening'], obj['RiskRating'], obj['CummulativeCerditTurnover'], obj['CummulativeDebitTurnover'], obj['CummulativeCashDepositTurnover'], obj['CummulativeCashWithdrawalTurnover'], obj['NoOfTransactionsToBeReported'], obj['TransactionDate'], obj['TransactionsID'], obj['TransactionMode'], obj['DebitCredit'], obj['amount'], obj['TransactionsCurrency'], obj['ProductType'], obj['ProductIdentifiers'],obj['TransactionType'], obj['unit'], obj['Date'], obj['DispositionOfFunds'], obj['RelatedAccountNumber'], obj['RelatedInstitutionName'], obj['RelatedInstitutionRefNum'], obj['Remark'], obj['offlineArc'], obj['Created_By'], obj['offlineArc'], obj['ticket_id']))
#                     conn.commit()
#                     return redirect(url_for("archived"))
#                 else:
#                     # update_archive_query1 = "UPDATE [dbo].[offline_collection] SET [ROS_cmt] = ?, [DGM_cmt] = ? WHERE [ticket_id] = ?"
#                     # cursor.execute(update_archive_query1,(None,None,obj['ticket_id']))
#                     # conn.commit()
#                     # return redirect(url_for("branchmakers"))
#                     select_get_doj_data = "SELECT [DateofOpening] FROM [dbo].[offline_scenarios] WHERE [ticket_id] = ?"
#                     conn.execute(select_get_doj_data,(obj['ticket_id'],))
#                     doj_data = conn.fetchone()
#                     doj_data = doj_data[0]
#                     print('doaadata:',doj_data)

#                     if obj['DateofOpening'] is None or obj['DateofOpening'] == '':
#                         obj['DateofOpening'] = doj_data
#                     print('DateofOpening:', obj['DateofOpening'])

#                     update_archive_query1 = "UPDATE [dbo].[offline_scenarios] SET [Created_Date] = ?,[Customerno] = ?,[casename] = ?,[scenario] = ?,[Guidance] = ?,[RuleScenario] = ?,[personname] = ?,[SourceofAlert] = ?,[alertindicator] = ?,[SuspiciousDueToproceedofCrime] = ?,[SuspiciousDueToComplexTranscaction] = ?,[SuspiciousDueToNoecoRational] = ?,[SuspiciousDueToFinancingTerrorism] = ?,[AttemptedTranscaction] = ?,[LEAInformed] = ?,[PriorityRating] = ?,[ReportCoverage] = ?,[leadetails] = ?,[AdditionalDocument] = ?,[Aroundofsuspision] = ?,[DetailsofInvestigation] = ?,[AccountNumber] = ?,[AccountType] = ?,[holdername] = ?,[AccountHolderType] = ?,[AccountStatus] = ?,[DateofOpening] = ?,[RiskRating] = ?,[CummulativeCerditTurnover] = ?,[CummulativeDebitTurnover] = ?,[CummulativeCashDepositTurnover] = ?,[CummulativeCashWithdrawalTurnover] = ?,[NoOfTransactionsToBeReported] = ?,[TransactionDate] = ?,[TransactionsID] = ?,[TransactionMode] = ?,[DebitCredit] = ?,[amount] = ?,[TransactionsCurrency] = ?,[ProductType] = ?,[ProductIdentifiers] = ?,[TransactionType] = ?,[unit] = ?,[Date] = ?,[DispositionOfFunds] = ?,[RelatedAccountNumber] = ?,[RelatedInstitutionName] = ?,[RelatedInstitutionRefNum] = ?,[Remark] = ?,[ROS_cmt] = ?,[Created_By] = ?,[DGM_cmt] = ? WHERE [ticket_id] = ?"
#                     conn.execute(update_archive_query1, (obj['Created_Date'], obj['Customerno'], obj['casename'], obj['scenario'],obj['Guidance'], obj['RuleScenario'], obj['personname'], obj['SourceofAlert'], obj['alertindicator'], obj['SuspiciousDueToproceedofCrime'], obj['SuspiciousDueToComplexTranscaction'],obj['SuspiciousDueToNoecoRational'], obj['SuspiciousDueToFinancingTerrorism'], obj['AttemptedTranscaction'], obj['LEAInformed'], obj['PriorityRating'], obj['ReportCoverage'], obj['leadetails'], obj['AdditionalDocument'], obj['Aroundofsuspision'], obj['DetailsofInvestigation'], obj['AccountNumber'], obj['AccountType'], obj['holdername'], obj['AccountHolderType'], obj['AccountStatus'], obj['DateofOpening'], obj['RiskRating'], obj['CummulativeCerditTurnover'], obj['CummulativeDebitTurnover'], obj['CummulativeCashDepositTurnover'], obj['CummulativeCashWithdrawalTurnover'], obj['NoOfTransactionsToBeReported'], obj['TransactionDate'], obj['TransactionsID'], obj['TransactionMode'], obj['DebitCredit'], obj['amount'], obj['TransactionsCurrency'], obj['ProductType'], obj['ProductIdentifiers'],obj['TransactionType'], obj['unit'], obj['Date'], obj['DispositionOfFunds'], obj['RelatedAccountNumber'], obj['RelatedInstitutionName'], obj['RelatedInstitutionRefNum'], obj['Remark'], None, obj['Created_By'], None, obj['ticket_id']))
#                     conn.commit()
#                     return redirect(url_for("branchmakers"))
#             else:
#                 if obj['offlineArc'] == 'archive':
#                     select_get_doj_data = "SELECT [DateofOpening] FROM [dbo].[offline_scenarios] WHERE [ticket_id] = ?"
#                     conn.execute(select_get_doj_data,(obj['ticket_id'],))
#                     doj_data = conn.fetchone()
#                     doj_data = doj_data[0]
#                     print('doaadata:',doj_data)

#                     if obj['DateofOpening'] is None or obj['DateofOpening'] == '':
#                         obj['DateofOpening'] = doj_data
#                     print('DateofOpening:', obj['DateofOpening'])

#                     update_archive_query = "UPDATE [dbo].[offline_scenarios] SET [Created_Date] = ?,[Customerno] = ?,[casename] = ?,[scenario] = ?,[Guidance] = ?,[RuleScenario] = ?,[personname] = ?,[SourceofAlert] = ?,[alertindicator] = ?,[SuspiciousDueToproceedofCrime] = ?,[SuspiciousDueToComplexTranscaction] = ?,[SuspiciousDueToNoecoRational] = ?,[SuspiciousDueToFinancingTerrorism] = ?,[AttemptedTranscaction] = ?,[LEAInformed] = ?,[PriorityRating] = ?,[ReportCoverage] = ?,[leadetails] = ?,[AdditionalDocument] = ?,[Aroundofsuspision] = ?,[DetailsofInvestigation] = ?,[AccountNumber] = ?,[AccountType] = ?,[holdername] = ?,[AccountHolderType] = ?,[AccountStatus] = ?,[DateofOpening] = ?,[RiskRating] = ?,[CummulativeCerditTurnover] = ?,[CummulativeDebitTurnover] = ?,[CummulativeCashDepositTurnover] = ?,[CummulativeCashWithdrawalTurnover] = ?,[NoOfTransactionsToBeReported] = ?,[TransactionDate] = ?,[TransactionsID] = ?,[TransactionMode] = ?,[DebitCredit] = ?,[amount] = ?,[TransactionsCurrency] = ?,[ProductType] = ?,[ProductIdentifiers] = ?,[TransactionType] = ?,[unit] = ?,[Date] = ?,[DispositionOfFunds] = ?,[RelatedAccountNumber] = ?,[RelatedInstitutionName] = ?,[RelatedInstitutionRefNum] = ?,[Remark] = ?,[ROS_cmt] = ?,[Created_By] = ?,[DGM_cmt] = ? WHERE [ticket_id] = ?"
#                     conn.execute(update_archive_query, (obj['Created_Date'], obj['Customerno'], obj['casename'], obj['scenario'],obj['Guidance'], obj['RuleScenario'], obj['personname'], obj['SourceofAlert'], obj['alertindicator'], obj['SuspiciousDueToproceedofCrime'], obj['SuspiciousDueToComplexTranscaction'],obj['SuspiciousDueToNoecoRational'], obj['SuspiciousDueToFinancingTerrorism'], obj['AttemptedTranscaction'], obj['LEAInformed'], obj['PriorityRating'], obj['ReportCoverage'], obj['leadetails'], obj['AdditionalDocument'], obj['Aroundofsuspision'], obj['DetailsofInvestigation'], obj['AccountNumber'], obj['AccountType'], obj['holdername'], obj['AccountHolderType'], obj['AccountStatus'], obj['DateofOpening'], obj['RiskRating'], obj['CummulativeCerditTurnover'], obj['CummulativeDebitTurnover'], obj['CummulativeCashDepositTurnover'], obj['CummulativeCashWithdrawalTurnover'], obj['NoOfTransactionsToBeReported'], obj['TransactionDate'], obj['TransactionsID'], obj['TransactionMode'], obj['DebitCredit'], obj['amount'], obj['TransactionsCurrency'], obj['ProductType'], obj['ProductIdentifiers'],obj['TransactionType'], obj['unit'], obj['Date'], obj['DispositionOfFunds'], obj['RelatedAccountNumber'], obj['RelatedInstitutionName'], obj['RelatedInstitutionRefNum'], obj['Remark'], obj['offlineArc'], obj['Created_By'], obj['offlineArc'], obj['ticket_id']))
#                     conn.commit()
#                     return redirect(url_for("archived"))
#                 else:
#                     select_get_doj_data = "SELECT [DateofOpening] FROM [dbo].[offline_scenarios] WHERE [ticket_id] = ?"
#                     conn.execute(select_get_doj_data,(obj['ticket_id'],))
#                     doj_data = cursor.fetchone()
#                     doj_data = doj_data[0]
#                     print('doaadata:',doj_data)

#                     if obj['DateofOpening'] is None or obj['DateofOpening'] == '':
#                         obj['DateofOpening'] = doj_data
#                     print('DateofOpening:', obj['DateofOpening'])

#                     update_archive_query1 = "UPDATE [dbo].[offline_scenarios] SET [Created_Date] = ?,[Customerno] = ?,[casename] = ?,[scenario] = ?,[Guidance] = ?,[RuleScenario] = ?,[personname] = ?,[SourceofAlert] = ?,[alertindicator] = ?,[SuspiciousDueToproceedofCrime] = ?,[SuspiciousDueToComplexTranscaction] = ?,[SuspiciousDueToNoecoRational] = ?,[SuspiciousDueToFinancingTerrorism] = ?,[AttemptedTranscaction] = ?,[LEAInformed] = ?,[PriorityRating] = ?,[ReportCoverage] = ?,[leadetails] = ?,[AdditionalDocument] = ?,[Aroundofsuspision] = ?,[DetailsofInvestigation] = ?,[AccountNumber] = ?,[AccountType] = ?,[holdername] = ?,[AccountHolderType] = ?,[AccountStatus] = ?,[DateofOpening] = ?,[RiskRating] = ?,[CummulativeCerditTurnover] = ?,[CummulativeDebitTurnover] = ?,[CummulativeCashDepositTurnover] = ?,[CummulativeCashWithdrawalTurnover] = ?,[NoOfTransactionsToBeReported] = ?,[TransactionDate] = ?,[TransactionsID] = ?,[TransactionMode] = ?,[DebitCredit] = ?,[amount] = ?,[TransactionsCurrency] = ?,[ProductType] = ?,[ProductIdentifiers] = ?,[TransactionType] = ?,[unit] = ?,[Date] = ?,[DispositionOfFunds] = ?,[RelatedAccountNumber] = ?,[RelatedInstitutionName] = ?,[RelatedInstitutionRefNum] = ?,[Remark] = ?,[ROS_cmt] = ?,[Created_By] = ?,[DGM_cmt] = ? WHERE [ticket_id] = ?"
#                     conn.execute(update_archive_query1, (obj['Created_Date'], obj['Customerno'], obj['casename'], obj['scenario'],obj['Guidance'], obj['RuleScenario'], obj['personname'], obj['SourceofAlert'], obj['alertindicator'], obj['SuspiciousDueToproceedofCrime'], obj['SuspiciousDueToComplexTranscaction'],obj['SuspiciousDueToNoecoRational'], obj['SuspiciousDueToFinancingTerrorism'], obj['AttemptedTranscaction'], obj['LEAInformed'], obj['PriorityRating'], obj['ReportCoverage'], obj['leadetails'], obj['AdditionalDocument'], obj['Aroundofsuspision'], obj['DetailsofInvestigation'], obj['AccountNumber'], obj['AccountType'], obj['holdername'], obj['AccountHolderType'], obj['AccountStatus'], obj['DateofOpening'], obj['RiskRating'], obj['CummulativeCerditTurnover'], obj['CummulativeDebitTurnover'], obj['CummulativeCashDepositTurnover'], obj['CummulativeCashWithdrawalTurnover'], obj['NoOfTransactionsToBeReported'], obj['TransactionDate'], obj['TransactionsID'], obj['TransactionMode'], obj['DebitCredit'], obj['amount'], obj['TransactionsCurrency'], obj['ProductType'], obj['ProductIdentifiers'],obj['TransactionType'], obj['unit'], obj['Date'], obj['DispositionOfFunds'], obj['RelatedAccountNumber'], obj['RelatedInstitutionName'], obj['RelatedInstitutionRefNum'], obj['Remark'], None, obj['Created_By'], None, obj['ticket_id']))
#                     conn.commit()
#                     return redirect(url_for("archived"))
#                 # Notify the branchmaker that the ticket has been created
#                 # notify = notification(branchmaker_email)
        
#     except Exception as e:
#         # Log the error traceback for debugging
#         print("An error occurred:", e)
#         traceback.print_exc()  # Print the traceback to identify the issue
#         return "An error occurred while processing your request", 500  # Return an error response


# Here, We are storing the data in the database
@app.route('/post_manual_str', methods=['POST'])
@secure_route(required_role=['BranchMakers'])
def post_manual_str():

    conn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
    cursor = conn.cursor()

    Tick_id = request.args.get('account_number')
    print('offid:',Tick_id)
    try:
        branchmaker_email = session.get('email_id')
        if request.method == 'POST':
            current_datetime = datetime.now()
            current_date = current_datetime.date()
            current_date_str = current_date.isoformat()
            Created_Date = current_date_str
            customerno = request.form.get('Customerno')
            casename = request.form.get('casename')
            personname = request.form.get('personname')
            RuleScenario = request.form.get('RuleScenario')
            Guidance = request.form.get('Guidance')
            scenario = request.form.get('scenario')
            SourceofAlert = request.form.get('sourceofalert')
            alertindicator = request.form.get('alertindicator')
            SuspiciousDueToproceedofCrime = request.form.get('SuspiciousDueToproceedofCrime')
            SuspiciousDueToComplexTranscaction = request.form.get('SuspiciousDueToComplexTranscaction')
            SuspiciousDueToNoecoRational = request.form.get('SuspiciousDueToNoecoRational')
            SuspiciousDueToFinancingTerrorism = request.form.get('SuspiciousDueToFinancingTerrorism')
            AttemptedTranscaction = request.form.get('AttemptedTranscaction')
            LEAInformed = request.form.get('LEAInformed')
            PriorityRating = request.form.get('PriorityRating')
            ReportCoverage = request.form.get('ReportCoverage')
            leadetails = request.form.get('leadetails')
            AdditionalDocument = request.form.get('AdditionalDocument')
            Aroundofsuspision = request.form.get('Aroundofsuspision')
            DetailsofInvestigation = request.form.get('DetailsofInvestigation')
            AccountNumber = request.form.get('AccountNumber')
            AccountType = request.form.get('AccountType')
            holdername = request.form.get('holdername')
            AccountHolderType = request.form.get('AccountHolderType')
            AccountStatus = request.form.get('AccountStatus')
            DateofOpening = request.form.get('DateofOpening')
            RiskRating = request.form.get('RiskRating')
            CummulativeCerditTurnover = request.form.get('CummulativeCerditTurnover')
            CummulativeDebitTurnover = request.form.get('CummulativeDebitTurnover')
            CummulativeCashDepositTurnover = request.form.get('CummulativeCashDepositTurnover')
            CummulativeCashWithdrawalTurnover = request.form.get('CummulativeCashWithdrawalTurnover')
            NoOfTransactionsToBeReported = request.form.get('NoOfTransactionsToBeReported')
            address = request.form.get('Address')
            pincode = request.form.get('Pincode')
            city = request.form.get('City')
            dob = request.form.get('DOB')
            mobilenumber = request.form.get('MobileNumber')
            pan = request.form.get('PAN')
            transactionamount = request.form.get('TransactionAmount')
            transactiontypeacc = request.form.get('TransactionTypeacc')
            transactioncategory = request.form.get('TransactionCategory')
            transactioncurrency = request.form.get('TransactionCurrency')
            bankname = request.form.get('BankName')
            bankstate = request.form.get('BankState')
            TransactionDate = request.form.get('TransactionDate')
            TransactionsID = request.form.get('TransactionsID')
            TransactionMode = request.form.get('TransactionMode')
            DebitCredit = request.form.get('DebitCredit')
            amount = request.form.get('amount')
            TransactionsCurrency = request.form.get('TransactionsCurrency')
            ProductType = request.form.get('ProductType')
            ProductIdentifiers = request.form.get('ProductIdentifiers')
            TransactionType = request.form.get('TransactionType')
            unit = request.form.get('unit')
            Date = request.form.get('Date')
            DispositionOfFunds = request.form.get('DispositionOfFunds')
            RelatedAccountNumber = request.form.get('RelatedAccountNumber')
            RelatedInstitutionName = request.form.get('RelatedInstitutionName')
            RelatedInstitutionRefNum = request.form.get('RelatedInstitutionRefNum')
            Remark = request.form.get('Remark')
            offlineArc = request.form.get('offlineArc')

            ticket_id = request.form.get('ticket_id')

            # Check if any of the fields are empty and set them to None
            if (leadetails == ''):
                leadetails = None
            if( Aroundofsuspision == ''):
                Aroundofsuspision = None
            if(DetailsofInvestigation == ''):
                DetailsofInvestigation = None
            if(holdername == ''):
                holdername = None
            if(DateofOpening == ''):
                DateofOpening = None
            if(CummulativeCerditTurnover == ''):
                CummulativeCerditTurnover = None 
            if(CummulativeDebitTurnover == ''):
                CummulativeDebitTurnover = None
            if(CummulativeCashDepositTurnover == ''):
                CummulativeCashDepositTurnover = None
            if(CummulativeCashWithdrawalTurnover == ''):
                CummulativeCashWithdrawalTurnover = None
            if(NoOfTransactionsToBeReported == ''):
                NoOfTransactionsToBeReported = None
            if (address) == '':
                address = None
            if (pincode) == '':
                pincode = None
            if (city) == '':
                city = None
            if (dob) == '':
                dob = None
            if (mobilenumber) == '':
                mobilenumber = None
            if (pan) == '':
                pan = None
            if (transactionamount) == '':
                transactionamount = None
            if (transactiontypeacc) == '':
                transactiontypeacc = None
            if (transactioncategory) == '':
                transactioncategory = None
            if (transactioncurrency) == '':
                transactioncurrency = None
            if (bankname) == '':
                bankname = None
            if (bankstate) == '':
                bankstate = None
            if(TransactionsID == ''):
                TransactionsID = None
            if(amount == ''):
                amount = None
            if(ProductIdentifiers == ''):
                ProductIdentifiers = None
            if(unit == ''):
                unit = None 
            if(DispositionOfFunds == ''):
                DispositionOfFunds = None
            if(RelatedInstitutionName == ''):
                RelatedInstitutionName = None
            if(RelatedInstitutionRefNum == ''):
                RelatedInstitutionRefNum = None
            if(Remark == ''):
                Remark = None                
            obj = {
                'ticket_id': request.form.get('ticket_id') or f"ARM-VRV-{str(uuid.uuid4())}",
                'Created_Date': current_date_str,
                'Customerno': customerno,
                'casename': casename,
                'scenario': scenario,
                'Guidance': Guidance,
                'RuleScenario': RuleScenario,
                'personname': personname,
                'SourceofAlert': SourceofAlert,
                'alertindicator': alertindicator,
                'SuspiciousDueToproceedofCrime': SuspiciousDueToproceedofCrime,
                'SuspiciousDueToComplexTranscaction': SuspiciousDueToComplexTranscaction,
                'SuspiciousDueToNoecoRational': SuspiciousDueToNoecoRational,
                'SuspiciousDueToFinancingTerrorism': SuspiciousDueToFinancingTerrorism,
                'AttemptedTranscaction': AttemptedTranscaction,
                'LEAInformed': LEAInformed,
                'PriorityRating': PriorityRating,
                'ReportCoverage': ReportCoverage,
                'leadetails': leadetails,
                'AdditionalDocument': AdditionalDocument,
                'Aroundofsuspision': Aroundofsuspision,
                'DetailsofInvestigation': DetailsofInvestigation,
                'AccountNumber': AccountNumber,
                'AccountType': AccountType,
                'holdername': holdername,
                'AccountHolderType': AccountHolderType,
                'AccountStatus': AccountStatus,
                'DateofOpening': DateofOpening,
                'RiskRating': RiskRating,
                'CummulativeCerditTurnover': CummulativeCerditTurnover,
                'CummulativeDebitTurnover': CummulativeDebitTurnover,
                'CummulativeCashDepositTurnover': CummulativeCashDepositTurnover,
                'CummulativeCashWithdrawalTurnover': CummulativeCashWithdrawalTurnover,
                'NoOfTransactionsToBeReported': NoOfTransactionsToBeReported,
                'Address':address,
                'Pincode':pincode,
                'City':city,
                'DOB':dob,
                'MobileNumber': mobilenumber,
                'PAN':pan,
                'TransactionAmount':transactionamount,
                'TransactionTypeacc':transactiontypeacc,
                'TransactionCategory': transactioncategory,
                'TransactionCurrency':transactioncurrency,
                'BankName':bankname,
                'BankState' : bankstate,
                'TransactionDate': TransactionDate,
                'TransactionsID': TransactionsID,
                'TransactionMode': TransactionMode,
                'DebitCredit': DebitCredit,
                'amount': amount,
                'TransactionsCurrency': TransactionsCurrency,
                'ProductType': ProductType,
                'ProductIdentifiers': ProductIdentifiers,
                'TransactionType': TransactionType,
                'unit': unit,
                'Date': Date,
                'DispositionOfFunds': DispositionOfFunds,
                'RelatedAccountNumber': RelatedAccountNumber,
                'RelatedInstitutionName': RelatedInstitutionName,
                'RelatedInstitutionRefNum': RelatedInstitutionRefNum,
                'Remark': Remark,
                'offlineArc': offlineArc,
                'Created_By': branchmaker_email
            }
            if obj['offlineArc'] != 'archive' and not obj['ticket_id']:
                obj['ticket_id'] = f"ARM-VRV-{str(uuid.uuid4())}"
                        
            if Tick_id is None:
                cursor.execute("SELECT * FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'offline_scenarios'")
                table_exists = cursor.fetchone()
                
                 # Getting the user id of the branchmaker
                select_query = "SELECT [id] FROM [user] WHERE [EmailId] = ?"
                cursor.execute(select_query, (branchmaker_email,))
                result = cursor.fetchone()
                obj['id'] = result[0]
                print('ID:', obj['id'])

                # if result is not None:
                #     obj['id'] = result[0]
                #     print('ID:', obj['id'])
                # else:
                #     print('No user found with the provided email ID')

                table_exists_query = """
                    SELECT CASE WHEN OBJECT_ID('dbo.offline_scenarios', 'U') IS NOT NULL THEN 1 ELSE 0 END AS TableExists
                    """
                cursor.execute(table_exists_query)
                table_exists = cursor.fetchone()[0]
                # If table does not exist, create it
                if not table_exists:
                    create_table_query = ("""
                        CREATE TABLE offline_scenarios(
                            [ticket_id] [varchar](max) NULL, 
                            allocatedTicket int NULL,
                            BranchmakerCasesTicket int NULL,
                            ROSCasesTicket int NULL,
                            BranchmakerClosedTicket int NULL,
                            ROSClosedTicket int NULL,
                            sentBackClosedTicket int NULL,
                            unsatisfiedTicket int NULL,
                            approved int NULL,
                            rejected int NULL,
                            DGMCasesTicket int NULL,
                            [DiffCount] [varchar](max) NULL, 
                            [offlineCasesCount] [varchar](max) NULL, 
                            [PendingCount] [varchar](max) NULL, 
                            [Created_Date] [varchar](max) NULL, 
                            [Customerno] [varchar](max) NULL, 
                            [casename] [varchar](max) NULL, 
                            [scenario] [varchar](max) NULL, 
                            [Guidance] [varchar](max) NULL, 
                            [RuleScenario] [varchar](max) NULL, 
                            [personname] [varchar](max) NULL, 
                            [SourceofAlert] [varchar](max) NULL, 
                            [alertindicator] [varchar](max) NULL, 
                            [SuspiciousDueToproceedofCrime] [varchar](max) NULL, 
                            [SuspiciousDueToComplexTranscaction] [varchar](max) NULL, 
                            [SuspiciousDueToNoecoRational] [varchar](max) NULL, 
                            [SuspiciousDueToFinancingTerrorism] [varchar](max) NULL, 
                            [AttemptedTranscaction] [varchar](max) NULL, 
                            [LEAInformed] [varchar](max) NULL, 
                            [PriorityRating] [varchar](max) NULL, 
                            [ReportCoverage] [varchar](max) NULL, 
                            [leadetails] [varchar](max) NULL, 
                            [AdditionalDocument] [varchar](max) NULL, 
                            [Aroundofsuspision] [varchar](max) NULL, 
                            [DetailsofInvestigation] [varchar](max) NULL, 
                            [AccountNumber] [varchar](max) NULL, 
                            [AccountType] [varchar](max) NULL, 
                            [holdername] [varchar](max) NULL, 
                            [AccountHolderType] [varchar](max) NULL, 
                            [AccountStatus] [varchar](max) NULL, 
                            [DateofOpening] [varchar](max) NULL, 
                            [RiskRating] [varchar](max) NULL, 
                            [CummulativeCerditTurnover] [varchar](max) NULL, 
                            [CummulativeDebitTurnover] [varchar](max) NULL, 
                            [CummulativeCashDepositTurnover] [varchar](max) NULL, 
                            [CummulativeCashWithdrawalTurnover] [varchar](max) NULL, 
                            [NoOfTransactionsToBeReported] [varchar](max) NULL, 
                            [Address] VARCHAR(MAX) NULL,
                            [Pincode] VARCHAR(MAX) NULL,
                            [City] VARCHAR(MAX) NULL,
                            [DOB] VARCHAR(MAX) NULL,
                            [MobileNumber] VARCHAR(MAX) NULL,
                            [PAN] VARCHAR(MAX) NULL,
                            [TransactionAmount] VARCHAR(MAX) NULL,
                            [TransactionTypeacc] VARCHAR(MAX) NULL,
                            [TransactionCategory] VARCHAR(MAX) NULL,
                            [TransactionCurrency] VARCHAR(MAX) NULL,
                            [BankName] VARCHAR(MAX) NULL,
                            [BankState] VARCHAR(MAX) NULL,
                            [TransactionDate] [varchar](max) NULL, 
                            [TransactionsID] [varchar](max) NULL, 
                            [TransactionMode] [varchar](max) NULL, 
                            [DebitCredit] [varchar](max) NULL, 
                            [amount] [varchar](max) NULL, 
                            [TransactionsCurrency] [varchar](max) NULL,
                            [ProductType] [varchar](max) NULL, 
                            [ProductIdentifiers] [varchar](max) NULL, 
                            [TransactionType] [varchar](max) NULL, 
                            [unit] [varchar](max) NULL, 
                            [Date] [varchar](max) NULL, 
                            [DispositionOfFunds] [varchar](max) NULL,
                            [RelatedAccountNumber] [varchar](max) NULL,
                            [RelatedInstitutionName] [varchar](max) NULL,
                            [RelatedInstitutionRefNum] [varchar](max) NULL,
                            [Remark] [varchar](max) NULL,
                            [ROS_cmt] [varchar](max) NULL,
                            [Created_By] [varchar](max) NULL,
                            [DGM_cmt] [varchar](max) NULL
                        )
                    """)
                    cursor.execute(create_table_query)
                    conn.commit()
                    print("Table 'offline_scenarios' created.")
                    insert_query = """
                        INSERT INTO [dbo].[offline_scenarios] (
                            [ticket_id], [allocatedTicket], [BranchmakerCasesTicket], [ROSCasesTicket], 
                            [BranchmakerClosedTicket], [ROSClosedTicket], [sentBackClosedTicket], [unsatisfiedTicket], 
                            [approved], [rejected], [DGMCasesTicket], [DiffCount], [offlineCasesCount], [PendingCount], 
                            [Created_Date], [Customerno], [casename], [scenario], [Guidance], [RuleScenario], [personname], 
                            [SourceofAlert], [alertindicator], [SuspiciousDueToproceedofCrime], [SuspiciousDueToComplexTranscaction], 
                            [SuspiciousDueToNoecoRational], [SuspiciousDueToFinancingTerrorism], [AttemptedTranscaction], [LEAInformed], 
                            [PriorityRating], [ReportCoverage], [leadetails], [AdditionalDocument], [Aroundofsuspision], [DetailsofInvestigation], 
                            [AccountNumber], [AccountType], [holdername], [AccountHolderType], [AccountStatus], [DateofOpening], [RiskRating], 
                            [CummulativeCerditTurnover], [CummulativeDebitTurnover], [CummulativeCashDepositTurnover], [CummulativeCashWithdrawalTurnover], 
                            [NoOfTransactionsToBeReported], [Address], [Pincode], [City], [DOB], [MobileNumber], [PAN], [TransactionAmount], 
                            [TransactionTypeacc], [TransactionCategory], [TransactionCurrency], [BankName], [BankState], [TransactionDate], 
                            [TransactionsID], [TransactionMode], [DebitCredit], [amount], [TransactionsCurrency], [ProductType], [ProductIdentifiers], 
                            [TransactionType], [unit], [Date], [DispositionOfFunds], [RelatedAccountNumber], [RelatedInstitutionName], 
                            [RelatedInstitutionRefNum], [Remark], [ROS_cmt], [Created_By], [DGM_cmt]
                        ) VALUES (
                            ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?,?, 
                            ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?,?
                        )
                        """
                    # Ensure all keys exist in obj and adjust if necessary
                    params = (
                        obj['ticket_id'], obj.get('allocatedTicket', None), obj.get('BranchmakerCasesTicket', None), obj.get('ROSCasesTicket', None),
                        obj.get('BranchmakerClosedTicket', None), obj.get('ROSClosedTicket', None), obj.get('sentBackClosedTicket', None), obj.get('unsatisfiedTicket', None),
                        obj.get('approved', None), obj.get('rejected', None), obj.get('DGMCasesTicket', None), obj.get('DiffCount', None), obj.get('offlineCasesCount', None),
                        obj.get('PendingCount', None), obj.get('Created_Date', None), obj.get('Customerno', None), obj.get('casename', None), obj.get('scenario', None),
                        obj.get('Guidance', None), obj.get('RuleScenario', None), obj.get('personname', None), obj.get('SourceofAlert', None), obj.get('alertindicator', None),
                        obj.get('SuspiciousDueToproceedofCrime', None), obj.get('SuspiciousDueToComplexTranscaction', None), obj.get('SuspiciousDueToNoecoRational', None),
                        obj.get('SuspiciousDueToFinancingTerrorism', None), obj.get('AttemptedTranscaction', None), obj.get('LEAInformed', None), obj.get('PriorityRating', None),
                        obj.get('ReportCoverage', None), obj.get('leadetails', None), obj.get('AdditionalDocument', None), obj.get('Aroundofsuspision', None),
                        obj.get('DetailsofInvestigation', None), obj.get('AccountNumber', None), obj.get('AccountType', None), obj.get('holdername', None),
                        obj.get('AccountHolderType', None), obj.get('AccountStatus', None), obj.get('DateofOpening', None), obj.get('RiskRating', None),
                        obj.get('CummulativeCerditTurnover', None), obj.get('CummulativeDebitTurnover', None), obj.get('CummulativeCashDepositTurnover', None),
                        obj.get('CummulativeCashWithdrawalTurnover', None), obj.get('NoOfTransactionsToBeReported', None), obj.get('Address', None),
                        obj.get('Pincode', None), obj.get('City', None), obj.get('DOB', None), obj.get('MobileNumber', None), obj.get('PAN', None),
                        obj.get('TransactionAmount', None), obj.get('TransactionTypeacc', None), obj.get('TransactionCategory', None),
                        obj.get('TransactionCurrency', None), obj.get('BankName', None), obj.get('BankState', None), obj.get('TransactionDate', None),
                        obj.get('TransactionsID', None), obj.get('TransactionMode', None), obj.get('DebitCredit', None), obj.get('amount', None),
                        obj.get('TransactionsCurrency', None), obj.get('ProductType', None), obj.get('ProductIdentifiers', None),
                        obj.get('TransactionType', None), obj.get('unit', None), obj.get('Date', None), obj.get('DispositionOfFunds', None),
                        obj.get('RelatedAccountNumber', None), obj.get('RelatedInstitutionName', None), obj.get('RelatedInstitutionRefNum', None),
                        obj.get('Remark', None), obj.get('ROS_cmt', None), obj.get('Created_By', None), obj.get('DGM_cmt', None)
                    )

                    # Execute the query with the correct number of parameters
                    cursor.execute(insert_query, params)

                    # Commit the transaction
                    conn.commit()
                    # insert_query = "INSERT INTO [dbo].[offline_scenarios]([ticket_id],[allocatedTicket],[BranchmakerCasesTicket],[ROSCasesTicket],[BranchmakerClosedTicket],[ROSClosedTicket],[sentBackClosedTicket],[unsatisfiedTicket],[approved],[rejected],[Created_Date],[Customerno],[casename],[scenario],[Guidance],[RuleScenario],[personname],[SourceofAlert],[alertindicator],[SuspiciousDueToproceedofCrime],[SuspiciousDueToComplexTranscaction],[SuspiciousDueToNoecoRational],[SuspiciousDueToFinancingTerrorism],[AttemptedTranscaction],[LEAInformed],[PriorityRating],[ReportCoverage],[leadetails],[AdditionalDocument],[Aroundofsuspision],[DetailsofInvestigation],[AccountNumber],[AccountType],[holdername],[AccountHolderType],[AccountStatus],[DateofOpening],[RiskRating],[CummulativeCerditTurnover],[CummulativeDebitTurnover],[CummulativeCashDepositTurnover],[CummulativeCashWithdrawalTurnover],[NoOfTransactionsToBeReported],[Address],[Pincode],[City],[DOB],[MobileNumber],[PAN],[TransactionAmount],[TransactionTypeacc],[TransactionCategory],[TransactionCurrency],[BankName],[BankState] ,[TransactionDate],[TransactionsID],[TransactionMode],[DebitCredit],[amount],[TransactionsCurrency],[ProductType],[ProductIdentifiers],[TransactionType],[unit],[Date],[DispositionOfFunds],[RelatedAccountNumber],[RelatedInstitutionName],[RelatedInstitutionRefNum],[Remark],[ROS_cmt],[Created_By],[DGM_cmt]) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)"
                    # cursor.execute(insert_query, (obj['ticket_id'],obj['id'], obj['id'], None, None, None, None, None, None, None, obj['Created_Date'], 
                    # obj['Customerno'], obj['casename'], obj['scenario'],obj['Guidance'], obj['RuleScenario'], obj['personname'], 
                    # obj['SourceofAlert'], obj['alertindicator'], obj['SuspiciousDueToproceedofCrime'], obj['SuspiciousDueToComplexTranscaction'],
                    # obj['SuspiciousDueToNoecoRational'], obj['SuspiciousDueToFinancingTerrorism'], obj['AttemptedTranscaction'], obj['LEAInformed'], 
                    # obj['PriorityRating'], obj['ReportCoverage'], obj['leadetails'], obj['AdditionalDocument'], obj['Aroundofsuspision'], 
                    # obj['DetailsofInvestigation'], obj['AccountNumber'], obj['AccountType'], obj['holdername'], obj['AccountHolderType'], 
                    # 
                    # obj['AccountStatus'], obj['DateofOpening'], obj['RiskRating'], obj['CummulativeCerditTurnover'], obj['CummulativeDebitTurnover'], 
                    # obj['CummulativeCashDepositTurnover'], obj['CummulativeCashWithdrawalTurnover'], obj['NoOfTransactionsToBeReported'],
                    #  obj['Address'],obj['Pincode'],obj['City'], obj['DOB'],obj['MobileNumber'],obj['PAN'],obj['TransactionAmount'],
                    # obj['TransactionTypeacc'],obj['TransactionCategory'],obj['TransactionCurrency'],obj['BankName'],obj['BankState'], 
                    # obj['TransactionDate'], obj['TransactionsID'], obj['TransactionMode'], obj['DebitCredit'], obj['amount'], 
                    # obj['TransactionsCurrency'], obj['ProductType'], obj['ProductIdentifiers'],obj['TransactionType'],
                    #  obj['unit'], obj['Date'], obj['DispositionOfFunds'], obj['RelatedAccountNumber'], obj['RelatedInstitutionName'],
                    #  obj['RelatedInstitutionRefNum'], obj['Remark'], None, obj['Created_By'], None))
                    # conn.commit()
                else:
                    print("Table 'offline_scenarios' already exists.")
                    # insert_query = "INSERT INTO [dbo].[offline_scenarios]([ticket_id],[allocatedTicket],[BranchmakerCasesTicket],[ROSCasesTicket],[BranchmakerClosedTicket],[ROSClosedTicket],[sentBackClosedTicket],[unsatisfiedTicket],[approved],[rejected],[Created_Date],[Customerno],[casename],[scenario],[Guidance],[RuleScenario],[personname],[SourceofAlert],[alertindicator],[SuspiciousDueToproceedofCrime],[SuspiciousDueToComplexTranscaction],[SuspiciousDueToNoecoRational],[SuspiciousDueToFinancingTerrorism],[AttemptedTranscaction],[LEAInformed],[PriorityRating],[ReportCoverage],[leadetails],[AdditionalDocument],[Aroundofsuspision],[DetailsofInvestigation],[AccountNumber],[AccountType],[holdername],[AccountHolderType],[AccountStatus],[DateofOpening],[RiskRating],[CummulativeCerditTurnover],[CummulativeDebitTurnover],[CummulativeCashDepositTurnover],[CummulativeCashWithdrawalTurnover],[NoOfTransactionsToBeReported],[Address],[Pincode],[City],[DOB],[MobileNumber],[PAN],[TransactionAmount],[TransactionTypeacc],[TransactionCategory],[TransactionCurrency],[BankName],[BankState] ,[TransactionDate],[TransactionsID],[TransactionMode],[DebitCredit],[amount],[TransactionsCurrency],[ProductType],[ProductIdentifiers],[TransactionType],[unit],[Date],[DispositionOfFunds],[RelatedAccountNumber],[RelatedInstitutionName],[RelatedInstitutionRefNum],[Remark],[ROS_cmt],[Created_By],[DGM_cmt]) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)"
                    # cursor.execute(insert_query, (obj['ticket_id'],obj['id'], obj['id'], None, None, None, None, None, None, None, obj['Created_Date'], obj['Customerno'], obj['casename'], obj['scenario'],obj['Guidance'], obj['RuleScenario'], obj['personname'], obj['SourceofAlert'], obj['alertindicator'], obj['SuspiciousDueToproceedofCrime'], obj['SuspiciousDueToComplexTranscaction'],obj['SuspiciousDueToNoecoRational'], obj['SuspiciousDueToFinancingTerrorism'], obj['AttemptedTranscaction'], obj['LEAInformed'], obj['PriorityRating'], obj['ReportCoverage'], obj['leadetails'], obj['AdditionalDocument'], obj['Aroundofsuspision'], obj['DetailsofInvestigation'], obj['AccountNumber'], obj['AccountType'], obj['holdername'], obj['AccountHolderType'], obj['AccountStatus'], obj['DateofOpening'], obj['RiskRating'], obj['CummulativeCerditTurnover'], obj['CummulativeDebitTurnover'], obj['CummulativeCashDepositTurnover'], obj['CummulativeCashWithdrawalTurnover'], obj['NoOfTransactionsToBeReported'], obj['Address'],obj['Pincode'],obj['City'], obj['DOB'],obj['MobileNumber'],obj['PAN'],obj['TransactionAmount'],obj['TransactionTypeacc'],obj['TransactionCategory'],obj['TransactionCurrency'],obj['BankName'],obj['BankState'], obj['TransactionDate'], obj['TransactionsID'], obj['TransactionMode'], obj['DebitCredit'], obj['amount'], obj['TransactionsCurrency'], obj['ProductType'], obj['ProductIdentifiers'],obj['TransactionType'], obj['unit'], obj['Date'], obj['DispositionOfFunds'], obj['RelatedAccountNumber'], obj['RelatedInstitutionName'], obj['RelatedInstitutionRefNum'], obj['Remark'], None, obj['Created_By'], None))
                    # conn.commit()
                    insert_query = """
                        INSERT INTO [dbo].[offline_scenarios] (
                            [ticket_id], [allocatedTicket], [BranchmakerCasesTicket], [ROSCasesTicket], 
                            [BranchmakerClosedTicket], [ROSClosedTicket], [sentBackClosedTicket], [unsatisfiedTicket], 
                            [approved], [rejected], [DGMCasesTicket], [DiffCount], [offlineCasesCount], [PendingCount], 
                            [Created_Date], [Customerno], [casename], [scenario], [Guidance], [RuleScenario], [personname], 
                            [SourceofAlert], [alertindicator], [SuspiciousDueToproceedofCrime], [SuspiciousDueToComplexTranscaction], 
                            [SuspiciousDueToNoecoRational], [SuspiciousDueToFinancingTerrorism], [AttemptedTranscaction], [LEAInformed], 
                            [PriorityRating], [ReportCoverage], [leadetails], [AdditionalDocument], [Aroundofsuspision], [DetailsofInvestigation], 
                            [AccountNumber], [AccountType], [holdername], [AccountHolderType], [AccountStatus], [DateofOpening], [RiskRating], 
                            [CummulativeCerditTurnover], [CummulativeDebitTurnover], [CummulativeCashDepositTurnover], [CummulativeCashWithdrawalTurnover], 
                            [NoOfTransactionsToBeReported], [Address], [Pincode], [City], [DOB], [MobileNumber], [PAN], [TransactionAmount], 
                            [TransactionTypeacc], [TransactionCategory], [TransactionCurrency], [BankName], [BankState], [TransactionDate], 
                            [TransactionsID], [TransactionMode], [DebitCredit], [amount], [TransactionsCurrency], [ProductType], [ProductIdentifiers], 
                            [TransactionType], [unit], [Date], [DispositionOfFunds], [RelatedAccountNumber], [RelatedInstitutionName], 
                            [RelatedInstitutionRefNum], [Remark], [ROS_cmt], [Created_By], [DGM_cmt]
                        ) VALUES (
                            ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?,?, 
                            ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?,?
                        )
                        """
                    # Ensure all keys exist in obj and adjust if necessary
                    params = (
                        obj['ticket_id'], obj.get('allocatedTicket', None), obj.get('BranchmakerCasesTicket', None), obj.get('ROSCasesTicket', None),
                        obj.get('BranchmakerClosedTicket', None), obj.get('ROSClosedTicket', None), obj.get('sentBackClosedTicket', None), obj.get('unsatisfiedTicket', None),
                        obj.get('approved', None), obj.get('rejected', None), obj.get('DGMCasesTicket', None), obj.get('DiffCount', None), obj.get('offlineCasesCount', None),
                        obj.get('PendingCount', None), obj.get('Created_Date', None), obj.get('Customerno', None), obj.get('casename', None), obj.get('scenario', None),
                        obj.get('Guidance', None), obj.get('RuleScenario', None), obj.get('personname', None), obj.get('SourceofAlert', None), obj.get('alertindicator', None),
                        obj.get('SuspiciousDueToproceedofCrime', None), obj.get('SuspiciousDueToComplexTranscaction', None), obj.get('SuspiciousDueToNoecoRational', None),
                        obj.get('SuspiciousDueToFinancingTerrorism', None), obj.get('AttemptedTranscaction', None), obj.get('LEAInformed', None), obj.get('PriorityRating', None),
                        obj.get('ReportCoverage', None), obj.get('leadetails', None), obj.get('AdditionalDocument', None), obj.get('Aroundofsuspision', None),
                        obj.get('DetailsofInvestigation', None), obj.get('AccountNumber', None), obj.get('AccountType', None), obj.get('holdername', None),
                        obj.get('AccountHolderType', None), obj.get('AccountStatus', None), obj.get('DateofOpening', None), obj.get('RiskRating', None),
                        obj.get('CummulativeCerditTurnover', None), obj.get('CummulativeDebitTurnover', None), obj.get('CummulativeCashDepositTurnover', None),
                        obj.get('CummulativeCashWithdrawalTurnover', None), obj.get('NoOfTransactionsToBeReported', None), obj.get('Address', None),
                        obj.get('Pincode', None), obj.get('City', None), obj.get('DOB', None), obj.get('MobileNumber', None), obj.get('PAN', None),
                        obj.get('TransactionAmount', None), obj.get('TransactionTypeacc', None), obj.get('TransactionCategory', None),
                        obj.get('TransactionCurrency', None), obj.get('BankName', None), obj.get('BankState', None), obj.get('TransactionDate', None),
                        obj.get('TransactionsID', None), obj.get('TransactionMode', None), obj.get('DebitCredit', None), obj.get('amount', None),
                        obj.get('TransactionsCurrency', None), obj.get('ProductType', None), obj.get('ProductIdentifiers', None),
                        obj.get('TransactionType', None), obj.get('unit', None), obj.get('Date', None), obj.get('DispositionOfFunds', None),
                        obj.get('RelatedAccountNumber', None), obj.get('RelatedInstitutionName', None), obj.get('RelatedInstitutionRefNum', None),
                        obj.get('Remark', None), obj.get('ROS_cmt', None), obj.get('Created_By', None), obj.get('DGM_cmt', None)
                    )

                    # Execute the query with the correct number of parameters
                    cursor.execute(insert_query, params)

                    # Commit the transaction
                    conn.commit()


                if obj['offlineArc'] == 'archive':
                    select_get_doj_data = "SELECT [DateofOpening] FROM [dbo].[offline_scenarios] WHERE [ticket_id] = ?"
                    cursor.execute(select_get_doj_data,(obj['ticket_id'],))
                    doj_data = cursor.fetchone()
                    doj_data = doj_data[0]
                    print('doaadata:',doj_data)

                    if obj['DateofOpening'] is None or obj['DateofOpening'] == '':
                        obj['DateofOpening'] = doj_data
                    print('DateofOpening:', obj['DateofOpening'])

                    update_archive_query = "UPDATE [dbo].[offline_scenarios] SET [Created_Date] = ?,[Customerno] = ?,[casename] = ?,[scenario] = ?,[Guidance] = ?,[RuleScenario] = ?,[personname] = ?,[SourceofAlert] = ?,[alertindicator] = ?,[SuspiciousDueToproceedofCrime] = ?,[SuspiciousDueToComplexTranscaction] = ?,[SuspiciousDueToNoecoRational] = ?,[SuspiciousDueToFinancingTerrorism] = ?,[AttemptedTranscaction] = ?,[LEAInformed] = ?,[PriorityRating] = ?,[ReportCoverage] = ?,[leadetails] = ?,[AdditionalDocument] = ?,[Aroundofsuspision] = ?,[DetailsofInvestigation] = ?,[AccountNumber] = ?,[AccountType] = ?,[holdername] = ?,[AccountHolderType] = ?,[AccountStatus] = ?,[DateofOpening] = ?,[RiskRating] = ?,[CummulativeCerditTurnover] = ?,[CummulativeDebitTurnover] = ?,[CummulativeCashDepositTurnover] = ?,[CummulativeCashWithdrawalTurnover] = ?,[NoOfTransactionsToBeReported] = ?,[Address] = ?,[Pincode] = ?,[City] = ?,[DOB] = ?,[MobileNumber] = ?,[PAN] = ?,[TransactionAmount] = ?,[TransactionTypeacc] = ?,[TransactionCategory] = ?,[TransactionCurrency] = ?,[BankName] = ?,[BankState] = ?,[TransactionDate] = ?,[TransactionsID] = ?,[TransactionMode] = ?,[DebitCredit] = ?,[amount] = ?,[TransactionsCurrency] = ?,[ProductType] = ?,[ProductIdentifiers] = ?,[TransactionType] = ?,[unit] = ?,[Date] = ?,[DispositionOfFunds] = ?,[RelatedAccountNumber] = ?,[RelatedInstitutionName] = ?,[RelatedInstitutionRefNum] = ?,[Remark] = ?,[ROS_cmt] = ?,[Created_By] = ?,[DGM_cmt] = ? WHERE [ticket_id] = ?"
                    cursor.execute(update_archive_query, (obj['Created_Date'], obj['Customerno'], obj['casename'], obj['scenario'],obj['Guidance'], obj['RuleScenario'], obj['personname'], obj['SourceofAlert'], obj['alertindicator'], obj['SuspiciousDueToproceedofCrime'], obj['SuspiciousDueToComplexTranscaction'],obj['SuspiciousDueToNoecoRational'], obj['SuspiciousDueToFinancingTerrorism'], obj['AttemptedTranscaction'], obj['LEAInformed'], obj['PriorityRating'], obj['ReportCoverage'], obj['leadetails'], obj['AdditionalDocument'], obj['Aroundofsuspision'], obj['DetailsofInvestigation'], obj['AccountNumber'], obj['AccountType'], obj['holdername'], obj['AccountHolderType'], obj['AccountStatus'], obj['DateofOpening'], obj['RiskRating'], obj['CummulativeCerditTurnover'], obj['CummulativeDebitTurnover'], obj['CummulativeCashDepositTurnover'], obj['CummulativeCashWithdrawalTurnover'], obj['NoOfTransactionsToBeReported'],obj['Address'],obj['Pincode'],obj['City'], obj['DOB'],obj['MobileNumber'],obj['PAN'],obj['TransactionAmount'],obj['TransactionTypeacc'],obj['TransactionCategory'],obj['TransactionCurrency'],obj['BankName'],obj['BankState'],obj['TransactionDate'], obj['TransactionsID'], obj['TransactionMode'], obj['DebitCredit'], obj['amount'], obj['TransactionsCurrency'], obj['ProductType'], obj['ProductIdentifiers'],obj['TransactionType'], obj['unit'], obj['Date'], obj['DispositionOfFunds'], obj['RelatedAccountNumber'], obj['RelatedInstitutionName'], obj['RelatedInstitutionRefNum'], obj['Remark'], obj['offlineArc'], obj['Created_By'], obj['offlineArc'], obj['ticket_id']))
                    conn.commit()
                    return redirect(url_for("archived"))
                else:
                    select_get_doj_data = "SELECT [DateofOpening] FROM [dbo].[offline_scenarios] WHERE [ticket_id] = ?"
                    cursor.execute(select_get_doj_data,(obj['ticket_id'],))
                    doj_data = cursor.fetchone()
                    doj_data = doj_data[0]
                    print('doaadata:',doj_data)

                    if obj['DateofOpening'] is None or obj['DateofOpening'] == '':
                        obj['DateofOpening'] = doj_data
                    print('DateofOpening:', obj['DateofOpening'])

                    update_archive_query1 = "UPDATE [dbo].[offline_scenarios] SET [Created_Date] = ?,[Customerno] = ?,[casename] = ?,[scenario] = ?,[Guidance] = ?,[RuleScenario] = ?,[personname] = ?,[SourceofAlert] = ?,[alertindicator] = ?,[SuspiciousDueToproceedofCrime] = ?,[SuspiciousDueToComplexTranscaction] = ?,[SuspiciousDueToNoecoRational] = ?,[SuspiciousDueToFinancingTerrorism] = ?,[AttemptedTranscaction] = ?,[LEAInformed] = ?,[PriorityRating] = ?,[ReportCoverage] = ?,[leadetails] = ?,[AdditionalDocument] = ?,[Aroundofsuspision] = ?,[DetailsofInvestigation] = ?,[AccountNumber] = ?,[AccountType] = ?,[holdername] = ?,[AccountHolderType] = ?,[AccountStatus] = ?,[DateofOpening] = ?,[RiskRating] = ?,[CummulativeCerditTurnover] = ?,[CummulativeDebitTurnover] = ?,[CummulativeCashDepositTurnover] = ?,[CummulativeCashWithdrawalTurnover] = ?,[NoOfTransactionsToBeReported] = ?,[Address] = ?,[Pincode] = ?,[City] = ?,[DOB] = ?,[MobileNumber] = ?,[PAN] = ?,[TransactionAmount] = ?,[TransactionTypeacc] = ?,[TransactionCategory] = ?,[TransactionCurrency] = ?,[BankName] = ?,[BankState] = ?,[TransactionDate] = ?,[TransactionsID] = ?,[TransactionMode] = ?,[DebitCredit] = ?,[amount] = ?,[TransactionsCurrency] = ?,[ProductType] = ?,[ProductIdentifiers] = ?,[TransactionType] = ?,[unit] = ?,[Date] = ?,[DispositionOfFunds] = ?,[RelatedAccountNumber] = ?,[RelatedInstitutionName] = ?,[RelatedInstitutionRefNum] = ?,[Remark] = ?,[ROS_cmt] = ?,[Created_By] = ?,[DGM_cmt] = ? WHERE [ticket_id] = ?"
                    cursor.execute(update_archive_query1, (obj['Created_Date'], obj['Customerno'], obj['casename'], obj['scenario'],obj['Guidance'], obj['RuleScenario'], obj['personname'], obj['SourceofAlert'], obj['alertindicator'], obj['SuspiciousDueToproceedofCrime'], obj['SuspiciousDueToComplexTranscaction'],obj['SuspiciousDueToNoecoRational'], obj['SuspiciousDueToFinancingTerrorism'], obj['AttemptedTranscaction'], obj['LEAInformed'], obj['PriorityRating'], obj['ReportCoverage'], obj['leadetails'], obj['AdditionalDocument'], obj['Aroundofsuspision'], obj['DetailsofInvestigation'], obj['AccountNumber'], obj['AccountType'], obj['holdername'], obj['AccountHolderType'], obj['AccountStatus'], obj['DateofOpening'], obj['RiskRating'], obj['CummulativeCerditTurnover'], obj['CummulativeDebitTurnover'], obj['CummulativeCashDepositTurnover'], obj['CummulativeCashWithdrawalTurnover'], obj['NoOfTransactionsToBeReported'],obj['Address'],obj['Pincode'],obj['City'], obj['DOB'],obj['MobileNumber'],obj['PAN'],obj['TransactionAmount'],obj['TransactionTypeacc'],obj['TransactionCategory'],obj['TransactionCurrency'],obj['BankName'],obj['BankState'],obj['TransactionDate'], obj['TransactionsID'], obj['TransactionMode'], obj['DebitCredit'], obj['amount'], obj['TransactionsCurrency'], obj['ProductType'], obj['ProductIdentifiers'],obj['TransactionType'], obj['unit'], obj['Date'], obj['DispositionOfFunds'], obj['RelatedAccountNumber'], obj['RelatedInstitutionName'], obj['RelatedInstitutionRefNum'], obj['Remark'], None, obj['Created_By'], None, obj['ticket_id']))
                    conn.commit()
                    return redirect(url_for("branchmakers"))
        else:
            if obj['offlineArc'] == 'archive':
                select_get_doj_data = "SELECT [DateofOpening] FROM [dbo].[offline_scenarios] WHERE [ticket_id] = ?"
                cursor.execute(select_get_doj_data,(obj['ticket_id'],))
                doj_data = cursor.fetchone()
                doj_data = doj_data[0]
                print('doaadata:',doj_data)

                if obj['DateofOpening'] is None or obj['DateofOpening'] == '':
                    obj['DateofOpening'] = doj_data
                print('DateofOpening:', obj['DateofOpening'])

                update_archive_query = "UPDATE [dbo].[offline_scenarios] SET [Created_Date] = ?,[Customerno] = ?,[casename] = ?,[scenario] = ?,[Guidance] = ?,[RuleScenario] = ?,[personname] = ?,[SourceofAlert] = ?,[alertindicator] = ?,[SuspiciousDueToproceedofCrime] = ?,[SuspiciousDueToComplexTranscaction] = ?,[SuspiciousDueToNoecoRational] = ?,[SuspiciousDueToFinancingTerrorism] = ?,[AttemptedTranscaction] = ?,[LEAInformed] = ?,[PriorityRating] = ?,[ReportCoverage] = ?,[leadetails] = ?,[AdditionalDocument] = ?,[Aroundofsuspision] = ?,[DetailsofInvestigation] = ?,[AccountNumber] = ?,[AccountType] = ?,[holdername] = ?,[AccountHolderType] = ?,[AccountStatus] = ?,[DateofOpening] = ?,[RiskRating] = ?,[CummulativeCerditTurnover] = ?,[CummulativeDebitTurnover] = ?,[CummulativeCashDepositTurnover] = ?,[CummulativeCashWithdrawalTurnover] = ?,[NoOfTransactionsToBeReported] = ?,[Address] = ?,[Pincode] = ?,[City] = ?,[DOB] = ?,[MobileNumber] = ?,[PAN] = ?,[TransactionAmount] = ?,[TransactionTypeacc] = ?,[TransactionCategory] = ?,[TransactionCurrency] = ?,[BankName] = ?,[BankState] = ?,[TransactionDate] = ?,[TransactionsID] = ?,[TransactionMode] = ?,[DebitCredit] = ?,[amount] = ?,[TransactionsCurrency] = ?,[ProductType] = ?,[ProductIdentifiers] = ?,[TransactionType] = ?,[unit] = ?,[Date] = ?,[DispositionOfFunds] = ?,[RelatedAccountNumber] = ?,[RelatedInstitutionName] = ?,[RelatedInstitutionRefNum] = ?,[Remark] = ?,[ROS_cmt] = ?,[Created_By] = ?,[DGM_cmt] = ? WHERE [ticket_id] = ?"
                cursor.execute(update_archive_query, (obj['Created_Date'], obj['Customerno'], obj['casename'], obj['scenario'],obj['Guidance'], obj['RuleScenario'], obj['personname'], obj['SourceofAlert'], obj['alertindicator'], obj['SuspiciousDueToproceedofCrime'], obj['SuspiciousDueToComplexTranscaction'],obj['SuspiciousDueToNoecoRational'], obj['SuspiciousDueToFinancingTerrorism'], obj['AttemptedTranscaction'], obj['LEAInformed'], obj['PriorityRating'], obj['ReportCoverage'], obj['leadetails'], obj['AdditionalDocument'], obj['Aroundofsuspision'], obj['DetailsofInvestigation'], obj['AccountNumber'], obj['AccountType'], obj['holdername'], obj['AccountHolderType'], obj['AccountStatus'], obj['DateofOpening'], obj['RiskRating'], obj['CummulativeCerditTurnover'], obj['CummulativeDebitTurnover'], obj['CummulativeCashDepositTurnover'], obj['CummulativeCashWithdrawalTurnover'], obj['NoOfTransactionsToBeReported'],obj['Address'],obj['Pincode'],obj['City'], obj['DOB'],obj['MobileNumber'],obj['PAN'],obj['TransactionAmount'],obj['TransactionTypeacc'],obj['TransactionCategory'],obj['TransactionCurrency'],obj['BankName'],obj['BankState'], obj['TransactionDate'], obj['TransactionsID'], obj['TransactionMode'], obj['DebitCredit'], obj['amount'], obj['TransactionsCurrency'], obj['ProductType'], obj['ProductIdentifiers'],obj['TransactionType'], obj['unit'], obj['Date'], obj['DispositionOfFunds'], obj['RelatedAccountNumber'], obj['RelatedInstitutionName'], obj['RelatedInstitutionRefNum'], obj['Remark'], obj['offlineArc'], obj['Created_By'], obj['offlineArc'], obj['ticket_id']))
                conn.commit()
                return redirect(url_for("archived"))
            else:
                select_get_doj_data = "SELECT [DateofOpening] FROM [dbo].[offline_scenarios] WHERE [ticket_id] = ?"
                cursor.execute(select_get_doj_data,(obj['ticket_id'],))
                doj_data = cursor.fetchone()
                doj_data = doj_data[0]
                print('doaadata:',doj_data)

                if obj['DateofOpening'] is None or obj['DateofOpening'] == '':
                    obj['DateofOpening'] = doj_data
                print('DateofOpening:', obj['DateofOpening'])

                update_archive_query1 = "UPDATE [dbo].[offline_scenarios] SET [Created_Date] = ?,[Customerno] = ?,[casename] = ?,[scenario] = ?,[Guidance] = ?,[RuleScenario] = ?,[personname] = ?,[SourceofAlert] = ?,[alertindicator] = ?,[SuspiciousDueToproceedofCrime] = ?,[SuspiciousDueToComplexTranscaction] = ?,[SuspiciousDueToNoecoRational] = ?,[SuspiciousDueToFinancingTerrorism] = ?,[AttemptedTranscaction] = ?,[LEAInformed] = ?,[PriorityRating] = ?,[ReportCoverage] = ?,[leadetails] = ?,[AdditionalDocument] = ?,[Aroundofsuspision] = ?,[DetailsofInvestigation] = ?,[AccountNumber] = ?,[AccountType] = ?,[holdername] = ?,[AccountHolderType] = ?,[AccountStatus] = ?,[DateofOpening] = ?,[RiskRating] = ?,[CummulativeCerditTurnover] = ?,[CummulativeDebitTurnover] = ?,[CummulativeCashDepositTurnover] = ?,[CummulativeCashWithdrawalTurnover] = ?,[NoOfTransactionsToBeReported] = ?,[Address] = ?,[Pincode] = ?,[City] = ?,[DOB] = ?,[MobileNumber] = ?,[PAN] = ?,[TransactionAmount] = ?,[TransactionTypeacc] = ?,[TransactionCategory] = ?,[TransactionCurrency] = ?,[BankName] = ?,[BankState] = ?,[TransactionDate] = ?,[TransactionsID] = ?,[TransactionMode] = ?,[DebitCredit] = ?,[amount] = ?,[TransactionsCurrency] = ?,[ProductType] = ?,[ProductIdentifiers] = ?,[TransactionType] = ?,[unit] = ?,[Date] = ?,[DispositionOfFunds] = ?,[RelatedAccountNumber] = ?,[RelatedInstitutionName] = ?,[RelatedInstitutionRefNum] = ?,[Remark] = ?,[ROS_cmt] = ?,[Created_By] = ?,[DGM_cmt] = ? WHERE [ticket_id] = ?"
                cursor.execute(update_archive_query1, (obj['Created_Date'], obj['Customerno'], obj['casename'], obj['scenario'],obj['Guidance'], obj['RuleScenario'], obj['personname'], obj['SourceofAlert'], obj['alertindicator'], obj['SuspiciousDueToproceedofCrime'], obj['SuspiciousDueToComplexTranscaction'],obj['SuspiciousDueToNoecoRational'], obj['SuspiciousDueToFinancingTerrorism'], obj['AttemptedTranscaction'], obj['LEAInformed'], obj['PriorityRating'], obj['ReportCoverage'], obj['leadetails'], obj['AdditionalDocument'], obj['Aroundofsuspision'], obj['DetailsofInvestigation'], obj['AccountNumber'], obj['AccountType'], obj['holdername'], obj['AccountHolderType'], obj['AccountStatus'], obj['DateofOpening'], obj['RiskRating'], obj['CummulativeCerditTurnover'], obj['CummulativeDebitTurnover'], obj['CummulativeCashDepositTurnover'], obj['CummulativeCashWithdrawalTurnover'], obj['NoOfTransactionsToBeReported'],obj['Address'],obj['Pincode'],obj['City'], obj['DOB'],obj['MobileNumber'],obj['PAN'],obj['TransactionAmount'],obj['TransactionTypeacc'],obj['TransactionCategory'],obj['TransactionCurrency'],obj['BankName'],obj['BankState'], obj['TransactionDate'], obj['TransactionsID'], obj['TransactionMode'], obj['DebitCredit'], obj['amount'], obj['TransactionsCurrency'], obj['ProductType'], obj['ProductIdentifiers'],obj['TransactionType'], obj['unit'], obj['Date'], obj['DispositionOfFunds'], obj['RelatedAccountNumber'], obj['RelatedInstitutionName'], obj['RelatedInstitutionRefNum'], obj['Remark'], None, obj['Created_By'], None, obj['ticket_id']))
                conn.commit()
                return redirect(url_for("archived"))            
            #     # Notify the branchmaker that the ticket has been created
            #     # notify = notification(branchmaker_email)
        
    except Exception as e:
        print("An error occurred:", e)
        traceback.print_exc()  # Print the traceback to identify the issue
        return "An error occurred while processing your request", 500  # Return an error response



# --------------------------------- OFFLINE CASES UPDATED SUBMIT TO NEXT LEVEL OFFICERS END - POINT --------------------------------------------


# @app.route('/update_offline_str_ros', methods=['POST'])
# @secure_route(required_role=['ROS','DGM/PO','BranchMakers'])
# def update_offline_str_ros():
#     conn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
#     cursor = conn.cursor()
#     ros_email_updated = session['email_id']
#     print('RoS update Email id:',ros_email_updated)
#     if request.method == 'POST':
#         try:
#             current_datetime = datetime.now()
#             current_date = current_datetime.date()
#             midnight_datetime = datetime.combine(current_date, datetime.min.time())

#             comments = request.form.get('comments')
#             # ticket_id = request.form.get('ticket_id')
#             approval_status = request.form.get('finalReport')
            
#             ticket_id = request.args.get('account_number')
#             mailid = session['email_id']
#             # print("seesion id:",mailid)

#             cursor = mysql2.connection.cursor()
#             select_get_doj_data = "SELECT [DateofOpening] FROM [dbo].[offline_collection] WHERE [ticket_id] = ?"
#             cursor.execute(select_get_doj_data,(ticket_id,))
#             doj_data = cursor.fetchone()
#             doj_data = doj_data[0]
#             print('doaadata:',doj_data)

#             customerno = request.form.get('Customerno')
#             casename = request.form.get('casename')
#             personname = request.form.get('personname')
#             RuleScenario = request.form.get('RuleScenario')
#             Guidance = request.form.get('Guidance')
#             scenario = request.form.get('scenario')
#             SourceofAlert = request.form.get('SourceofAlert')
#             alertindicator = request.form.get('alertindicator')
#             SuspiciousDueToproceedofCrime = request.form.get('SuspiciousDueToproceedofCrime')
#             SuspiciousDueToComplexTranscaction = request.form.get('SuspiciousDueToComplexTranscaction')
#             SuspiciousDueToNoecoRational = request.form.get('SuspiciousDueToNoecoRational')
#             SuspiciousDueToFinancingTerrorism = request.form.get('SuspiciousDueToFinancingTerrorism')
#             AttemptedTranscaction = request.form.get('AttemptedTranscaction')
#             LEAInformed = request.form.get('LEAInformed')
#             PriorityRating = request.form.get('PriorityRating')
#             ReportCoverage = request.form.get('ReportCoverage')
#             leadetails = request.form.get('leadetails')
#             AdditionalDocument = request.form.get('AdditionalDocument')
#             Aroundofsuspision = request.form.get('Aroundofsuspision')
#             DetailsofInvestigation = request.form.get('DetailsofInvestigation')
#             AccountNumber = request.form.get('AccountNumber')
#             AccountType = request.form.get('AccountType')
#             holdername = request.form.get('holdername')
#             AccountHolderType = request.form.get('AccountHolderType')
#             AccountStatus = request.form.get('AccountStatus')
#             DateofOpening = request.form.get('DateofOpening')
#             RiskRating = request.form.get('RiskRating')
#             CummulativeCerditTurnover = request.form.get('CummulativeCerditTurnover')
#             CummulativeDebitTurnover = request.form.get('CummulativeDebitTurnover')
#             CummulativeCashDepositTurnover = request.form.get('CummulativeCashDepositTurnover')
#             CummulativeCashWithdrawalTurnover = request.form.get('CummulativeCashWithdrawalTurnover')
#             NoOfTransactionsToBeReported = request.form.get('NoOfTransactionsToBeReported')
#             TransactionDate = request.form.get('TransactionDate')
#             TransactionsID = request.form.get('TransactionsID')
#             TransactionMode = request.form.get('TransactionMode')
#             DebitCredit = request.form.get('DebitCredit')
#             amount = request.form.get('amount')
#             TransactionsCurrency = request.form.get('TransactionsCurrency')
#             ProductType = request.form.get('ProductType')
#             ProductIdentifiers = request.form.get('ProductIdentifiers')
#             TransactionType = request.form.get('TransactionType')
#             unit = request.form.get('unit')
#             Date = request.form.get('Date')
#             DispositionOfFunds = request.form.get('DispositionOfFunds')
#             RelatedAccountNumber = request.form.get('RelatedAccountNumber')
#             RelatedInstitutionName = request.form.get('RelatedInstitutionName')
#             RelatedInstitutionRefNum = request.form.get('RelatedInstitutionRefNum')
#             Remark = request.form.get('Remark')

#             obj = {
#              'Customerno': customerno,
#                 'casename': casename,
#                 'scenario': scenario,
#                 'Guidance': Guidance,
#                 'RuleScenario': RuleScenario,
#                 'personname': personname,
#                 'SourceofAlert': SourceofAlert,
#                 'alertindicator': alertindicator,
#                 'SuspiciousDueToproceedofCrime': SuspiciousDueToproceedofCrime,
#                 'SuspiciousDueToComplexTranscaction': SuspiciousDueToComplexTranscaction,
#                 'SuspiciousDueToNoecoRational': SuspiciousDueToNoecoRational,
#                 'SuspiciousDueToFinancingTerrorism': SuspiciousDueToFinancingTerrorism,
#                 'AttemptedTranscaction': AttemptedTranscaction,
#                 'LEAInformed': LEAInformed,
#                 'PriorityRating': PriorityRating,
#                 'ReportCoverage': ReportCoverage,
#                 'leadetails': leadetails,
#                 'AdditionalDocument': AdditionalDocument,
#                 'Aroundofsuspision': Aroundofsuspision,
#                 'DetailsofInvestigation': DetailsofInvestigation,
#                 'AccountNumber': AccountNumber,
#                 'AccountType': AccountType,
#                 'holdername': holdername,
#                 'AccountHolderType': AccountHolderType,
#                 'AccountStatus': AccountStatus,
#                 'DateofOpening': DateofOpening,
#                 'RiskRating': RiskRating,
#                 'CummulativeCerditTurnover': CummulativeCerditTurnover,
#                 'CummulativeDebitTurnover': CummulativeDebitTurnover,
#                 'CummulativeCashDepositTurnover': CummulativeCashDepositTurnover,
#                 'CummulativeCashWithdrawalTurnover': CummulativeCashWithdrawalTurnover,
#                 'NoOfTransactionsToBeReported': NoOfTransactionsToBeReported,
#                 'TransactionDate': TransactionDate,
#                 'TransactionsID': TransactionsID,
#                 'TransactionMode': TransactionMode,
#                 'DebitCredit': DebitCredit,
#                 'amount': amount,
#                 'TransactionsCurrency': TransactionsCurrency,
#                 'ProductType': ProductType,
#                 'ProductIdentifiers': ProductIdentifiers,
#                 'TransactionType': TransactionType,
#                 'unit': unit,
#                 'Date': Date,
#                 'DispositionOfFunds': DispositionOfFunds,
#                 'RelatedAccountNumber': RelatedAccountNumber,
#                 'RelatedInstitutionName': RelatedInstitutionName,
#                 'RelatedInstitutionRefNum': RelatedInstitutionRefNum,
#                 'Remark': Remark,
#                }
#             if obj['DateofOpening'] is None or obj['DateofOpening'] == '':
#                 obj['DateofOpening'] = doj_data
#                 print('DateofOpening:', obj['DateofOpening'])

#             update_data_for_ros_dgm_query = "UPDATE [dbo].[offline_collection] SET[SuspiciousDueToproceedofCrime] = ?,[SuspiciousDueToComplexTranscaction] = ?,[SuspiciousDueToNoecoRational] = ?,[SuspiciousDueToFinancingTerrorism] = ?,[AttemptedTranscaction] = ?,[LEAInformed] = ?,[PriorityRating] = ?,[ReportCoverage] = ?,[leadetails] = ?,[AdditionalDocument] = ?,[Aroundofsuspision] = ?,[DetailsofInvestigation] = ? WHERE [ticket_id] = ?"
#             cursor.execute(update_data_for_ros_dgm_query, (obj['SuspiciousDueToproceedofCrime'], obj['SuspiciousDueToComplexTranscaction'],obj['SuspiciousDueToNoecoRational'], obj['SuspiciousDueToFinancingTerrorism'], obj['AttemptedTranscaction'], obj['LEAInformed'], obj['PriorityRating'], obj['ReportCoverage'], obj['leadetails'], obj['AdditionalDocument'], obj['Aroundofsuspision'], obj['DetailsofInvestigation'], ticket_id))
#             conn.commit()

#             search_query = "SELECT [Role] FROM [dbo].[user] WHERE EmailId = ?"
#             cursor.execute(search_query, (mailid,))
#             result = cursor.fetchone()
#             print("data:",result[0])
#             # cursor.close()
#             if result[0] == 'ROS':
#                 if comments:
#                     select_check_query = "SELECT [ROS_cmt], [DGM_cmt] FROM [dbo].[offline_collection] WHERE [ticket_id] = ?"
#                     cursor.execute(select_check_query, (ticket_id,))
#                     result = cursor.fetchone()
#                     print('result:',result[0])
#                     print('result:',result[1])
#                     if result[0] == None and result[1] == None:
#                         # print('Approved:',approval_status)  
#                         # update the ROS Comment based on Account Number(Account Number is getting for the Privious Page)
#                         update_query = "UPDATE [dbo].[offline_collection] SET ROS_cmt = ?, DGM_cmt = ? WHERE ticket_id = ?"
#                         cursor = mysql2.connection.cursor()
#                         cursor.execute(update_query, (comments, None, ticket_id))
#                         mysql2.connection.commit()

#                         # Getting the Ticket ID based on Account Number(Account Number is getting for the Privious Page)
#                         # select_query = "SELECT [ticket_id] FROM [dbo].[offline_collection] WHERE [AccountNumber] = ?"
#                         # cursor.execute(select_query, (ticket_id,))
#                         # result = cursor.fetchone()
#                         # Tid = result[0]

#                         # Getting the User ID based on User Email id(Email id is getting for the Session)
#                         select_query = "SELECT [id] FROM [dbo].[user] WHERE [EmailId] = ?"
#                         cursor.execute(select_query, (ros_email_updated,))
#                         result1 = cursor.fetchone()
#                         id = result1[0]

#                         print(' ID: ',id)
#                         print('TID:',ticket_id)

#                         # Update the User ID in allocated_ticket & ROSCasesTicket in the offline_tickets table based on Ticket ID
#                         update_query1 = "UPDATE [dbo].[offline_tickets] SET [allocatedTicket] = ?, [ROSCasesTicket] = ? WHERE [ticketid] = ?"
#                         cursor.execute(update_query1, (id, id, ticket_id))
#                         mysql2.connection.commit()

#                         cursor.close()
#                         print(comments)
#                         print(ticket_id)
#                         session['success_message'] = 'ROS Commented successfully.'
#                         return redirect(url_for("ros"))
#                     else:
#                         # print('Approved:',approval_status)  
#                         # update the ROS Comment based on Account Number(Account Number is getting for the Privious Page)
#                         update_query = "UPDATE [dbo].[offline_collection] SET ROS_cmt = ?, DGM_cmt = ? WHERE ticket_id = ?"
#                         cursor = mysql2.connection.cursor()
#                         cursor.execute(update_query, (comments, None, ticket_id))
#                         mysql2.connection.commit()

#                         # Getting the Ticket ID based on Account Number(Account Number is getting for the Privious Page)
#                         # select_query = "SELECT [ticket_id] FROM [dbo].[offline_collection] WHERE [AccountNumber] = ?"
#                         # cursor.execute(select_query, (ticket_id,))
#                         # result = cursor.fetchone()
#                         # Tid = result[0]

#                         # Getting the User ID based on User Email id(Email id is getting for the Session)
#                         select_query = "SELECT [id] FROM [dbo].[user] WHERE [EmailId] = ?"
#                         cursor.execute(select_query, (ros_email_updated,))
#                         result1 = cursor.fetchone()
#                         id = result1[0]

#                         print(' ID: ',id)
#                         print('TID:',ticket_id)

#                         # Update the User ID in allocated_ticket & ROSCasesTicket in the offline_tickets table based on Ticket ID
#                         update_query1 = "UPDATE [dbo].[offline_tickets] SET [allocatedTicket] = ?, [ROSCasesTicket] = ? WHERE [ticketid] = ?"
#                         cursor.execute(update_query1, (id, id, ticket_id))
#                         mysql2.connection.commit()

#                         cursor.close()
#                         print(comments)
#                         print(ticket_id)
#                         # session['success_message'] = 'ROS Commented successfully.'
#                         return redirect(url_for("archived"))
#                 else:
#                     update_archive_cmt = "UPDATE [dbo].[offline_collection] SET [ROS_cmt] = ?, [DGM_cmt] = ?  WHERE ticket_id = ?"
#                     cursor.execute(update_archive_cmt,(None, 'archive', ticket_id))
#                     conn.commit()

#                     # Getting the User ID based on User Email id(Email id is getting for the Session)
#                     select_query = "SELECT [id] FROM [dbo].[user] WHERE [EmailId] = ?"
#                     cursor.execute(select_query, (ros_email_updated,))
#                     result1 = cursor.fetchone()
#                     id = result1[0]

#                     print(' ID: ',id)
#                     print('TID:',ticket_id)

#                     # Update the User ID in allocated_ticket & ROSCasesTicket in the offline_tickets table based on Ticket ID
#                     update_query1 = "UPDATE [dbo].[offline_tickets] SET [allocatedTicket] = ? WHERE [ticketid] = ?"
#                     cursor.execute(update_query1, (id, ticket_id))
#                     mysql2.connection.commit()

#                     cursor.close()
#                     print(comments)
#                     print(ticket_id)
#                     # session['success_message'] = 'Successfully Archived in archives.'
#                     return redirect(url_for("archived"))
                
#             elif result[0] == 'DGM/PO':
#                 if comments:
#                     select_check_query = "SELECT [ROS_cmt], [DGM_cmt] FROM [dbo].[offline_collection] WHERE [ticket_id] = ?"
#                     cursor.execute(select_check_query, (ticket_id,))
#                     result = cursor.fetchone()
#                     print('result:',result[0])
#                     print('result:',result[1])
#                     if result[1] == None:
#                         if approval_status == 'Approved':
#                             # update the DGM/PO Comment based on Account Number(Account Number is getting for the Privious Page)
#                             update_query = "UPDATE [dbo].[offline_collection] SET DGM_cmt = ? WHERE ticket_id = ?"
#                             cursor = mysql2.connection.cursor()
#                             cursor.execute(update_query, (comments, ticket_id))
#                             mysql2.connection.commit()

#                             # Getting the Ticket ID based on Account Number(Account Number is getting for the Privious Page)
#                             # select_query = "SELECT [ticket_id] FROM [dbo].[offline_collection] WHERE [ticket_id] = ?"
#                             # cursor.execute(select_query, (ticket_id,))
#                             # result = cursor.fetchone()
#                             # Tid = result[0]

#                             # Getting the User ID based on Role(Because here i have only one DGM so no need check the email id)
#                             select_query = "SELECT [id] FROM [dbo].[user] WHERE [Role] = ?"
#                             cursor.execute(select_query, ('DGM/PO',))
#                             result1 = cursor.fetchone()
#                             id = result1[0]

#                             print(' ID: ',id)
#                             print('TID:',ticket_id)

#                             # 0 represents false and 1 represents true.    
#                             # Update the User ID in allocated_ticket & DGMCasesTicket in the offline_tickets table based on Ticket ID
#                             update_query1 = "UPDATE [dbo].[offline_tickets] SET [allocatedTicket] = ?, [DGMCasesTicket] = ?, [rejected] = ?, [approved] = ? WHERE [ticketid] = ?"
#                             cursor.execute(update_query1, (id, id, 0, 1, ticket_id))
#                             mysql2.connection.commit()

#                             cursor.close()
#                             print(comments)
#                             print(ticket_id)
#                             session['success_message'] = 'DGM Commented successfully.'
#                             return redirect(url_for("offline_dgm_Str"))
#                         elif approval_status == 'Rejected':
#                             # update the DGM/PO Comment based on Account Number(Account Number is getting for the Privious Page)
#                             update_query = "UPDATE [dbo].[offline_collection] SET DGM_cmt = ? WHERE ticket_id = ?"
#                             cursor = mysql2.connection.cursor()
#                             cursor.execute(update_query, (comments, ticket_id))
#                             mysql2.connection.commit()

#                             # Getting the User ID based on Role(Because here i have only one DGM so no need check the email id)
#                             select_query = "SELECT [id] FROM [dbo].[user] WHERE [Role] = ?"
#                             cursor.execute(select_query, ('DGM/PO',))
#                             result1 = cursor.fetchone()
#                             id = result1[0]

#                             print(' ID: ',id)
#                             print('TID:',ticket_id)

#                             # 0 represents false and 1 represents true. 
#                             # Update the User ID in allocated_ticket & DGMCasesTicket in the offline_tickets table based on Ticket ID
#                             update_query1 = "UPDATE [dbo].[offline_tickets] SET [allocatedTicket] = ?, [DGMCasesTicket] = ?, [rejected] = ?, [approved] = ? WHERE [ticketid] = ?"
#                             cursor.execute(update_query1, (id, id, 1, 0, ticket_id))
#                             mysql2.connection.commit()

#                             cursor.close()
#                             print(comments)
#                             print(ticket_id)
                                
#                             return redirect(url_for("offline_dgm_Str"))
#                     else:
#                         if approval_status == 'Approved':
#                             # update the DGM/PO Comment based on Account Number(Account Number is getting for the Privious Page)
#                             update_query = "UPDATE [dbo].[offline_collection] SET DGM_cmt = ? WHERE ticket_id = ?"
#                             cursor = mysql2.connection.cursor()
#                             cursor.execute(update_query, (comments, ticket_id))
#                             mysql2.connection.commit()

#                             # Getting the Ticket ID based on Account Number(Account Number is getting for the Privious Page)
#                             # select_query = "SELECT [ticket_id] FROM [dbo].[offline_collection] WHERE [ticket_id] = ?"
#                             # cursor.execute(select_query, (ticket_id,))
#                             # result = cursor.fetchone()
#                             # Tid = result[0]

#                             # Getting the User ID based on Role(Because here i have only one DGM so no need check the email id)
#                             select_query = "SELECT [id] FROM [dbo].[user] WHERE [Role] = ?"
#                             cursor.execute(select_query, ('DGM/PO',))
#                             result1 = cursor.fetchone()
#                             id = result1[0]

#                             print(' ID: ',id)
#                             print('TID:',ticket_id)

#                             # 0 represents false and 1 represents true.    
#                             # Update the User ID in allocated_ticket & DGMCasesTicket in the offline_tickets table based on Ticket ID
#                             update_query1 = "UPDATE [dbo].[offline_tickets] SET [allocatedTicket] = ?, [DGMCasesTicket] = ?, [rejected] = ?, [approved] = ? WHERE [ticketid] = ?"
#                             cursor.execute(update_query1, (id, id, 0, 1, ticket_id))
#                             mysql2.connection.commit()

#                             cursor.close()
#                             print(comments)
#                             print(ticket_id)
#                             # session['success_message'] = 'DGM Commented successfully.'
#                             return redirect(url_for("archived"))
#                         elif approval_status == 'Rejected':
#                                 # update the DGM/PO Comment based on Account Number(Account Number is getting for the Privious Page)
#                                 update_query = "UPDATE [dbo].[offline_collection] SET DGM_cmt = ? WHERE ticket_id = ?"
#                                 cursor = mysql2.connection.cursor()
#                                 cursor.execute(update_query, (comments, ticket_id))
#                                 mysql2.connection.commit()

#                                 # Getting the User ID based on Role(Because here i have only one DGM so no need check the email id)
#                                 select_query = "SELECT [id] FROM [dbo].[user] WHERE [Role] = ?"
#                                 cursor.execute(select_query, ('DGM/PO',))
#                                 result1 = cursor.fetchone()
#                                 id = result1[0]

#                                 print(' ID: ',id)
#                                 print('TID:',ticket_id)

#                                 # 0 represents false and 1 represents true. 
#                                 # Update the User ID in allocated_ticket & DGMCasesTicket in the offline_tickets table based on Ticket ID
#                                 update_query1 = "UPDATE [dbo].[offline_tickets] SET [allocatedTicket] = ?, [DGMCasesTicket] = ?, [rejected] = ?, [approved] = ? WHERE [ticketid] = ?"
#                                 cursor.execute(update_query1, (id, id, 1, 0, ticket_id))
#                                 mysql2.connection.commit()

#                                 cursor.close()
#                                 print(comments)
#                                 print(ticket_id)
                                    
#                                 return redirect(url_for("archived"))
#                 else:
#                     update_archive_cmt = "UPDATE [dbo].[offline_collection] SET [DGM_cmt] = ?  WHERE ticket_id = ?"
#                     cursor.execute(update_archive_cmt,('archive', ticket_id))
#                     conn.commit()

#                     # Getting the User ID based on Role(Because here i have only one DGM so no need check the email id)
#                     select_query = "SELECT [id] FROM [dbo].[user] WHERE [Role] = ?"
#                     cursor.execute(select_query, ('DGM/PO',))
#                     result1 = cursor.fetchone()
#                     id = result1[0]

#                     print(' ID: ',id)
#                     print('TID:',ticket_id)

#                     # 0 represents false and 1 represents true. 
#                     # Update the User ID in allocated_ticket & DGMCasesTicket in the offline_tickets table based on Ticket ID
#                     update_query1 = "UPDATE [dbo].[offline_tickets] SET [allocatedTicket] = ? WHERE [ticketid] = ?"
#                     cursor.execute(update_query1, (id, ticket_id))
#                     mysql2.connection.commit()

#                     # session['success_message'] = 'Successfully Archived in archives.'
#                     return redirect(url_for("archived"))
#             else:
#                 return "Unable to find the Details..."
#         except Exception as e:
#             # Print or log the error message
#             print("Error:", e)
#             # Rollback the transaction if an error occurs
#             mysql2.connection.rollback()
#             # Close the cursor
#             # cursor.close()
#             # Redirect or return an error message
#             return 'An error occurred while updating the data'
#     else:
#         return 'Method Not Allowed', 405

@app.route('/update_offline_str_ros', methods=['POST'])
@secure_route(required_role=['ROS','AGM','DGM/PO','BranchMakers'])
def update_offline_str_ros():
    conn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
    cursor = conn.cursor()
    ros_email_updated = session['email_id']
    print('RoS update Email id:',ros_email_updated)
    if request.method == 'POST':
        try:
            current_datetime = datetime.now()
            current_date = current_datetime.date()
            midnight_datetime = datetime.combine(current_date, datetime.min.time())

            comments = request.form.get('comments')
            # ticket_id = request.form.get('ticket_id')
            approval_status = request.form.get('finalReport')
            
            ticket_id = request.args.get('account_number')
            mailid = session['email_id']
            # print("seesion id:",mailid)

            cursor = mysql2.connection.cursor()
            select_get_doj_data = "SELECT [DateofOpening] FROM [dbo].[offline_scenarios] WHERE [ticket_id] = ?"
            cursor.execute(select_get_doj_data,(ticket_id,))
            doj_data = cursor.fetchone()
            doj_data = doj_data[0]
            print('doaadata:',doj_data)

            customerno = request.form.get('Customerno')
            casename = request.form.get('casename')
            personname = request.form.get('personname')
            RuleScenario = request.form.get('RuleScenario')
            Guidance = request.form.get('Guidance')
            scenario = request.form.get('scenario')
            SourceofAlert = request.form.get('SourceofAlert')
            alertindicator = request.form.get('alertindicator')
            SuspiciousDueToproceedofCrime = request.form.get('SuspiciousDueToproceedofCrime')
            SuspiciousDueToComplexTranscaction = request.form.get('SuspiciousDueToComplexTranscaction')
            SuspiciousDueToNoecoRational = request.form.get('SuspiciousDueToNoecoRational')
            SuspiciousDueToFinancingTerrorism = request.form.get('SuspiciousDueToFinancingTerrorism')
            AttemptedTranscaction = request.form.get('AttemptedTranscaction')
            LEAInformed = request.form.get('LEAInformed')
            PriorityRating = request.form.get('PriorityRating')
            ReportCoverage = request.form.get('ReportCoverage')
            leadetails = request.form.get('leadetails')
            AdditionalDocument = request.form.get('AdditionalDocument')
            Aroundofsuspision = request.form.get('Aroundofsuspision')
            DetailsofInvestigation = request.form.get('DetailsofInvestigation')
            AccountNumber = request.form.get('AccountNumber')
            AccountType = request.form.get('AccountType')
            holdername = request.form.get('holdername')
            AccountHolderType = request.form.get('AccountHolderType')
            AccountStatus = request.form.get('AccountStatus')
            DateofOpening = request.form.get('DateofOpening')
            RiskRating = request.form.get('RiskRating')
            CummulativeCerditTurnover = request.form.get('CummulativeCerditTurnover')
            CummulativeDebitTurnover = request.form.get('CummulativeDebitTurnover')
            CummulativeCashDepositTurnover = request.form.get('CummulativeCashDepositTurnover')
            CummulativeCashWithdrawalTurnover = request.form.get('CummulativeCashWithdrawalTurnover')
            NoOfTransactionsToBeReported = request.form.get('NoOfTransactionsToBeReported')
            address = request.form.get('Address')
            pincode = request.form.get('Pincode')
            city = request.form.get('City')
            dob = request.form.get('DOB')
            mobilenumber = request.form.get('MobileNumber')
            pan = request.form.get('PAN')
            transactionamount = request.form.get('TransactionAmount')
            transactiontypeacc = request.form.get('TransactionTypeacc')
            transactioncategory = request.form.get('TransactionCategory')
            transactioncurrency = request.form.get('TransactionCurrency')
            bankname = request.form.get('BankName')
            bankstate = request.form.get('BankState')
            TransactionDate = request.form.get('TransactionDate')
            TransactionsID = request.form.get('TransactionsID')
            TransactionMode = request.form.get('TransactionMode')
            DebitCredit = request.form.get('DebitCredit')
            amount = request.form.get('amount')
            TransactionsCurrency = request.form.get('TransactionsCurrency')
            ProductType = request.form.get('ProductType')
            ProductIdentifiers = request.form.get('ProductIdentifiers')
            TransactionType = request.form.get('TransactionType')
            unit = request.form.get('unit')
            Date = request.form.get('Date')
            DispositionOfFunds = request.form.get('DispositionOfFunds')
            RelatedAccountNumber = request.form.get('RelatedAccountNumber')
            RelatedInstitutionName = request.form.get('RelatedInstitutionName')
            RelatedInstitutionRefNum = request.form.get('RelatedInstitutionRefNum')
            Remark = request.form.get('Remark')

            obj = {
             'Customerno': customerno,
                'casename': casename,
                'scenario': scenario,
                'Guidance': Guidance,
                'RuleScenario': RuleScenario,
                'personname': personname,
                'SourceofAlert': SourceofAlert,
                'alertindicator': alertindicator,
                'SuspiciousDueToproceedofCrime': SuspiciousDueToproceedofCrime,
                'SuspiciousDueToComplexTranscaction': SuspiciousDueToComplexTranscaction,
                'SuspiciousDueToNoecoRational': SuspiciousDueToNoecoRational,
                'SuspiciousDueToFinancingTerrorism': SuspiciousDueToFinancingTerrorism,
                'AttemptedTranscaction': AttemptedTranscaction,
                'LEAInformed': LEAInformed,
                'PriorityRating': PriorityRating,
                'ReportCoverage': ReportCoverage,
                'leadetails': leadetails,
                'AdditionalDocument': AdditionalDocument,
                'Aroundofsuspision': Aroundofsuspision,
                'DetailsofInvestigation': DetailsofInvestigation,
                'AccountNumber': AccountNumber,
                'AccountType': AccountType,
                'holdername': holdername,
                'AccountHolderType': AccountHolderType,
                'AccountStatus': AccountStatus,
                'DateofOpening': DateofOpening,
                'RiskRating': RiskRating,
                'CummulativeCerditTurnover': CummulativeCerditTurnover,
                'CummulativeDebitTurnover': CummulativeDebitTurnover,
                'CummulativeCashDepositTurnover': CummulativeCashDepositTurnover,
                'CummulativeCashWithdrawalTurnover': CummulativeCashWithdrawalTurnover,
                'NoOfTransactionsToBeReported': NoOfTransactionsToBeReported,
                'address':address,
                'pincode':pincode,
                'city':city,
                'dob':dob,
                'mobilenumber': mobilenumber,
                'pan':pan,
                'transactionamount':transactionamount,
                'transactiontypeacc':transactiontypeacc,
                'transactioncategory': transactioncategory,
                'transactioncurrency':transactioncurrency,
                'bankname':bankname,
                'bankstate' : bankstate,
                'TransactionDate': TransactionDate,
                'TransactionsID': TransactionsID,
                'TransactionMode': TransactionMode,
                'DebitCredit': DebitCredit,
                'amount': amount,
                'TransactionsCurrency': TransactionsCurrency,
                'ProductType': ProductType,
                'ProductIdentifiers': ProductIdentifiers,
                'TransactionType': TransactionType,
                'unit': unit,
                'Date': Date,
                'DispositionOfFunds': DispositionOfFunds,
                'RelatedAccountNumber': RelatedAccountNumber,
                'RelatedInstitutionName': RelatedInstitutionName,
                'RelatedInstitutionRefNum': RelatedInstitutionRefNum,
                'Remark': Remark,
               }
            if obj['DateofOpening'] is None or obj['DateofOpening'] == '':
                obj['DateofOpening'] = doj_data
                print('DateofOpening:', obj['DateofOpening'])

            update_data_for_ros_dgm_query = "UPDATE [dbo].[offline_scenarios] SET[SuspiciousDueToproceedofCrime] = ?,[SuspiciousDueToComplexTranscaction] = ?,[SuspiciousDueToNoecoRational] = ?,[SuspiciousDueToFinancingTerrorism] = ?,[AttemptedTranscaction] = ?,[LEAInformed] = ?,[PriorityRating] = ?,[ReportCoverage] = ?,[leadetails] = ?,[AdditionalDocument] = ?,[Aroundofsuspision] = ?,[DetailsofInvestigation] = ? WHERE [ticket_id] = ?"
            cursor.execute(update_data_for_ros_dgm_query, (obj['SuspiciousDueToproceedofCrime'], obj['SuspiciousDueToComplexTranscaction'],obj['SuspiciousDueToNoecoRational'], obj['SuspiciousDueToFinancingTerrorism'], obj['AttemptedTranscaction'], obj['LEAInformed'], obj['PriorityRating'], obj['ReportCoverage'], obj['leadetails'], obj['AdditionalDocument'], obj['Aroundofsuspision'], obj['DetailsofInvestigation'], ticket_id))
            conn.commit()

            search_query = "SELECT [Role] FROM [dbo].[user] WHERE EmailId = ?"
            cursor.execute(search_query, (mailid,))
            result = cursor.fetchone()
            print("data:",result[0])
            # cursor.close()
            if result[0] == 'ROS':
                if comments:
                    select_check_query = "SELECT [ROS_cmt], [DGM_cmt] FROM [dbo].[offline_scenarios] WHERE [ticket_id] = ?"
                    cursor.execute(select_check_query, (ticket_id,))
                    result = cursor.fetchone()
                    print('result:',result[0])
                    print('result:',result[1])
                    if result[0] == None and result[1] == None:
                        # print('Approved:',approval_status)  
                        # update the ROS Comment based on Account Number(Account Number is getting from the Privious Page)
                        update_query = "UPDATE [dbo].[offline_scenarios] SET ROS_cmt = ?, DGM_cmt = ? WHERE ticket_id = ?"
                        cursor = mysql2.connection.cursor()
                        cursor.execute(update_query, (comments, None, ticket_id))
                        mysql2.connection.commit()

                        # Getting the Ticket ID based on Account Number(Account Number is getting for the Privious Page)
                        # select_query = "SELECT [ticket_id] FROM [dbo].[offline_collection] WHERE [AccountNumber] = ?"
                        # cursor.execute(select_query, (ticket_id,))
                        # result = cursor.fetchone()
                        # Tid = result[0]

                        # Getting the User ID based on User Email id(Email id is getting from the Session)
                        select_query = "SELECT [id] FROM [dbo].[user] WHERE [EmailId] = ?"
                        cursor.execute(select_query, (ros_email_updated,))
                        result1 = cursor.fetchone()
                        id = result1[0]

                        print(' ID: ',id)
                        print('TID:',ticket_id)

                        # Update the User ID in allocated_ticket & ROSCasesTicket in the offline_tickets table based on Ticket ID
                        update_query1 = "UPDATE [dbo].[offline_scenarios] SET [allocatedTicket] = ?, [ROSCasesTicket] = ? WHERE [ticket_id] = ?"
                        cursor.execute(update_query1, (id, id, ticket_id))
                        mysql2.connection.commit()

                        cursor.close()
                        print(comments)
                        print(ticket_id)
                        session['success_message'] = 'ROS Commented successfully.'
                        return redirect(url_for("ros"))
                    else:
                        # print('Approved:',approval_status)  
                        # update the ROS Comment based on Account Number(Account Number is getting for the Privious Page)
                        update_query = "UPDATE [dbo].[offline_scenarios] SET ROS_cmt = ?, DGM_cmt = ? WHERE [ticket_id] = ?"
                        cursor = mysql2.connection.cursor()
                        cursor.execute(update_query, (comments, None, ticket_id))
                        mysql2.connection.commit()

                        # Getting the Ticket ID based on Account Number(Account Number is getting for the Privious Page)
                        # select_query = "SELECT [ticket_id] FROM [dbo].[offline_collection] WHERE [AccountNumber] = ?"
                        # cursor.execute(select_query, (ticket_id,))
                        # result = cursor.fetchone()
                        # Tid = result[0]

                        # Getting the User ID based on User Email id(Email id is getting for the Session)
                        select_query = "SELECT [id] FROM [dbo].[user] WHERE [EmailId] = ?"
                        cursor.execute(select_query, (ros_email_updated,))
                        result1 = cursor.fetchone()
                        id = result1[0]

                        print(' ID: ',id)
                        print('TID:',ticket_id)

                        # Update the User ID in allocated_ticket & ROSCasesTicket in the offline_tickets table based on Ticket ID
                        update_query1 = "UPDATE [dbo].[offline_scenarios] SET [allocatedTicket] = ?, [ROSCasesTicket] = ? WHERE [ticket_id] = ?"
                        cursor.execute(update_query1, (id, id, ticket_id))
                        mysql2.connection.commit()

                        cursor.close()
                        print(comments)
                        print(ticket_id)
                        # session['success_message'] = 'ROS Commented successfully.'
                        return redirect(url_for("archived"))
                else:
                    update_archive_cmt = "UPDATE [dbo].[offline_scenarios] SET [ROS_cmt] = ?, [DGM_cmt] = ?  WHERE [ticket_id] = ?"
                    cursor.execute(update_archive_cmt,(None, 'archive', ticket_id))
                    conn.commit()

                    # Getting the User ID based on User Email id(Email id is getting for the Session)
                    select_query = "SELECT [id] FROM [dbo].[user] WHERE [EmailId] = ?"
                    cursor.execute(select_query, (ros_email_updated,))
                    result1 = cursor.fetchone()
                    id = result1[0]

                    print(' ID: ',id)
                    print('TID:',ticket_id)

                    # Update the User ID in allocated_ticket & ROSCasesTicket in the offline_tickets table based on Ticket ID
                    update_query1 = "UPDATE [dbo].[offline_scenarios] SET [allocatedTicket] = ? WHERE [ticket_id] = ?"
                    cursor.execute(update_query1, (id, ticket_id))
                    mysql2.connection.commit()

                    cursor.close()
                    print(comments)
                    print(ticket_id)
                    # session['success_message'] = 'Successfully Archived in archives.'
                    return redirect(url_for("archived"))
                
            elif result[0] == 'DGM/PO':
                if comments:
                    select_check_query = "SELECT [ROS_cmt], [DGM_cmt] FROM [dbo].[offline_scenarios] WHERE [ticket_id] = ?"
                    cursor.execute(select_check_query, (ticket_id,))
                    result = cursor.fetchone()
                    print('result:',result[0])
                    print('result:',result[1])
                    if result[1] == None:
                        if approval_status == 'Approved':
                            # update the DGM/PO Comment based on Account Number(Account Number is getting for the Privious Page)
                            update_query = "UPDATE [dbo].[offline_scenarios] SET DGM_cmt = ? WHERE [ticket_id] = ?"
                            cursor = mysql2.connection.cursor()
                            cursor.execute(update_query, (comments, ticket_id))
                            mysql2.connection.commit()

                            # Getting the Ticket ID based on Account Number(Account Number is getting for the Privious Page)
                            # select_query = "SELECT [ticket_id] FROM [dbo].[offline_collection] WHERE [ticket_id] = ?"
                            # cursor.execute(select_query, (ticket_id,))
                            # result = cursor.fetchone()
                            # Tid = result[0]

                            # Getting the User ID based on Role(Because here i have only one DGM so no need check the email id)
                            select_query = "SELECT [id] FROM [dbo].[user] WHERE [Role] = ?"
                            cursor.execute(select_query, ('DGM/PO',))
                            result1 = cursor.fetchone()
                            id = result1[0]

                            print(' ID: ',id)
                            print('TID:',ticket_id)

                            # 0 represents false and 1 represents true.    
                            # Update the User ID in allocated_ticket & DGMCasesTicket in the offline_tickets table based on Ticket ID
                            update_query1 = "UPDATE [dbo].[offline_scenarios] SET [allocatedTicket] = ?, [DGMCasesTicket] = ?, [rejected] = ?, [approved] = ? WHERE [ticket_id] = ?"
                            cursor.execute(update_query1, (id, id, 0, 1, ticket_id))
                            mysql2.connection.commit()

                            cursor.close()
                            print(comments)
                            print(ticket_id)
                            session['success_message'] = 'DGM Commented successfully.'
                            return redirect(url_for("offline_dgm_Str"))
                        elif approval_status == 'Rejected':
                            # update the DGM/PO Comment based on Account Number(Account Number is getting for the Privious Page)
                            update_query = "UPDATE [dbo].[offline_scenarios] SET DGM_cmt = ? WHERE [ticket_id] = ?"
                            cursor = mysql2.connection.cursor()
                            cursor.execute(update_query, (comments, ticket_id))
                            mysql2.connection.commit()

                            # Getting the User ID based on Role(Because here i have only one DGM so no need check the email id)
                            select_query = "SELECT [id] FROM [dbo].[user] WHERE [Role] = ?"
                            cursor.execute(select_query, ('DGM/PO',))
                            result1 = cursor.fetchone()
                            id = result1[0]

                            print(' ID: ',id)
                            print('TID:',ticket_id)

                            # 0 represents false and 1 represents true. 
                            # Update the User ID in allocated_ticket & DGMCasesTicket in the offline_tickets table based on Ticket ID
                            update_query1 = "UPDATE [dbo].[offline_scenarios] SET [allocatedTicket] = ?, [DGMCasesTicket] = ?, [rejected] = ?, [approved] = ? WHERE [ticket_id] = ?"
                            cursor.execute(update_query1, (id, id, 1, 0, ticket_id))
                            mysql2.connection.commit()

                            cursor.close()
                            print(comments)
                            print(ticket_id)
                                
                            return redirect(url_for("offline_dgm_Str"))
                    else:
                        if approval_status == 'Approved':
                            # update the DGM/PO Comment based on Account Number(Account Number is getting for the Privious Page)
                            update_query = "UPDATE [dbo].[offline_scenarios] SET DGM_cmt = ? WHERE [ticket_id] = ?"
                            cursor = mysql2.connection.cursor()
                            cursor.execute(update_query, (comments, ticket_id))
                            mysql2.connection.commit()

                            # Getting the Ticket ID based on Account Number(Account Number is getting for the Privious Page)
                            # select_query = "SELECT [ticket_id] FROM [dbo].[offline_collection] WHERE [ticket_id] = ?"
                            # cursor.execute(select_query, (ticket_id,))
                            # result = cursor.fetchone()
                            # Tid = result[0]

                            # Getting the User ID based on Role(Because here i have only one DGM so no need check the email id)
                            select_query = "SELECT [id] FROM [dbo].[user] WHERE [Role] = ?"
                            cursor.execute(select_query, ('DGM/PO',))
                            result1 = cursor.fetchone()
                            id = result1[0]

                            print(' ID: ',id)
                            print('TID:',ticket_id)

                            # 0 represents false and 1 represents true.    
                            # Update the User ID in allocated_ticket & DGMCasesTicket in the offline_tickets table based on Ticket ID
                            update_query1 = "UPDATE [dbo].[offline_scenarios] SET [allocatedTicket] = ?, [DGMCasesTicket] = ?, [rejected] = ?, [approved] = ? WHERE [ticket_id] = ?"
                            cursor.execute(update_query1, (id, id, 0, 1, ticket_id))
                            mysql2.connection.commit()

                            cursor.close()
                            print(comments)
                            print(ticket_id)
                            # session['success_message'] = 'DGM Commented successfully.'
                            return redirect(url_for("archived"))
                        elif approval_status == 'Rejected':
                                # update the DGM/PO Comment based on Account Number(Account Number is getting for the Privious Page)
                                update_query = "UPDATE [dbo].[offline_scenarios] SET DGM_cmt = ? WHERE [ticket_id] = ?"
                                cursor = mysql2.connection.cursor()
                                cursor.execute(update_query, (comments, ticket_id))
                                mysql2.connection.commit()

                                # Getting the User ID based on Role(Because here i have only one DGM so no need check the email id)
                                select_query = "SELECT [id] FROM [dbo].[user] WHERE [Role] = ?"
                                cursor.execute(select_query, ('DGM/PO',))
                                result1 = cursor.fetchone()
                                id = result1[0]

                                print(' ID: ',id)
                                print('TID:',ticket_id)

                                # 0 represents false and 1 represents true. 
                                # Update the User ID in allocated_ticket & DGMCasesTicket in the offline_tickets table based on Ticket ID
                                update_query1 = "UPDATE [dbo].[offline_scenarios] SET [allocatedTicket] = ?, [DGMCasesTicket] = ?, [rejected] = ?, [approved] = ? WHERE [ticket_id] = ?"
                                cursor.execute(update_query1, (id, id, 1, 0, ticket_id))
                                mysql2.connection.commit()

                                cursor.close()
                                print(comments)
                                print(ticket_id)
                                    
                                return redirect(url_for("archived"))
                else:
                    update_archive_cmt = "UPDATE [dbo].[offline_scenarios] SET [DGM_cmt] = ?  WHERE [ticket_id] = ?"
                    cursor.execute(update_archive_cmt,('archive', ticket_id))
                    conn.commit()

                    # Getting the User ID based on Role(Because here i have only one DGM so no need check the email id)
                    select_query = "SELECT [id] FROM [dbo].[user] WHERE [Role] = ?"
                    cursor.execute(select_query, ('DGM/PO',))
                    result1 = cursor.fetchone()
                    id = result1[0]

                    print(' ID: ',id)
                    print('TID:',ticket_id)

                    # 0 represents false and 1 represents true. 
                    # Update the User ID in allocated_ticket & DGMCasesTicket in the offline_tickets table based on Ticket ID
                    update_query1 = "UPDATE [dbo].[offline_scenarios] SET [allocatedTicket] = ? WHERE [ticket_id] = ?"
                    cursor.execute(update_query1, (id, ticket_id))
                    mysql2.connection.commit()

                    # session['success_message'] = 'Successfully Archived in archives.'
                    return redirect(url_for("archived"))
            else:
                return "Unable to find the Details..."
        except Exception as e:
            # Print or log the error message
            print("Error:", e)
            # Rollback the transaction if an error occurs
            mysql2.connection.rollback()
            # Close the cursor
            # cursor.close()
            # Redirect or return an error message
            return 'An error occurred while updating the data'
    else:
        return 'Method Not Allowed', 405


# ============================== ALL FORM PAGES SUBMITE THE DATA END - POINT CODE END's HERE ===================================================


# ------------------------------------------------------------------------------------------------------------------------------------------------------------



# ============================== OFFLINE ARCHIVE ROW  END - POINT CODE START's HERE ===================================================


# @app.route('/archived',methods=['GET','POST'])
# @secure_route(required_role=['ROS','DGM/PO','BranchMakers'])
# def archived():
#     conn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
#     cursor = conn.cursor()
#     emailId = session['email_id']

#     # notify = notification(emailId)

#     # Retrieve user data from the MS SQL database
#     select_query = "SELECT * FROM [user] WHERE [EmailId] = ?"
#     cursor.execute(select_query, (emailId,))
#     user = cursor.fetchone()
#     print(user)
#     print(user[6])
#     # Check if user data is found
#     if user is None:
#         return "User data not found. Please log in again."

#     # Check if 'image' is present in user and encode it as base64
#     if 'image' in user:
#         # Encode the image data as a base64 string
#         user['image'] = base64.b64encode(user['image']).decode('utf-8')

#     offlineArc = 'archive'

#     communuser = {}
#     if 'image' in user:
#         communuser['image'] = base64.b64encode(user['image']).decode('utf-8')

#     for cols in user:
#         if cols == 'DGM/PO':
#             dashBoardType = 'DGMdashboard'
#             select_query_dgm = "SELECT * FROM [dbo].[offline_scenarios] WHERE [DGM_cmt] = 'archive' and [ROS_cmt] != 'archive'"
#             cursor.execute(select_query_dgm)
#             info = cursor.fetchall()
#             print(info)
#             return render_template('archivedPage.html', data=info, allImages=communuser, role='DGM/PO', type='archived', dashBoardType=dashBoardType)
#         elif cols == 'ROS':
#             dashBoardType = 'ros'
#             select_query1 = "SELECT [BranchCode] FROM [user] WHERE [EmailId] = ?"
#             cursor.execute(select_query1, (emailId,))
#             branch_code = cursor.fetchone()
#             B_code = branch_code[0]
#             # print(branch_code[0])

#             # Getting the Email Id of Branchmaker based on Branch Code and Role
#             select_query2 = "SELECT [EmailId] FROM [user] WHERE [BranchCode] = ? AND [Role] = 'BranchMakers'"
#             cursor.execute(select_query2, (B_code,))
#             ros_email_id = cursor.fetchall()
#             # print(ros_email_id)

#             # In the above method will give us tuple of Email Id so we have to convert it into list
#             for tup in ros_email_id:
#                 B_email = tup
#             print(B_email[0])
#             createdby = B_email[0]
            
#             # Getting the data from offline_collection table based on Created By
#             select_query3 = "SELECT * FROM [dbo].[offline_scenarios] WHERE [Created_By] = ? AND [DGM_cmt] = 'archive' AND [ROS_cmt] IS NULL"
#             cursor.execute(select_query3, (createdby,))
#             info = cursor.fetchall()
#             print(info)

#             return render_template('archivedPage.html', data=info, allImages=communuser, role='ROS', type='archived', dashBoardType=dashBoardType)
#         elif cols == 'BranchMakers':
#             dashBoardType = 'branchmakers'
#             # Query to retrieve archived data based on archived ticket numbers
#             select_query1 = "SELECT * FROM [dbo].[offline_scenarios] WHERE [Created_By] = ? AND [ROS_cmt] = ? AND [DGM_cmt] = ?"
#             cursor.execute(select_query1, (emailId, offlineArc, offlineArc))
#             info = cursor.fetchall()
#             print(info)
#             return render_template('archivedPage.html', data=info, allImages=communuser, role='BranchMakers', type='archived', dashBoardType=dashBoardType)


@app.route('/archived',methods=['GET','POST'])
@secure_route(required_role=['ROS','DGM/PO','BranchMakers'])
def archived():
    conn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")
    cursor = conn.cursor()
    emailId = session['email_id']

    # notify = notification(emailId)

    # Retrieve user data from the MS SQL database
    select_query = "SELECT * FROM [dbo].[user] WHERE [EmailId] = ?"
    cursor.execute(select_query, (emailId,))
    user = cursor.fetchone()
    print(user)
    print(user[6])
    # Check if user data is found
    if user is None:
        return "User data not found. Please log in again."

    # Check if 'image' is present in user and encode it as base64
    if 'image' in user:
        # Encode the image data as a base64 string
        user['image'] = base64.b64encode(user['image']).decode('utf-8')

    offlineArc = 'archive'

    communuser = {}
    if 'image' in user:
        communuser['image'] = base64.b64encode(user['image']).decode('utf-8')

    for cols in user:
        if cols == 'DGM/PO':
            dashBoardType = 'DGMdashboard'
            # Here we are try to the data arrengment for DGM 
            # Getting the user id based on Email Id(Email Id stored in Session)
            # select_query1 = "SELECT id from [dbo].[user] where [EmailId] = ?"
            # cursor.execute(select_query1, (emailId,))
            # branch_code = cursor.fetchone()
            # B_id = branch_code[0]
            # # print(branch_code[0])

            # # Getting the ticket id in the offline_tickets table based on the user id
            # select_query2 = "SELECT ticketid FROM [dbo].[offline_tickets] where DGMCasesTicket = ?"
            # cursor.execute(select_query2, (B_id,))
            # ros_email_id = cursor.fetchall()
            # # print(ros_email_id)

            # # checking the data like if data is present or not
            # if ros_email_id:
            #     T_id = []
            #     # In the above method will give us tuple of Email Id so we have to convert it into list
            #     for tup in ros_email_id:
            #         B_email = tup
            #         T_id.append(B_email[0])
            # info = []
            # for i in T_id:
            #     select_query3 = "SELECT * FROM [dbo].[offline_collection] where ticket_id = ? and [DGM_cmt] = 'archive'"
            #     cursor.execute(select_query3, (i,))
            #     info.append(cursor.fetchall())
            #     print(info)
            select_query_dgm = "SELECT * FROM [dbo].[offline_scenarios] WHERE [DGM_cmt] = 'archive' and [ROS_cmt] != 'archive'"
            try:
                cursor.execute(select_query_dgm)
                info = cursor.fetchall()
                print(info)
                return render_template('archivedPage.html', data=info, allImages=communuser, role='DGM/PO', type='archived', dashBoardType=dashBoardType)
            except pyodbc.Error as e:
                print(f"Database error occurred: {e}")
                return render_template( 'archivedPage.html', data=[], allImages=communuser, role='DGM/PO', type='archived', dashBoardType=dashBoardType)
            
            # return render_template('archivedPage.html', data=info, allImages=communuser, role='DGM/PO', type='archived', dashBoardType=dashBoardType)
        elif cols == 'ROS':
            dashBoardType = 'ros'
            select_query1 = "SELECT [BranchCode] FROM [user] WHERE [EmailId] = ?"
            cursor.execute(select_query1, (emailId,))
            branch_code = cursor.fetchone()
            B_code = branch_code[0]
            # print(branch_code[0])

            # Getting the Email Id of Branchmaker based on Branch Code and Role
            select_query2 = "SELECT [EmailId] FROM [user] WHERE [BranchCode] = ? AND [Role] = 'BranchMakers'"
            cursor.execute(select_query2, (B_code,))
            ros_email_id = cursor.fetchall()
            # print(ros_email_id)

            # In the above method will give us tuple of Email Id so we have to convert it into list
            for tup in ros_email_id:
                B_email = tup
            print(B_email[0])
            createdby = B_email[0]
            select_query3 = "SELECT * FROM [dbo].[offline_scenarios] WHERE [Created_By] = ? AND [DGM_cmt] = 'archive' AND [ROS_cmt] IS NULL"
            try:
                cursor.execute(select_query3, (createdby,))
                info = cursor.fetchall()
                print(info)
                return render_template('archivedPage.html', data=info, allImages=communuser, role='ROS', type='archived', dashBoardType=dashBoardType)
            except pyodbc.Error as e:
                print(f"Database error occurred: {e}")
                return render_template('archivedPage.html', data=[], allImages=communuser, role='ROS', type='archived', dashBoardType=dashBoardType)
            
        elif cols == 'BranchMakers':
            dashBoardType = 'branchmakers'
            select_query1 = "SELECT * FROM [dbo].[offline_scenarios] WHERE [Created_By] = ? AND [ROS_cmt] = ? AND [DGM_cmt] = ?"
            try:
                cursor.execute(select_query1, (emailId, offlineArc, offlineArc))
                info = cursor.fetchall()
                print(info)
                return render_template('archivedPage.html', data=info, allImages=communuser, role='BranchMakers', type='archived', dashBoardType=dashBoardType)
            except pyodbc.Error as e:
                print(f"Database error occurred: {e}")
                return render_template('archivedPage.html', data=[], allImages=communuser, role='BranchMakers', type='archived', dashBoardType=dashBoardType)
            

# ============================== OFFLINE ARCHIVE ROW  END - POINT CODE END's HERE ===================================================



# --------------------------------------------------------------------------------------------------------------------------------------------------



# =============================== ACCOUNT HISTORY PAGE DISPLAY END - POINT START's HERE ======================================================



@app.route('/acc_holder_history', methods=['POST', 'GET'])
@secure_route(required_role=['MLRO', 'AGM', 'CM/SM', 'ROS', 'DGM/PO'])
def acc_holder_history():    
    # accConn = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    accConn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlAccHis = accConn.cursor() 

    accnum = request.form.get('accnum')
    custid = request.form.get('custid')
    txdate = request.form.get('txdate')
    scenarioType = request.form.get('scenarioType')
    txtype = request.form.get('txtype')
    trnsflowtype = request.form.get('trnsflowtype')
    cashflowtype = request.form.get('cashflowtype')
    redirectedURL = request.form.get('typee')

    success_message = session.pop('success_message', None)

    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    
    mlro_email = session['email_id']

    try:

        query = "SELECT * FROM [user] WHERE EmailId = ?"
        mysqlAccHis.execute(query, (mlro_email,))
        
        rows = mysqlAccHis.fetchall()
        columns = [col[0] for col in mysqlAccHis.description]
        mlro = [{columns[i]: row[i] for i in range(len(columns))} for row in rows]
        
        if not mlro:
            return "User data not found. Please log in again."

        mlro = mlro[0]
        if 'image' in mlro:
            mlro['image'] = base64.b64encode(mlro['image']).decode('utf-8')

        role = mlro.get('Role')
        customerData, accountsOfTheCustomer, allTransactionData, allTransactionDataMultiDrMultiCr = [], [], [], []
        scenarioTransactionData, scenarioTransactionDataMultiDrMultiCr, scenarioCodesEffected = [], [], []

        if role in ['MLRO', 'CM/SM', 'DGM/PO']:

            query_scenarios = "SELECT DISTINCT scenario_code FROM scenarios WHERE CUSTCD = ?"
            mysqlAccHis.execute(query_scenarios, (custid,))
            scenarioCodesEffected = [row[0] for row in mysqlAccHis.fetchall()]


            placeholders = ','.join('?' for _ in scenarioCodesEffected)

            scenariosSumariquire = f"SELECT Alert_title FROM Thresholds WHERE code IN ({placeholders})"


            mysqlAccHis.execute(scenariosSumariquire,scenarioCodesEffected)

            scenariosSumari = [row[0] for row in mysqlAccHis.fetchall()]


            query_customer_account_details = "SELECT * FROM CUSTOMERS WHERE CUSTCD = ?"
            mysqlAccHis.execute(query_customer_account_details, (custid,))
            accountsOfTheCustomer = mysqlAccHis.fetchall()

            # query = """
            #     SELECT mc.TXDATE,mc.TXTYPE, t.TRNFLOWTYPE, mc.TXAMT, mc.TXFRCURRCD, mc.TXNNO, mc.ACCTNO AS TRNFACCTNO
            #     FROM TRANSACTIONS t
            #     JOIN TRANSACTIONS mc ON mc.TRNFACCTNO = t.ACCTNO AND mc.TXNNO = t.TXNNO AND mc.TXDATE = t.TXDATE
            #     WHERE t.ACCTNO = ?
            #     UNION
            #     SELECT TXDATE,TXTYPE, TRNFLOWTYPE, TXAMT, TXFRCURRCD, TXNNO, TRNFACCTNO
            #     FROM TRANSACTIONS
            #     WHERE ACCTNO = ? AND TRNFACCTNO IS NOT NULL AND TRNFACCTNO != 'M'
            # """
            # cur.execute(query, (accnum, accnum))
            # allTransactionData = cur.fetchall()

            # multiDrmultiCRquery = """
            #     WITH UniqueTransactions AS (
            #         SELECT DISTINCT mc.TXDATE, mc.TXNNO
            #         FROM TRANSACTIONS mc
            #         WHERE mc.ACCTNO = ? AND mc.TRNFACCTNO = 'M'
            #     )
            #     SELECT t.ACCTNO, t.TXDATE, t.TXTYPE, t.TRNFLOWTYPE, t.TXAMT, t.TXFRCURRCD, t.TXNNO
            #     FROM TRANSACTIONS t
            #     JOIN UniqueTransactions ut ON t.TXDATE = ut.TXDATE AND t.TXNNO = ut.TXNNO
            #     WHERE t.TRNFACCTNO = 'M'
            #     ORDER BY t.TXNNO ASC
            # """
            # cur.execute(multiDrmultiCRquery, (accnum,))
            # allTransactionDataMultiDrMultiCr = cur.fetchall()

            try:
                divDate = txdate.split(' to ')
                presentDate1 = datetime.strptime(divDate[1], '%Y-%m-%d').strftime('%Y-%m-%d')
                pastDate1 = datetime.strptime(divDate[0], '%Y-%m-%d').strftime('%Y-%m-%d')


                query1 = """
                    SELECT mc.TXDATE,mc.TXTYPE, t.TRNFLOWTYPE, mc.TXAMT, mc.TXFRCURRCD, mc.TXNNO, mc.ACCTNO AS TRNFACCTNO
                    FROM TRANSACTIONS t
                    JOIN TRANSACTIONS mc ON mc.TRNFACCTNO = t.ACCTNO AND mc.TXNNO = t.TXNNO AND mc.TXDATE = t.TXDATE
                    WHERE t.ACCTNO = ? AND t.TXTYPE = ? AND t.TRNFLOWTYPE = ? AND CONVERT(DATE, t.TXDATE, 105) BETWEEN ? AND ?
                    UNION
                    SELECT TXDATE,TXTYPE, TRNFLOWTYPE, TXAMT, TXFRCURRCD, TXNNO, TRNFACCTNO
                    FROM TRANSACTIONS
                    WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TRNFACCTNO IS NOT NULL AND TRNFACCTNO != 'M' AND CONVERT(DATE, TXDATE, 105) BETWEEN ? AND ?
                """
                mysqlAccHis.execute(query1, (accnum,txtype,trnsflowtype,pastDate1, presentDate1, accnum,txtype,trnsflowtype,pastDate1, presentDate1))
                scenarioTransactionData = mysqlAccHis.fetchall()

                scenariosmultiDrmultiCRquery = """
                    WITH UniqueTransactions AS (
                        SELECT DISTINCT mc.TXDATE, mc.TXNNO
                        FROM TRANSACTIONS mc
                        WHERE mc.ACCTNO = ? AND mc.TXTYPE = ? AND mc.TRNFACCTNO = 'M' AND mc.TRNFLOWTYPE = ? AND CONVERT(DATE, mc.TXDATE, 105) BETWEEN ? AND ?
                    )
                    SELECT t.ACCTNO, t.TXDATE,t.TXTYPE, t.TRNFLOWTYPE, t.TXAMT, t.TXFRCURRCD, t.TXNNO
                    FROM TRANSACTIONS t
                    JOIN UniqueTransactions ut ON t.TXDATE = ut.TXDATE AND t.TXNNO = ut.TXNNO
                    WHERE t.TRNFACCTNO = 'M'
                    ORDER BY t.TXNNO ASC
                """
                mysqlAccHis.execute(scenariosmultiDrmultiCRquery, (accnum,txtype,trnsflowtype, pastDate1, presentDate1))
                scenarioTransactionDataMultiDrMultiCr = mysqlAccHis.fetchall()
            except:
                query1 = """
                    SELECT mc.TXDATE,mc.TXTYPE, t.TRNFLOWTYPE, mc.TXAMT, mc.TXFRCURRCD, mc.TXNNO, mc.ACCTNO AS TRNFACCTNO
                    FROM TRANSACTIONS t
                    JOIN TRANSACTIONS mc ON mc.TRNFACCTNO = t.ACCTNO AND mc.TXNNO = t.TXNNO AND mc.TXDATE = t.TXDATE
                    WHERE t.ACCTNO = ? AND t.TXTYPE = ? AND t.TRNFLOWTYPE = ? AND t.TXDATE = ?
                    UNION
                    SELECT TXDATE,TXTYPE, TRNFLOWTYPE, TXAMT, TXFRCURRCD, TXNNO, TRNFACCTNO
                    FROM TRANSACTIONS
                    WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TXDATE = ? AND TRNFACCTNO IS NOT NULL AND TRNFACCTNO != 'M'
                """
                mysqlAccHis.execute(query1, (accnum,txtype,trnsflowtype,txdate, accnum,txtype,trnsflowtype,txdate))
                scenarioTransactionData = mysqlAccHis.fetchall()

                scenariosmultiDrmultiCRquery = """
                    WITH UniqueTransactions AS (
                        SELECT DISTINCT mc.TXDATE, mc.TXNNO
                        FROM TRANSACTIONS mc
                        WHERE mc.ACCTNO = ? AND mc.TXTYPE = ? AND mc.TRNFACCTNO = 'M' AND mc.TRNFLOWTYPE = ? AND mc.TXDATE = ? 
                    )
                    SELECT t.ACCTNO, t.TXDATE,t.TXTYPE, t.TRNFLOWTYPE, t.TXAMT, t.TXFRCURRCD, t.TXNNO
                    FROM TRANSACTIONS t
                    JOIN UniqueTransactions ut ON t.TXDATE = ut.TXDATE AND t.TXNNO = ut.TXNNO
                    WHERE t.TRNFACCTNO = 'M'
                    ORDER BY t.TXNNO ASC
                """
                mysqlAccHis.execute(scenariosmultiDrmultiCRquery, (accnum,txtype,trnsflowtype, txdate))
                scenarioTransactionDataMultiDrMultiCr = mysqlAccHis.fetchall()

        accConn.close()

        return render_template('account_holder_details.html',
                            scenarioTransactionData=scenarioTransactionData,
                            scenarioTransactionDataMultiDrMultiCr=scenarioTransactionDataMultiDrMultiCr,
                            scenarioCodesEffected=scenarioCodesEffected,
                            accountsOfTheCustomer=accountsOfTheCustomer,
                            scenariosSumari=scenariosSumari,
                            accnum=accnum,
                            custid = custid,
                                txdate = txdate,
                                txtype = txtype,
                                trnsflowtype = trnsflowtype,
                                cashflowtype = cashflowtype,
                            scenarioType=scenarioType,
                            role=role,
                            allImages=mlro,
                            type=redirectedURL)
    except Exception as e:

        mysqlAccHis.rollback()
        accConn.close()

        return f"Somthing Went Wrong {e} Please Re-Login Again",500




# =============================== ACCOUNT HISTORY PAGE DISPLAY END - POINT END's HERE ========================================================




@app.route('/offlineactiveCases', methods=['POST', 'GET'])
def offlineactiveCases():
    print('###################')
    data = request.get_json()
    
    ticket_numbers = data.get('tickets', [])
    print('data:', data)

    # Create a temporary directory
    temp_dir = tempfile.mkdtemp()

    # Create the individual files
    pdf_filename = os.path.join(temp_dir, 'AlltableData.pdf')
    pdf_content = create_pdf1(data, pdf_filename)

    csv_filename = os.path.join(temp_dir, 'AlltableData.csv')
    csv_content = create_csv1(data, csv_filename)

    text_filename = os.path.join(temp_dir, 'AlltableData.txt')
    text_content = create_text1(data, text_filename)

    word_filename = os.path.join(temp_dir, 'AlltableData.docx')
    create_word1(data, word_filename)

    # Create a zip file
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.write(pdf_filename, arcname='AlltableData.pdf')
        zip_file.write(csv_filename, arcname='AlltableData.csv')
        zip_file.write(text_filename, arcname='AlltableData.txt')
        zip_file.write(word_filename, arcname='AlltableData.docx')

    # Close and delete the temporary directory
    shutil.rmtree(temp_dir)

    # Move the pointer to the beginning of the buffer
    zip_buffer.seek(0)

    # Send the zip file as an attachment
    return send_file(zip_buffer, as_attachment=True, mimetype='application/zip', download_name='AlltableData.zip')

def create_pdf1(data, filename):
       
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    list_of_dicts= data['tickets']
    # Create table data
    table_data = []
    for obj in list_of_dicts:
        for key, value in obj.items():
            wrapper = textwrap.TextWrapper(width=70)
            wrapped_lines = wrapper.wrap(text=str(value))
            wrapped_text = "\n".join(wrapped_lines)
            table_data.append([key, wrapped_text])

    # Create table
    table = Table(table_data, colWidths=[200, 400],  # Adjust the column widths as needed
                  style=[
                      ('GRID', (0, 0), (-1, -1), 1, colors.black),  # Add grid lines
                      ('BACKGROUND', (0, 0), (-1, 0), colors.white),  # Add background color to header row
                      ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),  # Set text color in header row
                      ('ALIGN', (0, 0), (-1, -1), 'LEFT'),  # Center align all cells
                      ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Middle align all cells
                      ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),  # Inner grid lines
                      ('BOX', (0, 0), (-1, -1), 0.25, colors.black),  # Add box around the table
                  ])

    # Build the story

    all_names = ["MLRO", "CM", "GM/PO \n (Approved)"]
    names_table_data = [[all_names[0], all_names[1], all_names[2]]]
    names_table = Table(names_table_data, colWidths=[180, 180, 180],
                        style=[
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                            ('ALIGN', (2, 0), (-1, -1), 'RIGHT')

                        ])

    spacer = Spacer(1, 70)

    doc.build([table, spacer, names_table])
    buffer.seek(0)
    with open(filename, 'wb') as f:
        f.write(buffer.read())

def create_csv1(data, filename):
    df = pd.DataFrame(data['tickets'])
    
    # Write DataFrame to CSV file
    df.to_csv(filename, index=False)

    # Alternatively, if you want to return CSV content as bytes
    csv_buffer = StringIO()  # Create a StringIO object to hold CSV content
    df.to_csv(csv_buffer, index=False)
    
    # Get the CSV content as bytes
    csv_content = csv_buffer.getvalue().encode('utf-8')
    return csv_content

def create_text1(data, filename):
    ticket_numbers = data.get('tickets', [])
    if not ticket_numbers:
        return b''  # Return an empty bytes object if ticket_numbers is empty
    
    # Create text content with each key-value pair separated by commas for all ticket numbers
    text_content = ''
    for ticket_data in ticket_numbers:
        text_content += ','.join([f"{key}: {value}" for key, value in ticket_data.items()])
        text_content += '\n'  # Add new line between ticket details
    
    # Encode the text content to bytes
    text_bytes = text_content.encode('utf-8')
    
    # Write text content to file
    with open(filename, 'wb') as text_file:
        text_file.write(text_bytes)
    
    return text_bytes

def create_word1(data, filename):
    ticket_numbers = data.get('tickets', [])
    if not ticket_numbers:
        return b''  # Return an empty bytes object if ticket_numbers is empty
    
    # Create a new Document object
    doc = Document()
    
    # Add a table with two columns for key-value pairs
    table = doc.add_table(rows=1, cols=2)
    
    # Iterate over each ticket and add its key-value pairs to the table
    for ticket in ticket_numbers:
        for key, value in ticket.items():
            # Add key-value pairs to the table
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
        
        # Add an empty row between tickets
        table.add_row().cells
    
    # Save the document
    doc.save(filename)



@app.route('/offlineactiveCases2', methods=['POST', 'GET'])
def offlineactiveCases2():
    data = request.get_json()
    print('dataaaa:',data)
    ticket_numbers = data.get('tickets', [])

    # Create a temporary directory
    temp_dir = tempfile.mkdtemp()

    try:
        # Create the individual files
        pdf_filename = os.path.join(temp_dir, 'AlltableData.pdf')
        create_pdf2(data, pdf_filename)

        csv_filename = os.path.join(temp_dir, 'AlltableData.csv')
        create_csv2(data, csv_filename)

        text_filename = os.path.join(temp_dir, 'AlltableData.txt')
        create_text2(data, text_filename)

        word_filename = os.path.join(temp_dir, 'AlltableData.docx')
        create_word2(data, word_filename)

        # Create a zip file
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            zip_file.write(pdf_filename, arcname='AlltableData.pdf')
            zip_file.write(csv_filename, arcname='AlltableData.csv')
            zip_file.write(text_filename, arcname='AlltableData.txt')
            zip_file.write(word_filename, arcname='AlltableData.docx')

        # Move the pointer to the beginning of the buffer
        zip_buffer.seek(0)

        # Send the zip file as an attachment
        return send_file(zip_buffer, as_attachment=True, mimetype='application/zip', download_name='AlltableData.zip')
    finally:
        # Close and delete the temporary directory
        shutil.rmtree(temp_dir)

def create_pdf2(data, filename):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    
    # Extract list of tickets from the 'data' dictionary
    tickets = data.get('tickets', [])
    
    # Create table data
    table_data = []
    for ticket_list in tickets:
        for ticket_data in ticket_list:
            ticket_table_data = []
            for key, value in ticket_data.items():
                ticket_table_data.append([key, value])
            table_data.append(ticket_table_data)

    # Check if table_data is empty
    if not table_data:
        return None  # Exit function if there's no data
    
    # Build tables
    tables = []
    for ticket_table_data in table_data:
        # Create table
        table = Table(ticket_table_data, colWidths=[200, 400],  # Adjust the column widths as needed
                      style=[
                          ('GRID', (0, 0), (-1, -1), 1, colors.black),  # Add grid lines
                          ('BACKGROUND', (0, 0), (-1, 0), colors.white),  # Add background color to header row
                          ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),  # Set text color in header row
                          ('ALIGN', (0, 0), (-1, -1), 'LEFT'),  # Center align all cells
                          ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Middle align all cells
                          ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),  # Inner grid lines
                          ('BOX', (0, 0), (-1, -1), 0.25, colors.black),  # Add box around the table
                      ])
        tables.append(table)

    # Build the story
    all_names = ["MLRO", "CM", "GM/PO \n (Approved)"]
    names_table_data = [[all_names[0], all_names[1], all_names[2]]]
    names_table = Table(names_table_data, colWidths=[180, 180, 180],
                        style=[
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                            ('ALIGN', (2, 0), (-1, -1), 'RIGHT')
                        ])

    spacer = Spacer(1, 70)

    story = [spacer, names_table]
    story.extend(tables)
    
    doc.build(story)
    buffer.seek(0)
    with open(filename, 'wb') as f:
        f.write(buffer.read())

    return filename

def create_csv2(data, filename):
    tickets = data.get('tickets', [])
    if not tickets:
        return b''  # Return an empty bytes object if there are no tickets
    
    with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        # Write column headers
        writer.writerow(tickets[0][0].keys())
        # Write ticket data
        for ticket_list in tickets:
            for ticket_data in ticket_list:
                writer.writerow(ticket_data.values())
    return filename

def create_text2(data, filename):
    tickets = data.get('tickets', [])
    if not tickets:
        return b''  # Return an empty bytes object if there are no tickets
    
    with open(filename, 'w', encoding='utf-8') as text_file:
        for ticket_list in tickets:
            for ticket_data in ticket_list:
                text_file.write(','.join([f"{key}: {value}" for key, value in ticket_data.items()]) + '\n')
    return filename

def create_word2(data, filename):
    tickets = data.get('tickets', [])
    if not tickets:
        return b''  # Return an empty bytes object if there are no tickets
    
    doc = Document()
    for ticket_list in tickets:
        for ticket_data in ticket_list:
            table = doc.add_table(rows=len(ticket_data), cols=2)
            for row_index, (key, value) in enumerate(ticket_data.items()):
                table.cell(row_index, 0).text = key
                table.cell(row_index, 1).text = str(value)
            doc.add_paragraph()  # Add empty paragraph between tickets
    doc.save(filename)
    return filename





# ==================================== 10% of Closed data Auto Distribution Functions======================================================


def verify_tenpercent_data():
    current_datetime = datetime.now()
    current_date = current_datetime.date()
    midnight_datetime = datetime.combine(current_date, datetime.min.time())

    # connverifyTen = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    connverifyTen = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlverifyTen = connverifyTen.cursor()
    try:
        query = "SELECT id FROM [user] WHERE Role = 'CM/SM'"
        mysqlverifyTen.execute(query)
        rows = mysqlverifyTen.fetchone()
        

        cmId = rows[0]

    

        res = mysqlverifyTen.execute("SELECT ticketid FROM scenarios WHERE  mlroClosedTicket IS NOT NULL AND mlroCasesTicket IS NULL AND ten_percent_ticket IS NULL AND currentDate = ? ",(str(midnight_datetime),)).fetchall()
        accounts = [doc for doc in res]

        print("accounts : ",accounts)


        if accounts:


            total_accounts = len(accounts)
            print(total_accounts)
            ten_percent = max(1, int(total_accounts * 0.1)) 
            print(ten_percent) 
            data = random.sample(accounts, ten_percent)
            print(data)

            
            for row in data:
                
                updateQuery = "UPDATE scenarios SET ten_percent_ticket = ?,mlroClosedTicket = NULL WHERE ticketid = ?"

                mysqlverifyTen.execute(updateQuery, (cmId, row[0]))
                mysqlverifyTen.commit()
            connverifyTen.close()
            print("10% Closed Alerts where Distributed......")
    except Exception as e:
        mysqlverifyTen.rollback()
        connverifyTen.close()

        print(e)
        print("10% not done..........")
      





# ====================== rised  10% closed data send back to MLRO form CM/SM to create a form with comment=============================








# ============== profile section ===============

@app.route('/profile', methods=['GET', 'POST'])
@secure_route(required_role='MLRO')
def profile():
     email = session['email_id']

     cur = mysql2.connection.cursor()
    #  user = users_collection.find_one({'emailid':email})
     userr = cur.execute("SELECT * FROM [user] WHERE EmailId = ?", (email,))
     user = cur.fetchone()
     return render_template('profileMLRO.html',mlrouser=user,role='MLRO',type='profile')


@app.route('/profileIT', methods=['GET', 'POST'])
@secure_route(required_role='IT OFFICER')
def profileIT():
     email = session['email_id']
     cur = mysql2.connection.cursor()
    #  user = users_collection.find_one({'emailid':email})
     userr = cur.execute("SELECT * FROM [user] WHERE EmailId = ?", (email,))
     user = cur.fetchone()
     
     return render_template('profileIT.html',ituser=user,role='IT OFFICER',type='profileIT')


@app.route('/profileDGM', methods=['GET', 'POST'])
@secure_route(required_role='DGM/PO')
def profileDGM():
     email = session['email_id']

     cur = mysql2.connection.cursor()
    #  user = users_collection.find_one({'emailid':email})
     userr = cur.execute("SELECT * FROM [user] WHERE EmailId = ?", (email,))
     user = cur.fetchone()
     return render_template('profileDGM.html',dgmuser=user,role='DGM/PO',type='profileDGM')

@app.route('/profile_CM_SM', methods=['GET', 'POST'])
@secure_route(required_role='CM/SM')
def profile_CM_SM():
     email = session['email_id']

     cur = mysql2.connection.cursor()
    #  user = users_collection.find_one({'emailid':email})
     userr = cur.execute("SELECT * FROM [user] WHERE EmailId = ?", (email,))
     user = cur.fetchone()
     return render_template('profile_CM_SM.html',cmuser=user,role='CM/SM',type='profile_CM_SM')



@app.route('/profileBranchMaker', methods=['GET', 'POST'])
@secure_route(required_role='BranchMakers')
def profileBranchMaker():
    email = session['email_id']
    # print(email)
    cur = mysql2.connection.cursor()
    #  user = users_collection.find_one({'emailid':email})
    user = cur.execute("SELECT * FROM [user] WHERE EmailId = ?", (email,))
    user = cur.fetchone()
    return render_template('profileBranchMaker.html',branchmakeruser=user,role='BranchMakers',type='profileBranchMaker')



@app.route('/profileROS', methods=['GET', 'POST'])
@secure_route(required_role='ROS')
def profileROS():
    email = session['email_id']
    # print(email)
    cur = mysql2.connection.cursor()
    #  user = users_collection.find_one({'emailid':email})
    user = cur.execute("SELECT * FROM [user] WHERE EmailId = ?", (email,))
    user = cur.fetchone()
    return render_template('profileROS.html',rosuser=user,role='ROS',type='profileROS')

@app.route('/SDNprofile', methods=['GET', 'POST'])
@secure_route(required_role='PINACA_ADMIN')
def SDNprofile():
    email = session['email_id']
    # print(email)
    cur = mysql2.connection.cursor()
    #  user = users_collection.find_one({'emailid':email})
    user = cur.execute("SELECT * FROM [user] WHERE EmailId = ?", (email,))
    user = cur.fetchone()
    return render_template('SDN_User.html',sdnuser=user,role='PINACA_ADMIN',type='SDNprofile')




#-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# ROS Pending Cases part







@app.route('/AllTransactionDetails', methods=['POST', 'GET'])
@secure_route(required_role=['MLRO','AGM','CM/SM','ROS','DGM/PO'])
def AllTransactionDetails():
    accnum = request.form.get('accnum')
    custid = request.form.get('custid')
    txdate = request.form.get('txdate')

    cur = mysql2.connection.cursor()
    data_scenarios=[]
    query1 = "SELECT * FROM merged_collection WHERE CUSTCD = ?  "
    cur.execute(query1, str(custid))
    columns_scenarios = [desc[0] for desc in cur.description]
    unique_txnno1 = set()
    for row in cur.fetchall():
        scenario_object = {}
        for i, value in enumerate(row): 
            scenario_object[columns_scenarios[i]] = value
        if scenario_object['TXNNO'] not in unique_txnno1:
            data_scenarios.append(scenario_object)
            unique_txnno1.add(scenario_object['TXNNO'])

    temp_dir = tempfile.mkdtemp()

    # Create the individual files
    pdf_filename = os.path.join(temp_dir, 'Transactiondata.pdf')
    create_pdf4({'ticket_numbers': data_scenarios}, pdf_filename)

    csv_filename = os.path.join(temp_dir, 'Transactiondata.csv')
    create_csv4({'ticket_numbers': data_scenarios}, csv_filename)

    text_filename = os.path.join(temp_dir, 'Transactiondata.txt')
    create_text4({'ticket_numbers': data_scenarios}, text_filename)

    word_filename = os.path.join(temp_dir, 'Transactiondata.docx')
    create_word4({'ticket_numbers': data_scenarios}, word_filename)

    # Create a zip file
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.write(pdf_filename, arcname='Transactiondata.pdf')
        zip_file.write(csv_filename, arcname='Transactiondata.csv')
        zip_file.write(text_filename, arcname='Transactiondata.txt')
        zip_file.write(word_filename, arcname='Transactiondata.docx')

    # Close and delete the temporary directory
    shutil.rmtree(temp_dir)

    # Move the pointer to the beginning of the buffer
    zip_buffer.seek(0)

    # Send the zip file as an attachment
    return send_file(zip_buffer, as_attachment=True, mimetype='application/zip', download_name='Transactiondata.zip')

def create_pdf4(data, filename):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    list_of_dicts = data['ticket_numbers']
    
    # Extract specific keys from the data
    keys = ['TXDATE', 'TXTYPE','TXAMT', 'TXFRCURRCD','TXNNO', 'TRNFACCTNO', 'BENEFICIARYNAME', 'BENEFIADDRESS']
    
    # Create table data with headers
    table_data = [[' ', ' ']]
    
    # Generate rows for each object in list_of_dicts
    for obj in list_of_dicts:
        for key in keys:
            value = obj.get(key, '')
            table_data.append([key, value])

    # Calculate the number of columns
    num_cols = len(keys)
    
    # Initialize column widths with minimum width
    colWidths = [100, 300]  # Adjust as needed
    
    # Create table
    table = Table(table_data,
                  colWidths=colWidths,
                  style=[
                      ('GRID', (0, 0), (-1, -1), 1, colors.black),  # Add grid lines
                      ('BACKGROUND', (0, 0), (-1, 0), colors.white),  # Add background color to header row
                      ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),  # Set text color in header row
                      ('ALIGN', (0, 0), (-1, -1), 'LEFT'),  # Center align all cells
                      ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Middle align all cells
                      ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),  # Inner grid lines
                      ('BOX', (0, 0), (-1, -1), 0.25, colors.black),  # Add box around the table
                  ])

    # Build the story
    story = []
    story.append(table)
    doc.build(story)
    
    # Move buffer pointer to beginning and write to file
    buffer.seek(0)
    with open(filename, 'wb') as f:
        f.write(buffer.read())

def create_csv4(data, filename):
    df = pd.DataFrame(data['ticket_numbers'])
    selected_fields = ['TXDATE', 'TXTYPE', 'TXAMT', 'TXFRCURRCD', 'TXNNO', 'TRNFACCTNO', 'BENEFICIARYNAME', 'BENEFIADDRESS']
    df_selected = df[selected_fields]
    
    # Write DataFrame to CSV file
    df_selected.to_csv(filename, index=False)

    # Alternatively, if you want to return CSV content as bytes
    csv_buffer = StringIO()  # Create a StringIO object to hold CSV content
    df_selected.to_csv(csv_buffer, index=False)
    
    # Get the CSV content as bytes
    csv_content = csv_buffer.getvalue().encode('utf-8')
    return csv_content

def create_text4(data, filename):
    ticket_numbers = data.get('ticket_numbers', [])
    if not ticket_numbers:
        return b''  # Return an empty bytes object if ticket_numbers is empty
    
    # Create text content with each key-value pair separated by commas for all ticket numbers
    text_content = ''
    selected_fields = ['TXDATE', 'TXTYPE', 'TXAMT', 'TXFRCURRCD', 'TXNNO', 'TRNFACCTNO', 'BENEFICIARYNAME', 'BENEFIADDRESS']
    for ticket_data in ticket_numbers:
        text_content += ','.join([f"{field}: {ticket_data.get(field, '')}" for field in selected_fields])
        text_content += '\n'  # Add new line between ticket details
    
    # Encode the text content to bytes
    text_bytes = text_content.encode('utf-8')
    
    # Write text content to file
    with open(filename, 'wb') as text_file:
        text_file.write(text_bytes)
    
    return text_bytes

def create_word4(data, filename):
    ticket_numbers = data.get('ticket_numbers', [])
    if not ticket_numbers:
        return b''  # Return an empty bytes object if ticket_numbers is empty
    
    # Create a new Document object
    doc = Document()
    
    # Add a table with two columns for key-value pairs
    table = doc.add_table(rows=1, cols=8)
    
    # Add table header
    hdr_cells = table.rows[0].cells
    header = ['TXDATE', 'TXTYPE', 'TXAMT', 'TXFRCURRCD', 'TXNNO', 'TRNFACCTNO', 'BENEFICIARYNAME', 'BENEFIADDRESS']
    for i, field in enumerate(header):
        hdr_cells[i].text = field
    
    # Iterate over each ticket and add its key-value pairs to the table
    for ticket in ticket_numbers:
        row_cells = table.add_row().cells
        for i, field in enumerate(header):
            row_cells[i].text = ticket.get(field, '')
        
        # Add an empty row between tickets
        table.add_row().cells
    
    # Save the document
    doc.save(filename)



@app.route('/TransactionDetails', methods=['POST', 'GET'])
@secure_route(required_role=['MLRO','AGM','CM/SM','ROS','DGM/PO'])
def TransactionDetails():
    accnum = request.form.get('accnum')
    custid = request.form.get('custid')
    txdate = request.form.get('txdate')
    scenarioType = request.form.get('scenarioType')
    txtype = request.form.get('txtype')
    trnsflowtype = request.form.get('trnsflowtype')
    cashflowtype = request.form.get('cashflowtype')
    # accConn = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
    accConn = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

    mysqlAccHis = accConn.cursor() 



    try:
        divDate = txdate.split(' to ')
        presentDate1 = datetime.strptime(divDate[1], '%Y-%m-%d').strftime('%Y-%m-%d')
        pastDate1 = datetime.strptime(divDate[0], '%Y-%m-%d').strftime('%Y-%m-%d')


        query1 = """
            SELECT mc.TXDATE, t.TRNFLOWTYPE, mc.TXAMT, mc.TXFRCURRCD, mc.TXNNO, mc.ACCTNO AS TRNFACCTNO
            FROM TRANSACTIONS t
            JOIN TRANSACTIONS mc ON mc.TRNFACCTNO = t.ACCTNO AND mc.TXNNO = t.TXNNO AND mc.TXDATE = t.TXDATE
            WHERE t.ACCTNO = ? AND t.TXTYPE = ? AND t.TRNFLOWTYPE = ? AND CONVERT(DATE, t.TXDATE, 105) BETWEEN ? AND ?
            UNION
            SELECT TXDATE, TRNFLOWTYPE, TXAMT, TXFRCURRCD, TXNNO, TRNFACCTNO
            FROM TRANSACTIONS
            WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TRNFACCTNO IS NOT NULL AND TRNFACCTNO != 'M' AND CONVERT(DATE, TXDATE, 105) BETWEEN ? AND ?
        """
        mysqlAccHis.execute(query1, (accnum,txtype,trnsflowtype,pastDate1, presentDate1, accnum,txtype,trnsflowtype,pastDate1, presentDate1))
        scenarioTransactionData = mysqlAccHis.fetchall()
        scenarioTransactionData = [(accnum,txtype) + tuple(record) for record in scenarioTransactionData]



        scenariosmultiDrmultiCRquery = """
            WITH UniqueTransactions AS (
                SELECT DISTINCT mc.TXDATE, mc.TXNNO
                FROM TRANSACTIONS mc
                WHERE mc.ACCTNO = ? AND mc.TXTYPE = ? AND mc.TRNFACCTNO = 'M' AND mc.TRNFLOWTYPE = ? AND CONVERT(DATE, mc.TXDATE, 105) BETWEEN ? AND ?
            )
            SELECT t.ACCTNO,t.TXTYPE, t.TXDATE, t.TRNFLOWTYPE, t.TXAMT, t.TXFRCURRCD, t.TXNNO,NULL
            FROM TRANSACTIONS t
            JOIN UniqueTransactions ut ON t.TXDATE = ut.TXDATE AND t.TXNNO = ut.TXNNO
            WHERE t.TRNFACCTNO = 'M'
            ORDER BY t.TXNNO ASC
        """
        mysqlAccHis.execute(scenariosmultiDrmultiCRquery, (accnum,txtype,trnsflowtype, pastDate1, presentDate1))
        scenarioTransactionDataMultiDrMultiCr = mysqlAccHis.fetchall()
    except:
        query1 = """
            SELECT mc.TXDATE, t.TRNFLOWTYPE, mc.TXAMT, mc.TXFRCURRCD, mc.TXNNO, mc.ACCTNO AS TRNFACCTNO
            FROM TRANSACTIONS t
            JOIN TRANSACTIONS mc ON mc.TRNFACCTNO = t.ACCTNO AND mc.TXNNO = t.TXNNO AND mc.TXDATE = t.TXDATE
            WHERE t.ACCTNO = ? AND t.TXTYPE = ? AND t.TRNFLOWTYPE = ? AND t.TXDATE = ?
            UNION
            SELECT TXDATE, TRNFLOWTYPE, TXAMT, TXFRCURRCD, TXNNO, TRNFACCTNO
            FROM TRANSACTIONS
            WHERE ACCTNO = ? AND TXTYPE = ? AND TRNFLOWTYPE = ? AND TXDATE = ? AND TRNFACCTNO IS NOT NULL AND TRNFACCTNO != 'M'
        """
        mysqlAccHis.execute(query1, (accnum,txtype,trnsflowtype,txdate, accnum,txtype,trnsflowtype,txdate))
        scenarioTransactionData = mysqlAccHis.fetchall()
        scenarioTransactionData = [(accnum,txtype) + tuple(record) for record in scenarioTransactionData]


        scenariosmultiDrmultiCRquery = """
            WITH UniqueTransactions AS (
                SELECT DISTINCT mc.TXDATE, mc.TXNNO
                FROM TRANSACTIONS mc
                WHERE mc.ACCTNO = ? AND mc.TXTYPE = ? AND mc.TRNFACCTNO = 'M' AND mc.TRNFLOWTYPE = ? AND mc.TXDATE = ? 
            )
            SELECT t.ACCTNO,t.TXTYPE,t.TXDATE, t.TRNFLOWTYPE, t.TXAMT, t.TXFRCURRCD, t.TXNNO,NULL
            FROM TRANSACTIONS t
            JOIN UniqueTransactions ut ON t.TXDATE = ut.TXDATE AND t.TXNNO = ut.TXNNO
            WHERE t.TRNFACCTNO = 'M'
            ORDER BY t.TXNNO ASC
        """
        mysqlAccHis.execute(scenariosmultiDrmultiCRquery, (accnum,txtype,trnsflowtype, txdate))
        scenarioTransactionDataMultiDrMultiCr = mysqlAccHis.fetchall()


    scenarios_extened = scenarioTransactionData + scenarioTransactionDataMultiDrMultiCr

    
    data_scenarios=[]
    unique_txnno1 = set()


    merged_collection_data = [list(row) for row in scenarios_extened]
    merged_collection_columns = ['ACCTNO','TXTYPE','TXDATE', 'TRNFLOWTYPE', 'TXAMT','TXFRCURRCD','TXNNO','TRNFACCTNO']
    df_merged_collection = pd.DataFrame(merged_collection_data, columns=merged_collection_columns)
    df_merged_collection['ACCTNO'] = "' " + df_merged_collection['ACCTNO'].astype(str)
    df_merged_collection['TRNFACCTNO'] = "' " + df_merged_collection['TRNFACCTNO'].astype(str)
    
    for index, row in df_merged_collection.iterrows():
        scenario_object = {}
        for column in df_merged_collection.columns:
            scenario_object[column] = row[column]
        
        # Check if TXNNO is unique before adding to data_scenarios
        if scenario_object['TXNNO'] not in unique_txnno1:
            data_scenarios.append(scenario_object)
            unique_txnno1.add(scenario_object['TXNNO'])

    temp_dir = tempfile.mkdtemp()

    # Create the individual files
    pdf_filename = os.path.join(temp_dir, 'scenarioTransactionData.pdf')
    create_pdf3({'ticket_numbers': data_scenarios}, pdf_filename)

    csv_filename = os.path.join(temp_dir, 'scenarioTransactionData.csv')
    create_csv3({'ticket_numbers': data_scenarios}, csv_filename)

    text_filename = os.path.join(temp_dir, 'scenarioTransactionData.txt')
    create_text3({'ticket_numbers': data_scenarios}, text_filename)

    word_filename = os.path.join(temp_dir, 'scenarioTransactionData.docx')
    create_word3({'ticket_numbers': data_scenarios}, word_filename)

    # Create a zip file
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.write(pdf_filename, arcname='scenarioTransactionData.pdf')
        zip_file.write(csv_filename, arcname='scenarioTransactionData.csv')
        zip_file.write(text_filename, arcname='scenarioTransactionData.txt')
        zip_file.write(word_filename, arcname='scenarioTransactionData.docx')

    # Close and delete the temporary directory
    shutil.rmtree(temp_dir)

    # Move the pointer to the beginning of the buffer
    zip_buffer.seek(0)

    # Send the zip file as an attachment
    return send_file(zip_buffer, as_attachment=True, mimetype='application/zip', download_name='scenarioTransactionData.zip')


def create_pdf3(data, filename):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    list_of_dicts = data['ticket_numbers']
    
    # Extract specific keys from the data
    keys = ['ACCTNO','TXTYPE','TXDATE', 'TRNFLOWTYPE', 'TXAMT','TXFRCURRCD','TXNNO','TRNFACCTNO']
    
    # Create table data with headers
    table_data = [[' ', ' ']]
    
    # Generate rows for each object in list_of_dicts
    for obj in list_of_dicts:
        for key in keys:
            value = obj.get(key, '')
            table_data.append([key, value])

    # Calculate the number of columns
    num_cols = len(keys)
    
    # Initialize column widths with minimum width
    colWidths = [100, 300]  # Adjust as needed
    
    # Create table
    table = Table(table_data,
                  colWidths=colWidths,
                  style=[
                      ('GRID', (0, 0), (-1, -1), 1, colors.black),  # Add grid lines
                      ('BACKGROUND', (0, 0), (-1, 0), colors.white),  # Add background color to header row
                      ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),  # Set text color in header row
                      ('ALIGN', (0, 0), (-1, -1), 'LEFT'),  # Center align all cells
                      ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Middle align all cells
                      ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),  # Inner grid lines
                      ('BOX', (0, 0), (-1, -1), 0.25, colors.black),  # Add box around the table
                  ])

    # Build the story
    story = []
    story.append(table)
    doc.build(story)
    
    # Move buffer pointer to beginning and write to file
    buffer.seek(0)
    with open(filename, 'wb') as f:
        f.write(buffer.read())

def create_csv3(data, filename):
    df = pd.DataFrame(data['ticket_numbers'])
    selected_fields = ['ACCTNO','TXTYPE','TXDATE', 'TRNFLOWTYPE', 'TXAMT','TXFRCURRCD','TXNNO','TRNFACCTNO']
    df_selected = df[selected_fields]
    
    # Write DataFrame to CSV file
    df_selected.to_csv(filename, index=False)

    # Alternatively, if you want to return CSV content as bytes
    csv_buffer = StringIO()  # Create a StringIO object to hold CSV content
    df_selected.to_csv(csv_buffer, index=False)
    
    # Get the CSV content as bytes
    csv_content = csv_buffer.getvalue().encode('utf-8')
    return csv_content

def create_text3(data, filename):
    ticket_numbers = data.get('ticket_numbers', [])
    if not ticket_numbers:
        return b''  # Return an empty bytes object if ticket_numbers is empty
    
    # Create text content with each key-value pair separated by commas for all ticket numbers
    text_content = ''
    selected_fields = ['ACCTNO','TXTYPE','TXDATE', 'TRNFLOWTYPE', 'TXAMT','TXFRCURRCD','TXNNO','TRNFACCTNO']
    for ticket_data in ticket_numbers:
        text_content += ','.join([f"{field}: {ticket_data.get(field, '')}" for field in selected_fields])
        text_content += '\n'  # Add new line between ticket details
    
    # Encode the text content to bytes
    text_bytes = text_content.encode('utf-8')
    
    # Write text content to file
    with open(filename, 'wb') as text_file:
        text_file.write(text_bytes)
    
    return text_bytes

def create_word3(data, filename):
    ticket_numbers = data.get('ticket_numbers', [])
    if not ticket_numbers:
        return b''  # Return an empty bytes object if ticket_numbers is empty
    
    # Create a new Document object
    doc = Document()
    
    # Add a table with two columns for key-value pairs
    table = doc.add_table(rows=1, cols=8)
    
    # Add table header
    hdr_cells = table.rows[0].cells
    header = ['ACCTNO','TXTYPE','TXDATE', 'TRNFLOWTYPE', 'TXAMT','TXFRCURRCD','TXNNO','TRNFACCTNO']
    for i, field in enumerate(header):
        hdr_cells[i].text = field
    
    # Iterate over each ticket and add its key-value pairs to the table
    for ticket in ticket_numbers:
        row_cells = table.add_row().cells
        for i, field in enumerate(header):
            row_cells[i].text = ticket.get(field, '')
        
        # Add an empty row between tickets
        table.add_row().cells
    
    # Save the document
    doc.save(filename)





# After login ROS can see his profile newly added

#--------------------------------------------------------------------------------------------

#--------------------------------------------------------------------------------------------
#--------------------------------------------------------------------------------------------


# ==================== OFFLINE FLOW ENDS  HERE USING MS SQL ==================================


# ============== Downloading all STR Starts ===============
    
# ===================== table data download starts==============
def flatten_dict(d):
    flattened = {}
    for key, value in d.items():
        if isinstance(value, dict):
            flattened.update(flatten_dict(value))
        else:
            flattened[key] = value
    return flattened

@app.route('/activeCases', methods=['POST', 'GET'])
def activeCases():
    data = request.get_json()
    
    ticket_numbers = data.get('ticket_numbers', [])

    # Create a temporary directory
    temp_dir = tempfile.mkdtemp()

    # Create the individual files
    pdf_filename = os.path.join(temp_dir, 'AlltableData.pdf')
    pdf_content = create_pdf(data, pdf_filename)

    csv_filename = os.path.join(temp_dir, 'AlltableData.csv')
    csv_content = create_csv(data, csv_filename)

    text_filename = os.path.join(temp_dir, 'AlltableData.txt')
    text_content = create_text(data, text_filename)

    word_filename = os.path.join(temp_dir, 'AlltableData.docx')
    create_word(data, word_filename)

    # Create a zip file
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.write(pdf_filename, arcname='AlltableData.pdf')
        zip_file.write(csv_filename, arcname='AlltableData.csv')
        zip_file.write(text_filename, arcname='AlltableData.txt')
        zip_file.write(word_filename, arcname='AlltableData.docx')

    # Close and delete the temporary directory
    shutil.rmtree(temp_dir)

    # Move the pointer to the beginning of the buffer
    zip_buffer.seek(0)

    # Send the zip file as an attachment
    return send_file(zip_buffer, as_attachment=True, mimetype='application/zip', download_name='AlltableData.zip')

def create_pdf(data, filename):
       
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    list_of_dicts= data['ticket_numbers']
    # Create table data
    table_data = []
    for obj in list_of_dicts:
        for key, value in obj.items():
            wrapper = textwrap.TextWrapper(width=70)
            wrapped_lines = wrapper.wrap(text=str(value))
            wrapped_text = "\n".join(wrapped_lines)
            table_data.append([key, wrapped_text])

    # Create table
    table = Table(table_data, colWidths=[200, 400],  # Adjust the column widths as needed
                  style=[
                      ('GRID', (0, 0), (-1, -1), 1, colors.black),  # Add grid lines
                      ('BACKGROUND', (0, 0), (-1, 0), colors.white),  # Add background color to header row
                      ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),  # Set text color in header row
                      ('ALIGN', (0, 0), (-1, -1), 'LEFT'),  # Center align all cells
                      ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Middle align all cells
                      ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),  # Inner grid lines
                      ('BOX', (0, 0), (-1, -1), 0.25, colors.black),  # Add box around the table
                  ])

    # Build the story

    all_names = ["MLRO", "CM", "GM/PO \n (Approved)"]
    names_table_data = [[all_names[0], all_names[1], all_names[2]]]
    names_table = Table(names_table_data, colWidths=[180, 180, 180],
                        style=[
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                            ('ALIGN', (2, 0), (-1, -1), 'RIGHT')

                        ])

    spacer = Spacer(1, 70)

    doc.build([table, spacer, names_table])
    buffer.seek(0)
    with open(filename, 'wb') as f:
        f.write(buffer.read())

def create_csv(data, filename):
    df = pd.DataFrame(data['ticket_numbers'])
    
    # Write DataFrame to CSV file
    df.to_csv(filename, index=False)

    # Alternatively, if you want to return CSV content as bytes
    csv_buffer = StringIO()  # Create a StringIO object to hold CSV content
    df.to_csv(csv_buffer, index=False)
    
    # Get the CSV content as bytes
    csv_content = csv_buffer.getvalue().encode('utf-8')
    return csv_content

def create_text(data, filename):
    ticket_numbers = data.get('ticket_numbers', [])
    if not ticket_numbers:
        return b''  # Return an empty bytes object if ticket_numbers is empty
    
    # Create text content with each key-value pair separated by commas for all ticket numbers
    text_content = ''
    for ticket_data in ticket_numbers:
        text_content += ','.join([f"{key}: {value}" for key, value in ticket_data.items()])
        text_content += '\n'  # Add new line between ticket details
    
    # Encode the text content to bytes
    text_bytes = text_content.encode('utf-8')
    
    # Write text content to file
    with open(filename, 'wb') as text_file:
        text_file.write(text_bytes)
    
    return text_bytes

def create_word(data, filename):
    ticket_numbers = data.get('ticket_numbers', [])
    if not ticket_numbers:
        return b''  # Return an empty bytes object if ticket_numbers is empty
    
    # Create a new Document object
    doc = Document()
    
    # Add a table with two columns for key-value pairs
    table = doc.add_table(rows=1, cols=2)
    
    # Iterate over each ticket and add its key-value pairs to the table
    for ticket in ticket_numbers:
        for key, value in ticket.items():
            # Add key-value pairs to the table
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
        
        # Add an empty row between tickets
        table.add_row().cells
    
    # Save the document
    doc.save(filename)





# =====================   RATING ENDS =================================================================

@app.route('/Ratings', methods=['GET', 'POST'])
@secure_route(required_role='DGM/PO')
def Ratings():
    if 'email_id' not in session:
        return redirect(url_for('post_login'))
    
    email = session['email_id']


    # Define the SQL query to fetch distinct scenarios
    query = "SELECT DISTINCT scenario_code FROM scenarios;"

    # Execute the query
    with mysql2.connection.cursor() as cur:
        cur.execute(query)

        # Fetch all distinct scenarios
        distinct_scenarios = [row[0] for row in cur.fetchall()]

        # Initialize a list to store scenario data
        scenario_data = []
        
        #Query to get the maximum count of alerts for any scenario on the fixed date
        max_count_query = """
        SELECT MAX(alert_count) FROM (
            SELECT COUNT(*) as alert_count 
            FROM scenarios 
            GROUP BY scenario_code
        ) as counts;
        """

        # Execute the query to get the max count
        cur.execute(max_count_query, ())
        max_count = cur.fetchone()[0] or 1  # Default to 1 to avoid division by zero
        print("max_count", max_count)

        
        for scenario_code in distinct_scenarios:
            # Construct the query to count alerts for the current scenario on the current date
            count_query = "SELECT COUNT(*) FROM scenarios WHERE scenario_code = ?;"

            # Execute the count query with the current scenario code
            cur.execute(count_query, (str(scenario_code)))

            # Fetch the count result
            count = cur.fetchone()[0]

            # Calculate rating
#             # rating = min((count / 5), 4.9)
            rating = min((count / max_count) * 5, 4.9)
            rounded_rating = round(rating, 1)

            # Append scenario data to the list
            scenario_data.append({
                'scenario_code': str(scenario_code),  # Ensure scenario_code is converted to string
                'total_count': count,
                'rating': rounded_rating
            })

        sorted_scenario_data = sorted(scenario_data, key=lambda x: x['rating'], reverse=True)


       
        return render_template('Ratings.html', role='DGM/PO', data=sorted_scenario_data)
    
# =====================   RATING ENDS =================================================================







# # =====================   KAMAL CODE STARTS =================================================================

def extract_letters(text):
    # Define a Unicode regex pattern that includes Unicode letters, numbers, combining characters, and specified punctuation
    pattern = rx.compile(r'[^\p{L}\p{N}\p{M}&\-.(),]', re.UNICODE)
    return pattern.sub('', text)

main_docs_list=[]

def docs_matched(documents, keyword_names_list, rangeInputIndex, keyword_dob,table_name):
    docs_list = []
    for doc in documents:
        if all(keyword.lower() in doc['name'].lower() for keyword in keyword_names_list):
            similarity_ratio = fuzz.WRatio(" ".join(keyword_names_list), doc['name'].lower())
            if keyword_dob:
                document_dob_str = doc.get("dob")
                normalized_dob_str = normalize_date(document_dob_str)
                if normalized_dob_str is not None:
                    if normalized_dob_str in keyword_dob:
                        if similarity_ratio >= int(rangeInputIndex):
                            doc['fuzzy_value'] = similarity_ratio
                            new_list = [name.capitalize() for name in keyword_names_list]
                            doc['key_name'] = " ".join(new_list)
                            doc['table_name']=table_name
                            docs_list.append(doc)
            else:
                if similarity_ratio >= int(rangeInputIndex):
                    doc['fuzzy_value'] = similarity_ratio
                    new_list = [name.capitalize() for name in keyword_names_list]
                    doc['key_name'] = " ".join(new_list)
                    doc['table_name']=table_name
                    docs_list.append(doc)
    return docs_list

MONTHS = {
    'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
    'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
}

def extract_unique_values(docs_list, key):
    seen = set()
    unique_list = []
    for doc in docs_list:
        value = doc[key]
        if value not in seen:
            seen.add(value)
            unique_list.append(doc)
    return unique_list

def normalize_date(date_str):
    if isinstance(date_str, str):
        date_str = re.sub(r'[/\-.]', '-', date_str)
        date_parts = date_str.split('-')
        if len(date_parts) == 3 and len(date_parts[0]) == 4:
            return '-'.join(date_parts)
        elif len(date_parts) == 3 and len(date_parts[-1]) == 4:
            return '-'.join(date_parts[::-1])
        elif len(date_parts) == 3 and len(date_parts[-2]) == 4:
            return '-'.join([date_parts[-1], date_parts[0], date_parts[1]])
        elif len(date_parts) == 3 and any(part[:3].lower() in MONTHS for part in date_parts):
            month_number = str(MONTHS[date_parts[1][:3].lower()])
            return '-'.join([date_parts[-1], month_number, date_parts[0]])
        else:
            return None
    elif isinstance(date_str, list):
        if len(date_str) == 3:
            if any(len(part) == 4 for part in date_str):
                return '-'.join(date_str)
            else:
                return None
        else:
            return None
    else:
        return None

def get_table_columns(table_name):
    """Helper function to get column names for a given table."""
    try:
        columns_query = f"SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = '{table_name}'"
        columns = conn_kamal.execute(columns_query).fetchall()
        return [column[0] for column in columns]
    except Exception as e:
        print(f"Error fetching columns for table {table_name}: {e}")
        return []

sub_cats = defaultdict(set)
categories_sl_wl = defaultdict(set)
cats_cols_wl_sl = defaultdict(set)

for table_info in kamal_tables:
    table_name = table_info[0]
    if table_name in ['AdverseMedia', 'PEP']:
        continue
    
    category_query = f"SELECT DISTINCT category FROM {table_name}"
    sub_category_query = f"SELECT DISTINCT sub_category FROM {table_name}"
    try:
        category_result = conn_kamal.execute(category_query).fetchall()
        sub_category_result = conn_kamal.execute(sub_category_query).fetchall()

        for category_row in category_result:
            category = category_row[0]
            sub_cats[category].update([row[0] for row in sub_category_result])
            cats_cols_wl_sl[category].update([table_name])
            categories_sl_wl[table_name].update([row[0] for row in sub_category_result])
    except Exception as e:
        print(f"Error Fetching Categories and sub-Categories from {table_name} : {e}")


def parse_datetime(date_str):
    """Tries to parse a datetime string with multiple formats."""
    for fmt in ("%Y-%m-%dT%H:%M:%S.%fZ", "%Y-%m-%dT%H:%M:%S.%f", "%Y-%m-%d %H:%M:%S.%f", "%Y-%m-%d %H:%M:%S.%f"):
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    return None


main_docs_list=[]
@app.route("/sdn", methods=["GET", "POST"])
@secure_route(required_role=['MLRO', 'ROS', 'BranchMakers'])
def sdn():
    role = session.get('user_role')
    print(role,"role")
    if role not in ['MLRO','ROS','BranchMakers']:
        return "Access denied", 403
    msg = None
    if 'upload_success' in session:
        msg = session.pop('upload_success','')
    elif 'upload_error' in session:
        msg = session.pop('upload_error','')
    elif 'upload_error_format' in session:
        msg = session.pop('upload_error_format', '')
    elif 'upload_error_reading' in session:
        msg = session.pop('upload_error_reading', '')
    elif 'upload_error_empty' in session:
        msg=session.pop('upload_error_empty','')

    global main_docs_list
    if request.method == "POST":
        main_docs_list = []
        try:
            keyword_name = request.form.get("name")
            rangeInputIndex = request.form.get('rangeInputIndex')
            id_type = request.form.get('id_type')
            entity_select = request.form.get("entitySelect")
            pob = request.form.get('pob')
            address = request.form.get('address')
            nationality = request.form.get('nationality')

            category_values = request.form.getlist('category')
            keyword_dob = request.form.get('dob')
            
            base_condition = "1=1"
            base_params = []
            if id_type:
                id_number = request.form.get('id_number')
            if keyword_name:
                keyword_names = re.split(r'[,\s]+', keyword_name.strip().lower())
                keyword_names_list = [extract_letters(name) for name in keyword_names]
                print("keyword_names_list:",keyword_names_list)
                base_condition += " AND ("
                for name in keyword_names:
                    base_condition += "LOWER(name) LIKE ? OR "
                    base_params.append(f"%{name}%")
                base_condition = base_condition[:-4]
                base_condition += ")"
                
                if entity_select:
                    base_condition += " AND entity_type = ?"
                    base_params.append(entity_select)
                    
                if id_type and id_number:
                    base_condition += f" AND {id_type.lower()} = ?"
                    base_params.append(id_number)
                    
                if pob:
                    pob_terms = pob.strip().lower().split()
                    pob_where_clause = " AND ("
                    for term in pob_terms:
                        pob_where_clause += "LOWER(pob) LIKE ? OR "
                        base_params.append(f"%{term}%")
                    pob_where_clause = pob_where_clause[:-4]
                    pob_where_clause += ")"
                    base_condition += pob_where_clause


                if address:
                    address_terms = address.strip().lower().split()
                    address_where_clause = " AND ("
                    for term in address_terms:
                        if term.isdigit() and len(term) == 6:
                            # If the term consists of 6 digits, assume it's a postal code
                            address_where_clause += "address LIKE ? OR "
                            base_params.append(f"%{term}%")
                        else:
                            # If the term contains characters or does not have 6 digits, treat it as part of the address
                            address_where_clause += "LOWER(address) LIKE ? OR "
                            base_params.append(f"%{term}%")
                    address_where_clause = address_where_clause[:-4]  # Remove the last " OR "
                    address_where_clause += ")"
                    base_condition += address_where_clause
                if nationality:
                    nationality_terms = nationality.strip().lower().split()
                    nationality_where_clause = " AND ("
                    for term in nationality_terms:
                        nationality_where_clause += "LOWER(nationality) LIKE ? OR "
                        base_params.append(f"%{term}%")
                    nationality_where_clause = nationality_where_clause[:-4]
                    nationality_where_clause += ")"
                    base_condition += nationality_where_clause
                

                
                if len(category_values) != 0:
                    read_cat_values=[]
                    for category in category_values:
                        adjusted_base_condition = base_condition 
                        if category == 'Adverse_Media':
                            columns = get_table_columns('AdverseMedia')
                            query = f"SELECT * FROM AdverseMedia WHERE {base_condition}"
                            if 'dob' not in columns:
                                adjusted_base_condition = base_condition.replace(' AND dob = ?', '')
                                adjusted_base_params = [param for param in base_params if param != keyword_dob]
                            else:
                                adjusted_base_condition = base_condition
                                adjusted_base_params = base_params

                            try:
                                cursor = conn_kamal.cursor()
                                cursor.execute(query, adjusted_base_params)
                                columns = [column[0] for column in cursor.description]
                                data = [dict(zip(columns, row)) for row in cursor.fetchall()]
                                matched_docs = docs_matched(data, keyword_names_list, rangeInputIndex, keyword_dob,'AdverseMedia')
                                main_docs_list.extend(matched_docs)
                            except Exception as e:
                                print(f"Error executing query on Adverse Media: {e}")
                    
                        elif category == 'PEP':
                            columns = get_table_columns('PEP')
                            query = f"SELECT * FROM PEP WHERE {base_condition}"
                            if 'dob' not in columns:
                                adjusted_base_condition = base_condition.replace(' AND dob = ?', '')
                                adjusted_base_params = [param for param in base_params if param != keyword_dob]
                            else:
                                adjusted_base_condition = base_condition
                                adjusted_base_params = base_params

                            try:
                                cursor = conn_kamal.cursor()
                                cursor.execute(query, adjusted_base_params)
                                columns = [column[0] for column in cursor.description]
                                data = [dict(zip(columns, row)) for row in cursor.fetchall()]
                                matched_docs = docs_matched(data, keyword_names_list, rangeInputIndex, keyword_dob,category)
                                main_docs_list.extend(matched_docs)
                            except Exception as e:
                                print(f"Error executing query on PEP : {e}")
                        
                        else:
                            for main_cat in list(sub_cats.keys()):
                                if main_cat in category_values:
                                    selected_subcategories = list(set(category_values) & set(sub_cats[main_cat]))
                                    if len(selected_subcategories) != 0:
                                        for val in selected_subcategories:
                                            if val not in read_cat_values:
                                                read_cat_values.append(val)
                                                collections_list = cats_cols_wl_sl[main_cat]
                                                for col in collections_list:
                                                    columns = get_table_columns(col)
                                                    query = f"SELECT * FROM {col} WHERE category = ? AND sub_category = ? AND {adjusted_base_condition}"
                                                    if 'dob' not in columns:
                                                        adjusted_base_condition = base_condition.replace(' AND dob = ?', '')
                                                        adjusted_base_params = [param for param in base_params if param != keyword_dob]
                                                    else:
                                                        adjusted_base_condition = base_condition
                                                        adjusted_base_params = base_params

                                                    try:
                                                        cursor = conn_kamal.cursor()
                                                        cursor.execute(query, (main_cat, val, *adjusted_base_params))
                                                        columns = [column[0] for column in cursor.description]
                                                        data = [dict(zip(columns, row)) for row in cursor.fetchall()]
                                                        matched_docs = docs_matched(data, keyword_names_list, rangeInputIndex, keyword_dob,col)
                                                        main_docs_list.extend(matched_docs)
                                                    except Exception as e:
                                                        print(f"Error executing query on {col}: {e}")
                                    else:
                                        collections_list = cats_cols_wl_sl[main_cat]
                                        for col in collections_list:
                                            columns = get_table_columns(col)
                                            query = f"SELECT * FROM {col} WHERE category = ? AND {adjusted_base_condition}"
                                            if 'dob' not in columns:
                                                adjusted_base_condition = base_condition.replace(' AND dob = ?', '')
                                                adjusted_base_params = [param for param in base_params if param != keyword_dob]
                                            else:
                                                adjusted_base_condition = base_condition
                                                adjusted_base_params = base_params

                                            try:
                                                cursor = conn_kamal.cursor()
                                                cursor.execute(query, (main_cat, *adjusted_base_params))
                                                columns = [column[0] for column in cursor.description]
                                                data = [dict(zip(columns, row)) for row in cursor.fetchall()]
                                                matched_docs = docs_matched(data, keyword_names_list, rangeInputIndex, keyword_dob,col)
                                                main_docs_list.extend(matched_docs)
                                            except Exception as e:
                                                print(f"Error executing query on {col}: {e}")
                
                else:
                    for table in kamal_tables:
                        table_name = table[0]
                        columns = get_table_columns(table_name)
                        query = f"SELECT * FROM {table_name} WHERE {base_condition}"
                        if 'dob' not in columns:
                            adjusted_base_condition = base_condition.replace(' AND dob = ?', '')
                            adjusted_base_params = [param for param in base_params if param != keyword_dob]
                        else:
                            adjusted_base_condition = base_condition
                            adjusted_base_params = base_params

                        try:
                            cursor = conn_kamal.cursor()
                            cursor.execute(query, adjusted_base_params)
                            columns = [column[0] for column in cursor.description]
                            data = [dict(zip(columns, row)) for row in cursor.fetchall()]
                            matched_docs = docs_matched(data, keyword_names_list, rangeInputIndex, keyword_dob,table_name)
                            main_docs_list.extend(matched_docs)
                        except Exception as e:
                            print(f"Error executing query on {table_name}: {e}")
            else:
                session['no_input'] = "No input provided."
                        
            if not main_docs_list:
                session["no_records"] = "No records found."
            else:
                session['search_performed'] = True  
                    
        except Exception as e:
            print(f"An error occurred: {e}")
        return redirect(url_for("sdn"))
    else:
        
        if role in session:
            print(role,"12345678")
        if 'upload_success' in session:
            msg = session.pop('upload_success','')
        elif 'upload_error' in session:
            msg = session.pop('upload_error','')
        if 'no_input' in session:
            msg = session.pop('no_input', '')
        if 'no_records' in session:
            msg = session.pop('no_records', '')
        elif 'upload_error_format' in session:
            msg = session.pop('upload_error_format', '')
        elif 'upload_error_reading' in session:
            msg = session.pop('upload_error_reading', '')
        elif 'upload_error_empty' in session:
            msg=session.pop('upload_error_empty','')

        search_performed = session.pop('search_performed', False)  
        if not search_performed:
            main_docs_list = []  


        if len(main_docs_list) != 0:
            for doc in main_docs_list:
                created = doc.get('created')
                doc['original_created'] = created 

                
                if isinstance(created, str):
                    if created.lower() not in ['none', 'null', 'nan']:
                        parsed_date = parse_custom_datetime(created)
                        if parsed_date is not None:
                            doc['created'] = parsed_date
                        else:
                            print(f"Error converting 'created' field to datetime: Invalid format for {created}")
                    else:
                        doc['created'] = None  
                elif created is None:
                    doc['created'] = None  

            main_docs_list = [
                doc for doc in main_docs_list if doc['created'] is None or isinstance(doc['created'], datetime)
            ]

            for doc in main_docs_list:
                if doc['created'] is None:
                    doc['created'] = datetime.min

            main_docs_list.sort(key=operator.itemgetter('fuzzy_value', 'created'), reverse=True)
            print('======SDN SEARCHED MATCH DATA =======',main_docs_list)
            conn_kamal.execute("""
            IF NOT EXISTS (SELECT * FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'SDN_Searched_Data')
            BEGIN
                CREATE TABLE SDN_Searched_Data (
                    [Index] INT,
                    _id NVARCHAR(250),
                    category NVARCHAR(250),
                    sub_category NVARCHAR(250),
                    created DATETIME,
                    modified NVARCHAR(250),
                    entity_type NVARCHAR(250),
                    country NVARCHAR(250),
                    source NVARCHAR(MAX),
                    name NVARCHAR(MAX),
                    first_name NVARCHAR(250),
                    last_name NVARCHAR(250),
                    title NVARCHAR(250),
                    primary_name NVARCHAR(250),
                    alias NVARCHAR(250),
                    dob NVARCHAR(250),
                    pob NVARCHAR(250),
                    gender NVARCHAR(50),
                    nationality NVARCHAR(250),
                    position NVARCHAR(250),
                    address NVARCHAR(250),
                    photo NVARCHAR(250),
                    remarks NVARCHAR(250),
                    contact_number NVARCHAR(250),
                    email NVARCHAR(250),
                    passport NVARCHAR(250),
                    pan NVARCHAR(250),
                    cin NVARCHAR(250),
                    din NVARCHAR(250),
                    linked_to NVARCHAR(250),
                    description NVARCHAR(MAX),
                    published_date DATETIME,
                    domain_name NVARCHAR(250),
                    news_title NVARCHAR(MAX),
                    title_translated NVARCHAR(MAX),
                    summary NVARCHAR(MAX),
                    text NVARCHAR(MAX),
                    text_translated NVARCHAR(MAX),
                    person NVARCHAR(MAX),
                    entity_sentiment NVARCHAR(MAX),
                    fuzzy_value INT,
                    key_name NVARCHAR(250),
                    table_name NVARCHAR(250),
                    original_created NVARCHAR(250)
                );
            END
            """)
            # conn_kamal.commit()

            # Insert data into the table
            try:
                for item in main_docs_list:
                    conn_kamal.execute("""
                    INSERT INTO SDN_Searched_Data (
                        [Index], _id, category, sub_category, created, modified, entity_type, country, source, name, first_name, last_name, title, primary_name, alias, dob, pob, gender, nationality, position, address, photo, remarks, contact_number, email, passport, pan, cin, din, linked_to, description, published_date, domain_name, news_title, title_translated, summary, text, text_translated, person, entity_sentiment, fuzzy_value, key_name, table_name, original_created
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        item["Index"], item["_id"], item["category"], item["sub_category"], item["created"], item["modified"], item["entity_type"], item["country"], item["source"], item["name"], item["first_name"], item["last_name"], item["title"], item["primary_name"], item["alias"], item["dob"], item["pob"], item["gender"], item["nationality"], item["position"], item["address"], item["photo"], item["remarks"], item["contact_number"], item["email"], item["passport"], item["pan"], item["cin"], item["din"], item["linked_to"], item["description"], item["published_date"], item["domain_name"], item["news_title"], item["title_translated"], item["summary"], item["text"], item["text_translated"], item["person"], item["entity_sentiment"], item["fuzzy_value"], item["key_name"], item["table_name"], item["original_created"]
                    ))
                conn_kamal.commit()
            except Exception as e:
                print("An error occurred:", e)
                # conn_kamal.rollback()  # Rollback the transaction if an error occurs
            finally:
                print('hiiii')
                # conn_kamal.close()  # Close the connection

            # conn_kamal.execute(category_query)

            for doc in main_docs_list:
                # print('==========SDN SEARCHED DATA DOCS=====',doc)
                doc['created'] = doc['original_created']
                del doc['original_created']
            print('=====SDN SEARCHED 1ST MATCH DATA=======',main_docs_list[0])
        return render_template("sdndashboard.html", data=main_docs_list, msg=msg, unique_entities=['Individual', 'Organization', 'Country', 'Vessel', 'Aircraft', 'Entity', 'N/A'], sub_cats=sub_cats, type='sdn', role=role)

main_docs_list=[]
@app.route('/upload', methods=['POST', 'GET'])
@secure_route(required_role=['IT OFFICER', 'MLRO', 'ROS', 'BranchMakers'])
def upload():
    role = session.get('user_role')
    print(role,"role")
    if role not in ['IT OFFICER', 'MLRO','ROS','BranchMakers']:
        return "Access denied", 403

    global main_docs_list
    if request.method == 'POST':
        main_docs_list = []  
        try:
            if 'file_upload' in request.files:
                uploaded_file = request.files['file_upload']
                if uploaded_file.filename != '' and uploaded_file:
                    session['uploaded_file'] = "Yes"
                    file_extension = uploaded_file.filename.split('.')[-1].lower()
                    
                    if file_extension in ['csv', 'xls', 'xlsx', 'txt']:
                        if file_extension == 'csv':
                            df = pd.read_csv(uploaded_file)
                        elif file_extension in ['xls', 'xlsx']:
                            df = pd.read_excel(uploaded_file)
                        elif file_extension == 'txt':
                            df = pd.read_csv(uploaded_file, delimiter=',', encoding='utf-8')

                        df = df.fillna("")  

                        if df.empty:
                            session['upload_error_empty'] = "Uploaded file is empty or does not contain valid data."

                        data_dict = df.to_dict(orient='records')

                        for value in data_dict:
                            rangeInputIndex = '80'
                            id_type = None
                            id_number = None
                            pob = None
                            address = None
                            dob = None
                            category_values = []
                            base_condition = "1=1"
                            base_params = []

                            # Construct base condition and parameters based on uploaded data
                            if 'Name/Keyword' in value and value['Name/Keyword'] != '':
                                keyword_name = value['Name/Keyword']
                                keyword_names = re.split(r'[,\s]+', keyword_name.strip().lower())
                                keyword_names_list = [extract_letters(name) for name in keyword_names]
                                base_condition += " AND ("
                                for name in keyword_names:
                                    base_condition += "LOWER(name) LIKE ? OR "
                                    base_params.append(f"%{name}%")
                                base_condition = base_condition[:-4]
                                base_condition += ")"

                                if 'Fuzzy Value' in value and value['Fuzzy Value'] != '':
                                    rangeInputIndex = value['Fuzzy Value']

                                if 'ID Type' in value and value['ID Type'].strip().lower() in ['passport', 'pan', 'din', 'cin']:
                                    id_type = value['ID Type']
                                    if 'ID Number' in value and value['ID Number'].strip():
                                        id_number = value['ID Number']
                                        base_condition += f" AND {id_type.lower()} = ?"
                                        base_params.append(id_number)

                                if 'Entity' in value and value['Entity'].strip():
                                    base_condition += " AND entity_type = ?"
                                    base_params.append(value['Entity'])

                                if 'Place of Birth' in value and value['Place of Birth'] != '':
                                    pob = value['Place of Birth']
                                    pob_terms = pob.strip().lower().split()
                                    pob_where_clause = " AND ("
                                    for term in pob_terms:
                                        pob_where_clause += "LOWER(pob) LIKE ? OR "
                                        base_params.append(f"%{term}%")
                                    pob_where_clause = pob_where_clause[:-4]
                                    pob_where_clause += ")"
                                    base_condition += pob_where_clause

                                if 'Address' in value and value['Address'] != '':
                                    address = value['Address']
                                    address_terms = address.strip().lower().split()
                                    address_condition = " AND ("
                                    for term in address_terms:
                                        address_condition += "LOWER(address) LIKE ? OR "
                                        base_params.append(f"%{term}%")
                                    base_condition += address_condition[:-4] + ")"

                                if 'Date of Birth' in value and value['Date of Birth'] != '':
                                    dob = value['Date of Birth']

                                for key in value.keys():
                                    if "Category" in key and "multiple categories search" in key:
                                        category_column = key

                                        if value[category_column].strip():
                                            category_list = value[category_column].split(',')
                                            for category in category_list:
                                                category = category.strip().lower()
                                                if 'sanction' in category:
                                                    category_values.append('Sanction List')
                                                elif 'watch' in category:
                                                    category_values.append('Watch List')
                                                elif 'adverse' in category or 'media' in category:
                                                    category_values.append('AdverseMedia')
                                                elif 'pep' in category:
                                                    category_values.append('PEP')

                                if category_values:
                                    for cat in category_values:
                                        if cat in ['PEP', 'AdverseMedia']:
                                            columns = get_table_columns(cat)
                                            if 'dob' not in columns:
                                                adjusted_base_condition = base_condition.replace(' AND dob = ?', '')
                                                adjusted_base_params = [param for param in base_params if param != dob]
                                            else:
                                                adjusted_base_condition = base_condition
                                                adjusted_base_params = base_params

                                            query = f"SELECT * FROM {cat} WHERE {adjusted_base_condition}"

                                            try:
                                                cursor = conn_kamal.cursor()
                                                cursor.execute(query, adjusted_base_params)
                                                columns = [column[0] for column in cursor.description]
                                                data = [dict(zip(columns, row)) for row in cursor.fetchall()]
                                                matched_docs = docs_matched(data, keyword_names_list, rangeInputIndex, dob, cat)
                                                main_docs_list.extend(matched_docs)
                                            except Exception as e:
                                                session['upload_error'] = f"Error executing query on {cat}: {e}"

                                        elif cat in ['Sanction List', 'Watch List']:
                                            for main_cat in list(sub_cats.keys()):
                                                if main_cat == cat:
                                                    collections_list = cats_cols_wl_sl[main_cat]
                                                    for col in collections_list:
                                                        columns = get_table_columns(col)
                                                        if 'dob' not in columns:
                                                            adjusted_base_condition = base_condition.replace(' AND dob = ?', '')
                                                            adjusted_base_params = [param for param in base_params if param != dob]
                                                        else:
                                                            adjusted_base_condition = base_condition
                                                            adjusted_base_params = base_params

                                                        query = f"SELECT * FROM {col} WHERE category = ? AND {adjusted_base_condition}"

                                                        try:
                                                            cursor = conn_kamal.cursor()
                                                            cursor.execute(query, (main_cat, *adjusted_base_params))
                                                            columns = [column[0] for column in cursor.description]
                                                            data = [dict(zip(columns, row)) for row in cursor.fetchall()]
                                                            matched_docs = docs_matched(data, keyword_names_list, rangeInputIndex, dob, col)
                                                            main_docs_list.extend(matched_docs)
                                                        except Exception as e:
                                                            session['upload_error'] = f"Error executing query on {col}: {e}"

                                else:
                                    for table in kamal_tables:
                                        table_name = table[0]
                                        columns = get_table_columns(table_name)
                                        if 'dob' not in columns:
                                            adjusted_base_condition = base_condition.replace(' AND dob = ?', '')
                                            adjusted_base_params = [param for param in base_params if param != dob]
                                        else:
                                            adjusted_base_condition = base_condition
                                            adjusted_base_params = base_params

                                        query = f"SELECT * FROM {table_name} WHERE {adjusted_base_condition}"

                                        try:
                                            cursor = conn_kamal.cursor()
                                            cursor.execute(query, adjusted_base_params)
                                            columns = [column[0] for column in cursor.description]
                                            data = [dict(zip(columns, row)) for row in cursor.fetchall()]
                                            matched_docs = docs_matched(data, keyword_names_list, rangeInputIndex, dob, table_name)
                                            main_docs_list.extend(matched_docs)
                                        except Exception as e:
                                            session['upload_error'] = f"Error executing query on {table_name}: {e}"

                        session['upload_success'] = "File uploaded and processed successfully"
                        session['search_performed'] = True 

                    else:
                        session['upload_error_format'] = "Unsupported file format"
                        print(f"Unsupported file format: {uploaded_file.filename}")
                else:
                    session['upload_error'] = "No file uploaded"
            else:
                session['upload_error'] = "No file uploaded"
        except Exception as e:
            session['upload_error_reading'] = f"Error reading file: {e}"
            print(f"Error reading file: {e}")
        return redirect(url_for('sdn'))


def query_for_search(keyword_names_list,entitySelect,id_type,id_number,pob):
    regex_pattern = '|'.join(map(re.escape, keyword_names_list))
    query = {"name": {"$regex": regex_pattern, "$options": "i"}}
    if entitySelect:
        query["entity_type"] = entitySelect
    if id_number:
        if id_type:
            if id_type.lower() in {"passport", "pan", "din", "cin"}:
                query[id_type.lower()] = id_number
        if id_type is None:
            for id in ["passport", "pan", "din", "cin"]:
                id_query = {id:id_number}
                if "$or" in query:
                    query["$or"].append(id_query)
                else:
                    query["$or"] = [id_query]
    if pob:
        pob_regex_pattern = '|'.join(map(re.escape, pob.strip().lower().split()))
        pob_query = {"pob": {"$regex": pob_regex_pattern, "$options": "i"}}
        if "$or" in query:
            query["$or"].append(pob_query)
        else:
            query["$or"] = [pob_query]
    return query

@app.route("/view/<table_name>/<unique_id>/<value>")
@secure_route(required_role=[ 'MLRO', 'ROS', 'BranchMakers'])
def view(table_name, unique_id, value):
    role = session.get('user_role')
    if role not in ['MLRO','ROS','BranchMakers']:
        return "Access denied", 403


    print("Entered the view function")

    query = f"SELECT * FROM {table_name} WHERE _id LIKE '{unique_id}'"
    data = []

    try:
        cursor = conn_kamal.cursor()

        cursor.execute(query)

        columns = [column[0] for column in cursor.description]

        data = [dict(zip(columns, row)) for row in cursor.fetchall()]
       
    except Exception as e:
        print("Error encountered:", e)
        
    for doc in data:
        print("Processing document:", doc)
        for key, val in doc.items():
            if isinstance(val, datetime):

                doc[key] = val.isoformat()
                print(f"Converted {key} to ISO format")

    print("Rendering template")
    return render_template("view.html", value=value, document=data, type='view', role=role)

@app.route("/sdndashboard")
@secure_route(required_role=['IT OFFICER', 'MLRO', 'ROS', 'BranchMakers'])
def sdndashboard():
    role = session.get('user_role')
    print(role,"role")
    if role not in ['IT OFFICER', 'MLRO','ROS','BranchMakers']:
        return "Access denied", 403  


    return render_template("sdndashboard.html", unique_entities=['Individual', 'Organization', 'Country', 'Vessel', 'Aircraft', 'Entity', 'N/A'], sub_cats=sub_cats,type='sdn' ,role=role)



def infer_sql_type(dtype):
    return 'NVARCHAR(MAX)'
def table_exists(cursor, table_name):
    check_table_query = "SELECT COUNT(*) FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = ?"
    cursor.execute(check_table_query, (table_name,))
    return cursor.fetchone()[0] > 0

def create_table(cursor, table_name, df):
    columns = df.columns
    column_defs = ", ".join([f"[{col}] {infer_sql_type(df[col].dtype)}" for col in columns])
    create_table_query = f"CREATE TABLE [{table_name}] ([Index] INT IDENTITY(1,1) PRIMARY KEY, {column_defs})"
    cursor.execute(create_table_query)
    index_query = f"CREATE INDEX idx_{table_name}_index ON [{table_name}] ([Index])"
    cursor.execute(index_query)

def add_columns_if_not_exist(cursor, table_name, df):
    existing_columns_query = f"SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = '{table_name}'"
    cursor.execute(existing_columns_query)
    existing_columns = set(row[0] for row in cursor.fetchall())
    for column in df.columns:
        if column not in existing_columns:
            alter_table_query = f"ALTER TABLE [{table_name}] ADD [{column}] {infer_sql_type(df[column].dtype)}"
            cursor.execute(alter_table_query)

def insert_or_update_data(cursor, table_name, df, batch_size=25000):
    unique_id_column = '_id'
    if unique_id_column not in df.columns:
        df.insert(0, unique_id_column, range(1, 1 + len(df)))
    columns = df.columns
    add_columns_if_not_exist(cursor, table_name, df)
    for batch_start in range(0, len(df), batch_size):
        batch_df = df.iloc[batch_start:batch_start + batch_size]
        unique_ids = tuple(batch_df[unique_id_column])
        check_query = f"SELECT [{unique_id_column}] FROM [{table_name}] WHERE [{unique_id_column}] IN ({', '.join(['?']*len(unique_ids))})"
        cursor.execute(check_query, unique_ids)
        existing_ids = set(row[0] for row in cursor.fetchall())
        insert_rows = []
        update_rows = []
        for _, row in batch_df.iterrows():
            row_values = [str(row[col]) for col in columns]
            if row[unique_id_column] in existing_ids:
                update_rows.append(row_values)
            else:
                insert_rows.append(row_values)
        if insert_rows:
            placeholders = ", ".join(['?'] * len(columns))
            insert_query = f"INSERT INTO [{table_name}] ({', '.join([f'[{col}]' for col in columns])}) VALUES ({placeholders})"
            cursor.executemany(insert_query, insert_rows)
            print(f"Inserted {len(insert_rows)} rows into {table_name}.")
        if update_rows:
            set_clause = ", ".join([f"[{col}] = ?" for col in columns if col != unique_id_column])
            update_query = f"UPDATE [{table_name}] SET {set_clause} WHERE [{unique_id_column}] = ?"
            update_data = []
            for row in update_rows:
                update_data.append(row[1:] + [row[0]])  
            cursor.executemany(update_query, update_data)
            print(f"Updated {len(update_rows)} rows in {table_name}.")

warnings.simplefilter("ignore", UnknownTimezoneWarning)
csv.field_size_limit(10**7)
default_datetime = datetime.now() - timedelta(days=365 * 10)

def parse_custom_datetime(date_str):
    date_formats = [
        '%b %d %Y %I:%M%p',            # Jun 24 2022 3:00PM
        '%Y-%m-%dT%H:%M:%S.%f',        # 2022-06-04T12:20:43.565648
        '%Y-%m-%dT%H:%M:%S.%fZ',       # 2017-12-28T11:17:50.000Z (UTC)
        '%Y-%m-%d %H:%M:%S.%f',        # 2021-12-20 08:29:15.666472
        '%Y-%m-%d %H:%M:%S',           # 2022-06-04 12:20:43
        '%Y-%m-%d',                    # 2022-06-04
        '%d-%b-%Y %H:%M:%S',           # 04-Jun-2022 12:20:43
        '%m/%d/%Y %H:%M:%S',           # 06/24/2022 15:00:00
        '%b %d, %Y %I:%M %p',
    ]
    for fmt in date_formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    return None



def truncate_or_default(value, max_length):
    if not value:
        return None
    return str(value)[:max_length]

def create_table_adverse(cursor, table_name, df, keys):
    column_defs = []
    for col in df.columns:
        column_defs.append(f"[{col}] {infer_sql_type(df[col].dtype)}")
    
    data_type_map = {
        'id': 'UNIQUEIDENTIFIER',
        'created': 'NVARCHAR(MAX)',
        'modified': 'NVARCHAR(MAX)',
        'published_date': 'NVARCHAR(MAX)',
        'entity_sentiment': 'NVARCHAR(MAX)',
        'keywords_p': 'NVARCHAR(MAX)',
        'text_translated': 'NVARCHAR(MAX)',
        'news_title': 'NVARCHAR(MAX)',
        'title_translated': 'NVARCHAR(MAX)',
    }
    
    for key in keys:
        if key not in df.columns: 
            if key in ['person', 'name']:
                column_type = 'NVARCHAR(MAX)'
            elif key == 'source':
                column_type = 'NVARCHAR(MAX)'
            elif key in ['summary', 'text', 'description', 'title_translated']:
                column_type = 'NVARCHAR(MAX)'
            else:
                column_type = data_type_map.get(key, 'NVARCHAR(255)')
            column_defs.append(f"[{key}] {column_type}")
    
    create_table_query = f"CREATE TABLE [{table_name}] ([Index] INT IDENTITY(1,1) PRIMARY KEY, {', '.join(column_defs)})"
    
    cursor.execute(create_table_query)


def process_dataframe(df, filename):
    df.drop_duplicates(inplace=True)
    for column in df.columns:
        if df[column].dtype == 'object':
            try:
                df[column] = pd.to_numeric(df[column])
            except ValueError:
                df[column] = df[column].astype(str)

    conn = get_db_connection(conn_str)
    if conn:
        cursor = conn.cursor()
        if 'category' in df.columns:
            df['category'] = df['category'].astype(str)
            categories = df['category'].unique()
            for category in categories:
                category_df = df[df['category'] == category]
                if category.lower().startswith("adverse"):
                    table_name = "AdverseMedia"
                    if not table_exists(cursor, table_name):
                        create_table_adverse(cursor, table_name, category_df)
                    insert_or_update_data(cursor, table_name, category_df)
                elif category.lower().startswith("pep"):
                    table_name = "PEP"
                    if not table_exists(cursor, table_name):
                        create_table(cursor, table_name, category_df)
                    insert_or_update_data(cursor, table_name, category_df)
                
                elif category.lower().startswith("sanction"):
                    sanctionlist_category_df = df[df['category'] == category]
                    if 'sub_category' in sanctionlist_category_df.columns and sanctionlist_category_df['sub_category'].notna().any():
                        sanctionlist_category_df['sub_category'] = sanctionlist_category_df['sub_category'].astype(str)
                        subcategories = sanctionlist_category_df['sub_category'].unique()
                        for subcategory in subcategories:
                            subcategory_df = sanctionlist_category_df[sanctionlist_category_df['sub_category'] == subcategory]
                            subcategory_table_name = ''
                            if subcategory.lower() == 'un sanction list':
                                subcategory_table_name = 'UN_Sanction_List'
                            elif subcategory.lower() == 'ofsi uk sanction list':
                                subcategory_table_name = 'UK_Sanction_List'
                            elif subcategory.lower() == 'ofac sdn criminal individuals list':
                                subcategory_table_name = 'OFAC_SDN'
                            elif subcategory.lower() == 'ofac':
                                subcategory_table_name = 'OFAC_NONSDN'
                            elif subcategory.lower() == 'mha uapa':
                                subcategory_table_name = 'MHA_UAPA'
                            elif subcategory.lower() == 'eu sanction list':
                                subcategory_table_name = 'EU_Sanction_List'
                            if not table_exists(cursor, subcategory_table_name):
                                create_table(cursor, subcategory_table_name, subcategory_df)
                            insert_or_update_data(cursor, subcategory_table_name, subcategory_df)
                    
                elif category.lower().startswith("watch"):
                    watchlist_category_df = df[df['category'] == category]
                    if 'sub_category' in watchlist_category_df.columns and watchlist_category_df['sub_category'].notna().any():
                        watchlist_category_df['sub_category'] = watchlist_category_df['sub_category'].astype(str)
                        subcategories = watchlist_category_df['sub_category'].unique()
                        for subcategory in subcategories:
                            subcategory_df = watchlist_category_df[watchlist_category_df['sub_category'] == subcategory]
                            subcategory_table_name = ''
                            if subcategory.lower() in ['nia arrested person', 'nia most wanted', 'absconding criminals']:
                                subcategory_table_name = 'Criminal_Activity_Watchlist'
                            elif subcategory.lower() in ['sfio conviction', 'sfio proclaimed', 'sfio prosecution']:
                                subcategory_table_name = 'SFIO_Watchlist'
                            elif subcategory.lower() in ['blacklisted doctors', 'blacklisted ngos']:
                                subcategory_table_name = 'Blacklisted'
                            elif subcategory.lower() in ['interpol', 'eocn uae']:
                                subcategory_table_name = 'International_Enforcement_Network'
                            elif subcategory.lower() in ['ncdex cessation members', 'ncdex defaulter members', 'ncdex expelled members']:
                                subcategory_table_name = 'NCDEX_Integrity_Watchlist'
                            elif subcategory.lower() in ['mca defaulter company', 'mca defaulter director', 'mca disqualified directors', 'mca dormant directors', 'mca mlm company', 'mca proclaimed offenders']:
                                subcategory_table_name = 'MCA_Integrity_Watchlist'
                            elif subcategory.lower() in ['cibil', 'fcra cancelled', 'fiu defaulter list', 'tax defaulters', 'world bank']:
                                subcategory_table_name = 'Financial_Integrity_Watchlist'
                            elif subcategory.lower() in ['cbi', 'cvc', 'fatf', 'irda', 'sebi']:
                                subcategory_table_name = 'Regulatory_Enforcement'
                            elif subcategory.lower() in ['bse defaulters expelled', 'bse delist companies', 'bse suspended companies']:
                                subcategory_table_name = 'BSE'
                            elif subcategory.lower() in ['mcx defaulters', 'msei defaulters', 'nbfc', 'nse','nse defaulter', 'nse expelled']:
                                subcategory_table_name = 'Market_Compliance'
                            if not table_exists(cursor, subcategory_table_name):
                                create_table(cursor, subcategory_table_name, subcategory_df)
                            insert_or_update_data(cursor, subcategory_table_name, subcategory_df)
                else:
                    if category.lower() in ['null', 'nan', 'none'] or not category.strip():
                        table_name = filename.split('.')[0]  
                        print(f"Using filename as table name: {table_name}")
                    else:
                        table_name = category.replace(" ", "_") 
                        print(f"Using category as table name: {table_name}")
                                        
                    if not table_exists(cursor, table_name):
                        create_table(cursor, table_name, category_df)
                    insert_or_update_data(cursor, table_name, category_df)
        
        else:
            table_name = filename.split('.')[0]
            if not table_exists(cursor, table_name):
                create_table(cursor, table_name, df)
            insert_or_update_data(cursor, table_name, df)
            
        conn.commit()
        cursor.close()
        print(f"File processed successfully.", "success")

def process_csv_file(df, filename):

    conn = get_db_connection(conn_str)
    if conn:
        cursor = conn.cursor()
        if 'category' in df.columns:
            df['category'] = df['category'].astype(str)
            categories = df['category'].unique()
            for category in categories:
                category_df = df[df['category'] == category]
                if category.lower().startswith("adverse"):
                    keys = category_df.columns.tolist()
                    batch = []
                    batch_size = 5000
                    record_count = 0
                    table_name = "AdverseMedia"
                    if not table_exists(cursor, table_name):
                        create_table_adverse(cursor, table_name, category_df, keys)
                        print("here")
                    add_columns_if_not_exist(cursor, table_name, category_df)
                    insert_query = f"INSERT INTO [{table_name}] ({','.join([f'[{key}]' for key in keys])}) VALUES ({','.join(['?' for _ in keys])})"
                    for index, row in category_df.iterrows():
                        record_count += 1
                        row_data = []
                        max_column_length = 10000
                        for key in keys:
                            if key in ['created', 'modified', 'published_date']:
                                dt_str = row.get(key, '')
                                row_data.append(parse_custom_datetime(dt_str) if dt_str else default_datetime)
                            elif key in ['summary', 'text', 'entity_sentiment', 'keywords_p', 'description', 'text_translated', 'name']:
                                row_data.append(truncate_or_default(row.get(key), 50000))
                            elif key in ['person', 'source']:
                                row_data.append(truncate_or_default(row.get(key), max_column_length))
                            else:
                                row_data.append(truncate_or_default(row.get(key), 255))
                        check_query = f"SELECT COUNT(*) FROM [{table_name}] WHERE [_id] = ?"
                        cursor.execute(check_query, (row['_id'],))
                        if cursor.fetchone()[0] == 0:
                            batch.append(tuple(row_data))
                            if len(batch) >= batch_size:
                                try:
                                    cursor.executemany(insert_query, batch)
                                    conn.commit()
                                except Exception as insert_error:
                                    print(f"Error during insertion: {insert_error}")
                                batch = []
                        else:
                            print(f"Row with unique ID {row['_id']} already exists in {table_name}.")
                    if batch:
                        try:
                            cursor.executemany(insert_query, batch)
                            conn.commit()
                            print("inserted record")
                        except Exception as final_insert_error:
                            print(f"Error inserting remaining records: {final_insert_error}")
                

        cursor.close()
        print(f"File processed successfully.")
    else:
        print("Database connection failed.")


@app.route('/SDN_USER', methods=['GET', 'POST'])
@secure_route(required_role='PINACA_ADMIN')
def SDN_USER():
    msg = None
    if 'upload_error' in session:
        msg = session.pop('upload_error', '')
    elif 'upload_success' in session:
        msg = session.pop('upload_success', '')
    if request.method == 'POST':
        files = request.files.getlist('file')
        if not files:
            session['upload_error'] = "No file selected."
            return redirect(request.url)
        for uploaded_file in files:
            if uploaded_file and uploaded_file.filename:
                filename = secure_filename(uploaded_file.filename)
                try:
                    rawdata = uploaded_file.read()
                    result = chardet.detect(rawdata)
                    encoding = result['encoding'] if result['encoding'] else 'utf-8'
                    uploaded_file.seek(0)
                    if filename.endswith('.csv'):
                        df = pd.read_csv(uploaded_file, encoding=encoding, low_memory=False)
                        if 'category' not in df.columns:
                            session['upload_error'] = f"category column not found in the CSV file: {filename}"
                            return redirect(request.url)

                        elif df['category'].dropna().empty:
                            session['upload_error'] = f"category column is empty in the CSV file: {filename}"
                            return redirect(request.url)

                        elif df['category'].str.lower().str.startswith('adverse').any():
                            process_csv_file(df, filename)
                        else:
                            process_dataframe(df, filename)
                    elif filename.endswith(('.xls', '.xlsx')):
                        xls = pd.ExcelFile(uploaded_file)
                        if not xls.sheet_names:
                            session['upload_error'] = f"No sheets found in Excel file: {filename}"
                            return redirect(request.url)
                        for sheet_name in xls.sheet_names:
                            df = pd.read_excel(xls, sheet_name=sheet_name)
                            process_dataframe(df, filename)
                    else:
                        session['upload_error'] = f"Unsupported file format: {filename}"
                        return redirect(request.url)
                except Exception as e:
                    session['upload_error'] = f"Failed to read file {uploaded_file.filename}: {e}"
                    return redirect(request.url)
        session['upload_success'] = "Files uploaded and processed successfully."
        return redirect(request.url)
    return render_template('SDN_User.html', role='PINACA_ADMIN', msg=msg)



cust_str = f"Driver={{SQL Server}};SERVER=Charan\\MSSQLSERVER04;Database=ticketid;Trusted_Connection=yes;MARS_Connection=yes"

def get_cust_db_connection():
    sql_conn = None
    try:
        sql_conn = pyodbc.connect(cust_str)
        print("Connection to SQL Server established successfully.")
    except Exception as e:
        print("Error connecting to SQL Server:", e)
    return sql_conn



cust_conn=get_cust_db_connection()

@app.route('/customer_sdn', methods=['GET', 'POST'])
@secure_route(required_role='IT OFFICER')
def customer_sdn():
    cust_conn = get_cust_db_connection()
    if cust_conn:
        try:
            cursor = cust_conn.cursor()
            select_query = "SELECT CUSTCD, CustomerName, SanctionsOutput FROM OutputTable"
            cursor.execute(select_query)
            rows = cursor.fetchall()

            customer_data = defaultdict(lambda: {
                'CUSTCD': '',
                'CustomerName': '',
                'SanctionsOutput': []
            })

            for row in rows:
                custcd = row[0]
                customer_name = row[1]
                sanctions_output = eval(row[2]) 

                customer_data[custcd]['CUSTCD'] = custcd
                customer_data[custcd]['CustomerName'] = customer_name
                customer_data[custcd]['SanctionsOutput'].append(sanctions_output)

            data = list(customer_data.values())
            return render_template("customer_sdn.html", role='IT OFFICER', data=data, message=session.pop('message', None))
        except Exception as e:
            print(f"Error fetching customer data: {e}")
            session['message'] = "Error fetching customer data."
            return render_template("customer_sdn.html", role='IT OFFICER', data=[], message=session.pop('message', None))
        
    else:
        session['message'] = "Error connecting to the customer database."
        return render_template("customer_sdn.html", role='IT OFFICER', data=[], message=session.pop('message', None))

@app.route('/view_sanctions/<int:customer_id>', methods=['GET'])
@secure_route(required_role='IT OFFICER')
def view_sanctions(customer_id):
    cust_conn = get_cust_db_connection()
    if cust_conn:
        try:
            cursor = cust_conn.cursor()
            select_query = "SELECT SanctionsOutput FROM OutputTable WHERE CUSTCD = ?"
            cursor.execute(select_query, (customer_id,))
            rows = cursor.fetchall()
            sanctions_output = []

            for row in rows:
                sanctions_output.extend(eval(row[0]))  

            session['message'] = f"Sanctions data fetched for customer {customer_id}."
            return render_template("display_customer.html", role='IT OFFICER', customer_id=customer_id, sanctions_output=sanctions_output, message=session.pop('message', None))
        except Exception as e:
            print(f"Error fetching sanctions output for customer {customer_id}: {e}")
            session['message'] = f"Error fetching sanctions output for customer {customer_id}."
            return render_template("display_customer.html", role='IT OFFICER', customer_id=customer_id, sanctions_output=[], message=session.pop('message', None))
    session['message'] = f"Error connecting to the customer database."
    return render_template("display_customer.html", role='IT OFFICER', customer_id=customer_id, sanctions_output=[], message=session.pop('message', None))

# @app.route('/view_sanctions/<int:customer_id>', methods=['GET'])
# @secure_route(required_role='IT OFFICER')
# def view_sanctions(customer_id):
#     cust_conn = get_cust_db_connection()
#     if cust_conn:
#         try:
#             cursor = cust_conn.cursor()
#             select_query = "SELECT SanctionsOutput FROM OutputTable WHERE CUSTCD = ?"
#             cursor.execute(select_query, (customer_id,))
#             rows = cursor.fetchall()
#             sanctions_output = []

#             for row in rows:
#                 sanctions_output.extend(eval(row[0]))  

#             return render_template("display_customer.html", role='IT OFFICER', customer_id=customer_id, sanctions_output=sanctions_output)
#         except Exception as e:
#             print(f"Error fetching sanctions output for customer {customer_id}: {e}")
#             return render_template("display_customer.html", role='IT OFFICER', customer_id=customer_id, sanctions_output=[])
#     return render_template("display_customer.html", role='IT OFFICER', customer_id=customer_id, sanctions_output=[])






# # =====================   KAMAL CODE ENDS =================================================================

 
    
schedule.every().day.at("00:27").do(verify_tenpercent_data)


# schedule.every(5).minutes.do(detltaADupdate,For_AD_USERNAME_update,For_AD_PASSWORD_update)


def run_closed_alerts_loop():
    while True:
        schedule.run_pending()
        time.sleep(1)


def handle_upload(sftp, filename, file_contents):
    remote_directory = 'D:/OGBDATA'
    remote_path = remote_directory + '/' + filename
    with sftp.file(remote_path, 'w') as file:
        file.write(file_contents)
    print(f"File '{filename}' uploaded successfully.")


def sftp_server():
    hostname = '10.40.16.190'
    port = 22
    username = 'ogbuser'
    password = 'Technology@2022'

    transport = paramiko.Transport((hostname, port))
    transport.connect(username=username, password=password)
    sftp = paramiko.SFTPClient.from_transport(transport)

    while True:
        print("Waiting for client to connect...")
        # Accept a connection from a client
        client_transport, client_address = transport.accept()
        print(f"Connection established with {client_address}")
        client_sftp = paramiko.SFTPClient.from_transport(client_transport)

        # Receive file from client
        for filename in client_sftp.listdir():
            with client_sftp.file(filename, 'r') as file:
                file_contents = file.read()
            handle_upload(sftp, filename, file_contents)



def verifyFilesExists():
    while True:
        files = os.listdir("D:/OGBDATA")
        # files = False

        # if files and len(files) == 2:
        if files:
            print("Files found, scheduling DataInsertionProcess in 20 minutes...")
            time.sleep(60)  
            DataInsertionProcess()
        else:
            print("No files found, checking again in 5 minutes...")
            # TM11()
            # print("Completed.....................")
            # TM12()
            # print("Completed.....................")
            # TM13()
            # print("Completed.....................")
            # TM14()
            # print("Completed.....................")
            # TM21()
            # print("Completed.....................")
            # TM22()
            # print("Completed.....................")
            # TM23()
            # print("Completed.....................")
            # TM24()
            # print("Completed.....................")
            # insert_into_tickets()
            # print("Inserted to Tickets.....................")
            # allocate()
            # print("Allocated..........")
            time.sleep(60)  
        

def DataInsertionProcess():
    files = os.listdir("D:/OGBDATA")

    if files:
        if "AML_TRANSACTIONS.lst" in files:
            # TRDBInstance = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
            TRDBInstance = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

            TRDBInstanceConn = TRDBInstance.cursor() 
            print("DB CONNECTED for Transcation DB" ,TRDBInstance)
            try:
                chunk_size = 1000000
                pandsImples = pd.read_csv("D:/OGBDATA/AML_TRANSACTIONS.lst",sep='|',dtype={"ACCTNO":str,'TRNFACCTNO':str,'CUSTOMERNO':str,'TRF_SOL_ID':str,'BENEFICIARYNAME':str,'BENEFIADDRESS':str,'SENDERNAME':str,'SENDERADDRESS':str},quoting=3, chunksize=chunk_size)

                for pandsImple in pandsImples:
                    pandsImple = pandsImple.astype(str)

                    pandsImple = pandsImple.replace('nan', None)


                    columns = pandsImple.columns


                    tableInfo = TRDBInstanceConn.execute("SELECT 1 FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'TRANSACTIONS'")
                    table_exists = tableInfo.fetchone()

                    if not table_exists:
                        column_definitions = ", ".join([f"{col} NVARCHAR(50)" if col in ['ACCTNO', 'TXTYPE', 'TRNFLOWTYPE', 'TRNFACCTNO'] else f"{col} NVARCHAR(MAX)" for col in columns])
                        create_table_query = f"CREATE TABLE TRANSACTIONS ({column_definitions})"
                        TRDBInstanceConn.execute(create_table_query)
                        TRDBInstanceConn.commit()



                    placeholders = ", ".join(["?"] * len(columns))
                    insert_query = f"INSERT INTO TRANSACTIONS ({', '.join(columns)}) VALUES ({placeholders})"
                


                    data = [tuple(row) for row in pandsImple.to_numpy()]


                    batch_size = 1000000

                    for i in range(0, len(data), batch_size):
                        batch = data[i:i+batch_size]
                        TRDBInstanceConn.executemany(insert_query, batch)
                        TRDBInstanceConn.commit()

             
            

                print("completed...........")
            
                date = datetime.now().strftime("%Y_%m_%d")
                os.rename('D:/OGBDATA/AML_TRANSACTIONS.lst',f'D:/OGBCBSHISTORY/AML_TRANSACTIONS_{date}.lst')
            except Exception as e:
                TRDBInstanceConn.rollback()
                print(f"Error processing AML_TRANSACTIONS.lst: {e}")
        if "AML_CRED.lst" in files:
            pandsImple = pd.read_csv("D:/OGBDATA/AML_CRED.lst",sep='|',dtype={"ACCTNO":str,'CUSTCD':str,'PINCODE':str,'TYPEOFOPRN':str,'Occupation':str,'RIP':str,'RLTLamount':str,'PRIMARY_SOL_ID':str,'Guardian_CUSTCD':str,'CUST_CONSTITUTION':str},quoting=3)

            pandsImple = pandsImple.astype(str)
            pandsImple = pandsImple.replace('nan', None)

            # CUSTInstance = pyodbc.connect(f"Driver={{SQL Server}};SERVER={serverIP2};Database=ticketid;UID={userSQL2};PWD={pwdSQL2}")
            CUSTInstance = pyodbc.connect("Driver={SQL Server};SERVER=Charan\\MSSQLSERVER04;Database=ticketid; MARS_Connection=YES")

            CUSTInstanceConn = CUSTInstance.cursor() 
            print("DB CONNECTED Customer Data db" ,CUSTInstance)


            columns = pandsImple.columns
            try:
                tableInfo = CUSTInstanceConn.execute("SELECT 1 FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'CUSTOMERS'")
                table_exists = tableInfo.fetchone()

                if not table_exists:
                    column_definitions = ", ".join([f"{col} NVARCHAR(50)" if col in ['ACCTNO', 'CUSTCD','ACCT_OPENDATE','CUSTID_OPENDATE'] else f"{col} NVARCHAR(MAX)" for col in columns])
                    create_table_query = f"CREATE TABLE CUSTOMERS ({column_definitions})"
                    CUSTInstanceConn.execute(create_table_query)
                    CUSTInstanceConn.commit()

                
                placeholders = ", ".join(["?"] * len(columns))
                insert_query = f"INSERT INTO CUSTOMERS ({', '.join(columns)}) VALUES ({placeholders})"

                update_query = f"UPDATE CUSTOMERS SET {', '.join([f'{col} = ?' for col in columns])} WHERE ACCTNO = ?"

                data = [tuple(row) for row in pandsImple.to_numpy()]

                batch_size = 100000

                for i in range(0, len(data), batch_size):
                    batch = data[i:i+batch_size]

                    CUSTInstanceConn.executemany(insert_query, batch)
                    # for row in batch:
                    #     CUSTInstanceConn.execute(update_query, row + (row[0],))  
                    CUSTInstanceConn.commit()


                # data = [tuple(row) for row in pandsImple.to_numpy()]


                # batch_size = 1000000;

                # for i in range(0, len(data), batch_size):
                #     batch = data[i:i+batch_size]
                #     CUSTInstanceConn.executemany(insert_query, batch)
                #     CUSTInstanceConn.commit()




                # for index, row in pandsImple.iterrows():
                #     updateExistingUser = CUSTInstanceConn.execute(f"SELECT ACCTNO FROM CUSTOMERS WHERE ACCTNO = ?", (row['ACCTNO'],)).fetchone()


                #     if updateExistingUser:
                #         set_clause = ", ".join([f"{col} = ?" for col in columns])
                #         update_query = f"UPDATE CUSTOMERS SET {set_clause} WHERE ACCTNO = ?"
                #         CUSTInstanceConn.execute(update_query, tuple(row) + (row['ACCTNO'],))
                #     else:
                #         placeholders = ", ".join(["?"] * len(columns))
                #         insert_query = f"INSERT INTO CUSTOMERS ({', '.join(columns)}) VALUES ({placeholders})"
                #         CUSTInstanceConn.execute(insert_query, tuple(row))
                #     CUSTInstanceConn.commit()
            

                print("completed...........")
            
                date = datetime.now().strftime("%Y_%m_%d")
                os.rename('D:/OGBDATA/AML_CRED.lst',f'D:/OGBCBSHISTORY/AML_CRED_{date}.lst')
            except Exception as e:
                CUSTInstanceConn.rollback()
                print(f"Error processing AML_TRANSACTIONS.lst: {e}")

        # TM11()
        # print("Completed.....................")
        # TM12()
        # print("Completed.....................")
        # TM13()
        # print("Completed.....................")
        # TM14()
        # print("Completed.....................")
        # TM21()
        # print("Completed.....................")
        # TM22()
        # print("Completed.....................")
        # TM23()
        # print("Completed.....................")
        # TM24()
        # print("Completed.....................")
        # insert_into_tickets()
        # print("Inserted to Tickets.....................")
        # allocate()
        # print("Allocated..........")

    verifyFilesExists()

if __name__ == "__main__":
    closed_while_thread = threading.Thread(target=run_closed_alerts_loop)
    closed_while_thread.start()
    app_thread = threading.Thread(target=app.run)
    while True:
        schedule.run_pending()
        time.sleep(1)
        app.run()






    

